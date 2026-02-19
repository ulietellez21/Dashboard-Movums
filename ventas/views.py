from django.shortcuts import render, redirect, get_object_or_404
from django.views.generic import (
    ListView, DetailView, CreateView, UpdateView, View, DeleteView, TemplateView
)
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User 
from django.template.loader import render_to_string
from django.urls import reverse_lazy, reverse
from django.db.models.functions import Coalesce
from django.db.models import Sum, Max, Count, F, Q, Value, IntegerField, ExpressionWrapper, Prefetch
from django.db.models import DecimalField as ModelDecimalField
from django.db import transaction
from django.contrib import messages
from django.core.paginator import Paginator
from django.http import HttpResponse, HttpResponseRedirect, JsonResponse
from django.conf import settings
from django.utils import timezone 
from django.utils import formats  # Para formatear fechas en el PDF
# IMPORTACIÓN CLAVE: Necesaria para generar slugs automáticamente
from django.utils.text import slugify 
from datetime import timedelta, date
from collections import Counter
import math, re, logging, secrets, json, io, os
import datetime # Necesario para el contexto del PDF (campo now)
from decimal import Decimal, InvalidOperation # Importar Decimal para asegurar precisión en cálculos financieros

logger = logging.getLogger(__name__)

# Intento cargar WeasyPrint; si falla (por dependencias GTK), defino placeholders.
try:
    from weasyprint import HTML, CSS 
    WEASYPRINT_AVAILABLE = True
except ImportError:
    logger.warning("WeasyPrint no está disponible. La generación de PDF fallará.")
    class HTML:
        def __init__(self, string, base_url=None): pass
        def write_pdf(self): return b''
    class CSS: pass
    WEASYPRINT_AVAILABLE = False


from .models import (
    VentaViaje,
    AbonoPago,
    Logistica,
    LogisticaServicio,
    ContratoGenerado,
    Notificacion,
    Proveedor,
    ConfirmacionVenta,
    Ejecutivo,
    Oficina,
    PlantillaConfirmacion,
    Cotizacion,
    VentaPromocionAplicada,
    AbonoProveedor,
    ComisionVenta,
    ComisionMensual,
    SolicitudCancelacion,
)
from crm.models import Cliente
from crm.services import KilometrosService
from usuarios.models import Perfil
from usuarios import permissions as perm
from usuarios import mixins as usuarios_mixins
from .services.cancelacion import CancelacionService
from .validators import validate_uploaded_file, safe_int
from django.core.exceptions import ValidationError as DjangoValidationError
from django.forms import modelformset_factory
from .forms import (
    VentaViajeForm,
    LogisticaForm,
    LogisticaServicioForm,
    LogisticaServicioFormSet,
    AbonoPagoForm,
    ProveedorForm,
    ConfirmacionVentaForm,
    EjecutivoForm,
    OficinaForm,
    CotizacionForm,
    SolicitarAbonoProveedorForm,
    ConfirmarAbonoProveedorForm,
    SolicitudCancelacionForm,
)
from .utils import numero_a_texto
from .services.logistica import (
    build_financial_summary,
    build_service_rows,
    build_logistica_card,
)

# Función auxiliar para obtener el rol (delega a la capa centralizada de permisos)
# NOTA: Usar perm.get_user_role(user, request) directamente en las vistas para aprovechar cache
def get_user_role(user, request=None):
    """Wrapper para compatibilidad. Usar perm.get_user_role(user, request) directamente."""
    return perm.get_user_role(user, request)


def _get_logistica_servicio_formset(venta, request_POST=None, queryset=None, prefix='servicios'):
    """Formset de servicios logísticos. Solo filas existentes (extra=0). Nuevas filas se añaden con POST 'añadir_servicio_logistica'."""
    FormSetClass = modelformset_factory(
        LogisticaServicio,
        form=LogisticaServicioForm,
        extra=0,
        can_delete=False,
    )
    qs = queryset if queryset is not None else venta.servicios_logisticos.all().order_by('orden', 'pk')
    kwargs = {'queryset': qs, 'prefix': prefix}
    if request_POST is not None:
        kwargs['data'] = request_POST
    return FormSetClass(**kwargs)

class ContratoGeneradoDetailView(LoginRequiredMixin, DetailView):
    """
    Muestra los detalles de un contrato generado. 
    Se asume que el modelo 'VentaViaje' representa este contrato.
    """
    # Usamos VentaViaje como modelo por defecto. Ajustar si tienes un modelo ContratoGenerado dedicado.
    model = VentaViaje 
    template_name = 'ventas/contrato_generado_detalle.html' 
    context_object_name = 'contrato'

    def get_queryset(self):
        """Filtra para que solo el usuario autenticado pueda ver sus propias ventas/contratos."""
        # Se asume que el modelo VentaViaje tiene un campo 'vendedor' relacionado al usuario.
        return self.model.objects.filter(vendedor=self.request.user)

# ------------------- 1. DASHBOARD GERENCIAL -------------------

class DashboardView(LoginRequiredMixin, ListView):
    model = VentaViaje
    template_name = 'dashboard.html'
    context_object_name = 'ventas'
    paginate_by = 50  # ESCALABILIDAD: Limitar resultados por página

    def get_queryset(self):
        user = self.request.user
        queryset = perm.get_ventas_queryset_base(VentaViaje, user, self.request)

        # Aplicar filtro por fecha de viaje (soporta fecha única o rango)
        fecha_filtro = self.request.GET.get('fecha_filtro')
        fecha_desde = self.request.GET.get('fecha_desde')
        fecha_hasta = self.request.GET.get('fecha_hasta')
        
        # Prioridad: rango de fechas > fecha única
        if fecha_desde and fecha_hasta:
            try:
                from datetime import datetime
                fecha_desde_obj = datetime.strptime(fecha_desde, '%Y-%m-%d').date()
                fecha_hasta_obj = datetime.strptime(fecha_hasta, '%Y-%m-%d').date()
                # Asegurar que desde <= hasta
                if fecha_desde_obj <= fecha_hasta_obj:
                    queryset = queryset.filter(fecha_inicio_viaje__range=[fecha_desde_obj, fecha_hasta_obj])
            except ValueError:
                pass
        elif fecha_filtro:
            try:
                from datetime import datetime
                fecha_obj = datetime.strptime(fecha_filtro, '%Y-%m-%d').date()
                queryset = queryset.filter(fecha_inicio_viaje=fecha_obj)
            except ValueError:
                pass

        return queryset.select_related('cliente', 'vendedor', 'proveedor').order_by('-fecha_inicio_viaje', '-fecha_creacion')

    def get_context_data(self, **kwargs):
        # Asegurar que object_list esté definido antes de llamar a super()
        if not hasattr(self, 'object_list'):
            self.object_list = self.get_queryset()
        context = super().get_context_data(**kwargs)
        user = self.request.user
        
        user_rol = perm.get_user_role(user, self.request)
        context['user_rol'] = user_rol
        context['puede_aprobar_rechazar_cancelacion'] = perm.can_approve_reject_cancelacion(user, self.request)
        
        # Inicializar notificaciones vacías por defecto (para evitar errores en template)
        context['notificaciones'] = Notificacion.objects.none()
        context['notificaciones_count'] = 0

        # --- Lógica de KPIs (roles con acceso total y CONTADOR; Gerente/Directores ven según su scope en listado) ---
        if perm.has_full_access(user, self.request) or perm.is_contador(user, self.request):
            # KPI 1: Saldo Pendiente Global
            total_vendido_agg = VentaViaje.objects.aggregate(Sum('costo_venta_final'))['costo_venta_final__sum']
            total_vendido = total_vendido_agg if total_vendido_agg is not None else Decimal('0.00')
            
            # Total pagado incluye abonos + montos de apertura
            total_abonos_agg = AbonoPago.objects.aggregate(Sum('monto'))['monto__sum']
            total_abonos = total_abonos_agg if total_abonos_agg is not None else Decimal('0.00')
            
            total_apertura_agg = VentaViaje.objects.aggregate(Sum('cantidad_apertura'))['cantidad_apertura__sum']
            total_apertura = total_apertura_agg if total_apertura_agg is not None else Decimal('0.00')
            
            total_pagado = total_abonos + total_apertura
            
            context['saldo_total_pendiente'] = total_vendido - total_pagado

            # KPI 2: Servicios Pendientes (Logística)
            # Filtrar solo ventas que tienen logística asociada
            context['alertas_logistica_count'] = VentaViaje.objects.filter(
                logistica__isnull=False
            ).filter(
                Q(logistica__vuelo_confirmado=False) |
                Q(logistica__hospedaje_reservado=False) |
                Q(logistica__traslado_confirmado=False) |
                Q(logistica__tickets_confirmado=False)
            ).count()
            
            # INNOVACIÓN 2: Ranking de Vendedores
            context['ranking_ventas'] = VentaViaje.objects.filter(
                vendedor__isnull=False,
            ).values('vendedor__username').annotate(
                num_ventas=Count('id'),
                total_vendido=Sum('costo_venta_final')
            ).order_by('-num_ventas', '-total_vendido')[:5]
            
            # INNOVACIÓN 3: Notificaciones para JEFE/Director General y para quien puede aprobar cancelaciones (incl. Director Administrativo)
            if perm.has_full_access(user, self.request) or perm.can_approve_reject_cancelacion(user, self.request):
                context['notificaciones'] = Notificacion.objects.filter(
                    usuario=user,
                    vista=False  # Solo mostrar notificaciones no vistas
                ).select_related('venta', 'venta__cliente', 'abono', 'solicitud_cancelacion').order_by('-fecha_creacion')[:30]  # Últimas 30 no vistas
                context['notificaciones_count'] = Notificacion.objects.filter(
                    usuario=user,
                    vista=False
                ).count()  # Contador solo de no vistas
                
                # Solicitudes de cancelación pendientes (JEFE, Director General, Director Administrativo)
                if perm.can_approve_reject_cancelacion(user, self.request):
                    solicitudes_pendientes = SolicitudCancelacion.objects.filter(
                        estado='PENDIENTE'
                    ).select_related(
                        'venta', 'venta__cliente', 'venta__vendedor', 'solicitado_por'
                    ).order_by('-fecha_solicitud')[:20]
                    context['solicitudes_cancelacion_pendientes'] = solicitudes_pendientes
                    context['solicitudes_cancelacion_count'] = solicitudes_pendientes.count()
            # --- Lógica para CONTADOR (dentro del bloque JEFE/CONTADOR) ---
            elif user_rol == 'CONTADOR':
                # Notificaciones para CONTADOR: mostrar solo las no vistas
                # IMPORTANTE: Filtrar por vista=False para que desaparezcan al marcarlas
                notificaciones_contador = Notificacion.objects.filter(
                    usuario=user,
                    vista=False  # Solo mostrar notificaciones no vistas
                ).select_related('venta', 'venta__cliente', 'abono').prefetch_related('abono__confirmado_por').order_by('-fecha_creacion')
                
                # Convertir a lista para asegurar que se evalúe el QuerySet
                notificaciones_list = list(notificaciones_contador[:20])
                context['notificaciones'] = notificaciones_list
                context['notificaciones_count'] = len(notificaciones_list)
                
                # Para las secciones adicionales: obtener solo las pendientes
                notificaciones_pendientes_pagos = notificaciones_contador.filter(
                    tipo='PAGO_PENDIENTE',
                    confirmado=False
                )
                
                # Ventas con pagos pendientes de confirmación para el CONTADOR
                # Obtener todas las notificaciones pendientes y sus ventas relacionadas
                ventas_pendientes_ids = notificaciones_pendientes_pagos.values_list('venta_id', flat=True).distinct()
                if ventas_pendientes_ids:
                    context['ventas_pendientes'] = VentaViaje.objects.filter(
                        pk__in=ventas_pendientes_ids
                    ).select_related('cliente', 'vendedor').prefetch_related('abonos').order_by('-fecha_creacion')[:10]
                else:
                    context['ventas_pendientes'] = VentaViaje.objects.none()
                
                # Abonos pendientes de confirmación (para mostrar en el dashboard)
                # IMPORTANTE: Debe coincidir con la lógica de PagosPorConfirmarView
                # Solo mostrar abonos con comprobante subido
                context['abonos_pendientes'] = AbonoPago.objects.filter(
                    Q(forma_pago__in=['TRN', 'TAR', 'DEP']) & 
                    Q(confirmado=False) &
                    Q(comprobante_subido=True)
                ).select_related('venta', 'venta__cliente', 'registrado_por').order_by('-fecha_pago')[:10]
                
                # Ventas con apertura pendiente (para mostrar en el dashboard)
                # IMPORTANTE: Debe coincidir con la lógica de PagosPorConfirmarView
                # Para TRN/TAR/DEP: requiere comprobante subido
                # Para CRE: no requiere comprobante, solo estar en EN_CONFIRMACION
                ventas_apertura_pendiente = VentaViaje.objects.filter(
                    Q(estado_confirmacion='EN_CONFIRMACION') &  # Solo las que están en confirmación
                    (
                        # Transferencia, Tarjeta, Depósito: requieren comprobante
                        # Para NAC: cantidad_apertura > 0; para INT: cantidad_apertura_usd > 0
                        (Q(modo_pago_apertura__in=['TRN', 'TAR', 'DEP']) & 
                         Q(comprobante_apertura_subido=True) &
                         (
                             Q(cantidad_apertura__gt=0) |  # Ventas nacionales
                             (Q(tipo_viaje='INT') & Q(cantidad_apertura_usd__gt=0))  # Ventas internacionales
                         )) |
                        # Crédito: no requiere comprobante ni cantidad_apertura > 0
                        Q(modo_pago_apertura='CRE')
                    )
                ).exclude(
                    estado_confirmacion='COMPLETADO'  # Excluir explícitamente las ya confirmadas
                ).select_related('cliente', 'vendedor').order_by('-fecha_creacion')[:10]
                
                # Convertir ventas a notificaciones para mantener compatibilidad con el template
                # El template espera notificaciones, pero ahora usamos directamente las ventas
                # Clase auxiliar para crear objetos similares a Notificacion
                class NotificacionTemporal:
                    def __init__(self, venta, mensaje):
                        self.venta = venta
                        self.mensaje = mensaje
                        self.fecha_creacion = venta.fecha_creacion
                
                notificaciones_apertura_list = []
                for venta in ventas_apertura_pendiente:
                    # Crear una notificación temporal para el template (no se guarda en BD)
                    modo_pago_display = dict(VentaViaje.MODO_PAGO_CHOICES).get(venta.modo_pago_apertura, venta.modo_pago_apertura)
                    mensaje = f"Pago de apertura pendiente de confirmación: ${venta.cantidad_apertura:,.2f} ({modo_pago_display}) - Venta #{venta.pk} - Cliente: {venta.cliente.nombre_completo_display}"
                    notificaciones_apertura_list.append(NotificacionTemporal(venta, mensaje))
                
                context['notificaciones_apertura_pendiente'] = notificaciones_apertura_list
                
                # Contador de ventas con estado "En confirmación"
                context['ventas_en_confirmacion_count'] = VentaViaje.objects.filter(
                    estado_confirmacion='EN_CONFIRMACION'
                ).count()
                
                # Abonos a proveedor pendientes de confirmar (para mostrar en el dashboard)
                abonos_proveedor_pendientes_dashboard = AbonoProveedor.objects.filter(
                    estado='APROBADO'
                ).select_related('venta', 'venta__cliente', 'solicitud_por', 'aprobado_por').order_by('-fecha_aprobacion')[:10]
                context['abonos_proveedor_pendientes'] = abonos_proveedor_pendientes_dashboard
                
                # Notificaciones de abonos a proveedor para el contador
                notificaciones_abonos_proveedor = Notificacion.objects.filter(
                    usuario=user,
                    tipo__in=['SOLICITUD_ABONO_PROVEEDOR', 'ABONO_PROVEEDOR_APROBADO'],
                    vista=False
                ).select_related('venta', 'abono_proveedor').order_by('-fecha_creacion')[:10]
                context['notificaciones_abonos_proveedor'] = notificaciones_abonos_proveedor
            
        # --- Lógica de KPIs (se mantiene para vendedores) ---
        elif user_rol == 'VENDEDOR':
            # Notificaciones para VENDEDOR: solo las notificaciones creadas específicamente para este vendedor
            # IMPORTANTE: Simplificar la consulta - si la notificación está asignada al usuario, no hace falta filtrar por ventas
            # Filtrar por vista=False para que desaparezcan al marcarlas
            notificaciones_vendedor = Notificacion.objects.filter(
                usuario=user,  # Solo notificaciones creadas para este vendedor específico
                vista=False  # Solo mostrar notificaciones no vistas
            ).select_related('venta', 'venta__cliente', 'abono').order_by('-fecha_creacion')
            
            # Mostrar las últimas 30 no vistas
            context['notificaciones'] = notificaciones_vendedor[:30]
            context['notificaciones_count'] = notificaciones_vendedor.count()  # Contador solo de no vistas
            # OPTIMIZACIÓN N+1: Prefetch abonos confirmados para evitar consultas extra
            mis_ventas = VentaViaje.objects.filter(vendedor=user).prefetch_related(
                Prefetch(
                    'abonos',
                    queryset=AbonoPago.objects.filter(Q(confirmado=True) | Q(forma_pago='EFE'))
                )
            )
            
            # KPI 1: Mi Saldo Pendiente
            mi_total_vendido_agg = mis_ventas.aggregate(Sum('costo_venta_final'))['costo_venta_final__sum']
            mi_total_vendido = mi_total_vendido_agg if mi_total_vendido_agg is not None else Decimal('0.00')
            
            # Total pagado incluye abonos + montos de apertura
            mi_total_abonos_agg = AbonoPago.objects.filter(
                venta__vendedor=user
            ).filter(
                Q(confirmado=True) | Q(forma_pago='EFE')
            ).aggregate(Sum('monto'))['monto__sum']
            mi_total_abonos = mi_total_abonos_agg if mi_total_abonos_agg is not None else Decimal('0.00')
            
            mi_total_apertura_agg = mis_ventas.aggregate(Sum('cantidad_apertura'))['cantidad_apertura__sum']
            mi_total_apertura = mi_total_apertura_agg if mi_total_apertura_agg is not None else Decimal('0.00')
            
            mi_total_pagado = mi_total_abonos + mi_total_apertura
            
            context['mi_saldo_pendiente'] = mi_total_vendido - mi_total_pagado

            # KPI 2: Mis Ventas Cerradas (ventas donde el total pagado >= costo_venta_final)
            # El total pagado ahora incluye cantidad_apertura + abonos (calculado en la propiedad total_pagado)
            # OPTIMIZACIÓN: Los abonos ya están prefetched, no genera N+1
            ventas_cerradas = 0
            for venta in mis_ventas:
                if venta.total_pagado >= venta.costo_venta_final:
                    ventas_cerradas += 1
            context['mis_ventas_cerradas'] = ventas_cerradas

            # Cotizaciones propias del vendedor que aún no están convertidas en venta (con días desde realización)
            hoy = timezone.localdate()
            cotizaciones_pendientes_qs = Cotizacion.objects.filter(
                vendedor=user
            ).exclude(
                estado='CONVERTIDA'
            ).select_related('cliente').order_by('-creada_en')[:15]
            cotizaciones_con_dias = []
            for cot in cotizaciones_pendientes_qs:
                dias = (hoy - cot.creada_en.date()).days
                cotizaciones_con_dias.append({'cotizacion': cot, 'dias_desde_creacion': dias})
            context['cotizaciones_pendientes_vendedor'] = cotizaciones_con_dias

        # Agregar filtros de fecha al contexto
        context['fecha_filtro'] = self.request.GET.get('fecha_filtro', '')
        context['fecha_desde'] = self.request.GET.get('fecha_desde', '')
        context['fecha_hasta'] = self.request.GET.get('fecha_hasta', '')
        
        # Variables del template que pueden no estar definidas
        context.setdefault('total_ingresos_mtd', Decimal('0.00'))
        context.setdefault('total_ventas', 0)
        context.setdefault('nuevos_clientes_mtd', 0)
        context.setdefault('envios_pendientes', context.get('alertas_logistica_count', 0))
        context.setdefault('cotizaciones_pendientes_vendedor', [])
        
        # Preparar lista de ventas individuales para las cards (solo si hay filtro)
        if context['fecha_filtro'] or (context['fecha_desde'] and context['fecha_hasta']):
            context['ventas_filtradas'] = list(context['ventas'])[:50]  # Limitar a 50 para rendimiento
        else:
            context['ventas_filtradas'] = []

        return context


@login_required
def pagos_pendientes_count(request):
    """Endpoint ligero para polling: devuelve la cantidad de pagos pendientes del contador."""
    rol = perm.get_user_role(request.user, request)
    if rol != 'CONTADOR':
        return JsonResponse({'count': 0})
    count = AbonoPago.objects.filter(
        forma_pago__in=['TRN', 'TAR', 'DEP'],
        confirmado=False,
        comprobante_subido=True
    ).count()
    count += VentaViaje.objects.filter(
        estado_confirmacion='EN_CONFIRMACION'
    ).count()
    return JsonResponse({'count': count})


# ------------------- 2. LISTADO DE VENTAS - SOLUCIÓN AL ERROR DE ANOTACIÓN -------------------

class VentaViajeListView(LoginRequiredMixin, ListView):
    model = VentaViaje 
    template_name = 'ventas/venta_list.html'
    paginate_by = None  # Paginación se hace por separado para activas y cerradas (ver get_context_data)

    def get_queryset(self):
        user = self.request.user
        base_query = perm.get_ventas_queryset_base(self.model, user, self.request)

        # 2. Aplicar filtro por folio/ID
        busqueda_folio = self.request.GET.get('busqueda_folio', '').strip()
        if busqueda_folio:
            # Buscar por folio o por ID numérico
            if busqueda_folio.isdigit():
                base_query = base_query.filter(Q(folio__icontains=busqueda_folio) | Q(pk=int(busqueda_folio)))
            else:
                base_query = base_query.filter(folio__icontains=busqueda_folio)
        
        # 3. Aplicar filtro por servicio
        busqueda_servicio = self.request.GET.get('busqueda_servicio', '').strip()
        if busqueda_servicio:
            if busqueda_servicio == 'VAR':
                # Buscar ventas con múltiples servicios (folio empieza con VAR)
                base_query = base_query.filter(folio__startswith='VAR-')
            else:
                # Buscar ventas que contengan este servicio
                base_query = base_query.filter(
                    Q(servicios_seleccionados__icontains=busqueda_servicio) | 
                    Q(folio__startswith=f'{busqueda_servicio}-')
                )
        
        # 4. Aplicar filtro por fecha de viaje si se proporciona (soporta fecha única o rango)
        fecha_filtro = self.request.GET.get('fecha_filtro')
        fecha_desde = self.request.GET.get('fecha_desde')
        fecha_hasta = self.request.GET.get('fecha_hasta')
        
        # Prioridad: rango de fechas > fecha única
        if fecha_desde and fecha_hasta:
            try:
                from datetime import datetime
                fecha_desde_obj = datetime.strptime(fecha_desde, '%Y-%m-%d').date()
                fecha_hasta_obj = datetime.strptime(fecha_hasta, '%Y-%m-%d').date()
                # Asegurar que desde <= hasta
                if fecha_desde_obj <= fecha_hasta_obj:
                    base_query = base_query.filter(fecha_inicio_viaje__range=[fecha_desde_obj, fecha_hasta_obj])
            except ValueError:
                pass
        elif fecha_filtro:
            try:
                from datetime import datetime
                fecha_obj = datetime.strptime(fecha_filtro, '%Y-%m-%d').date()
                base_query = base_query.filter(fecha_inicio_viaje=fecha_obj)
            except ValueError:
                pass

        # 3. Optimizar el queryset con select_related, prefetch_related y anotaciones
        # OPTIMIZACIÓN N+1: Anotar total_pagado_anotado para evitar consultas en propiedades
        queryset = base_query.select_related(
            'cliente', 'vendedor', 'proveedor'
        ).prefetch_related(
            Prefetch(
                'abonos',
                queryset=AbonoPago.objects.filter(Q(confirmado=True) | Q(forma_pago='EFE'))
            )
        ).annotate(
            # Calcular total de abonos confirmados directamente en la consulta
            total_abonos_confirmados=Coalesce(
                Sum('abonos__monto', filter=Q(abonos__confirmado=True) | Q(abonos__forma_pago='EFE')),
                Value(Decimal('0.00')),
                output_field=ModelDecimalField()
            )
        ).order_by('-fecha_inicio_viaje', '-fecha_creacion')
        
        
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        ventas_list = list(context['object_list'])  # Lista completa (sin paginación global)

        # Separar ventas activas y cerradas
        # Las ventas canceladas van directamente a cerradas
        # Las ventas no canceladas se separan según si están pagadas o no (esta_pagada respeta INT/USD)
        ventas_activas_list = []
        ventas_cerradas_list = []
        
        for venta in ventas_list:
            if venta.estado == 'CANCELADA':
                ventas_cerradas_list.append(venta)
            else:
                if venta.esta_pagada:
                    ventas_cerradas_list.append(venta)
                else:
                    ventas_activas_list.append(venta)
        
        # Paginación por separado: así "Contratos Activos" muestra todos los activos (pag. 1, 2, ...)
        # y "Contratos Cerrados" muestra todos los cerrados (pag. 1, 2, ...)
        paginate_by_activas = 13
        paginate_by_cerradas = 13
        page_activas = self.request.GET.get('page_activas', '1')
        page_cerradas = self.request.GET.get('page_cerradas', '1')
        
        paginator_activas = Paginator(ventas_activas_list, paginate_by_activas)
        paginator_cerradas = Paginator(ventas_cerradas_list, paginate_by_cerradas)
        
        try:
            page_obj_activas = paginator_activas.page(int(page_activas))
        except (ValueError, TypeError):
            page_obj_activas = paginator_activas.page(1) if paginator_activas.num_pages > 0 else None
        except Exception:
            # EmptyPage o página inválida: usar página 1 solo si hay resultados
            page_obj_activas = paginator_activas.page(1) if paginator_activas.num_pages > 0 else None
        
        try:
            page_obj_cerradas = paginator_cerradas.page(int(page_cerradas))
        except (ValueError, TypeError):
            page_obj_cerradas = paginator_cerradas.page(1) if paginator_cerradas.num_pages > 0 else None
        except Exception:
            page_obj_cerradas = paginator_cerradas.page(1) if paginator_cerradas.num_pages > 0 else None
        
        if page_obj_activas is None:
            context['ventas_activas'] = []
            context['page_obj_activas'] = None
            context['current_page_activas'] = 1
        else:
            context['ventas_activas'] = page_obj_activas.object_list
            context['page_obj_activas'] = page_obj_activas
            context['current_page_activas'] = page_obj_activas.number
        if page_obj_cerradas is None:
            context['ventas_cerradas'] = []
            context['page_obj_cerradas'] = None
            context['current_page_cerradas'] = 1
        else:
            context['ventas_cerradas'] = page_obj_cerradas.object_list
            context['page_obj_cerradas'] = page_obj_cerradas
            context['current_page_cerradas'] = page_obj_cerradas.number
        context['user_rol'] = perm.get_user_role(self.request.user, self.request)
        context['ventas_para_cotizacion'] = ventas_list
        
        # Agregar filtros al contexto para mantenerlos en el formulario
        context['fecha_filtro'] = self.request.GET.get('fecha_filtro', '')
        context['fecha_desde'] = self.request.GET.get('fecha_desde', '')
        context['fecha_hasta'] = self.request.GET.get('fecha_hasta', '')
        context['busqueda_folio'] = self.request.GET.get('busqueda_folio', '')
        context['busqueda_servicio'] = self.request.GET.get('busqueda_servicio', '')
        
        # Para CONTADOR: agregar ventas con pagos pendientes de confirmación
        if perm.is_contador(self.request.user, self.request):
            # Ventas con estado "En confirmación" o con abonos pendientes
            ventas_pendientes_ids = list(AbonoPago.objects.filter(
                Q(forma_pago__in=['TRN', 'TAR', 'DEP']) & Q(confirmado=False)
            ).values_list('venta_id', flat=True).distinct())
            
            ventas_en_confirmacion_ids = list(VentaViaje.objects.filter(
                estado_confirmacion='EN_CONFIRMACION'
            ).values_list('pk', flat=True))
            
            # Combinar ambas listas y convertir a lista para el template
            todas_ids_pendientes = list(set(ventas_pendientes_ids + ventas_en_confirmacion_ids))
            context['ventas_pendientes_confirmacion_ids'] = todas_ids_pendientes

        cotizacion_payload = []
        for venta in ventas_list:
            cliente = venta.cliente
            fecha_inicio = venta.fecha_inicio_viaje.isoformat() if venta.fecha_inicio_viaje else ''
            fecha_fin = venta.fecha_fin_viaje.isoformat() if venta.fecha_fin_viaje else ''
            dias = None
            noches = None
            if venta.fecha_inicio_viaje and venta.fecha_fin_viaje:
                delta = (venta.fecha_fin_viaje - venta.fecha_inicio_viaje).days
                if delta >= 0:
                    dias = delta + 1
                    noches = max(delta, 0)

            cotizacion_payload.append({
                'id': venta.pk,
                'slug': venta.slug_safe,
                'cliente': cliente.nombre_completo_display,
                'cotizaciones': getattr(cliente, 'cotizaciones_generadas', 0),
                'fecha_inicio': fecha_inicio,
                'fecha_fin': fecha_fin,
                'dias': dias,
                'noches': noches,
                'increment_url': reverse('incrementar_cotizaciones_cliente', kwargs={'slug': venta.slug_safe, 'pk': venta.pk})
            })

        context['cotizacion_ventas_json'] = json.dumps(cotizacion_payload, ensure_ascii=False)
        context['cotizacion_fecha_hoy'] = timezone.localdate().isoformat()
        context['puede_generar_cotizacion'] = not perm.is_contador(self.request.user, self.request) and bool(cotizacion_payload)
        
        del context['object_list'] 
        
        return context


class _VentaViajeDetailViewPostMixin:
    # ----------------------------------------------------------------------
    # MÉTODO POST: Para gestionar los formularios: Logística y Abonos
    # ----------------------------------------------------------------------
    @transaction.atomic  # ✅ INTEGRIDAD: Transacción atómica para abonos y actualizaciones de estado
    def post(self, request, *args, **kwargs):
        # Es crucial establecer el objeto (VentaViaje) al inicio del POST
        # self.get_object() utiliza los nuevos kwargs (slug y pk)
        self.object = self.get_object() 
        context = self.get_context_data(object=self.object)

        # 0. Añadir un servicio extra (Vuelo/Hospedaje/Tour): crea una fila y redirige a logística
        if request.POST.get('añadir_servicio_logistica') and self._puede_gestionar_logistica_financiera(request.user, self.object):
            tipo = (request.POST.get('tipo') or '').strip().upper()
            servicios_codes = [c.strip() for c in (self.object.servicios_seleccionados or '').split(',') if c.strip()]
            choices = dict(VentaViaje.SERVICIOS_CHOICES)
            if tipo == 'TOU' and 'TOU' in servicios_codes:
                nombre = choices.get('TOU', 'Tour y Actividades')
                next_orden = (self.object.servicios_logisticos.filter(codigo_servicio='TOU').aggregate(mx=Max('orden'))['mx'] or 0) + 1
                LogisticaServicio.objects.create(venta=self.object, codigo_servicio='TOU', nombre_servicio=nombre, orden=next_orden)
                messages.success(request, "Se añadió una fila de Tour. Complete los datos y guarde.")
            elif tipo == 'VUE' and 'VUE' in servicios_codes:
                nombre = choices.get('VUE', 'Vuelo')
                next_orden = (self.object.servicios_logisticos.filter(codigo_servicio='VUE').aggregate(mx=Max('orden'))['mx'] or 0) + 1
                LogisticaServicio.objects.create(venta=self.object, codigo_servicio='VUE', nombre_servicio=nombre, orden=next_orden)
                messages.success(request, "Se añadió una fila de Vuelo. Complete los datos y guarde.")
            elif tipo == 'HOS' and 'HOS' in servicios_codes:
                nombre = choices.get('HOS', 'Hospedaje')
                next_orden = (self.object.servicios_logisticos.filter(codigo_servicio='HOS').aggregate(mx=Max('orden'))['mx'] or 0) + 1
                LogisticaServicio.objects.create(venta=self.object, codigo_servicio='HOS', nombre_servicio=nombre, orden=next_orden)
                messages.success(request, "Se añadió una fila de Hospedaje. Complete los datos y guarde.")
            else:
                messages.error(request, "No se pudo añadir el servicio. Verifique que el tipo sea correcto.")
            return redirect(reverse('detalle_venta', kwargs={'pk': self.object.pk, 'slug': self.object.slug_safe}) + '?tab=logistica')

        # 1. Manejo del control financiero por servicio (Guardar ajustes)
        if 'actualizar_servicios_logistica' in request.POST:
            if not self._puede_gestionar_logistica_financiera(request.user, self.object):
                messages.error(request, "No tienes permiso para actualizar el control financiero de esta venta.")
                return redirect(reverse('detalle_venta', kwargs={'pk': self.object.pk, 'slug': self.object.slug_safe}) + '?tab=logistica')

            # [NUEVO] Paso 1: Obtener una COPIA LIMPIA de los originales para referencia
            # Es vital hacer esto ANTES de validar el formset y en una variable separada
            # Usamos list() para asegurar que se traigan de la BD ahora mismo
            servicios_db_list = list(self.object.servicios_logisticos.all())
            originales_limpios = {s.pk: s for s in servicios_db_list}
            
            # Paso 2: Crear el formset (permite múltiples TOU: extra=1 y can_delete cuando hay Tour)
            servicios_qs = self.object.servicios_logisticos.all().order_by('orden', 'pk')
            formset = _get_logistica_servicio_formset(
                self.object,
                request_POST=request.POST,
                queryset=servicios_qs,
                prefix='servicios',
            )

            if formset.is_valid():
                total_pagado = self.object.total_pagado
                total_marcado_pagado = Decimal('0.00')
                # Snapshot de pagado/fecha_pagado antes de modificar (para revertir si la validación falla)
                pagado_snapshot = {s.pk: (s.pagado, s.fecha_pagado) for s in servicios_db_list}

                for form in formset.forms:
                    if not form.cleaned_data:
                        continue
                    
                    servicio_id = form.instance.pk if form.instance.pk else None
                    original = originales_limpios.get(servicio_id) if servicio_id else None
                    
                    if not original:
                        continue 
                    
                    # --- LÓGICA DE ACTUALIZACIÓN CON BLOQUEO POR ROL ---
                    # Monto planificado y pagado: cualquier rol puede llenar por primera vez (vacíos).
                    # Solo Director Admin, Director General y JEFE pueden EDITAR cuando ya tienen valores.
                    # Nombre del proveedor: solo Director Admin, Director General y JEFE siempre.
                    puede_editar_restringidos = perm.can_edit_logistica_campos_restringidos(request.user, request)

                    if puede_editar_restringidos:
                        nuevo_monto = form.cleaned_data.get('monto_planeado') or Decimal('0.00')
                        nuevo_pagado = form.cleaned_data.get('pagado', False)
                        nuevo_opcion_proveedor = (form.cleaned_data.get('opcion_proveedor', '') or '').strip()
                    else:
                        # 1. MONTO PLANIFICADO: permitir si está vacío (primera vez); si ya tiene valor, conservar
                        nuevo_monto = form.cleaned_data.get('monto_planeado')
                        campo_ya_llenado = original.monto_planeado and original.monto_planeado > Decimal('0.00')
                        if campo_ya_llenado:
                            nuevo_monto = original.monto_planeado
                        else:
                            nuevo_monto = nuevo_monto or Decimal('0.00')

                        # 2. PAGADO: permitir si está vacío (primera vez); si ya marcado, conservar
                        nuevo_pagado = form.cleaned_data.get('pagado', False)
                        if original.pagado:
                            nuevo_pagado = True

                        # 3. OPCIÓN PROVEEDOR: permitir si está vacío (primera vez); si ya tiene valor, conservar
                        proveedor_ya_llenado = (original.opcion_proveedor or '').strip()
                        if proveedor_ya_llenado:
                            nuevo_opcion_proveedor = proveedor_ya_llenado
                        else:
                            nuevo_opcion_proveedor = (form.cleaned_data.get('opcion_proveedor', '') or '').strip()
                    
                    # --- APLICAR CAMBIOS ---
                    original.monto_planeado = nuevo_monto
                    original.pagado = nuevo_pagado
                    original.opcion_proveedor = nuevo_opcion_proveedor
                    
                    # Gestión de fecha de pago
                    if nuevo_pagado and not original.fecha_pagado:
                         original.fecha_pagado = timezone.now()
                    elif not nuevo_pagado:
                         original.fecha_pagado = None
                    
                    # Calcular total marcado como pagado para validación
                    if nuevo_pagado:
                        total_marcado_pagado += nuevo_monto
                
                # Calcular suma de montos planificados (se usa para INT y para referencia)
                suma_planeados = sum(
                    (orig.monto_planeado or Decimal('0.00')) for orig in originales_limpios.values()
                )

                # Validar que el total marcado como pagado no exceda el total pagado
                # Excepción: Crédito (CRE) no requiere apertura ni abonos - se paga fuera del dashboard
                es_credito = getattr(self.object, 'modo_pago_apertura', None) == 'CRE'
                pagado_excede = not es_credito and total_marcado_pagado > total_pagado + Decimal('0.01')

                if pagado_excede:
                    # Revertir casillas de pagado: restaurar valores originales de BD
                    for pk, (pagado_orig, fecha_orig) in pagado_snapshot.items():
                        orig = originales_limpios.get(pk)
                        if orig:
                            orig.pagado = pagado_orig
                            orig.fecha_pagado = fecha_orig

                # Guardar siempre montos y proveedores (no se pierden al añadir servicios).
                # pagado/fecha_pagado solo se persisten si la validación pasa.
                for orig in originales_limpios.values():
                    orig.save(update_fields=['monto_planeado', 'pagado', 'fecha_pagado', 'opcion_proveedor'])

                if pagado_excede:
                    formset._non_form_errors = formset.error_class([
                        f"No puedes marcar como pagados ${total_marcado_pagado:,.2f} cuando solo hay ${total_pagado:,.2f} registrados en abonos y apertura."
                    ])
                    context = self.get_context_data()
                    self._prepare_logistica_finanzas_context(context, self.object, formset=formset, servicios_qs=servicios_qs)
                    return self.render_to_response(context)

                # Para INT: si costo_neto_usd no estaba definido, persistir la suma de la tabla como "servicios planificados"
                if self.object.tipo_viaje == 'INT' and (getattr(self.object, 'costo_neto_usd', None) is None or self.object.costo_neto_usd <= 0) and suma_planeados > 0:
                    self.object.costo_neto_usd = suma_planeados
                    self.object.save(update_fields=['costo_neto_usd'])

                # Sincronizar venta.proveedor desde la tabla de logística: si algún servicio tiene
                # "Nombre del proveedor" que coincide con un Proveedor con metodo_pago_preferencial,
                # asignar ese proveedor a la venta para que se muestre la sección Abonos a Proveedor
                if self.object.tipo_viaje in ('NAC', 'INT_MXN'):
                    for original in originales_limpios.values():
                        nombre_prov = (original.opcion_proveedor or '').strip()
                        if nombre_prov:
                            proveedor_pref = Proveedor.objects.filter(
                                nombre__iexact=nombre_prov,
                                metodo_pago_preferencial=True
                            ).first()
                            if proveedor_pref and (not self.object.proveedor or not self.object.proveedor.metodo_pago_preferencial):
                                self.object.proveedor = proveedor_pref
                                self.object.save(update_fields=['proveedor'])
                                break

                # Eliminar filas TOU y VUE vacías (monto 0 y sin nombre de proveedor) al guardar
                tou_vacios = list(
                    self.object.servicios_logisticos.filter(codigo_servicio='TOU')
                )
                for s in tou_vacios:
                    if (s.monto_planeado or Decimal('0.00')) <= Decimal('0.00') and not (s.opcion_proveedor or '').strip():
                        s.delete()
                
                vue_vacios = list(
                    self.object.servicios_logisticos.filter(codigo_servicio='VUE')
                )
                for s in vue_vacios:
                    if (s.monto_planeado or Decimal('0.00')) <= Decimal('0.00') and not (s.opcion_proveedor or '').strip():
                        s.delete()

                hos_vacios = list(
                    self.object.servicios_logisticos.filter(codigo_servicio='HOS')
                )
                for s in hos_vacios:
                    if (s.monto_planeado or Decimal('0.00')) <= Decimal('0.00') and not (s.opcion_proveedor or '').strip():
                        s.delete()

                messages.success(request, "Control por servicio actualizado correctamente.")
                # Re-renderizar en la misma petición (sin redirect) para que la tabla muestre la fila recién guardada.
                # No ejecutar sync aquí (skip_sync=True): evitar que se cree un servicio duplicado al guardar.
                if getattr(self.object, '_prefetched_objects_cache', None) and 'servicios_logisticos' in self.object._prefetched_objects_cache:
                    del self.object._prefetched_objects_cache['servicios_logisticos']
                servicios_qs_final = self.object.servicios_logisticos.all().order_by('orden', 'pk')
                context = self.get_context_data()
                self._prepare_logistica_finanzas_context(context, self.object, formset=None, servicios_qs=servicios_qs_final, skip_sync=True)
                context['activar_tab_logistica'] = True  # Mantener pestaña Logística visible tras guardar
                return self.render_to_response(context)
            else:
                messages.error(request, "Revisa los montos ingresados para cada servicio.")

            self._prepare_logistica_finanzas_context(context, self.object, formset=formset, servicios_qs=servicios_qs)

        # 3. Manejo del Formulario de Abono
        elif 'registrar_abono' in request.POST:
            # CONTADOR solo lectura, no puede registrar abonos
            user_rol = perm.get_user_role(request.user, request)
            if user_rol == 'CONTADOR':
                messages.error(request, "No tienes permiso para registrar abonos. Solo puedes visualizarlos.")
                return redirect(reverse('detalle_venta', kwargs={'pk': self.object.pk, 'slug': self.object.slug_safe}) + '?tab=abonos')
            
            # Para INT el usuario ingresa monto_usd en el template; el form requiere monto > 0, pasamos 1 y luego sobrescribimos
            post_abono = request.POST.copy()
            if self.object.tipo_viaje == 'INT' and not (post_abono.get('monto') or '').strip():
                post_abono['monto'] = '1'
            abono_form = AbonoPagoForm(post_abono)
            if abono_form.is_valid():
                abono = abono_form.save(commit=False)
                abono.venta = self.object
                abono.registrado_por = request.user

                # Ventas internacionales: monto ingresado es USD; guardar monto_usd y tipo_cambio (referencia)
                if self.object.tipo_viaje == 'INT':
                    monto_usd_str = (request.POST.get('monto_usd') or '').replace('$', '').replace(',', '').strip()
                    tc_abono_str = (request.POST.get('tipo_cambio_abono') or '').replace(',', '').strip()
                    try:
                        monto_usd_val = Decimal(monto_usd_str) if monto_usd_str else None
                        tc_abono_val = Decimal(tc_abono_str) if tc_abono_str else None
                    except (ValueError, InvalidOperation):
                        monto_usd_val = tc_abono_val = None
                    if monto_usd_val is not None and monto_usd_val > 0:
                        abono.monto_usd = monto_usd_val.quantize(Decimal('0.01'))
                        abono.tipo_cambio_aplicado = tc_abono_val or self.object.tipo_cambio
                        abono.monto = Decimal('0.00')  # INT: no usar MXN; fuente de verdad es monto_usd
                
                # Obtener la forma de pago del formulario
                forma_pago = abono_form.cleaned_data.get('forma_pago', 'EFE')
                
                # ⚠️ IMPORTANTE: Determinar si requiere confirmación del contador
                # Transferencia (TRN), Tarjeta (TAR) y Depósito (DEP) requieren confirmación
                # Solo Efectivo (EFE) se confirma automáticamente
                requiere_confirmacion = forma_pago in ['TRN', 'TAR', 'DEP']
                
                # Establecer el estado de confirmación ANTES de guardar
                if requiere_confirmacion:
                    # Transferencia, Tarjeta o Depósito: NO confirmado, requiere aprobación del contador
                    abono.confirmado = False
                    abono.confirmado_por = None
                    abono.confirmado_en = None
                else:
                    # Solo Efectivo se confirma automáticamente
                    abono.confirmado = True
                    abono.confirmado_por = None  # No requiere confirmación manual
                    abono.confirmado_en = timezone.now()  # Fecha de confirmación automática
                
                # ✅ INTEGRIDAD FINANCIERA: Validar que el abono no exceda el saldo restante
                saldo_restante = self.object.saldo_restante
                monto_abono = abono.monto_usd if (self.object.tipo_viaje == 'INT' and abono.monto_usd) else abono.monto
                
                if monto_abono > saldo_restante:
                    messages.error(
                        request, 
                        f"El monto del abono (${monto_abono:,.2f}) excede el saldo restante de la venta (${saldo_restante:,.2f}). "
                        f"Por favor, ajusta el monto del abono."
                    )
                    context = self.get_context_data(object=self.object)
                    context['abono_form'] = abono_form
                    return self.render_to_response(context)
                
                # Guardar el abono primero para obtener su PK
                abono.save()
                
                # Procesar según si requiere confirmación o no
                if requiere_confirmacion:
                    # ⚠️ FLUJO DE APROBACIÓN PARA TRANSFERENCIA/TARJETA/DEPÓSITO
                    # ✅ NUEVO FLUJO: NO se crean notificaciones automáticas
                    # Las notificaciones se crearán solo cuando se suba el comprobante
                    
                    # 1. Cambiar estado de venta a "En confirmación" SOLO si no está ya COMPLETADO.
                    # Si la apertura (o crédito) ya fue confirmada, la venta está COMPLETADO y no debe
                    # cambiarse: así evitamos que la apertura "reaparezca" en el dashboard del contador.
                    if self.object.estado_confirmacion != 'COMPLETADO':
                        self.object.estado_confirmacion = 'EN_CONFIRMACION'
                        self.object.save(update_fields=['estado_confirmacion'])
                    self.object.actualizar_estado_financiero()
                    
                    forma_pago_display = dict(AbonoPago.FORMA_PAGO_CHOICES).get(forma_pago, forma_pago)
                    monto_abono = abono.monto_usd if (self.object.tipo_viaje == 'INT' and abono.monto_usd is not None) else abono.monto
                    moneda_abono = 'USD' if self.object.tipo_viaje == 'INT' else 'MXN'
                    messages.success(request, f"Abono de ${monto_abono:,.2f} {moneda_abono} ({forma_pago_display}) registrado exitosamente. ⏳ Por favor, sube el comprobante para enviarlo al contador.")
                else:
                    # ⚠️ FLUJO AUTOMÁTICO SOLO PARA EFECTIVO
                    
                    # Recalcular el total pagado después de guardar el abono
                    self.object.actualizar_estado_financiero()
                    
                    forma_pago_display = dict(AbonoPago.FORMA_PAGO_CHOICES).get(forma_pago, forma_pago)
                    
                    # Crear notificación para el VENDEDOR de la venta (si existe y no es quien registra el abono)
                    if self.object.vendedor and self.object.vendedor != request.user:
                        mensaje_vendedor = f"Abono registrado y confirmado en tu venta: ${abono.monto:,.2f} ({forma_pago_display}) - Venta #{self.object.pk} - Cliente: {self.object.cliente.nombre_completo_display}"
                        Notificacion.objects.create(
                            usuario=self.object.vendedor,
                            tipo='ABONO',
                            mensaje=mensaje_vendedor,
                            venta=self.object,
                            abono=abono,
                            confirmado=True
                        )
                    
                    # Si la venta se liquidó completamente, crear notificación de liquidación
                    if self.object.esta_pagada and self.object.vendedor:
                        mensaje_liquidacion = f"¡Venta #{self.object.pk} completamente liquidada! - Cliente: {self.object.cliente.nombre_completo_display} - Total: ${self.object.costo_venta_final:,.2f}"
                        Notificacion.objects.create(
                            usuario=self.object.vendedor,
                            tipo='LIQUIDACION',
                            mensaje=mensaje_liquidacion,
                            venta=self.object,
                            confirmado=False
                        )
                    
                    monto_abono = abono.monto_usd if (self.object.tipo_viaje == 'INT' and abono.monto_usd is not None) else abono.monto
                    moneda_abono = 'USD' if self.object.tipo_viaje == 'INT' else 'MXN'
                    messages.success(request, f"Abono de ${monto_abono:,.2f} {moneda_abono} ({forma_pago_display}) registrado exitosamente. ✅ Confirmado automáticamente.")
                
                # Redirige a la pestaña de Abonos
                # ******************************************************************
                # IMPORTANTE: Se actualiza la redirección para usar SLUG y PK (YA ESTABA BIEN AQUÍ)
                # ******************************************************************
                return redirect(reverse('detalle_venta', kwargs={'pk': self.object.pk, 'slug': self.object.slug_safe}) + '?tab=abonos')
            else:
                messages.error(request, "Error al registrar el abono. Revisa el monto y la forma de pago. ⚠️")
                context['abono_form'] = abono_form # Muestra el formulario con errores
                
        # 3. Manejo del formulario de confirmaciones
        elif 'registrar_confirmacion' in request.POST:
            if not self._puede_subir_confirmaciones(request.user, self.object):
                messages.error(request, "No tienes permiso para agregar confirmaciones a esta venta.")
                return redirect(reverse('detalle_venta', kwargs={'pk': self.object.pk, 'slug': self.object.slug_safe}) + '?tab=confirmaciones')

            # Validar archivos directamente desde request.FILES
            archivos = request.FILES.getlist('archivos')
            nota = request.POST.get('nota', '').strip()

            if not archivos:
                messages.error(request, "Debes seleccionar al menos un archivo.")
                confirmacion_form = ConfirmacionVentaForm(request.POST, request.FILES)
                context['confirmacion_form'] = confirmacion_form
            else:
                # SEGURIDAD: Validar archivos antes de procesarlos
                archivos_validos = []
                errores_validacion = []
                for archivo in archivos:
                    try:
                        validate_uploaded_file(archivo)
                        archivos_validos.append(archivo)
                    except DjangoValidationError as e:
                        errores_validacion.append(str(e.message if hasattr(e, 'message') else e))
                
                if errores_validacion:
                    messages.error(request, "Archivos rechazados: " + " | ".join(errores_validacion))
                    if not archivos_validos:
                        confirmacion_form = ConfirmacionVentaForm(request.POST, request.FILES)
                        context['confirmacion_form'] = confirmacion_form
                        return self.render_to_response(context)
                
                # Procesar los archivos validados
                creadas = 0
                errores = []
                for archivo in archivos_validos:
                    try:
                        ConfirmacionVenta.objects.create(
                            venta=self.object,
                            archivo=archivo,
                            nota=nota,
                            subido_por=request.user if request.user.is_authenticated else None
                        )
                        creadas += 1
                    except Exception as e:
                        errores.append(f"Error al guardar {archivo.name}: {str(e)}")

                if creadas > 0:
                    mensaje = f"Se cargaron {creadas} archivo(s) de confirmación correctamente."
                    if errores:
                        mensaje += " " + " ".join(errores)
                    messages.success(request, mensaje)
                else:
                    messages.error(request, "No se pudo cargar ningún archivo. " + " ".join(errores))
                
                return redirect(reverse('detalle_venta', kwargs={'pk': self.object.pk, 'slug': self.object.slug_safe}) + '?tab=confirmaciones')

        # Si no hubo redirección exitosa, re-renderiza la respuesta con el contexto actualizado
        return self.render_to_response(context)

    def _puede_subir_confirmaciones(self, user, venta):
        """CONTADOR solo lectura, no puede subir confirmaciones."""
        if not user.is_authenticated:
            return False
        if user.is_superuser or user == venta.vendedor:
            return True
        rol = perm.get_user_role(user).upper()
        # CONTADOR no puede subir confirmaciones, solo visualizarlas
        return 'JEFE' in rol

    # ------------------- Utilidades internas para logística financiera -------------------

    def _puede_ver_logistica_tab(self, user, venta):
        if not user or not user.is_authenticated:
            return False
        if perm.has_full_access(user, self.request) or perm.is_contador(user, self.request):
            return True
        if perm.can_edit_datos_viaje(user, self.request):
            return True
        return perm.is_vendedor(user) and venta.vendedor == user

    def _puede_gestionar_logistica_financiera(self, user, venta):
        """
        Quién puede gestionar la pestaña Logística (habilitar el formulario).
        Cualquier rol que pueda ver la pestaña puede llenar monto/pagado por primera vez.
        La edición de campos ya llenados y nombre del proveedor se controla por campo.
        """
        if not user or not user.is_authenticated:
            return False
        return self._puede_ver_logistica_tab(user, venta)

    def _sync_logistica_servicios(self, venta):
        """Asegura al menos una fila por cada código en servicios_seleccionados. No toca monto ni pagado; no actualiza filas ya existentes (extras TOU/VUE/HOS)."""
        servicios_codes = []
        if venta.servicios_seleccionados:
            servicios_codes = [
                code.strip() for code in venta.servicios_seleccionados.split(',')
                if code.strip()
            ]
        choices = dict(VentaViaje.SERVICIOS_CHOICES)

        # Solo para crear la fila inicial si no existe ninguna de ese tipo
        nombres_proveedor = {}
        if venta.servicios_detalle:
            for linea in venta.servicios_detalle.split('\n'):
                linea = linea.strip()
                if not linea or ' - Proveedor: ' not in linea:
                    continue
                partes = linea.split(' - Proveedor: ', 1)
                if len(partes) != 2:
                    continue
                nombre_servicio = partes[0].strip()
                resto = partes[1].strip()
                nombre_proveedor = resto.split(' - Opción: ')[0].strip() if ' - Opción: ' in resto else resto
                if nombre_proveedor:
                    nombres_proveedor[nombre_servicio] = nombre_proveedor

        for idx, code in enumerate(servicios_codes):
            nombre = choices.get(code)
            if not nombre:
                continue
            existentes = LogisticaServicio.objects.filter(venta_id=venta.pk, codigo_servicio=code)
            if not existentes.exists():
                # Crear fila inicial si no existe ninguna de este tipo
                opcion_proveedor = nombres_proveedor.get(nombre, '')
                LogisticaServicio.objects.create(
                    venta=venta,
                    codigo_servicio=code,
                    nombre_servicio=nombre,
                    orden=idx,
                    opcion_proveedor=opcion_proveedor
                )
            elif nombres_proveedor.get(nombre):
                # Sincronizar nombre de proveedor desde el formulario de edición
                # en la PRIMERA fila (por defecto) de este servicio, sin importar
                # cuántas filas extra haya añadido el usuario.
                fila = existentes.order_by('orden', 'pk').first()
                nuevo_proveedor = nombres_proveedor[nombre]
                if fila and (fila.opcion_proveedor or '').strip() != nuevo_proveedor:
                    fila.opcion_proveedor = nuevo_proveedor
                    fila.save(update_fields=['opcion_proveedor'])

        # Eliminar servicios que ya no están contratados (consulta directa para no depender de relación en caché).
        # IMPORTANTE: Si servicios_codes está vacío (servicios_seleccionados vacío o None),
        # NO borrar los servicios existentes.
        if servicios_codes:
            LogisticaServicio.objects.filter(venta_id=venta.pk).exclude(codigo_servicio__in=servicios_codes).delete()

    def _prepare_logistica_finanzas_context(self, context, venta, formset=None, servicios_qs=None, skip_sync=False):
        # En pestaña Logística: monto planificado, pagado y nombre proveedor solo editables por JEFE, Gerente y 3 directores
        user_rol = perm.get_user_role(self.request.user, self.request)
        puede_editar_campos_bloqueados = perm.can_edit_campos_bloqueados(self.request.user, self.request)
        context['puede_editar_campos_bloqueados'] = puede_editar_campos_bloqueados
        context['puede_desbloquear_todos_los_campos'] = perm.can_edit_datos_viaje(self.request.user, self.request)

        if not context.get('mostrar_tab_logistica'):
            return

        # Forzar consulta fresca: limpiar prefetch para que _sync y la tabla vean todos los servicios (incl. recién añadidos)
        if getattr(venta, '_prefetched_objects_cache', None) and 'servicios_logisticos' in venta._prefetched_objects_cache:
            del venta._prefetched_objects_cache['servicios_logisticos']
        # No ejecutar sync al re-renderizar tras "Guardar ajustes": ya tenemos el conjunto correcto y sync podría crear un duplicado
        if not skip_sync:
            self._sync_logistica_servicios(venta)
        if servicios_qs is None:
            # Consulta directa por venta_id para no depender de la relación en memoria (asegura ver servicios recién creados tras redirect)
            servicios_qs = LogisticaServicio.objects.filter(venta_id=venta.pk).order_by('orden', 'pk')

        if formset is None:
            formset = _get_logistica_servicio_formset(venta, queryset=servicios_qs, prefix='servicios')

        # Determinar si el usuario tiene acceso total
        user_rol = perm.get_user_role(self.request.user, self.request)
        es_jefe = perm.has_full_access(self.request.user, self.request)
        context['es_jefe'] = es_jefe
        puede_editar_restringidos = perm.can_edit_logistica_campos_restringidos(self.request.user, self.request)
        context['puede_editar_logistica_restringidos'] = puede_editar_restringidos

        if not context.get('puede_editar_servicios_financieros'):
            for form in formset.forms:
                for field in form.fields.values():
                    field.widget.attrs['disabled'] = 'disabled'
        else:
            # Cualquier rol puede llenar campos por primera vez (vacíos).
            # Solo Director Admin, Director General y JEFE pueden EDITAR cuando ya tienen valores.
            for form in formset.forms:
                if form.instance and form.instance.pk:
                    if not puede_editar_restringidos:
                        # monto_planeado: bloquear solo si ya tiene valor
                        if form.instance.monto_planeado and form.instance.monto_planeado > Decimal('0.00'):
                            form.fields['monto_planeado'].widget.attrs['disabled'] = 'disabled'
                            form.fields['monto_planeado'].widget.attrs['readonly'] = 'readonly'

                        # opcion_proveedor: bloquear solo si ya tiene valor
                        if (form.instance.opcion_proveedor or '').strip():
                            form.fields['opcion_proveedor'].widget.attrs['disabled'] = 'disabled'
                            form.fields['opcion_proveedor'].widget.attrs['readonly'] = 'readonly'

                        # pagado: bloquear solo si ya está marcado
                        if form.instance.pagado:
                            form.fields['pagado'].widget.attrs['disabled'] = 'disabled'

        resumen = build_financial_summary(venta, servicios_qs)
        # Suma efectiva desde el formset (valores del formulario o POST) para que el alert muestre lo que el usuario ve
        if formset.is_bound and formset.forms:
            suma_efectiva = Decimal('0.00')
            for form in formset.forms:
                if form.cleaned_data is not None:
                    m = form.cleaned_data.get('monto_planeado')
                elif form.instance and form.instance.pk:
                    m = getattr(form.instance, 'monto_planeado', None)
                else:
                    m = None
                if m is None and form.data:
                    raw = form.data.get(form.add_prefix('monto_planeado'), '')
                    if raw:
                        try:
                            raw_limpio = str(raw).replace('$', '').replace(',', '').replace(' ', '').strip()
                            if raw_limpio:
                                m = Decimal(raw_limpio)
                        except (InvalidOperation, ValueError):
                            pass
                suma_efectiva += (m or Decimal('0.00'))
            resumen['suma_montos_planeados'] = suma_efectiva
            # Para INT: si el total de servicios planificados viene en 0 (BD) pero el usuario ya ingresó montos en el form, usar esa suma como objetivo
            total_objetivo = resumen.get('total_servicios_planeados') or Decimal('0.00')
            if venta.tipo_viaje == 'INT' and total_objetivo <= 0 and suma_efectiva > 0:
                resumen['total_servicios_planeados'] = suma_efectiva
                total_objetivo = suma_efectiva
            resumen['montos_cuadran'] = abs(suma_efectiva - total_objetivo) < Decimal('0.01')

        formset_forms = list(formset.forms)
        filas = build_service_rows(servicios_qs, resumen, formset_forms[: len(servicios_qs)], venta=venta)

        servicios_codes = [c.strip() for c in (venta.servicios_seleccionados or '').split(',') if c.strip()]
        has_tou = 'TOU' in servicios_codes
        has_vue = 'VUE' in servicios_codes
        has_hos = 'HOS' in servicios_codes

        context['servicios_financieros_formset'] = formset
        context['tiene_tou_para_otro_proveedor'] = has_tou
        context['tiene_vue_para_otro_proveedor'] = has_vue
        context['tiene_hos_para_otro_proveedor'] = has_hos
        context['logistica_finanzas'] = resumen
        context['servicios_logisticos_rows'] = filas
        context['servicios_logisticos_queryset'] = servicios_qs

# ------------------- 3. DETALLE DE VENTA MODIFICADA -------------------

class VentaViajeDetailView(LoginRequiredMixin, usuarios_mixins.VentaPermissionMixin, _VentaViajeDetailViewPostMixin, DetailView):
    model = VentaViaje
    template_name = 'ventas/venta_detail.html'
    context_object_name = 'venta'

    def get_queryset(self):
        """✅ PERFORMANCE: Optimizar queryset con prefetch de relaciones"""
        return super().get_queryset().select_related(
            'cliente', 'vendedor', 'proveedor'
        ).prefetch_related(
            'abonos',
            'servicios_logisticos'
        )

    # ******************************************************************
    # NUEVO: Implementación de get_object para usar SLUG y PK
    # ******************************************************************
    def get_object(self, queryset=None):
        # 1. Recupera los parámetros de la URL
        pk = self.kwargs.get('pk')
        slug = self.kwargs.get('slug')
        
        # 2. Define el queryset base si no se proporciona uno
        if queryset is None:
            queryset = self.get_queryset()
            
        # 3. Busca el objeto utilizando ambos parámetros para asegurar unicidad
        try:
            # Usamos get_queryset() que ya trae el .select_related() si lo tienes configurado
            obj = queryset.filter(pk=pk, slug=slug).first()
            if obj:
                return obj
        except VentaViaje.DoesNotExist:
            pass # Continúa al manejo de error 404
            
        # 4. Si el objeto no se encuentra, levanta un error 404
        from django.http import Http404
        raise Http404("No se encontró la Venta de Viaje que coincide con el ID y el slug.")
    
    # ******************************************************************
    # FIN de get_object
    # ******************************************************************

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        venta = self.object
        # Corregir venta marcada COMPLETADO/liquidada cuando en realidad no está pagada (apertura TRN/TAR/DEP sin confirmar)
        if venta.estado_confirmacion == 'COMPLETADO' and venta.saldo_restante > 0:
            venta.actualizar_estado_financiero(guardar=True)
        user_rol = perm.get_user_role(self.request.user, self.request)
        context['user_rol'] = user_rol
        context['es_jefe'] = perm.has_full_access(self.request.user, self.request)
        
        # Marcar notificaciones pendientes como vistas al entrar al detalle
        # Esto hace que al hacer clic en "Ver Venta" la notificación se marque automáticamente
        if self.request.user.is_authenticated:
            Notificacion.objects.filter(
                usuario=self.request.user,
                venta=venta,
                vista=False
            ).update(vista=True, fecha_vista=timezone.now())
        
        # Verificar si existen abonos reales pendientes
        # Esto ayuda a determinar si la apertura está realmente pendiente o si solo hay abonos nuevos pendientes
        abonos_pendientes_reales = venta.abonos.filter(confirmado=False).exists()
        context['tiene_abonos_pendientes'] = abonos_pendientes_reales
        
        # Inicialización del Formulario de Abono
        context['abono_form'] = AbonoPagoForm(initial={'venta': venta.pk}) 
        # Abonos existentes, ordenados por fecha
        context['abonos'] = venta.abonos.all().order_by('-fecha_pago')

        # Control de habilitación de documentos (contrato y comprobante de abonos)
        # Se desbloquean cuando hay al menos un abono confirmado O la apertura está confirmada
        tiene_abono_confirmado = venta.abonos.filter(confirmado=True).exists()
        
        # Apertura confirmada: para NAC cantidad_apertura > 0; para INT cantidad_apertura_usd (o resuelto) > 0
        apertura_confirmada = False
        tiene_apertura = False
        if venta.tipo_viaje == 'INT':
            apertura_usd = getattr(venta, 'cantidad_apertura_usd', None)
            # ✅ DIVISIÓN SEGURA: Validar tipo_cambio > 0
            if (apertura_usd is None or apertura_usd <= 0) and venta.tipo_cambio and venta.tipo_cambio > 0 and venta.cantidad_apertura:
                apertura_usd = (venta.cantidad_apertura / venta.tipo_cambio).quantize(Decimal('0.01'))
            tiene_apertura = apertura_usd and apertura_usd > 0
        else:
            tiene_apertura = (venta.cantidad_apertura or 0) > 0
        if tiene_apertura:
            if venta.modo_pago_apertura == 'EFE':
                apertura_confirmada = True
            else:
                apertura_confirmada = venta.estado_confirmacion == 'COMPLETADO'
        
        context['puede_generar_documentos'] = tiene_abono_confirmado or apertura_confirmada
        
        # Verificar si hay notificación de apertura pendiente para el CONTADOR
        if perm.is_contador(self.request.user, self.request):
            # Notificacion ya está importada al inicio del archivo
            context['notificacion_apertura_pendiente'] = Notificacion.objects.filter(
                usuario=self.request.user,
                venta=venta,
                tipo='PAGO_PENDIENTE',
                confirmado=False,
                abono__isnull=True  # Sin abono = apertura
            ).first()

        # Asegurar que exista registro de logística para sincronizar servicios
        try:
            venta.logistica
        except Logistica.DoesNotExist:
            Logistica.objects.create(venta=venta)
        mostrar_tab_logistica = self._puede_ver_logistica_tab(self.request.user, venta)
        context['mostrar_tab_logistica'] = mostrar_tab_logistica
        context['puede_editar_servicios_financieros'] = self._puede_gestionar_logistica_financiera(self.request.user, venta)
        # Permisos para editar campos bloqueados (JEFE/Director General y Gerente)
        context['puede_desbloquear_todos_los_campos'] = perm.can_edit_campos_bloqueados(self.request.user, self.request)
        # Botón "Editar datos del viaje": solo Gerente y los 3 directores
        context['puede_editar_datos_viaje'] = perm.can_edit_datos_viaje(self.request.user, self.request)
        if mostrar_tab_logistica:
            # Refrescar venta desde BD y limpiar prefetch para que la tabla use siempre datos frescos (evita que filas nuevas guardadas no se vean tras redirect)
            venta.refresh_from_db()
            if getattr(venta, '_prefetched_objects_cache', None) and 'servicios_logisticos' in venta._prefetched_objects_cache:
                del venta._prefetched_objects_cache['servicios_logisticos']
            # No ejecutar sync cuando estamos re-renderizando tras POST "Guardar ajustes" (evita duplicar filas VUE/HOS/TOU)
            skip_sync = (
                self.request.method == 'POST'
                and 'actualizar_servicios_logistica' in (self.request.POST or {})
            )
            self._prepare_logistica_finanzas_context(context, venta, skip_sync=skip_sync)
        
        # Abonos a Proveedor: usar SIEMPRE venta.puede_solicitar_abonos_proveedor (misma regla que POST)
        debe_mostrar_abonos = venta.puede_solicitar_abonos_proveedor
        if debe_mostrar_abonos:
            context['abonos_proveedor'] = venta.abonos_proveedor.all().select_related(
                'solicitud_por', 'aprobado_por', 'confirmado_por', 'cancelado_por'
            ).order_by('-fecha_solicitud')
            # Base para abonos = Servicios planificados (costo_neto en NAC; costo_neto_usd o suma tabla en INT)
            if venta.tipo_viaje == 'INT':
                base_abonos = venta.costo_neto_usd if (getattr(venta, 'costo_neto_usd', None) is not None and venta.costo_neto_usd > 0) else None
                if base_abonos is None or base_abonos <= 0:
                    from django.db.models import Sum
                    from django.db.models.functions import Coalesce
                    suma_logistica = venta.servicios_logisticos.aggregate(
                        total=Coalesce(Sum('monto_planeado'), Decimal('0.00'))
                    )['total']
                    base_abonos = suma_logistica or Decimal('0.00')
                context['servicios_planificados_abonos'] = base_abonos
            else:
                context['servicios_planificados_abonos'] = venta.costo_neto or Decimal('0.00')
            if venta.tipo_viaje == 'INT':
                context['total_abonado_proveedor'] = venta.total_abonado_proveedor
                context['saldo_pendiente_proveedor'] = venta.saldo_pendiente_proveedor
                context['moneda_abonos'] = 'USD'
            else:
                context['total_abonado_proveedor'] = venta.total_abonado_proveedor
                context['saldo_pendiente_proveedor'] = venta.saldo_pendiente_proveedor
                context['moneda_abonos'] = 'MXN'
            from .forms import SolicitarAbonoProveedorForm
            context['form_abono_proveedor'] = SolicitarAbonoProveedorForm(venta=venta, user=self.request.user)
            context['puede_solicitar_abono_proveedor'] = perm.can_solicitar_abono_proveedor(self.request.user, self.request)
            context['puede_aprobar_abono_proveedor'] = perm.can_approve_abono_proveedor(self.request.user, self.request)
            context['puede_confirmar_abono_proveedor'] = perm.can_confirm_abono_proveedor(self.request.user, self.request)
            context['puede_cancelar_abono_proveedor'] = perm.can_cancel_abono_proveedor(self.request.user, self.request)
        else:
            context['abonos_proveedor'] = []
            context['total_abonado_proveedor'] = Decimal('0.00')
            context['saldo_pendiente_proveedor'] = Decimal('0.00')
            context['servicios_planificados_abonos'] = Decimal('0.00')
            context['moneda_abonos'] = 'MXN'
            context['puede_solicitar_abono_proveedor'] = False
            context['puede_aprobar_abono_proveedor'] = False
            context['puede_confirmar_abono_proveedor'] = False
            context['puede_cancelar_abono_proveedor'] = False
        # El template debe usar esta variable (no duplicar la lógica): incluye proveedor en logística
        context['debe_mostrar_abonos'] = debe_mostrar_abonos
        
        # Inicialización del Formulario de Confirmaciones
        context['confirmaciones'] = venta.confirmaciones.select_related('subido_por').order_by('-fecha_subida')
        context['confirmacion_form'] = ConfirmacionVentaForm()
        context['puede_subir_confirmaciones'] = self._puede_subir_confirmaciones(self.request.user, venta)
        
        # Calcular descuentos y totales para el contexto (INT: usar propiedades en USD)
        if venta.tipo_viaje == 'INT':
            total_final = venta.costo_total_con_modificacion
            costo_base = total_final
            total_descuentos = Decimal('0.00')
            descuento_km = Decimal('0.00')
            descuento_promo = Decimal('0.00')
        else:
            descuento_km = venta.descuento_kilometros_mxn or Decimal('0.00')
            descuento_promo = venta.descuento_promociones_mxn or Decimal('0.00')
            total_descuentos = descuento_km + descuento_promo
            costo_base = (venta.costo_venta_final or Decimal('0.00')) + (venta.costo_modificacion or Decimal('0.00'))
            total_final = costo_base - total_descuentos

        context['total_descuentos'] = total_descuentos
        context['costo_base'] = costo_base
        context['total_final'] = total_final
        context['descuento_km'] = descuento_km
        context['descuento_promo'] = descuento_promo

        # Para INT: valores en USD para la sección Desglose (mostrar todo en USD, evitar repetir total)
        # Costo neto en INT viene del formulario en costo_neto_usd (costo_neto se deja en 0)
        if venta.tipo_viaje == 'INT' and venta.tipo_cambio and venta.tipo_cambio > 0:
            tc = venta.tipo_cambio
            context['fin_desglose_usd'] = True
            # Usar costo_neto_usd (dato del formulario nueva venta); fallback costo_neto/tc por legacy
            costo_neto_usd_val = getattr(venta, 'costo_neto_usd', None)
            if costo_neto_usd_val is not None and costo_neto_usd_val > 0:
                context['fin_costo_neto_usd'] = costo_neto_usd_val
            else:
                context['fin_costo_neto_usd'] = (venta.costo_neto or Decimal('0.00')) / tc
            context['fin_costo_base_usd'] = costo_base / tc
            context['fin_total_final_usd'] = total_final / tc
            context['fin_total_descuentos_usd'] = total_descuentos / tc
            context['fin_descuento_km_usd'] = descuento_km / tc
            context['fin_descuento_promo_usd'] = descuento_promo / tc
        else:
            context['fin_desglose_usd'] = False
        
        # ------------------- Contexto de Solicitud de Cancelación -------------------
        solicitud_cancelacion = getattr(venta, 'solicitud_cancelacion', None)
        context['solicitud_cancelacion'] = solicitud_cancelacion
        
        # Determinar si se puede solicitar cancelación
        puede_solicitar = (
            venta.estado == 'ACTIVA' and
            (not solicitud_cancelacion or solicitud_cancelacion.estado in ['RECHAZADA', 'CANCELADA']) and
            (venta.vendedor == self.request.user or perm.has_full_access(self.request.user, self.request))
        )
        context['puede_solicitar_cancelacion'] = puede_solicitar
        
        # Determinar si se puede cancelar definitivamente (requiere solicitud aprobada)
        puede_cancelar_definitivamente = (
            venta.estado == 'ACTIVA' and
            solicitud_cancelacion and
            solicitud_cancelacion.estado == 'APROBADA' and
            (venta.vendedor == self.request.user or perm.has_full_access(self.request.user, self.request))
        )
        context['puede_cancelar_definitivamente'] = puede_cancelar_definitivamente
        
        # Determinar si se puede reciclar (venta cancelada definitivamente)
        puede_reciclar = (
            venta.estado == 'CANCELADA' and
            (venta.vendedor == self.request.user or perm.has_full_access(self.request.user, self.request))
        )
        context['puede_reciclar_venta'] = puede_reciclar
        
        # Determinar si el usuario puede aprobar/rechazar solicitudes (JEFE, Director General, Director Administrativo)
        puede_aprobar_rechazar = (
            solicitud_cancelacion and
            solicitud_cancelacion.estado == 'PENDIENTE' and
            perm.can_approve_reject_cancelacion(self.request.user, self.request)
        )
        context['puede_aprobar_rechazar_cancelacion'] = puede_aprobar_rechazar
        
        # Formulario de solicitud de cancelación
        if puede_solicitar:
            context['solicitud_cancelacion_form'] = SolicitudCancelacionForm()
        
        # ------------------- Resumen Financiero (solo si está cancelada) -------------------
        if venta.estado == 'CANCELADA':
            # ✅ PERFORMANCE: Cargar abonos una sola vez (ya prefetched)
            abonos_list = list(venta.abonos.all().order_by('fecha_pago'))
            resumen_financiero = {
                'total_abonos': len(abonos_list),
                'monto_total_abonos': sum(abono.monto for abono in abonos_list),
                'monto_apertura': venta.cantidad_apertura or Decimal('0.00'),
                'total_pagado': venta.total_pagado,
                'abonos_detalle': [
                    {
                        'fecha': abono.fecha_pago,
                        'monto': abono.monto,
                        'forma_pago': abono.get_forma_pago_display(),
                        'confirmado': abono.confirmado,
                    }
                    for abono in abonos_list
                ],
            }
            context['resumen_financiero'] = resumen_financiero
        else:
            context['resumen_financiero'] = None
        
        return context


# ------------------- 3.1. ELIMINAR CONFIRMACIÓN DE VENTA -------------------

class EliminarConfirmacionView(LoginRequiredMixin, View):
    """Vista para eliminar una confirmación de venta."""
    
    def post(self, request, pk):
        try:
            confirmacion = get_object_or_404(ConfirmacionVenta, pk=pk)
            venta = confirmacion.venta
            
            # Verificar permisos (misma lógica que para subir)
            if not request.user.is_authenticated:
                messages.error(request, "Debes estar autenticado para realizar esta acción.")
                return redirect('detalle_venta', pk=venta.pk, slug=venta.slug_safe)
            
            puede_eliminar = False
            if request.user.is_superuser or request.user == venta.vendedor:
                puede_eliminar = True
            else:
                rol = perm.get_user_role(request.user, request).upper()
                # CONTADOR solo lectura, no puede eliminar confirmaciones
                puede_eliminar = 'JEFE' in rol
            
            if not puede_eliminar:
                messages.error(request, "No tienes permiso para eliminar confirmaciones de esta venta.")
                return redirect('detalle_venta', pk=venta.pk, slug=venta.slug_safe)
            
            # Guardar información antes de eliminar
            nombre_archivo = confirmacion.nombre_archivo
            venta_pk = venta.pk
            venta_slug = venta.slug_safe
            
            # Eliminar el archivo físico si existe
            if confirmacion.archivo:
                try:
                    confirmacion.archivo.delete(save=False)
                except Exception as e:
                    logger.warning(f"No se pudo eliminar el archivo físico: {e}")
            
            # Eliminar el registro
            confirmacion.delete()
            
            messages.success(request, f"Confirmación '{nombre_archivo}' eliminada correctamente.")
            return redirect(reverse('detalle_venta', kwargs={'pk': venta_pk, 'slug': venta_slug}) + '?tab=confirmaciones')
            
        except ConfirmacionVenta.DoesNotExist:
            messages.error(request, "La confirmación no existe.")
            return redirect('dashboard')
        except Exception as e:
            messages.error(request, f"Error al eliminar la confirmación: {str(e)}")
            if 'venta' in locals():
                return redirect(reverse('detalle_venta', kwargs={'pk': venta.pk, 'slug': venta.slug_safe}) + '?tab=confirmaciones')
            return redirect('dashboard')


# ------------------- 4. CREACIÓN Y EDICIÓN DE VENTA -------------------

class VentaViajeCreateView(LoginRequiredMixin, UserPassesTestMixin, CreateView):
    model = VentaViaje
    form_class = VentaViajeForm
    template_name = 'ventas/venta_form.html'
    
    def test_func(self):
        """Solo JEFE y VENDEDOR pueden crear ventas. CONTADOR solo lectura."""
        user_rol = perm.get_user_role(self.request.user, self.request)
        return perm.has_full_access(self.request.user, self.request) or perm.is_vendedor(self.request.user, self.request)
    
    def handle_no_permission(self):
        messages.error(self.request, "No tienes permiso para crear ventas. Solo puedes visualizarlas.")
        return redirect('lista_ventas')
    
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['request'] = self.request
        
        # Pre-llenar valores desde cotización si existe (tiene prioridad)
        cot_slug = self.request.GET.get('cotizacion')
        cotizacion_data = self.request.session.get('cotizacion_convertir', {})
        
        # Pre-seleccionar cliente si viene en la URL (desde detalle del cliente)
        # Solo si no hay cotización (la cotización tiene prioridad)
        cliente_pk = self.request.GET.get('cliente_pk')
        if cliente_pk and not (cot_slug or cotizacion_data):
            if 'initial' not in kwargs:
                kwargs['initial'] = {}
            try:
                cliente = Cliente.objects.get(pk=cliente_pk)
                kwargs['initial']['cliente'] = cliente
            except Cliente.DoesNotExist:
                pass
        
        # Función auxiliar para limpiar y convertir totales
        def limpiar_y_convertir_total(valor):
            """Convierte un string de total (puede tener comas) a Decimal"""
            if not valor:
                return Decimal('0.00')
            try:
                # Remover comas y espacios, luego convertir a Decimal
                valor_limpio = str(valor).replace(',', '').replace('$', '').strip()
                return Decimal(valor_limpio)
            except (ValueError, InvalidOperation):
                return Decimal('0.00')
        
        if cot_slug or cotizacion_data:
            slug_a_usar = cot_slug or cotizacion_data.get('cotizacion_slug')
            if slug_a_usar:
                qs_cot = perm.get_cotizaciones_queryset_base(Cotizacion, self.request.user, self.request)
                cot = qs_cot.filter(slug=slug_a_usar).first()
                if cot:
                    # Obtener el total de la sesión si está disponible (prioridad máxima)
                    total_cotizacion = Decimal('0.00')
                    
                    # 1. Intentar obtener TOTAL de la URL (Prioridad Máxima - Bypass de Sesión)
                    total_url = self.request.GET.get('total_b')
                    if total_url:
                        try:
                            # Limpiar y convertir
                            total_cotizacion = Decimal(str(total_url).replace(',', '').replace('$', ''))
                            logging.debug(f"Total recuperado VÍA URL (Bypass): {total_cotizacion}")
                        except (ValueError, InvalidOperation):
                            pass
                    
                    # 2. Si no vino en URL, intentar sesión (lógica original)
                    if total_cotizacion == Decimal('0.00') and cotizacion_data.get('total_cotizacion'):
                        try:
                            total_cotizacion = Decimal(cotizacion_data['total_cotizacion'])
                            logging.debug(f"Total recuperado de la sesión: {total_cotizacion}")
                        except (ValueError, InvalidOperation):
                            total_cotizacion = Decimal('0.00')
                    
                    # Calcular el costo final - SIEMPRE usar el total (URL tiene prioridad sobre sesión)
                    if total_cotizacion > 0:
                        costo_final = total_cotizacion
                        logging.debug(f"Usando total de la sesión: {costo_final}")
                    else:
                        # Solo si no hay total en la sesión, intentar obtenerlo de las propuestas
                        # Esto solo debería pasar si hay un error en el guardado de la sesión
                        propuestas = cot.propuestas if isinstance(cot.propuestas, dict) else {}
                        tipo_cot = propuestas.get('tipo', '')
                        # Intentar leer índice de la URL primero (Bypass de Sesión)
                        opcion_vuelo_index = self.request.GET.get('idx_b', '') or cotizacion_data.get('opcion_vuelo_index', '')
                        
                        if tipo_cot == 'vuelos' and propuestas.get('vuelos'):
                            vuelos = propuestas.get('vuelos', [])
                            if vuelos and len(vuelos) > 0:
                                # Intentar usar la opción seleccionada de la sesión
                                try:
                                    indice = int(opcion_vuelo_index) if opcion_vuelo_index else 0
                                    if indice < 0 or indice >= len(vuelos):
                                        indice = 0
                                except (ValueError, TypeError):
                                    indice = 0
                                
                                vuelo_seleccionado = vuelos[indice] if isinstance(vuelos, list) else vuelos.get(f'propuesta_{indice + 1}', {})
                                if isinstance(vuelo_seleccionado, dict) and vuelo_seleccionado.get('total'):
                                    total_primero = limpiar_y_convertir_total(vuelo_seleccionado.get('total'))
                                    if total_primero > 0:
                                        costo_final = total_primero
                                        logging.debug(f"Total obtenido de propuestas (índice {indice}): {costo_final}")
                                    else:
                                        costo_final = cot.total_estimado or Decimal('0.00')
                                else:
                                    costo_final = cot.total_estimado or Decimal('0.00')
                            elif tipo_cot == 'hospedaje' and propuestas.get('hoteles'):
                                hoteles = propuestas.get('hoteles', [])
                                if hoteles and len(hoteles) > 0:
                                    # Intentar leer índice de la URL primero (Bypass de Sesión)
                                    opcion_hotel_index = self.request.GET.get('idx_b', '') or cotizacion_data.get('opcion_hotel_index', '')
                                    try:
                                        indice = int(opcion_hotel_index) if opcion_hotel_index else 0
                                        if indice < 0 or indice >= len(hoteles):
                                            indice = 0
                                    except (ValueError, TypeError):
                                        indice = 0
                                    
                                    hotel_seleccionado = hoteles[indice] if isinstance(hoteles, list) else hoteles.get(f'propuesta_{indice + 1}', {})
                                    if isinstance(hotel_seleccionado, dict) and hotel_seleccionado.get('total'):
                                        total_primero = limpiar_y_convertir_total(hotel_seleccionado.get('total'))
                                        if total_primero > 0:
                                            costo_final = total_primero
                                            logging.debug(f"Total obtenido de propuestas hospedaje (get_form_kwargs, índice {indice}): {costo_final}")
                                        else:
                                            costo_final = cot.total_estimado or Decimal('0.00')
                                    else:
                                        costo_final = cot.total_estimado or Decimal('0.00')
                                else:
                                    costo_final = cot.total_estimado or Decimal('0.00')
                            else:
                                costo_final = cot.total_estimado or Decimal('0.00')
                        else:
                            costo_final = cot.total_estimado or Decimal('0.00')
                    
                    # Convertir Decimal a string para TextInput con formato de moneda
                    # El JavaScript aplicará el formato de moneda automáticamente
                    if isinstance(costo_final, Decimal):
                        # Convertir a string sin formato para que el JavaScript lo formatee
                        costo_final_value = str(costo_final)
                    else:
                        costo_final_value = str(costo_final) if costo_final else ''
                    
                    # Obtener edades de menores y servicios desde la sesión
                    edades_menores_valor = cotizacion_data.get('edades_menores', '')
                    # Si no está en la sesión, obtener directamente de la cotización y formatear
                    if not edades_menores_valor and cot.edades_menores and cot.edades_menores.strip():
                        edades = [e.strip() for e in cot.edades_menores.split(',') if e.strip()]
                        if edades:
                            # Crear formato con saltos de línea para cualquier cantidad de menores
                            edades_lista = [f"Menor {i+1} - {edad}" for i, edad in enumerate(edades)]
                            edades_menores_valor = '\n'.join(edades_lista)
                    
                    servicios_seleccionados_valor = cotizacion_data.get('servicios_seleccionados', [])
                    
                    # Establecer valores iniciales antes de inicializar el formulario
                    if 'initial' not in kwargs:
                        kwargs['initial'] = {}
                    kwargs['initial'].update({
                        'cliente': cot.cliente,
                        'fecha_inicio_viaje': cot.fecha_inicio,
                        'fecha_fin_viaje': cot.fecha_fin,
                        'edades_menores': edades_menores_valor,  # Prellenar con edades desde sesión
                        'costo_venta_final': costo_final_value,  # Usar string para TextInput con formato de moneda
                    })
                    
                    # Asegurar que el campo también tenga el valor directamente
                    if edades_menores_valor and 'edades_menores' in kwargs.get('initial', {}):
                        # El valor ya está en initial, Django lo usará automáticamente
                        pass
                    
                    # Prellenar servicios seleccionados en initial también
                    if servicios_seleccionados_valor:
                        kwargs['initial']['servicios_seleccionados'] = servicios_seleccionados_valor
        
        return kwargs

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        cliente = self._obtener_cliente_contextual(context.get('form'))
        context['cliente_preseleccionado'] = cliente
        if cliente:
            context['kilometros_cliente'] = KilometrosService.resumen_cliente(cliente)
        context['kilometros_meta'] = {
            'valor_por_km': KilometrosService.VALOR_PESO_POR_KM,
            'km_por_peso': KilometrosService.KM_POR_PESO,
            'max_porcentaje': KilometrosService.MAX_PORCENTAJE_REDENCION * 100,
        }
        # Prefill desde cotización
        cot_slug = self.request.GET.get('cotizacion')
        cotizacion_data = self.request.session.get('cotizacion_convertir', {})
        
        if cot_slug or cotizacion_data:
            # Usar el slug de la URL o de la sesión
            slug_a_usar = cot_slug or cotizacion_data.get('cotizacion_slug')
            if slug_a_usar:
                qs_cot = perm.get_cotizaciones_queryset_base(Cotizacion, self.request.user, self.request)
                cot = qs_cot.filter(slug=slug_a_usar).first()
                if cot:
                    form = context.get('form')
                    if form:
                        inicial = form.initial.copy()
                        
                        # Función auxiliar para limpiar y convertir totales
                        def limpiar_y_convertir_total(valor):
                            """Convierte un string de total (puede tener comas) a Decimal"""
                            if not valor:
                                return Decimal('0.00')
                            try:
                                # Remover comas y espacios, luego convertir a Decimal
                                valor_limpio = str(valor).replace(',', '').replace('$', '').strip()
                                return Decimal(valor_limpio)
                            except (ValueError, InvalidOperation):
                                return Decimal('0.00')
                        
                        # Obtener el total de la sesión si está disponible
                        total_cotizacion = Decimal('0.00')
                        if cotizacion_data.get('total_cotizacion'):
                            try:
                                total_cotizacion = Decimal(cotizacion_data['total_cotizacion'])
                            except (ValueError, InvalidOperation):
                                total_cotizacion = Decimal('0.00')
                        
                        # Calcular el costo final - SIEMPRE usar el total de la sesión si está disponible
                        if total_cotizacion > 0:
                            costo_final = total_cotizacion
                            logging.debug(f"Usando total de la sesión (get_context_data): {costo_final}")
                        else:
                            # Solo si no hay total en la sesión, intentar obtenerlo de las propuestas
                            # Esto solo debería pasar si hay un error en el guardado de la sesión
                            propuestas = cot.propuestas if isinstance(cot.propuestas, dict) else {}
                            tipo_cot = propuestas.get('tipo', '')
                            # Intentar leer índice de la URL primero (Bypass de Sesión)
                            opcion_vuelo_index = self.request.GET.get('idx_b', '') or cotizacion_data.get('opcion_vuelo_index', '')
                            
                            if tipo_cot == 'vuelos' and propuestas.get('vuelos'):
                                vuelos = propuestas.get('vuelos', [])
                                if vuelos and len(vuelos) > 0:
                                    # Intentar usar la opción seleccionada de la sesión
                                    try:
                                        indice = int(opcion_vuelo_index) if opcion_vuelo_index else 0
                                        if indice < 0 or indice >= len(vuelos):
                                            indice = 0
                                    except (ValueError, TypeError):
                                        indice = 0
                                    
                                    vuelo_seleccionado = vuelos[indice] if isinstance(vuelos, list) else vuelos.get(f'propuesta_{indice + 1}', {})
                                    if isinstance(vuelo_seleccionado, dict) and vuelo_seleccionado.get('total'):
                                        total_primero = limpiar_y_convertir_total(vuelo_seleccionado.get('total'))
                                        if total_primero > 0:
                                            costo_final = total_primero
                                            logging.debug(f"Total obtenido de propuestas (get_context_data, índice {indice}): {costo_final}")
                                        else:
                                            costo_final = cot.total_estimado or Decimal('0.00')
                                    else:
                                        costo_final = cot.total_estimado or Decimal('0.00')
                                else:
                                    costo_final = cot.total_estimado or Decimal('0.00')
                            elif tipo_cot == 'hospedaje' and propuestas.get('hoteles'):
                                hoteles = propuestas.get('hoteles', [])
                                if hoteles and len(hoteles) > 0:
                                    # Intentar leer índice de la URL primero (Bypass de Sesión)
                                    opcion_hotel_index = self.request.GET.get('idx_b', '') or cotizacion_data.get('opcion_hotel_index', '')
                                    try:
                                        indice = int(opcion_hotel_index) if opcion_hotel_index else 0
                                        if indice < 0 or indice >= len(hoteles):
                                            indice = 0
                                    except (ValueError, TypeError):
                                        indice = 0
                                    
                                    hotel_seleccionado = hoteles[indice] if isinstance(hoteles, list) else hoteles.get(f'propuesta_{indice + 1}', {})
                                    if isinstance(hotel_seleccionado, dict) and hotel_seleccionado.get('total'):
                                        total_primero = limpiar_y_convertir_total(hotel_seleccionado.get('total'))
                                        if total_primero > 0:
                                            costo_final = total_primero
                                            logging.debug(f"Total obtenido de propuestas hospedaje (get_context_data, índice {indice}): {costo_final}")
                                        else:
                                            costo_final = cot.total_estimado or Decimal('0.00')
                                    else:
                                        costo_final = cot.total_estimado or Decimal('0.00')
                                else:
                                    costo_final = cot.total_estimado or Decimal('0.00')
                            else:
                                costo_final = cot.total_estimado or Decimal('0.00')
                        
                        # Convertir Decimal a string para TextInput con formato de moneda
                        # El JavaScript aplicará el formato de moneda automáticamente
                        if isinstance(costo_final, Decimal):
                            # Convertir a string sin formato para que el JavaScript lo formatee
                            costo_final_value = str(costo_final)
                        else:
                            costo_final_value = str(costo_final) if costo_final else ''
                        
                        # Obtener edades de menores y servicios desde la sesión
                        edades_menores_valor = cotizacion_data.get('edades_menores', '')
                        # Si no está en la sesión, obtener directamente de la cotización y formatear
                        if not edades_menores_valor and cot.edades_menores and cot.edades_menores.strip():
                            edades = [e.strip() for e in cot.edades_menores.split(',') if e.strip()]
                            if edades:
                                # Crear formato con saltos de línea para cualquier cantidad de menores
                                edades_lista = [f"Menor {i+1} - {edad}" for i, edad in enumerate(edades)]
                                edades_menores_valor = '\n'.join(edades_lista)
                        
                        servicios_seleccionados_valor = cotizacion_data.get('servicios_seleccionados', [])
                        
                        inicial.update({
                            'cliente': cot.cliente,
                            'fecha_inicio_viaje': cot.fecha_inicio,
                            'fecha_fin_viaje': cot.fecha_fin,
                            'pasajeros': '',  # Campo vacío, se llena manualmente
                            'edades_menores': edades_menores_valor,  # Prellenar con edades desde sesión
                            'costo_venta_final': costo_final_value,  # Usar string para TextInput con formato de moneda
                        })
                        
                        # Establecer initial primero
                        form.initial = inicial
                        
                        # Luego asegurar que los campos específicos tengan el valor directamente
                        if edades_menores_valor and 'edades_menores' in form.fields:
                            form.fields['edades_menores'].initial = edades_menores_valor
                        
                        # Prellenar servicios seleccionados
                        if servicios_seleccionados_valor and 'servicios_seleccionados' in form.fields:
                            form.fields['servicios_seleccionados'].initial = servicios_seleccionados_valor
                        
                        # Asegurar que el valor también se establezca en el campo directamente
                        if 'costo_venta_final' in form.fields:
                            form.fields['costo_venta_final'].initial = costo_final_value
                            # También establecer el valor en el widget para asegurar que se muestre
                            if hasattr(form.fields['costo_venta_final'], 'widget'):
                                form.fields['costo_venta_final'].widget.attrs['value'] = str(costo_final_value)
                        
                        context['form'] = form
                    context['cotizacion_origen'] = cot
        return context

    def _obtener_cliente_contextual(self, form=None):
        cliente_pk = self.request.GET.get('cliente_pk') or self.request.POST.get('cliente')
        cliente = None
        if cliente_pk:
            cliente = Cliente.objects.filter(pk=cliente_pk).first()
        elif form and form.initial.get('cliente'):
            inicial = form.initial.get('cliente')
            if isinstance(inicial, Cliente):
                cliente = inicial
            else:
                cliente = Cliente.objects.filter(pk=inicial).first()
        return cliente
    
    def form_valid(self, form):
        # 1. Guarda temporalmente la instancia sin enviarla a la base de datos (commit=False)
        instance = form.save(commit=False)
        
        # 2. Asigna el vendedor (que es el usuario logueado)
        instance.vendedor = self.request.user
        
        # Si viene de cotización, enlazar (verificar GET parameter o sesión). Solo cotizaciones propias del vendedor.
        cot = None
        qs_cot = perm.get_cotizaciones_queryset_base(Cotizacion, self.request.user, self.request)
        cot_slug = self.request.GET.get('cotizacion')
        if cot_slug:
            cot = qs_cot.filter(slug=cot_slug).first()
        else:
            # Verificar si hay datos en la sesión
            cotizacion_data = self.request.session.get('cotizacion_convertir', {})
            if cotizacion_data.get('cotizacion_id'):
                try:
                    cot = qs_cot.filter(pk=cotizacion_data['cotizacion_id']).first()
                    # Limpiar la sesión después de usarla
                    if 'cotizacion_convertir' in self.request.session:
                        del self.request.session['cotizacion_convertir']
                except Exception as e:
                    logger.warning(f"Error al limpiar sesión de cotización: {e}")
        
        if cot:
            instance.cotizacion_origen = cot

        # 3. ¡IMPORTANTE! Eliminamos la lógica manual de generación de slug de aquí.
        # El modelo VentaViaje se encarga de generar y asegurar el slug único 
        # dentro de su método save() una vez que la PK ha sido asignada.
        
        # 4. Guarda la instancia, lo que dispara el método save() del modelo
        self.object = instance # Establece self.object para que get_success_url funcione
        self.object.save()
        
        # 4.1. Actualizar el estado de la cotización a CONVERTIDA cuando se crea la venta
        if cot:
            cot.estado = 'CONVERTIDA'
            cot.save(update_fields=['estado']) 
        
        # 5. Llama a save_m2m (necesario si hay campos ManyToMany en VentaViajeForm)
        form.save_m2m() 
        
        # 5.1. KILÓMETROS MOVUMS: Solo redimir (si aplica) al crear la venta
        # IMPORTANTE: La acumulación de kilómetros y bonos de promociones se aplicará
        # SOLO cuando la venta se liquide (a través del signal aplicar_promociones_al_liquidar)
        try:
            # PRIMERO: Redimir kilómetros si se aplicó descuento (esto SÍ se hace al crear la venta)
            if self.object.aplica_descuento_kilometros and self.object.descuento_kilometros_mxn > 0:
                # Calcular kilómetros a redimir: descuento_mxn / valor_por_km (0.05)
                km_a_redimir = (self.object.descuento_kilometros_mxn / KilometrosService.VALOR_PESO_POR_KM).quantize(Decimal('0.01'))
                if km_a_redimir > 0:
                    registro_redencion = KilometrosService.redimir(
                        cliente=self.object.cliente,
                        kilometros=km_a_redimir,
                        venta=self.object,
                        descripcion=f"Redención aplicada a venta #{self.object.pk}: ${self.object.descuento_kilometros_mxn:,.2f} MXN"
                    )
                    if registro_redencion:
                        logger.info(f"✅ Kilómetros redimidos para venta {self.object.pk}: {km_a_redimir} km (${self.object.descuento_kilometros_mxn:,.2f} MXN)")
                    else:
                        logger.warning(f"⚠️ No se pudieron redimir kilómetros para venta {self.object.pk} (posible saldo insuficiente)")
            
            # NOTA: La acumulación de kilómetros por compra y bonos de promociones
            # se aplicará automáticamente cuando la venta se liquide (a través del signal)
            logger.info(
                f"ℹ️ Promociones de kilómetros se aplicarán cuando la venta {self.object.pk} se liquide completamente."
            )
        except Exception:
            logger.exception(
                f"❌ Error procesando redención de kilómetros para la venta {self.object.pk} "
                f"(Cliente: {self.object.cliente.pk if self.object.cliente else 'N/A'})"
            )
    
        # 5.1. Lógica de notificaciones para apertura con Transferencia/Tarjeta
        # ✅ NUEVO FLUJO: NO se crean notificaciones automáticas
        # Las notificaciones se crearán solo cuando se suba el comprobante
        modo_pago = form.cleaned_data.get('modo_pago_apertura', 'EFE')
        cantidad_apertura = form.cleaned_data.get('cantidad_apertura', Decimal('0.00'))
        
        # ✅ NUEVO: Lógica para "Directo a Proveedor" (PRO)
        if modo_pago == 'PRO':
            # Calcular el total final con descuentos
            costo_base = (self.object.costo_venta_final or Decimal('0.00')) + (self.object.costo_modificacion or Decimal('0.00'))
            descuento_km = self.object.descuento_kilometros_mxn or Decimal('0.00')
            descuento_promo = self.object.descuento_promociones_mxn or Decimal('0.00')
            total_descuentos = descuento_km + descuento_promo
            total_final = costo_base - total_descuentos
            
            # Establecer cantidad_apertura igual al total final y estado como COMPLETADO
            self.object.cantidad_apertura = total_final
            self.object.estado_confirmacion = 'COMPLETADO'
            self.object.save(update_fields=['cantidad_apertura', 'estado_confirmacion'])
            
            # Crear notificación para el VENDEDOR (si existe)
            if self.object.vendedor:
                mensaje_vendedor_apertura = f"Venta #{self.object.pk} marcada como pagada (Directo a Proveedor): ${total_final:,.2f} - Cliente: {self.object.cliente.nombre_completo_display}"
                Notificacion.objects.create(
                    usuario=self.object.vendedor,
                    tipo='APERTURA',
                    mensaje=mensaje_vendedor_apertura,
                    venta=self.object,
                    confirmado=True
                )
        elif modo_pago == 'CRE':
            # Para crédito: estado EN_CONFIRMACION y cantidad_apertura = 0
            self.object.estado_confirmacion = 'EN_CONFIRMACION'
            self.object.cantidad_apertura = Decimal('0.00')
            self.object.save(update_fields=['estado_confirmacion', 'cantidad_apertura'])
        elif cantidad_apertura > 0 and modo_pago in ['TRN', 'TAR', 'DEP']:
            # Cambiar estado a "En confirmación"
            self.object.estado_confirmacion = 'EN_CONFIRMACION'
            self.object.save(update_fields=['estado_confirmacion'])
        elif cantidad_apertura > 0 and modo_pago == 'EFE':
            # Si es efectivo, se marca como completado automáticamente
            self.object.estado_confirmacion = 'COMPLETADO'
            self.object.save(update_fields=['estado_confirmacion'])
            
            # Crear notificación para el VENDEDOR (si existe)
            if self.object.vendedor:
                mensaje_vendedor_apertura = f"Apertura registrada y confirmada en tu venta #{self.object.pk}: ${cantidad_apertura:,.2f} (Efectivo) - Cliente: {self.object.cliente.nombre_completo_display}"
                Notificacion.objects.create(
                    usuario=self.object.vendedor,
                    tipo='APERTURA',
                    mensaje=mensaje_vendedor_apertura,
                    venta=self.object,
                    confirmado=True
                )
    
        messages.success(self.request, "Venta creada exitosamente. ¡No olvides gestionar la logística!")
        
        # 6. Retorna la respuesta de redirección usando la URL de éxito
        return HttpResponseRedirect(self.get_success_url())

    def get_success_url(self):
        # Redirección correcta a la vista de detalle (AHORA CON SLUG)
        return reverse_lazy('detalle_venta', kwargs={'pk': self.object.pk, 'slug': self.object.slug_safe})


class ClienteKilometrosResumenView(LoginRequiredMixin, View):
    """Devuelve el resumen de Kilómetros Movums del cliente para el formulario de ventas."""

    def get(self, request, cliente_id, *args, **kwargs):
        # SEGURIDAD: Validar acceso al cliente (prevenir IDOR)
        if not self._user_can_access_cliente(request.user, cliente_id):
            return JsonResponse({
                'success': False,
                'message': 'No tienes permiso para ver los datos de este cliente.'
            }, status=403)
        
        try:
            cliente = get_object_or_404(Cliente, pk=cliente_id)
            resumen = KilometrosService.resumen_cliente(cliente)
        except Exception as e:
            logger.error(f"Error al obtener resumen de kilómetros para cliente {cliente_id}: {str(e)}")
            return JsonResponse({
                'success': False,
                'message': 'Error al obtener los datos de Kilómetros.'
            }, status=500)

        if not resumen:
            return JsonResponse({
                'success': False,
                'message': 'No se encontraron datos de Kilómetros para este cliente.'
            }, status=404)

        def decimal_to_str(value):
            if value is None:
                return "0.00"
            return format(value, '.2f')

        data = {
            'cliente': cliente.nombre_completo_display,
            'participa': resumen.get('participa', False),
            'disponible': decimal_to_str(resumen.get('disponible')),
            'valor_equivalente': decimal_to_str(resumen.get('valor_equivalente')),
            'total_acumulado': decimal_to_str(resumen.get('total_acumulado')),
            'ultima_fecha': resumen.get('ultima_fecha').isoformat() if resumen.get('ultima_fecha') else None,
        }
        return JsonResponse({'success': True, 'data': data})

    def _user_can_access_cliente(self, user, cliente_id):
        """
        SEGURIDAD: Verifica si el usuario tiene permiso para ver datos del cliente.
        - JEFE y ADMIN pueden ver cualquier cliente
        - VENDEDOR solo puede ver clientes con los que tiene ventas asociadas
        """
        if not hasattr(user, 'perfil'):
            return False
        
        rol = user.perfil.rol
        
        # JEFE y ADMIN tienen acceso total
        if rol in ['JEFE', 'ADMIN']:
            return True
        
        # VENDEDOR: verificar que tenga ventas con este cliente
        if rol == 'VENDEDOR' and hasattr(user.perfil, 'ejecutivo'):
            return VentaViaje.objects.filter(
                cliente_id=cliente_id,
                vendedor=user.perfil.ejecutivo
            ).exists()
        
        # CONTADOR puede ver datos de kilómetros (solo lectura)
        if rol == 'CONTADOR':
            return True
        
        return False

class VentaViajeUpdateView(LoginRequiredMixin, UserPassesTestMixin, UpdateView):
    model = VentaViaje
    form_class = VentaViajeForm
    template_name = 'ventas/venta_form.html' # Usar el template de formulario si es UpdateView

    def test_func(self):
        # Botón "Editar datos del viaje": exclusivo de Gerente y los 3 directores (General, Administrativo, Ventas).
        # JEFE, Contador y Vendedor no pueden acceder a esta vista.
        return perm.can_edit_datos_viaje(self.request.user, self.request)
    
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['request'] = self.request
        return kwargs
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        venta = self.object
        user_rol = perm.get_user_role(self.request.user, self.request)
        
        # Contexto de solicitud de cancelación
        solicitud_cancelacion = getattr(venta, 'solicitud_cancelacion', None)
        context['solicitud_cancelacion'] = solicitud_cancelacion
        
        # Siempre mostrar el formulario si la venta está activa (la validación se hace en la vista)
        if venta.estado == 'ACTIVA':
            context['solicitud_cancelacion_form'] = SolicitudCancelacionForm()
        
        # Variable para verificar permisos (usada en la vista para validar, no en el template)
        puede_solicitar = (
            venta.estado == 'ACTIVA' and
            (not solicitud_cancelacion or solicitud_cancelacion.estado in ['RECHAZADA', 'CANCELADA']) and
            (venta.vendedor == self.request.user or perm.has_full_access(self.request.user, self.request))
        )
        context['puede_solicitar_cancelacion'] = puede_solicitar
        
        return context

    def handle_no_permission(self):
        pk = self.kwargs.get('pk')
        venta = get_object_or_404(VentaViaje, pk=pk) if pk else None
        messages.error(self.request, "Solo JEFE, Gerente y los directores pueden editar los datos del viaje.")
        if venta:
            return HttpResponseRedirect(reverse_lazy('detalle_venta', kwargs={'pk': venta.pk, 'slug': venta.slug_safe}))
        return redirect('lista_ventas')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        cliente = self.object.cliente if self.object else None
        context['cliente_preseleccionado'] = cliente
        if cliente:
            context['kilometros_cliente'] = KilometrosService.resumen_cliente(cliente)
        context['kilometros_meta'] = {
            'valor_por_km': KilometrosService.VALOR_PESO_POR_KM,
            'km_por_peso': KilometrosService.KM_POR_PESO,
            'max_porcentaje': KilometrosService.MAX_PORCENTAJE_REDENCION * 100,
        }
        return context

    def form_valid(self, form):
        """
        Maneja la validación del formulario y actualiza el costo total
        incluyendo el costo de modificación si se proporciona.
        """
        # Obtener valores anteriores ANTES de guardar
        venta_anterior = VentaViaje.objects.get(pk=form.instance.pk)
        aplica_descuento_anterior = venta_anterior.aplica_descuento_kilometros
        descuento_anterior = venta_anterior.descuento_kilometros_mxn or Decimal('0.00')
        
        # Obtener el costo de modificación del formulario
        costo_modificacion = form.cleaned_data.get('costo_modificacion', Decimal('0.00')) or Decimal('0.00')
        previo_modificacion = form.instance.costo_modificacion or Decimal('0.00')
        # SEGURIDAD: Usar permisos en lugar de username hardcodeado
        puede_ajustar_modificacion = self.request.user.is_superuser or perm.has_full_access(self.request.user, self.request)

        # Guardar la instancia
        self.object = form.save()

        # Protección de campos bloqueados: si el usuario no puede editar restringidos,
        # restaurar desde venta_anterior los campos financieros que no debe modificar
        puede_editar_bloqueados = perm.can_edit_campos_bloqueados(self.request.user, self.request)
        if not puede_editar_bloqueados:
            CAMPOS_RESTRINGIDOS = ['costo_neto', 'costo_venta_final', 'cantidad_apertura', 'costo_modificacion']
            CAMPOS_RESTRINGIDOS_INT = [
                'tarifa_base_usd', 'impuestos_usd', 'suplementos_usd', 'tours_usd',
                'tipo_cambio', 'cantidad_apertura_usd', 'costo_venta_final_usd',
                'costo_neto_usd', 'costo_modificacion_usd',
            ]
            for field_name in CAMPOS_RESTRINGIDOS:
                setattr(self.object, field_name, getattr(venta_anterior, field_name))
            if self.object.tipo_viaje == 'INT':
                for field_name in CAMPOS_RESTRINGIDOS_INT:
                    setattr(self.object, field_name, getattr(venta_anterior, field_name))
            update_fields = CAMPOS_RESTRINGIDOS + (CAMPOS_RESTRINGIDOS_INT if self.object.tipo_viaje == 'INT' else [])
            self.object.save(update_fields=update_fields)

        # Obtener valores nuevos DESPUÉS de guardar
        aplica_descuento_nuevo = self.object.aplica_descuento_kilometros
        descuento_nuevo = self.object.descuento_kilometros_mxn or Decimal('0.00')

        mensaje = "Venta actualizada correctamente."
        if puede_ajustar_modificacion:
            # JEFE/ADMIN: el valor del campo reemplaza el acumulado (puede restar/ajustar)
            self.object.costo_modificacion = max(Decimal('0.00'), costo_modificacion)
            self.object.save(update_fields=['costo_modificacion'])
            mensaje = f"Venta actualizada. Costo de modificación ajustado a ${self.object.costo_modificacion:,.2f}."
        elif costo_modificacion > 0:
            total_mod = previo_modificacion + costo_modificacion
            self.object.costo_modificacion = total_mod
            self.object.save(update_fields=['costo_modificacion'])
            mensaje = (
                f"Venta actualizada correctamente. "
                f"Costo adicional por modificación: ${costo_modificacion:,.2f}. "
                f"Total acumulado en modificaciones: ${total_mod:,.2f}."
            )
        else:
            if self.object.costo_modificacion != previo_modificacion:
                self.object.costo_modificacion = previo_modificacion
                self.object.save(update_fields=['costo_modificacion'])
        
        # Manejar cambios en descuento de kilómetros
        if self.object.cliente and self.object.cliente.participa_kilometros:
            try:
                # Si se aplicó descuento por primera vez o aumentó el descuento
                if aplica_descuento_nuevo and descuento_nuevo > 0:
                    if not aplica_descuento_anterior:
                        # Primera vez que se aplica descuento
                        km_a_redimir = (descuento_nuevo / KilometrosService.VALOR_PESO_POR_KM).quantize(Decimal('0.01'))
                        if km_a_redimir > 0:
                            registro_redencion = KilometrosService.redimir(
                                cliente=self.object.cliente,
                                kilometros=km_a_redimir,
                                venta=self.object,
                                descripcion=f"Redención aplicada a venta #{self.object.pk}: ${descuento_nuevo:,.2f} MXN"
                            )
                            if registro_redencion:
                                logger.info(f"✅ Kilómetros redimidos en actualización de venta {self.object.pk}: {km_a_redimir} km")
                    elif descuento_nuevo != descuento_anterior:
                        # El descuento cambió
                        diferencia = descuento_nuevo - descuento_anterior
                        if diferencia > 0:
                            # Descuento aumentó, redimir kilómetros adicionales
                            km_diferencia = (diferencia / KilometrosService.VALOR_PESO_POR_KM).quantize(Decimal('0.01'))
                            if km_diferencia > 0:
                                registro_redencion = KilometrosService.redimir(
                                    cliente=self.object.cliente,
                                    kilometros=km_diferencia,
                                    venta=self.object,
                                    descripcion=f"Redención adicional aplicada a venta #{self.object.pk}: ${diferencia:,.2f} MXN"
                                )
                                if registro_redencion:
                                    logger.info(f"✅ Kilómetros adicionales redimidos en actualización de venta {self.object.pk}: {km_diferencia} km")
                        # Si diferencia < 0, no podemos "devolver" kilómetros ya redimidos
            except Exception:
                logger.exception("❌ Error procesando redención de kilómetros en actualización de venta %s", self.object.pk)
        
        # Manejar cambios en promociones con bonos de kilómetros
        # IMPORTANTE: Los bonos SOLO se aplican cuando la venta se liquida.
        # Si la venta ya está liquidada y se agregan/modifican promociones, aplicar inmediatamente.
        # Si la venta NO está liquidada, los bonos se aplicarán automáticamente cuando se liquide (signal)
        if self.object.cliente and self.object.cliente.participa_kilometros:
            try:
                # Verificar si la venta está liquidada
                venta_liquidada = (
                    self.object.estado_confirmacion == 'COMPLETADO' or
                    self.object.total_pagado >= self.object.costo_total_con_modificacion or
                    self.object.esta_pagada
                )
                
                if venta_liquidada:
                    # Si está liquidada, verificar si ya se aplicaron bonos para promociones nuevas
                    # El signal aplicará automáticamente todas las promociones pendientes cuando se liquide
                    # Aquí solo verificamos si hay promociones nuevas que no se han aplicado aún
                    from crm.models import HistorialKilometros
                    
                    # Obtener promociones actuales con bonos
                    promociones_aplicadas_actuales_qs = self.object.promociones_aplicadas.filter(km_bono__gt=0)
                    
                    for vpa in promociones_aplicadas_actuales_qs:
                        # Verificar si ya se aplicó el bono para esta promoción específica
                        bono_ya_aplicado = HistorialKilometros.objects.filter(
                            venta=self.object,
                            tipo_evento='BONO_PROMOCION',
                            kilometros=vpa.km_bono
                        ).exists()
                        
                        if not bono_ya_aplicado and vpa.km_bono > 0:
                            # Bono nuevo que aún no se ha aplicado, aplicar ahora
                            registro_bono = KilometrosService.acumular_bono_promocion(
                                cliente=self.object.cliente,
                                kilometros=vpa.km_bono,
                                venta=self.object,
                                promocion=vpa.promocion,
                                descripcion=f"Bono de promoción al liquidar: {vpa.nombre_promocion or vpa.promocion.nombre}"
                            )
                            if registro_bono:
                                logger.info(
                                    f"✅ Bono de promoción acumulado en venta liquidada {self.object.pk}: "
                                    f"{vpa.km_bono} km ({vpa.nombre_promocion or vpa.promocion.nombre}, Cliente: {self.object.cliente.pk})"
                                )
                else:
                    # Venta no liquidada: los bonos se aplicarán cuando se liquide (a través del signal)
                    logger.info(
                        f"ℹ️ Bonos de promociones para venta {self.object.pk} se aplicarán cuando la venta se liquide."
                    )
            except Exception:
                logger.exception("❌ Error procesando bonos de promociones en actualización de venta %s", self.object.pk)
        
        # ✅ NUEVO: Lógica para "Directo a Proveedor" (PRO) en actualización
        modo_pago = form.cleaned_data.get('modo_pago_apertura', self.object.modo_pago_apertura)
        modo_pago_anterior = venta_anterior.modo_pago_apertura
        
        # Si se cambió a PRO o ya era PRO, aplicar la lógica
        if modo_pago == 'PRO':
            # Calcular el total final con descuentos
            costo_base = (self.object.costo_venta_final or Decimal('0.00')) + (self.object.costo_modificacion or Decimal('0.00'))
            descuento_km = self.object.descuento_kilometros_mxn or Decimal('0.00')
            descuento_promo = self.object.descuento_promociones_mxn or Decimal('0.00')
            total_descuentos = descuento_km + descuento_promo
            total_final = costo_base - total_descuentos
            
            # Establecer cantidad_apertura igual al total final y estado como COMPLETADO
            self.object.cantidad_apertura = total_final
            self.object.estado_confirmacion = 'COMPLETADO'
            self.object.save(update_fields=['cantidad_apertura', 'estado_confirmacion'])
            
            # Si cambió de otro modo a PRO, crear notificación
            if modo_pago_anterior != 'PRO' and self.object.vendedor:
                mensaje_vendedor = f"Venta #{self.object.pk} actualizada a pagada (Directo a Proveedor): ${total_final:,.2f} - Cliente: {self.object.cliente.nombre_completo_display}"
                Notificacion.objects.create(
                    usuario=self.object.vendedor,
                    tipo='APERTURA',
                    mensaje=mensaje_vendedor,
                    venta=self.object,
                    confirmado=True
                )
        
        # Si tiene apertura TRN/TAR/DEP y aún no la confirmó el contador, mantener EN_CONFIRMACION
        cantidad_apertura = getattr(self.object, 'cantidad_apertura', None) or Decimal('0.00')
        if (modo_pago in ['TRN', 'TAR', 'DEP'] and cantidad_apertura > 0 and
                not getattr(self.object, 'apertura_confirmada', False)):
            if self.object.estado_confirmacion != 'COMPLETADO':
                self.object.estado_confirmacion = 'EN_CONFIRMACION'
                self.object.save(update_fields=['estado_confirmacion'])
        
        self.object.actualizar_estado_financiero()
        messages.success(self.request, mensaje)
        return super().form_valid(form)
    
    def get_success_url(self):
        # Se asegura de usar 'detalle_venta' para la redirección de éxito (AHORA CON SLUG)
        return reverse_lazy('detalle_venta', kwargs={'pk': self.object.pk, 'slug': self.object.slug_safe})


# ------------------- VISTAS DE SOLICITUD DE CANCELACIÓN -------------------

class SolicitarCancelacionView(LoginRequiredMixin, UserPassesTestMixin, View):
    """Vista para solicitar la cancelación de una venta."""
    
    def test_func(self):
        """Solo el vendedor que creó la venta o el JEFE pueden solicitar cancelación."""
        venta = get_object_or_404(VentaViaje, pk=self.kwargs['pk'])
        user_rol = perm.get_user_role(self.request.user, self.request)
        if perm.is_contador(self.request.user, self.request):
            return False
        # No se puede solicitar si ya hay una solicitud pendiente o aprobada
        if hasattr(venta, 'solicitud_cancelacion'):
            solicitud = venta.solicitud_cancelacion
            if solicitud.estado in ['PENDIENTE', 'APROBADA']:
                return False
        return venta.vendedor == self.request.user or user_rol == 'JEFE'
    
    def handle_no_permission(self):
        messages.error(self.request, "No tienes permiso para solicitar la cancelación de esta venta.")
        return redirect('detalle_venta', pk=self.kwargs['pk'], slug=get_object_or_404(VentaViaje, pk=self.kwargs['pk']).slug_safe)
    
    def post(self, request, *args, **kwargs):
        venta = get_object_or_404(VentaViaje, pk=self.kwargs['pk'])
        
        # Verificar que no existe una solicitud pendiente o aprobada
        if hasattr(venta, 'solicitud_cancelacion'):
            solicitud_existente = venta.solicitud_cancelacion
            if solicitud_existente.estado in ['PENDIENTE', 'APROBADA']:
                messages.warning(request, "Ya existe una solicitud de cancelación pendiente o aprobada para esta venta.")
                return redirect('detalle_venta', pk=venta.pk, slug=venta.slug_safe)
        
        form = SolicitudCancelacionForm(request.POST)
        if form.is_valid():
            solicitud = form.save(commit=False)
            solicitud.venta = venta
            solicitud.solicitado_por = request.user
            solicitud.estado = 'PENDIENTE'
            solicitud.save()
            
            # Notificar a quienes pueden aprobar/rechazar: JEFE, Director General, Director Administrativo
            usuarios_autorizados = User.objects.filter(
                perfil__rol__in=['JEFE', 'DIRECTOR_GENERAL', 'DIRECTOR_ADMINISTRATIVO']
            ).distinct()
            mensaje = (
                f"Solicitud de cancelación de venta #{venta.folio or venta.pk} - "
                f"Cliente: {venta.cliente.nombre_completo_display} - "
                f"Vendedor: {venta.vendedor.get_full_name() or venta.vendedor.username if venta.vendedor else 'N/A'}\n\n"
                f"Motivo: {solicitud.motivo}"
            )
            for usuario in usuarios_autorizados:
                Notificacion.objects.create(
                    usuario=usuario,
                    tipo='SOLICITUD_CANCELACION',
                    mensaje=mensaje,
                    venta=venta,
                    solicitud_cancelacion=solicitud,
                    vista=False
                )
            
            messages.success(request, f"Solicitud de cancelación enviada. El director administrativo revisará tu solicitud.")
            logger.info(f"✅ Solicitud de cancelación creada para venta {venta.pk} por {request.user.username}")
        else:
            messages.error(request, "Error al enviar la solicitud. Por favor, verifica el formulario.")
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f"{field}: {error}")
        
        return redirect('detalle_venta', pk=venta.pk, slug=venta.slug_safe)


class AprobarCancelacionView(LoginRequiredMixin, UserPassesTestMixin, View):
    """Vista para aprobar una solicitud de cancelación (JEFE, Director General, Director Administrativo)."""
    
    def test_func(self):
        return perm.can_approve_reject_cancelacion(self.request.user, self.request)
    
    def handle_no_permission(self):
        messages.error(self.request, "No tienes permiso para aprobar solicitudes de cancelación.")
        solicitud = get_object_or_404(SolicitudCancelacion, pk=self.kwargs.get('pk'))
        return redirect('detalle_venta', pk=solicitud.venta.pk, slug=solicitud.venta.slug_safe)
    
    def post(self, request, *args, **kwargs):
        solicitud = get_object_or_404(SolicitudCancelacion, pk=self.kwargs['pk'])
        venta = solicitud.venta
        
        if solicitud.estado != 'PENDIENTE':
            messages.error(request, "Esta solicitud ya fue procesada.")
            return redirect('detalle_venta', pk=venta.pk, slug=venta.slug_safe)
        
        solicitud.estado = 'APROBADA'
        solicitud.aprobado_por = request.user
        solicitud.fecha_aprobacion = timezone.now()
        solicitud.save()
        
        # Crear notificación para el vendedor
        if venta.vendedor:
            mensaje = (
                f"Tu solicitud de cancelación para la venta #{venta.folio or venta.pk} "
                f"ha sido aprobada. Ahora puedes proceder con la cancelación definitiva."
            )
            Notificacion.objects.create(
                usuario=venta.vendedor,
                tipo='CANCELACION_APROBADA',
                mensaje=mensaje,
                venta=venta,
                solicitud_cancelacion=solicitud,
                vista=False
            )
        
        messages.success(request, f"Solicitud de cancelación aprobada. El vendedor podrá cancelar la venta definitivamente.")
        logger.info(f"✅ Solicitud de cancelación {solicitud.pk} aprobada por {request.user.username}")
        
        return redirect('detalle_venta', pk=venta.pk, slug=venta.slug_safe)


class RechazarCancelacionView(LoginRequiredMixin, UserPassesTestMixin, View):
    """Vista para rechazar una solicitud de cancelación (JEFE, Director General, Director Administrativo)."""
    
    def test_func(self):
        return perm.can_approve_reject_cancelacion(self.request.user, self.request)
    
    def handle_no_permission(self):
        messages.error(self.request, "No tienes permiso para rechazar solicitudes de cancelación.")
        solicitud = get_object_or_404(SolicitudCancelacion, pk=self.kwargs.get('pk'))
        return redirect('detalle_venta', pk=solicitud.venta.pk, slug=solicitud.venta.slug_safe)
    
    def post(self, request, *args, **kwargs):
        solicitud = get_object_or_404(SolicitudCancelacion, pk=self.kwargs['pk'])
        venta = solicitud.venta
        
        if solicitud.estado != 'PENDIENTE':
            messages.error(request, "Esta solicitud ya fue procesada.")
            return redirect('detalle_venta', pk=venta.pk, slug=venta.slug_safe)
        
        motivo_rechazo = request.POST.get('motivo_rechazo', '').strip()
        if not motivo_rechazo or len(motivo_rechazo) < 10:
            messages.error(request, "Debes proporcionar un motivo de rechazo de al menos 10 caracteres.")
            return redirect('detalle_venta', pk=venta.pk, slug=venta.slug_safe)
        
        solicitud.estado = 'RECHAZADA'
        solicitud.aprobado_por = request.user
        solicitud.fecha_aprobacion = timezone.now()
        solicitud.motivo_rechazo = motivo_rechazo
        solicitud.save()
        
        # Crear notificación para el vendedor
        if venta.vendedor:
            mensaje = (
                f"Tu solicitud de cancelación para la venta #{venta.folio or venta.pk} "
                f"ha sido rechazada.\n\nMotivo: {motivo_rechazo}"
            )
            Notificacion.objects.create(
                usuario=venta.vendedor,
                tipo='CANCELACION_RECHAZADA',
                mensaje=mensaje,
                venta=venta,
                solicitud_cancelacion=solicitud,
                vista=False
            )
        
        messages.success(request, f"Solicitud de cancelación rechazada.")
        logger.info(f"✅ Solicitud de cancelación {solicitud.pk} rechazada por {request.user.username}")
        
        return redirect('detalle_venta', pk=venta.pk, slug=venta.slug_safe)


class CancelarVentaView(LoginRequiredMixin, UserPassesTestMixin, View):
    """Vista para cancelar una venta definitivamente (requiere solicitud aprobada)."""
    
    def test_func(self):
        """Solo el vendedor que creó la venta o el JEFE pueden cancelarla, y debe tener solicitud aprobada."""
        venta = get_object_or_404(VentaViaje, pk=self.kwargs['pk'])
        user_rol = perm.get_user_role(self.request.user, self.request)
        if perm.is_contador(self.request.user, self.request):
            return False
        
        # Verificar que existe una solicitud aprobada
        if not hasattr(venta, 'solicitud_cancelacion'):
            return False
        
        solicitud = venta.solicitud_cancelacion
        if solicitud.estado != 'APROBADA':
            return False
        
        return venta.vendedor == self.request.user or user_rol == 'JEFE'
    
    def handle_no_permission(self):
        messages.error(self.request, "No tienes permiso para cancelar esta venta o la solicitud no está aprobada.")
        return redirect('detalle_venta', pk=self.kwargs['pk'], slug=get_object_or_404(VentaViaje, pk=self.kwargs['pk']).slug_safe)
    
    def post(self, request, *args, **kwargs):
        venta = get_object_or_404(VentaViaje, pk=self.kwargs['pk'])
        solicitud = venta.solicitud_cancelacion
        
        # Verificar que la solicitud está aprobada
        if solicitud.estado != 'APROBADA':
            messages.error(request, "La solicitud de cancelación debe estar aprobada antes de cancelar definitivamente.")
            return redirect('detalle_venta', pk=venta.pk, slug=venta.slug_safe)
        
        # Verificar si la venta ya estaba cancelada
        if venta.estado == 'CANCELADA':
            messages.info(request, f"La venta #{venta.pk} ya estaba cancelada.")
            return redirect('detalle_venta', pk=venta.pk, slug=venta.slug_safe)
        
        # Usar el servicio para cancelar definitivamente
        try:
            resultado = CancelacionService.cancelar_venta_definitivamente(venta, solicitud)
            
            if resultado['exito']:
                mensaje = f"La venta #{venta.folio or venta.pk} ha sido cancelada definitivamente."
                detalles = []
                
                if resultado['km_revertidos'] > 0:
                    detalles.append(f"Se revirtieron {resultado['km_revertidos']:,.2f} km acumulados")
                if resultado['km_devueltos'] > 0:
                    detalles.append(f"Se devolvieron {resultado['km_devueltos']:,.2f} km redimidos")
                if resultado['promociones_revertidas'] > 0:
                    detalles.append(f"Se revirtieron {resultado['promociones_revertidas']} promociones")
                if resultado['comisiones_canceladas'] > 0:
                    detalles.append(f"Se cancelaron {resultado['comisiones_canceladas']} comisiones")
                
                if detalles:
                    mensaje += " " + " y ".join(detalles) + "."
                
                if resultado['errores']:
                    mensaje += f" Advertencias: {'; '.join(resultado['errores'])}"
                
                messages.success(request, mensaje)
            else:
                messages.error(request, f"Error al cancelar la venta: {'; '.join(resultado['errores'])}")
        except Exception as e:
            logger.exception(f"❌ Error al cancelar venta {venta.pk}: {e}")
            messages.error(request, f"Error al cancelar la venta: {str(e)}")
        
        # Notificar a quienes pueden aprobar cancelaciones (JEFE, Director General, Director Administrativo)
        usuarios_autorizados = User.objects.filter(
            perfil__rol__in=['JEFE', 'DIRECTOR_GENERAL', 'DIRECTOR_ADMINISTRATIVO']
        ).distinct()
        mensaje_notif = (
            f"La venta #{venta.folio or venta.pk} - Cliente: {venta.cliente.nombre_completo_display} "
            f"ha sido cancelada definitivamente por {request.user.get_full_name() or request.user.username}"
        )
        for usuario in usuarios_autorizados:
            Notificacion.objects.create(
                usuario=usuario,
                tipo='CANCELACION_DEFINITIVA',
                mensaje=mensaje_notif,
                venta=venta,
                solicitud_cancelacion=solicitud,
                vista=False
            )
        
        return redirect('detalle_venta', pk=venta.pk, slug=venta.slug_safe)


class ReciclarVentaView(LoginRequiredMixin, UserPassesTestMixin, View):
    """Vista para reciclar una venta cancelada, creando una nueva venta con los datos de la cancelada."""
    
    def test_func(self):
        """Solo el vendedor de la venta cancelada o JEFE pueden reciclar."""
        venta = get_object_or_404(VentaViaje, pk=self.kwargs['pk'])
        user_rol = perm.get_user_role(self.request.user, self.request)
        
        # Solo se puede reciclar si la venta está cancelada
        if venta.estado != 'CANCELADA':
            return False
        
        return venta.vendedor == self.request.user or user_rol == 'JEFE'
    
    def handle_no_permission(self):
        messages.error(self.request, "No tienes permiso para reciclar esta venta o la venta no está cancelada.")
        return redirect('detalle_venta', pk=self.kwargs['pk'], slug=get_object_or_404(VentaViaje, pk=self.kwargs['pk']).slug_safe)
    
    def post(self, request, *args, **kwargs):
        venta_original = get_object_or_404(VentaViaje, pk=self.kwargs['pk'])
        
        if venta_original.estado != 'CANCELADA':
            messages.error(request, "Solo se pueden reciclar ventas canceladas.")
            return redirect('detalle_venta', pk=venta_original.pk, slug=venta_original.slug_safe)
        
        try:
            with transaction.atomic():
                # Crear nueva venta copiando datos de la original
                nueva_venta = VentaViaje.objects.create(
                    cliente=venta_original.cliente,
                    vendedor=venta_original.vendedor,
                    pasajeros=venta_original.pasajeros,
                    edades_menores=venta_original.edades_menores,
                    servicios_seleccionados=venta_original.servicios_seleccionados,
                    servicios_detalle=venta_original.servicios_detalle,
                    proveedor=venta_original.proveedor,
                    tipo_viaje=venta_original.tipo_viaje,
                    fecha_inicio_viaje=venta_original.fecha_inicio_viaje,
                    fecha_fin_viaje=venta_original.fecha_fin_viaje,
                    # NO copiar: folio (se genera nuevo), estado (siempre ACTIVA)
                    # NO copiar: cantidad_apertura, abonos, logística, etc.
                    costo_neto=Decimal('0.00'),
                    costo_venta_final=Decimal('0.00'),
                    estado='ACTIVA',
                    # Campos internacionales si aplica
                    tarifa_base_usd=venta_original.tarifa_base_usd if venta_original.tipo_viaje == 'INT' else None,
                    impuestos_usd=venta_original.impuestos_usd if venta_original.tipo_viaje == 'INT' else None,
                    suplementos_usd=venta_original.suplementos_usd if venta_original.tipo_viaje == 'INT' else None,
                    tours_usd=venta_original.tours_usd if venta_original.tipo_viaje == 'INT' else None,
                    tipo_cambio=venta_original.tipo_cambio if venta_original.tipo_viaje == 'INT' else None,
                )
                
                # ✅ PERFORMANCE: Prefetch abonos una sola vez
                abonos_originales = list(venta_original.abonos.all())
                
                # Copiar abonos de la venta original (referenciando la venta original)
                # Los abonos se referencian a la venta original para tener contexto
                for abono_original in abonos_originales:
                    AbonoPago.objects.create(
                        venta=nueva_venta,
                        monto=abono_original.monto,
                        forma_pago=abono_original.forma_pago,
                        fecha_pago=abono_original.fecha_pago,
                        registrado_por=abono_original.registrado_por,
                        monto_usd=abono_original.monto_usd,
                        tipo_cambio_aplicado=abono_original.tipo_cambio_aplicado,
                        confirmado=abono_original.confirmado,
                        confirmado_por=abono_original.confirmado_por,
                        confirmado_en=abono_original.confirmado_en,
                        requiere_factura=abono_original.requiere_factura,
                        # Nota: estos abonos referencian la nueva venta pero provienen de la original
                    )
                
                # Copiar cantidad de apertura si existe
                if venta_original.cantidad_apertura and venta_original.cantidad_apertura > 0:
                    nueva_venta.cantidad_apertura = venta_original.cantidad_apertura
                    nueva_venta.modo_pago_apertura = venta_original.modo_pago_apertura
                    nueva_venta.requiere_factura_apertura = venta_original.requiere_factura_apertura
                    nueva_venta.save(update_fields=['cantidad_apertura', 'modo_pago_apertura', 'requiere_factura_apertura'])
                
                messages.success(
                    request,
                    f"Venta reciclada exitosamente. Nueva venta #{nueva_venta.folio or nueva_venta.pk} creada."
                )
                logger.info(
                    f"✅ Venta {venta_original.pk} reciclada, nueva venta {nueva_venta.pk} creada por {request.user.username}"
                )
                
                return redirect('editar_venta', pk=nueva_venta.pk)
                
        except Exception as e:
            logger.exception(f"❌ Error al reciclar venta {venta_original.pk}: {e}")
            messages.error(request, f"Error al reciclar la venta: {str(e)}")
            return redirect('detalle_venta', pk=venta_original.pk, slug=venta_original.slug_safe)


# ------------------- 5. GESTIÓN DE ABONOS (REMOVIDA/INTEGRADA) -------------------
# NOTA: La vista AbonoPagoCreateView ha sido eliminada. La funcionalidad se maneja
# directamente en el método post de VentaViajeDetailView (Sección 3) para un
# flujo de trabajo más limpio con formularios anidados.

# ------------------- 6. GESTIÓN DE LOGÍSTICA (STANDALONE) -------------------

class LogisticaUpdateView(LoginRequiredMixin, UserPassesTestMixin, UpdateView):
    model = Logistica
    form_class = LogisticaForm
    template_name = 'ventas/logistica_form.html'
    # Espera un PK de la VentaViaje en la URL, no del objeto Logistica
    pk_url_kwarg = 'venta_pk' 

    def get_object(self, queryset=None):
        # Busca el objeto Logistica relacionado con el PK de la VentaViaje
        return get_object_or_404(Logistica, venta__pk=self.kwargs['venta_pk'])

    def test_func(self):
        # Solo JEFES o el VENDEDOR de la venta pueden acceder. CONTADOR solo lectura.
        venta = self.get_object().venta
        user_rol = perm.get_user_role(self.request.user, self.request)
        
        return perm.has_full_access(self.request.user) or venta.vendedor == self.request.user

    def get_success_url(self):
        messages.success(self.request, "Logística actualizada correctamente.")
        venta = self.object.venta
        # CORRECCIÓN: Ahora incluye el slug para la redirección de éxito.
        # Redirige al detalle de la venta, con el tab de Logística activo
        return reverse_lazy('detalle_venta', kwargs={'pk': venta.pk, 'slug': venta.slug_safe}) + '?tab=logistica'


# ------------------- 7. ALERTA LOGÍSTICA PENDIENTE (INNOVACIÓN 3) -------------------

class LogisticaPendienteView(LoginRequiredMixin, ListView):
    model = VentaViaje
    template_name = 'ventas/logistica_pendiente.html'
    context_object_name = 'ventas'
    paginate_by = 30  # ESCALABILIDAD: Limitar resultados por página

    STATUS_META = {
        'pending': ('danger', 'Pendiente'),
        'ready': ('warning text-dark', 'Listo para pagar'),
        'paid': ('success', 'Pagado'),
    }

    ESTADO_META = {
        'pendiente': ('danger', 'Servicios pendientes'),
        'ready': ('warning text-dark', 'Fondos disponibles'),
        'completo': ('success', 'Servicios cubiertos'),
        'sin_servicios': ('secondary', 'Sin servicios planificados'),
    }

    def dispatch(self, request, *args, **kwargs):
        if not request.user.is_authenticated:
            return super().dispatch(request, *args, **kwargs)

        self.user_role = perm.get_user_role(request.user, request)
        if self.user_role in ('JEFE', 'CONTADOR', 'VENDEDOR'):
            return super().dispatch(request, *args, **kwargs)

        messages.error(request, "No tienes permiso para acceder al tablero de logística.")
        return redirect(reverse('dashboard'))

    def get_queryset(self):
        queryset = (
            self.model.objects
            .filter(servicios_logisticos__isnull=False)
            .select_related('cliente', 'vendedor')
            .prefetch_related('servicios_logisticos')
            .distinct()
            .order_by('fecha_inicio_viaje')
        )

        if self.user_role == 'VENDEDOR':
            queryset = queryset.filter(vendedor=self.request.user)

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        ventas = context.get('ventas') or []

        tablero = []
        stats = Counter({'pending': 0, 'ready': 0, 'paid': 0})

        for venta in ventas:
            card = build_logistica_card(venta)
            tablero.append(card)
            for servicio in card['servicios']:
                stats[servicio['status']] += 1

        context.update({
            'user_rol': self.user_role,
            'tablero_logistica': tablero,
            'stats_servicios': stats,
            'mostrando_propias': self.user_role == 'VENDEDOR',
            'puede_ver_todos': self.user_role in ('JEFE', 'CONTADOR'),
        })
        return context

# ------------------- 8. REPORTE FINANCIERO -------------------

class ReporteFinancieroView(LoginRequiredMixin, usuarios_mixins.FinancialReportRequiredMixin, TemplateView):
    template_name = 'ventas/reporte_financiero.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        user = self.request.user
        base_ventas = perm.get_ventas_queryset_base(VentaViaje, user, self.request)

        # --- 1. CÁLCULOS PRINCIPALES DE AGREGACIÓN (filtrados por rol: vendedor=propias, gerente=oficina) ---
        # Usar las propiedades del modelo directamente para garantizar que coincidan con los valores reales
        # que se muestran en el detalle de cada venta (costo_total_con_modificacion, total_pagado, saldo_restante)
        total_ventas = Decimal('0.00')
        total_pagado = Decimal('0.00')
        saldo_pendiente = Decimal('0.00')
        
        # Optimizar consultas con prefetch para abonos
        ventas_con_datos = base_ventas.select_related('cliente').prefetch_related('abonos')
        
        # Iterar sobre las ventas usando las propiedades calculadas del modelo
        for venta in ventas_con_datos:
            # costo_total_con_modificacion ya incluye modificaciones y descuentos correctamente
            costo_venta = venta.costo_total_con_modificacion or Decimal('0.00')
            total_ventas += costo_venta
            
            # total_pagado ya considera solo abonos confirmados y apertura confirmada
            pagado_venta = venta.total_pagado or Decimal('0.00')
            total_pagado += pagado_venta
            
            # saldo_restante ya está calculado correctamente en el modelo
            saldo_venta = venta.saldo_restante or Decimal('0.00')
            saldo_pendiente += saldo_venta

        context['total_ventas'] = total_ventas
        context['total_pagado'] = total_pagado
        context['saldo_pendiente'] = saldo_pendiente

        pagos_esperados = total_ventas - saldo_pendiente
        diferencia = total_pagado - pagos_esperados
        context['consistencia'] = {
            'real': total_pagado,
            'esperado': pagos_esperados,
            'diferencia': diferencia,
            'es_consistente': abs(diferencia) < Decimal('0.01'),
        }

        # Lista de usuarios para el filtro del historial (vendedor solo se ve a sí mismo)
        if perm.is_vendedor(user, self.request):
            context['usuarios'] = User.objects.filter(pk=user.pk).order_by('username')
        else:
            context['usuarios'] = User.objects.filter(is_active=True).order_by('username')

        # Últimos movimientos solo de ventas que el usuario puede ver
        try:
            from auditoria.models import HistorialMovimiento
            ventas_ids = list(base_ventas.values_list('pk', flat=True))
            ultimos_movimientos = HistorialMovimiento.objects.select_related('usuario', 'content_type').order_by('-fecha_hora')
            movimientos_preview = []
            for mov in ultimos_movimientos[:50]:
                enlace_url = None
                enlace_texto = None
                if mov.content_type and mov.object_id:
                    try:
                        obj = mov.content_type.get_object_for_this_type(pk=mov.object_id)
                        if isinstance(obj, AbonoPago) and obj.venta_id in ventas_ids:
                            venta = obj.venta
                            enlace_url = reverse('detalle_venta', kwargs={'pk': venta.pk, 'slug': venta.slug_safe}) + '?tab=abonos'
                            enlace_texto = f"Ver Venta #{venta.pk}"
                        elif isinstance(obj, VentaViaje) and obj.pk in ventas_ids:
                            enlace_url = reverse('detalle_venta', kwargs={'pk': obj.pk, 'slug': obj.slug_safe})
                            if mov.tipo_evento in ['ABONO_REGISTRADO', 'ABONO_CONFIRMADO', 'ABONO_ELIMINADO']:
                                enlace_url += '?tab=abonos'
                            enlace_texto = f"Ver Venta #{obj.pk}"
                        elif isinstance(obj, Cotizacion):
                            enlace_url = reverse('detalle_cotizacion', kwargs={'slug': obj.slug})
                            enlace_texto = "Ver Cotización"
                        elif isinstance(obj, Cliente):
                            enlace_url = reverse('detalle_cliente', kwargs={'pk': obj.pk})
                            enlace_texto = "Ver Cliente"
                        else:
                            continue
                    except Exception:
                        continue
                movimientos_preview.append({'movimiento': mov, 'enlace_url': enlace_url, 'enlace_texto': enlace_texto})
                if len(movimientos_preview) >= 5:
                    break
            context['ultimos_movimientos'] = movimientos_preview
        except ImportError:
            context['ultimos_movimientos'] = []

        # Botón exportar Excel: solo quien tiene acceso total (JEFE, Director General)
        context['es_jefe'] = perm.has_full_access(user, self.request)

        # --- Datos para modales de detalle (Ingreso, Pagos, Saldo) ---
        ventas_para_modales = base_ventas.filter(estado='ACTIVA').select_related('cliente').prefetch_related('abonos')
        ventas_nac = ventas_para_modales.filter(tipo_viaje__in=['NAC', 'INT_MXN'])
        ventas_int = ventas_para_modales.filter(tipo_viaje='INT')

        # Modal 1: Detalle Ingreso Bruto
        total_nac = Decimal('0.00')
        total_int = Decimal('0.00')
        lista_ventas_ingreso = []
        for v in ventas_para_modales.select_related('cliente'):
            costo = v.costo_total_con_modificacion or Decimal('0.00')
            if v.tipo_viaje in ('NAC', 'INT_MXN'):
                total_nac += costo
            else:
                total_int += costo
            lista_ventas_ingreso.append({
                'folio': v.folio or f'#{v.pk}',
                'cliente': v.cliente.nombre_completo_display if v.cliente else '—',
                'monto': costo,
                'tipo': v.get_tipo_viaje_display(),
                'url': reverse('detalle_venta', kwargs={'pk': v.pk, 'slug': v.slug_safe}),
            })
        lista_ventas_ingreso.sort(key=lambda x: x['monto'], reverse=True)
        context['detalle_ingreso'] = {
            'num_ventas': ventas_para_modales.count(),
            'ventas_nac_count': ventas_nac.count(),
            'ventas_nac_monto': total_nac,
            'ventas_int_count': ventas_int.count(),
            'ventas_int_monto': total_int,
            'lista_ventas': lista_ventas_ingreso[:15],
        }

        # Modal 2: Detalle Total Pagos Recibidos
        # Para INT la fuente de verdad es monto_usd (monto se guarda en 0); para NAC es monto (MXN)
        abonos_confirmados_qs = AbonoPago.objects.filter(
            venta__in=base_ventas
        ).filter(Q(confirmado=True) | Q(forma_pago='EFE')).select_related('venta', 'venta__cliente').order_by('-fecha_pago')[:15]
        monto_desde_abonos_nac = Decimal('0.00')
        monto_desde_abonos_int = Decimal('0.00')
        ultimos_abonos = []
        for ab in abonos_confirmados_qs:
            venta = ab.venta
            if venta.tipo_viaje == 'INT':
                # INT: fuente de verdad es monto_usd; monto suele estar en 0
                monto_abono = ab.monto_usd if (ab.monto_usd is not None and ab.monto_usd > 0) else (ab.monto_usd_para_display or Decimal('0.00'))
                if monto_abono is None:
                    monto_abono = Decimal('0.00')
                monto_desde_abonos_int += monto_abono
                moneda = 'USD'
            else:
                monto_abono = ab.monto or Decimal('0.00')
                monto_desde_abonos_nac += monto_abono
                moneda = 'MXN'
            ultimos_abonos.append({
                'folio': venta.folio or f'#{venta.pk}',
                'cliente': venta.cliente.nombre_completo_display if venta.cliente else '—',
                'monto': monto_abono,
                'moneda': moneda,
                'fecha': ab.fecha_pago,
                'forma_pago': ab.get_forma_pago_display(),
                'url': reverse('detalle_venta', kwargs={'pk': venta.pk, 'slug': venta.slug_safe}) + '?tab=abonos',
            })
        # Total abonos (suma NAC + INT en sus monedas; para el resumen mostramos ambos)
        num_abonos_total = AbonoPago.objects.filter(
            venta__in=base_ventas
        ).filter(Q(confirmado=True) | Q(forma_pago='EFE')).count()
        monto_apertura = Decimal('0.00')
        for v in ventas_para_modales:
            if v._apertura_confirmada_para_conteo():
                monto_apertura += (v.cantidad_apertura or Decimal('0.00'))
        context['detalle_pagos'] = {
            'num_abonos': num_abonos_total,
            'monto_desde_abonos_nac': monto_desde_abonos_nac,
            'monto_desde_abonos_int': monto_desde_abonos_int,
            'monto_desde_apertura': monto_apertura,
            'ultimos_abonos': ultimos_abonos,
        }

        # Modal 3: Detalle Saldo Pendiente (CxC)
        ventas_con_saldo = []
        for v in ventas_para_modales.prefetch_related('abonos'):
            saldo = v.saldo_restante or Decimal('0.00')
            if saldo <= 0:
                continue
            ventas_con_saldo.append({
                'folio': v.folio or f'#{v.pk}',
                'cliente': v.cliente.nombre_completo_display if v.cliente else '—',
                'saldo': saldo,
                'url': reverse('detalle_venta', kwargs={'pk': v.pk, 'slug': v.slug_safe}),
            })
        ventas_con_saldo.sort(key=lambda x: x['saldo'], reverse=True)
        context['detalle_saldo'] = {
            'num_ventas_con_saldo': len(ventas_con_saldo),
            'ventas_pendientes': ventas_con_saldo[:20],
        }

        return context


class ExportarReporteFinancieroExcelView(LoginRequiredMixin, UserPassesTestMixin, View):
    """
    Exporta el reporte financiero completo a Excel.
    Solo accesible para JEFE o Director General (acceso total).
    """
    def test_func(self):
        return perm.has_full_access(self.request.user)

    def handle_no_permission(self):
        messages.error(self.request, "No tienes permiso para descargar el reporte financiero en Excel.")
        return redirect('reporte_financiero')

    def get(self, request):
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            messages.error(request, "Error: openpyxl no está instalado. Instala con: pip install openpyxl")
            return redirect('reporte_financiero')

        try:
            return self._generar_excel(request)
        except Exception as e:
            logger.exception("Error al generar reporte financiero Excel: %s", e)
            messages.error(
                request,
                "No se pudo generar el Excel. Revisa el registro del servidor o contacta al administrador."
            )
            return redirect('reporte_financiero')

    def _generar_excel(self, request):
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        # Colores Movums (gradiente morado/índigo usado en la web)
        header_fill = PatternFill(start_color="667EEA", end_color="667EEA", fill_type="solid")  # Índigo
        header_fill_secondary = PatternFill(start_color="764BA2", end_color="764BA2", fill_type="solid")  # Púrpura
        title_fill = PatternFill(start_color="5C0CD1", end_color="5C0CD1", fill_type="solid")  # Morado principal
        row_fill_even = PatternFill(start_color="F5F3FF", end_color="F5F3FF", fill_type="solid")  # Lavanda muy suave
        row_fill_odd = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=11)
        title_font = Font(bold=True, color="FFFFFF", size=14)
        border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )
        border_medium = Border(
            left=Side(style='medium'), right=Side(style='medium'),
            top=Side(style='medium'), bottom=Side(style='medium')
        )
        align_center = Alignment(horizontal='center', vertical='center', wrap_text=True)
        align_left = Alignment(horizontal='left', vertical='center', wrap_text=True)
        hoy = timezone.localdate()
        meses_nombres = ['', 'Enero', 'Febrero', 'Marzo', 'Abril', 'Mayo', 'Junio',
                         'Julio', 'Agosto', 'Septiembre', 'Octubre', 'Noviembre', 'Diciembre']

        # --- Hoja 1: Detalle de Ventas (Nueva primera hoja) ---
        ws_detalle_ventas = wb.active
        ws_detalle_ventas.title = "Detalle de Ventas"

        # Título de la hoja (estilo Movums)
        ws_detalle_ventas.merge_cells(start_row=1, start_column=1, end_row=1, end_column=17)
        titulo_cell = ws_detalle_ventas.cell(row=1, column=1, value="Reporte Financiero — Detalle de Ventas · Movums")
        titulo_cell.fill = title_fill
        titulo_cell.font = title_font
        titulo_cell.alignment = Alignment(horizontal='center', vertical='center')
        titulo_cell.border = border_medium
        ws_detalle_ventas.row_dimensions[1].height = 28

        # Encabezados de la nueva tabla (fila 2)
        headers_detalle = [
            'Oficina', 'Fecha', 'Folio', 'Usuario', 'Cliente', 'Costo Final',
            'Pagado', 'Pendiente', 'Estado', 'Nacional o Internacional',
            'Servicio', 'Servicios Planificados', 'Servicios Pagados',
            'Saldo disponible', 'Ganancia Estimada', 'Fecha límite de pago', 'Empresa'
        ]
        header_row = 2
        for col, h in enumerate(headers_detalle, 1):
            c = ws_detalle_ventas.cell(row=header_row, column=col, value=h)
            c.fill = header_fill
            c.font = header_font
            c.border = border
            c.alignment = align_center
        ws_detalle_ventas.row_dimensions[header_row].height = 22
        
        # Obtener todas las ventas con relaciones optimizadas
        ventas_todas = VentaViaje.objects.select_related(
            'cliente', 'vendedor', 'vendedor__ejecutivo_asociado__oficina'
        ).prefetch_related('servicios_logisticos').order_by('-fecha_inicio_viaje')

        row_detalle = 3
        for idx, venta in enumerate(ventas_todas):
            # Filas alternadas (lavanda suave / blanco)
            row_fill = row_fill_even if idx % 2 == 0 else row_fill_odd
            # Oficina: desde vendedor → ejecutivo_asociado → oficina
            oficina_nombre = ''
            if venta.vendedor and hasattr(venta.vendedor, 'ejecutivo_asociado'):
                ejecutivo = venta.vendedor.ejecutivo_asociado
                if ejecutivo and ejecutivo.oficina:
                    oficina_nombre = ejecutivo.oficina.nombre
            
            # Fecha: fecha_inicio_viaje
            fecha_str = venta.fecha_inicio_viaje.strftime('%d/%m/%Y') if venta.fecha_inicio_viaje else ''
            
            # Folio
            folio = venta.folio or ''
            
            # Usuario
            usuario = venta.vendedor.username if venta.vendedor else ''
            
            # Cliente
            cliente_nombre = venta.cliente.nombre_completo_display if venta.cliente else ''
            
            # Costo Final: costo_total_con_modificacion (ya incluye modificaciones y descuentos)
            costo_final = venta.costo_total_con_modificacion or Decimal('0.00')
            
            # Pagado
            total_pagado = venta.total_pagado or Decimal('0.00')
            
            # Pendiente
            saldo_pendiente = venta.saldo_restante or Decimal('0.00')
            
            # Estado
            estado = venta.get_estado_display()
            
            # Nacional o Internacional
            tipo_viaje = venta.get_tipo_viaje_display()
            
            # Servicio: códigos separados por coma
            servicio = venta.servicios_seleccionados_display or ''
            
            # Servicios Planificados: costo_neto (o costo_neto_usd para INT)
            if venta.tipo_viaje == 'INT':
                servicios_planificados = venta.costo_neto_usd if getattr(venta, 'costo_neto_usd', None) is not None else Decimal('0.00')
                if servicios_planificados == 0 and venta.costo_neto and venta.tipo_cambio and venta.tipo_cambio > 0:
                    servicios_planificados = (venta.costo_neto / venta.tipo_cambio).quantize(Decimal('0.01'))
            else:
                servicios_planificados = venta.costo_neto or Decimal('0.00')
            
            # Servicios Pagados: suma de monto_planeado donde pagado=True
            servicios_pagados = venta.servicios_logisticos.filter(pagado=True).aggregate(
                total=Coalesce(Sum('monto_planeado'), Decimal('0.00'))
            )['total']
            
            # Saldo disponible y Ganancia Estimada: usar build_financial_summary
            servicios_qs = venta.servicios_logisticos.all()
            resumen_financiero = build_financial_summary(venta, servicios_qs)
            saldo_disponible = resumen_financiero.get('saldo_disponible_servicios', Decimal('0.00'))
            ganancia_estimada = resumen_financiero.get('ganancia_estimada', Decimal('0.00'))
            
            # Fecha límite de pago
            fecha_limite = venta.fecha_vencimiento_pago.strftime('%d/%m/%Y') if venta.fecha_vencimiento_pago else ''
            
            # Empresa: nombre_empresa si es tipo EMPRESA, si no "N/A"
            empresa_nombre = 'N/A'
            if venta.cliente and venta.cliente.tipo_cliente == 'EMPRESA':
                empresa_nombre = venta.cliente.nombre_empresa or 'N/A'
            
            # Escribir fila con formato
            def set_cell(r, col, val, num_fmt=None):
                cell = ws_detalle_ventas.cell(row=r, column=col, value=val)
                cell.fill = row_fill
                cell.border = border
                cell.alignment = align_center if col in (2, 6, 7, 8, 9, 10, 12, 13, 14, 15, 16) else align_left
                if num_fmt:
                    cell.number_format = num_fmt
                return cell

            set_cell(row_detalle, 1, oficina_nombre)
            set_cell(row_detalle, 2, fecha_str)
            set_cell(row_detalle, 3, folio)
            set_cell(row_detalle, 4, usuario)
            set_cell(row_detalle, 5, cliente_nombre)
            set_cell(row_detalle, 6, float(costo_final), '#,##0.00')
            set_cell(row_detalle, 7, float(total_pagado), '#,##0.00')
            set_cell(row_detalle, 8, float(saldo_pendiente), '#,##0.00')
            set_cell(row_detalle, 9, estado)
            set_cell(row_detalle, 10, tipo_viaje)
            set_cell(row_detalle, 11, servicio)
            set_cell(row_detalle, 12, float(servicios_planificados), '#,##0.00')
            set_cell(row_detalle, 13, float(servicios_pagados), '#,##0.00')
            set_cell(row_detalle, 14, float(saldo_disponible), '#,##0.00')
            set_cell(row_detalle, 15, float(ganancia_estimada), '#,##0.00')
            set_cell(row_detalle, 16, fecha_limite)
            set_cell(row_detalle, 17, empresa_nombre)

            row_detalle += 1
        
        # Ajustar anchos de columnas
        ws_detalle_ventas.column_dimensions['A'].width = 20  # Oficina
        ws_detalle_ventas.column_dimensions['B'].width = 12  # Fecha
        ws_detalle_ventas.column_dimensions['C'].width = 15  # Folio
        ws_detalle_ventas.column_dimensions['D'].width = 15  # Usuario
        ws_detalle_ventas.column_dimensions['E'].width = 30  # Cliente
        ws_detalle_ventas.column_dimensions['F'].width = 15  # Costo Final
        ws_detalle_ventas.column_dimensions['G'].width = 15  # Pagado
        ws_detalle_ventas.column_dimensions['H'].width = 15  # Pendiente
        ws_detalle_ventas.column_dimensions['I'].width = 12  # Estado
        ws_detalle_ventas.column_dimensions['J'].width = 20  # Nacional o Internacional
        ws_detalle_ventas.column_dimensions['K'].width = 30  # Servicio
        ws_detalle_ventas.column_dimensions['L'].width = 20  # Servicios Planificados
        ws_detalle_ventas.column_dimensions['M'].width = 18  # Servicios Pagados
        ws_detalle_ventas.column_dimensions['N'].width = 18  # Saldo disponible
        ws_detalle_ventas.column_dimensions['O'].width = 18  # Ganancia Estimada
        ws_detalle_ventas.column_dimensions['P'].width = 18  # Fecha límite de pago
        ws_detalle_ventas.column_dimensions['Q'].width = 30  # Empresa
        ws_detalle_ventas.freeze_panes = 'A3'  # Congelar título y encabezados

        # --- Hoja 2: Resumen (ahora segunda hoja) ---
        ws_resumen = wb.create_sheet("Resumen", 1)
        row = 1
        ws_resumen.merge_cells(start_row=1, start_column=1, end_row=1, end_column=4)
        titulo_resumen = ws_resumen.cell(row=row, column=1, value="Reporte Financiero Consolidado · Movums")
        titulo_resumen.fill = title_fill
        titulo_resumen.font = title_font
        titulo_resumen.alignment = Alignment(horizontal='center', vertical='center')
        titulo_resumen.border = border_medium
        ws_resumen.row_dimensions[1].height = 26
        row += 2

        # Totales (igual que la pantalla Reporte Financiero: todas las ventas)
        total_ventas = VentaViaje.objects.aggregate(s=Sum('costo_venta_final')).get('s') or Decimal('0.00')
        total_abonos_sum = AbonoPago.objects.aggregate(s=Sum('monto')).get('s') or Decimal('0.00')
        total_apertura = VentaViaje.objects.aggregate(s=Sum('cantidad_apertura')).get('s') or Decimal('0.00')
        total_pagado = total_abonos_sum + total_apertura
        saldo_pendiente = total_ventas - total_pagado
        total_ventas = Decimal(total_ventas)
        total_pagado = Decimal(total_pagado)
        saldo_pendiente = Decimal(saldo_pendiente)

        # Relleno suave para bloque de totales
        totales_fill = PatternFill(start_color="E8E4F8", end_color="E8E4F8", fill_type="solid")
        for r in range(row, row + 3):
            for col in (1, 2):
                c = ws_resumen.cell(row=r, column=col, value=None)
                c.fill = totales_fill
                c.border = border
        ws_resumen.cell(row=row, column=1, value="Ingreso Bruto Total (Ventas activas):")
        ws_resumen.cell(row=row, column=1).font = Font(bold=True)
        ws_resumen.cell(row=row, column=2, value=float(total_ventas))
        ws_resumen.cell(row=row, column=2).number_format = '#,##0.00'
        ws_resumen.cell(row=row, column=2).fill = totales_fill
        ws_resumen.cell(row=row, column=2).border = border
        row += 1
        ws_resumen.cell(row=row, column=1, value="Total Pagos Recibidos:")
        ws_resumen.cell(row=row, column=1).font = Font(bold=True)
        ws_resumen.cell(row=row, column=2, value=float(total_pagado))
        ws_resumen.cell(row=row, column=2).number_format = '#,##0.00'
        ws_resumen.cell(row=row, column=2).fill = totales_fill
        ws_resumen.cell(row=row, column=2).border = border
        row += 1
        ws_resumen.cell(row=row, column=1, value="Saldo Pendiente (CxC):")
        ws_resumen.cell(row=row, column=1).font = Font(bold=True)
        ws_resumen.cell(row=row, column=2, value=float(saldo_pendiente))
        ws_resumen.cell(row=row, column=2).number_format = '#,##0.00'
        ws_resumen.cell(row=row, column=2).fill = totales_fill
        ws_resumen.cell(row=row, column=2).border = border
        row += 2

        # Liquidez por mes (próximos 12 meses)
        ws_resumen.merge_cells(start_row=row, start_column=1, end_row=row, end_column=4)
        liq_titulo = ws_resumen.cell(row=row, column=1, value="Liquidez por mes (próximos 12 meses)")
        liq_titulo.font = Font(bold=True, size=12, color="FFFFFF")
        liq_titulo.fill = header_fill_secondary
        liq_titulo.alignment = Alignment(horizontal='center', vertical='center')
        liq_titulo.border = border_medium
        row += 1
        headers_liq = ['Mes', 'Año', 'Esperado a cobrar (Vista A)', 'Liquidez acumulada (Vista B)']
        for col, h in enumerate(headers_liq, 1):
            c = ws_resumen.cell(row=row, column=col, value=h)
            c.fill = header_fill
            c.font = header_font
            c.border = border
            c.alignment = align_center
        row += 1

        liquidez_acum = total_pagado  # Lo ya cobrado
        from calendar import monthrange
        year = hoy.year
        month = hoy.month
        for i in range(12):
            month += 1
            if month > 12:
                month = 1
                year += 1
            _, ultimo = monthrange(year, month)
            inicio_mes = date(year, month, 1)
            fin_mes = date(year, month, ultimo)
            ventas_mes = VentaViaje.objects.filter(
                estado='ACTIVA',
                fecha_vencimiento_pago__gte=inicio_mes,
                fecha_vencimiento_pago__lte=fin_mes
            ).select_related('cliente')  # Solo activas para proyección de cobro
            esperado_mes = Decimal('0.00')
            for v in ventas_mes:
                esperado_mes += v.saldo_restante
            liquidez_acum += esperado_mes
            liq_fill = row_fill_even if (row - 1) % 2 == 0 else row_fill_odd
            for col in range(1, 5):
                cell = ws_resumen.cell(row=row, column=col)
                cell.border = border
                cell.fill = liq_fill
                cell.alignment = align_center if col > 1 else align_left
            ws_resumen.cell(row=row, column=1, value=meses_nombres[month])
            ws_resumen.cell(row=row, column=2, value=year)
            ws_resumen.cell(row=row, column=3, value=float(esperado_mes))
            ws_resumen.cell(row=row, column=3).number_format = '#,##0.00'
            ws_resumen.cell(row=row, column=4, value=float(liquidez_acum))
            ws_resumen.cell(row=row, column=4).number_format = '#,##0.00'
            row += 1

        # Ajustar anchos
        ws_resumen.column_dimensions['A'].width = 18
        ws_resumen.column_dimensions['B'].width = 10
        ws_resumen.column_dimensions['C'].width = 28
        ws_resumen.column_dimensions['D'].width = 28

        # --- Hoja 3: Detalle Ventas (activas + canceladas en sección aparte) ---
        ws_ventas = wb.create_sheet("Detalle Ventas", 2)
        headers_ventas = [
            'ID', 'Folio', 'Cliente', 'Fecha inicio viaje', 'Fecha fin viaje', 'Fecha límite pago',
            'Costo venta final', 'Total cobrado', 'Saldo pendiente', 'Descuento Km', 'Descuento promociones', 'Estado'
        ]
        for col, h in enumerate(headers_ventas, 1):
            c = ws_ventas.cell(row=1, column=col, value=h)
            c.fill = header_fill
            c.font = header_font
            c.border = border
            c.alignment = align_center
        row_v = 2
        ventas_activas_list = VentaViaje.objects.filter(estado='ACTIVA').select_related('cliente').order_by('fecha_vencimiento_pago')
        for idx, venta in enumerate(ventas_activas_list):
            _escribir_fila_venta(ws_ventas, row_v, venta, border, row_fill_even if idx % 2 == 0 else row_fill_odd)
            row_v += 1
        row_v += 1
        ws_ventas.cell(row=row_v, column=1, value="--- VENTAS CANCELADAS ---")
        ws_ventas.cell(row=row_v, column=1).font = Font(bold=True)
        ws_ventas.cell(row=row_v, column=1).fill = header_fill_secondary
        ws_ventas.cell(row=row_v, column=1).font = Font(bold=True, color="FFFFFF")
        row_v += 1
        ventas_canceladas = VentaViaje.objects.filter(estado='CANCELADA').select_related('cliente').order_by('-fecha_creacion')
        for idx, venta in enumerate(ventas_canceladas):
            _escribir_fila_venta(ws_ventas, row_v, venta, border, row_fill_even if idx % 2 == 0 else row_fill_odd)
            row_v += 1
        for col in range(1, len(headers_ventas) + 1):
            ws_ventas.column_dimensions[get_column_letter(col)].width = 16

        # --- Hoja 4: Detalle Abonos ---
        ws_abonos = wb.create_sheet("Detalle Abonos", 3)
        headers_abonos = ['ID Abono', 'Venta ID', 'Cliente', 'Monto', 'Forma pago', 'Fecha pago', 'Confirmado']
        for col, h in enumerate(headers_abonos, 1):
            c = ws_abonos.cell(row=1, column=col, value=h)
            c.fill = header_fill
            c.font = header_font
            c.border = border
            c.alignment = align_center
        row_ab = 2
        abonos = AbonoPago.objects.select_related('venta', 'venta__cliente').order_by('-fecha_pago')
        for ab in abonos:
            ws_abonos.cell(row=row_ab, column=1, value=ab.pk)
            ws_abonos.cell(row=row_ab, column=2, value=ab.venta_id)
            _cliente = getattr(ab.venta, 'cliente', None) if ab.venta_id else None
            _cliente_nombre = getattr(_cliente, 'nombre_completo_display', '') if _cliente else ''
            ws_abonos.cell(row=row_ab, column=3, value=_cliente_nombre)
            ws_abonos.cell(row=row_ab, column=4, value=float(ab.monto))
            ws_abonos.cell(row=row_ab, column=4).number_format = '#,##0.00'
            ws_abonos.cell(row=row_ab, column=5, value=ab.get_forma_pago_display())
            ws_abonos.cell(row=row_ab, column=6, value=ab.fecha_pago.strftime('%d/%m/%Y %H:%M') if ab.fecha_pago else '')
            ws_abonos.cell(row=row_ab, column=7, value='Sí' if ab.confirmado else 'No')
            for col in range(1, 8):
                ws_abonos.cell(row=row_ab, column=col).border = border
            row_ab += 1
        for col in range(1, 8):
            ws_abonos.column_dimensions[get_column_letter(col)].width = 16

        # --- Hoja 5: Abonos a Proveedores ---
        ws_prov = wb.create_sheet("Abonos a Proveedores", 4)
        headers_prov = ['ID', 'Venta ID', 'Proveedor', 'Monto', 'Monto USD', 'Estado', 'Fecha solicitud', 'Fecha aprobación', 'Fecha confirmación']
        for col, h in enumerate(headers_prov, 1):
            c = ws_prov.cell(row=1, column=col, value=h)
            c.fill = header_fill
            c.font = header_font
            c.border = border
            c.alignment = align_center
        row_p = 2
        abonos_prov = AbonoProveedor.objects.select_related('venta').order_by('-fecha_solicitud')
        for ap in abonos_prov:
            ws_prov.cell(row=row_p, column=1, value=ap.pk)
            ws_prov.cell(row=row_p, column=2, value=ap.venta_id)
            ws_prov.cell(row=row_p, column=3, value=ap.proveedor or '')
            ws_prov.cell(row=row_p, column=4, value=float(ap.monto))
            ws_prov.cell(row=row_p, column=4).number_format = '#,##0.00'
            ws_prov.cell(row=row_p, column=5, value=float(ap.monto_usd) if ap.monto_usd else '')
            if ap.monto_usd is not None:
                ws_prov.cell(row=row_p, column=5).number_format = '#,##0.00'
            ws_prov.cell(row=row_p, column=6, value=ap.get_estado_display())
            ws_prov.cell(row=row_p, column=7, value=ap.fecha_solicitud.strftime('%d/%m/%Y %H:%M') if ap.fecha_solicitud else '')
            ws_prov.cell(row=row_p, column=8, value=ap.fecha_aprobacion.strftime('%d/%m/%Y %H:%M') if ap.fecha_aprobacion else '')
            ws_prov.cell(row=row_p, column=9, value=ap.fecha_confirmacion.strftime('%d/%m/%Y %H:%M') if ap.fecha_confirmacion else '')
            for col in range(1, 10):
                ws_prov.cell(row=row_p, column=col).border = border
            row_p += 1
        for col in range(1, 10):
            ws_prov.column_dimensions[get_column_letter(col)].width = 18

        # --- Hoja 6: Comisiones ---
        ws_com = wb.create_sheet("Comisiones", 5)
        headers_com = ['Venta ID', 'Vendedor', 'Mes', 'Año', 'Tipo venta', 'Monto base', 'Porcentaje', 'Comisión calculada', 'Comisión pagada', 'Comisión pendiente', 'Estado pago', 'Cancelada']
        for col, h in enumerate(headers_com, 1):
            c = ws_com.cell(row=1, column=col, value=h)
            c.fill = header_fill
            c.font = header_font
            c.border = border
            c.alignment = align_center
        row_c = 2
        comisiones = ComisionVenta.objects.select_related('venta', 'vendedor').order_by('-anio', '-mes', 'venta_id')
        for cv in comisiones:
            ws_com.cell(row=row_c, column=1, value=cv.venta_id)
            ws_com.cell(row=row_c, column=2, value=cv.vendedor.get_full_name() or cv.vendedor.username)
            ws_com.cell(row=row_c, column=3, value=cv.mes)
            ws_com.cell(row=row_c, column=4, value=cv.anio)
            ws_com.cell(row=row_c, column=5, value=cv.get_tipo_venta_display())
            ws_com.cell(row=row_c, column=6, value=float(cv.monto_base_comision))
            ws_com.cell(row=row_c, column=6).number_format = '#,##0.00'
            ws_com.cell(row=row_c, column=7, value=float(cv.porcentaje_aplicado))
            ws_com.cell(row=row_c, column=8, value=float(cv.comision_calculada))
            ws_com.cell(row=row_c, column=8).number_format = '#,##0.00'
            ws_com.cell(row=row_c, column=9, value=float(cv.comision_pagada))
            ws_com.cell(row=row_c, column=9).number_format = '#,##0.00'
            ws_com.cell(row=row_c, column=10, value=float(cv.comision_pendiente))
            ws_com.cell(row=row_c, column=10).number_format = '#,##0.00'
            ws_com.cell(row=row_c, column=11, value=cv.get_estado_pago_venta_display())
            ws_com.cell(row=row_c, column=12, value='Sí' if cv.cancelada else 'No')
            for col in range(1, 13):
                ws_com.cell(row=row_c, column=col).border = border
            row_c += 1
        for col in range(1, 13):
            ws_com.column_dimensions[get_column_letter(col)].width = 16

        # Guardar en buffer para asegurar que el contenido se envíe correctamente
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        filename = f"reporte_financiero_{hoy.strftime('%Y-%m-%d')}.xlsx"
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        response['Content-Length'] = len(buffer.getvalue())
        return response


def _escribir_fila_venta(ws, row_v, venta, border, row_fill=None):
    """Escribe una fila de detalle de venta en la hoja Excel. row_fill opcional para filas alternadas."""
    cliente = getattr(venta, 'cliente', None)
    cliente_nombre = getattr(cliente, 'nombre_completo_display', '') if cliente else ''
    ws.cell(row=row_v, column=1, value=venta.pk)
    ws.cell(row=row_v, column=2, value=venta.folio or '')
    ws.cell(row=row_v, column=3, value=cliente_nombre)
    ws.cell(row=row_v, column=4, value=venta.fecha_inicio_viaje.strftime('%d/%m/%Y') if venta.fecha_inicio_viaje else '')
    ws.cell(row=row_v, column=5, value=venta.fecha_fin_viaje.strftime('%d/%m/%Y') if venta.fecha_fin_viaje else '')
    ws.cell(row=row_v, column=6, value=venta.fecha_vencimiento_pago.strftime('%d/%m/%Y') if venta.fecha_vencimiento_pago else '')
    ws.cell(row=row_v, column=7, value=float(venta.costo_venta_final))
    ws.cell(row=row_v, column=7).number_format = '#,##0.00'
    ws.cell(row=row_v, column=8, value=float(venta.total_pagado))
    ws.cell(row=row_v, column=8).number_format = '#,##0.00'
    ws.cell(row=row_v, column=9, value=float(venta.saldo_restante))
    ws.cell(row=row_v, column=9).number_format = '#,##0.00'
    desc_km = venta.descuento_kilometros_mxn or Decimal('0.00')
    desc_promo = venta.descuento_promociones_mxn or Decimal('0.00')
    ws.cell(row=row_v, column=10, value=float(desc_km))
    ws.cell(row=row_v, column=10).number_format = '#,##0.00'
    ws.cell(row=row_v, column=11, value=float(desc_promo))
    ws.cell(row=row_v, column=11).number_format = '#,##0.00'
    ws.cell(row=row_v, column=12, value=venta.get_estado_display())
    for col in range(1, 13):
        cell = ws.cell(row=row_v, column=col)
        cell.border = border
        if row_fill:
            cell.fill = row_fill


# ------------------- 9. GENERACIÓN DE PDF (INNOVACIÓN 4) -------------------

class ComprobanteAbonoPDFView(LoginRequiredMixin, DetailView):
    model = VentaViaje
    
    def get(self, request, *args, **kwargs):
        if not WEASYPRINT_AVAILABLE:
             # Si WeasyPrint no está cargado, devuelve un error 503 o un mensaje simple
             return HttpResponse("Error en la generación de PDF. Faltan dependencias (GTK3).", status=503)

        # Usamos el mismo método get_object que maneja PK y SLUG
        self.object = self.get_object() 
        venta = self.object
        
        # Calcular totales para el PDF (total_pagado/saldo_restante ya en USD para INT)
        total_pagado = venta.total_pagado
        saldo_restante = venta.saldo_restante
        moneda_pdf = 'USD' if venta.tipo_viaje == 'INT' else 'MXN'

        # Total final y descuentos (INT: costo_total_con_modificacion ya en USD)
        if venta.tipo_viaje == 'INT':
            total_final = venta.costo_total_con_modificacion
            costo_base = total_final
            total_descuentos = Decimal('0.00')
            descuento_km = Decimal('0.00')
            descuento_promo = Decimal('0.00')
        else:
            descuento_km = venta.descuento_kilometros_mxn or Decimal('0.00')
            descuento_promo = venta.descuento_promociones_mxn or Decimal('0.00')
            total_descuentos = descuento_km + descuento_promo
            costo_base = (venta.costo_venta_final or Decimal('0.00')) + (venta.costo_modificacion or Decimal('0.00'))
            total_final = costo_base - total_descuentos
        
        # Preparar ruta absoluta file:// para el membrete (WeasyPrint necesita URL absoluta)
        membrete_path = os.path.join(settings.BASE_DIR, 'static', 'img', 'membrete_movums.jpg')
        membrete_url = None
        if os.path.exists(membrete_path):
            # Crear URL file:// absoluta para WeasyPrint
            membrete_abs_path = os.path.abspath(membrete_path)
            # En Windows, necesitamos ajustar el formato de la ruta
            if os.name == 'nt':
                membrete_url = f"file:///{membrete_abs_path.replace(os.sep, '/')}"
            else:
                membrete_url = f"file://{membrete_abs_path}"
        
        # Obtener el contexto para la plantilla HTML
        context = {
            'venta': venta,
            'now': datetime.datetime.now(),
            'total_pagado': total_pagado,
            'saldo_restante': saldo_restante,
            'moneda_pdf': moneda_pdf,
            # Incluir TODOS los abonos para el detalle en el PDF (mostrar todos, incluso pendientes)
            'abonos': venta.abonos.all().order_by('fecha_pago'),
            'membrete_url': membrete_url,  # URL absoluta file:// para WeasyPrint
            'STATIC_URL': settings.STATIC_URL,
            # Información financiera completa igual que en detalle de venta
            'total_descuentos': total_descuentos,
            'costo_base': costo_base,
            'total_final': total_final,
            'descuento_km': descuento_km,
            'descuento_promo': descuento_promo,
        }

        # 1. Renderizar la plantilla HTML
        html_string = render_to_string('ventas/comprobante_abonos_pdf.html', context, request=request)
        
        # 2. Generar el PDF con WeasyPrint
        # Obtener rutas de archivos estáticos para recursos como imágenes
        static_dir = os.path.join(settings.BASE_DIR, 'static')
        static_dir_abs = os.path.abspath(static_dir)
        base_url = f"file://{static_dir_abs}/"
        
        html = HTML(string=html_string, base_url=base_url)
        pdf_file = html.write_pdf(stylesheets=[]) # Dejar la lista vacía a menos que se necesite un CSS específico

        # 3. Preparar la respuesta HTTP
        response = HttpResponse(pdf_file, content_type='application/pdf')
        filename = f"Comprobante_Venta_{venta.pk}.pdf"
        # Usar 'inline' si se quiere mostrar en el navegador o 'attachment' para forzar la descarga
        response['Content-Disposition'] = f'attachment; filename="{filename}"' 
        
        return response
    
# ------------------- 10. NUEVA VISTA PARA EL CONTRATO (LA QUE NECESITAS) -------------------
class ContratoVentaPDFView(LoginRequiredMixin, DetailView):
    """
    Vista para generar el Contrato de Venta en formato DOCX.
    Utiliza python-docx para generar un documento Word editable.
    Si la venta es de hospedaje, redirige a ContratoHospedajePDFView.
    """
    model = VentaViaje
    
    def get(self, request, *args, **kwargs):
        self.object = self.get_object() 
        venta = self.object
        
        # Si es hospedaje, usar el contrato específico
        if venta.servicios_seleccionados and 'HOS' in venta.servicios_seleccionados:
            # Redirigir a la vista específica de hospedaje
            from django.urls import reverse
            return HttpResponseRedirect(reverse('generar_contrato_hospedaje_pdf', kwargs={'pk': venta.pk, 'slug': venta.slug_safe}))
        
        # Verificar si es paquete nacional (directamente o por cotización)
        es_paquete_nacional = False
        if venta.servicios_seleccionados and 'PAQ' in venta.servicios_seleccionados and venta.tipo_viaje in ('NAC', 'INT_MXN'):
            es_paquete_nacional = True
        elif venta.tipo_viaje in ('NAC', 'INT_MXN') and venta.cotizacion_origen:
            # Si no tiene PAQ pero tiene cotización, verificar el tipo de cotización
            cotizacion = venta.cotizacion_origen
            if isinstance(cotizacion.propuestas, str):
                try:
                    import json
                    propuestas = json.loads(cotizacion.propuestas)
                except (ValueError, TypeError, json.JSONDecodeError) as e:
                    logger.warning(f"Error al decodificar propuestas de cotización: {e}")
                    propuestas = {}
            elif isinstance(cotizacion.propuestas, dict):
                propuestas = cotizacion.propuestas
            else:
                propuestas = {}
            
            tipo_cot = propuestas.get('tipo', '')
            if tipo_cot == 'paquete':
                es_paquete_nacional = True
        
        if es_paquete_nacional:
            # Redirigir a la vista específica de paquete nacional
            from django.urls import reverse
            return HttpResponseRedirect(reverse('generar_contrato_paquete_nacional_pdf', kwargs={'pk': venta.pk, 'slug': venta.slug_safe}))
        
        # TODAS las ventas internacionales usan el contrato internacional unificado
        if venta.tipo_viaje == 'INT':
            from django.urls import reverse
            return HttpResponseRedirect(reverse('generar_contrato_paquete_internacional_pdf', kwargs={'pk': venta.pk, 'slug': venta.slug_safe}))
        
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.oxml.ns import qn
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return HttpResponse("Error: python-docx no está instalado. Ejecuta: pip install python-docx", status=500)

        cliente = venta.cliente
        
        # Calcular saldo pendiente (asegurarse de que no sea negativo)
        from decimal import Decimal
        saldo_pendiente = max(Decimal('0.00'), venta.costo_total_con_modificacion - venta.total_pagado)
        
        # Usar plantilla con membrete si existe
        template_path = os.path.join(settings.BASE_DIR, 'static', 'docx', 'membrete.docx')
        if os.path.exists(template_path):
            doc = Document(template_path)
        else:
            doc = Document()
        
        # Configurar fuente predeterminada Arial 12
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        
        MOVUMS_BLUE = RGBColor(0, 74, 142)  # #004a8e
        TEXT_COLOR = RGBColor(47, 47, 47)  # #2f2f2f
        
        def set_run_font(run, size=12, bold=False, color=TEXT_COLOR):
            run.font.name = 'Arial'
            run.font.size = Pt(size)
            run.bold = bold
            run.font.color.rgb = color
        
        def add_paragraph(doc_obj, text='', size=12, bold=False, color=TEXT_COLOR, space_before=0, space_after=0, alignment=None):
            paragraph = doc_obj.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(space_before)
            paragraph.paragraph_format.space_after = Pt(space_after)
            if alignment:
                paragraph.alignment = alignment
            if text:
                run = paragraph.add_run(text)
                set_run_font(run, size=size, bold=bold, color=color)
            return paragraph
        
        def format_date(value):
            if not value:
                return '-'
            try:
                if isinstance(value, datetime.date):
                    return value.strftime('%d/%m/%Y')
                return str(value)
            except (ValueError, AttributeError) as e:
                logger.warning(f"Error al formatear fecha: {e}")
                return '-'
        
        def format_currency(value):
            if value in (None, '', 0):
                return '0.00'
            try:
                number = Decimal(str(value).replace(',', ''))
            except (ValueError, TypeError) as e:
                logger.warning(f"Error al formatear moneda: {e}")
                return str(value)
            return f"{number:,.2f}"
        
        # Título principal
        titulo = add_paragraph(doc, 'CONTRATO DE SERVICIOS TURÍSTICOS', size=18, bold=True, color=MOVUMS_BLUE, space_after=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Sección 1: DATOS GENERALES DEL CLIENTE
        add_paragraph(doc, '1. DATOS GENERALES DEL CLIENTE', size=14, bold=True, color=MOVUMS_BLUE, space_before=15, space_after=8)
        
        # Nombre completo
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run('Nombre completo: ')
        set_run_font(run1, size=12, bold=True)
        if cliente.nombre_completo_display:
            run2 = p.add_run(cliente.nombre_completo_display)
            set_run_font(run2, size=12)
        else:
            run2 = p.add_run('________________________')
            set_run_font(run2, size=12)
            run2.font.underline = True
        
        # Teléfono
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run('Teléfono: ')
        set_run_font(run1, size=12, bold=True)
        if cliente.telefono:
            run2 = p.add_run(cliente.telefono)
            set_run_font(run2, size=12)
        else:
            run2 = p.add_run('________________________')
            set_run_font(run2, size=12)
            run2.font.underline = True
        
        # Email
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run('Correo electrónico: ')
        set_run_font(run1, size=12, bold=True)
        if cliente.email:
            run2 = p.add_run(cliente.email)
            set_run_font(run2, size=12)
        else:
            run2 = p.add_run('________________________')
            set_run_font(run2, size=12)
            run2.font.underline = True
        
        # Identificación
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run('Identificación oficial: ')
        set_run_font(run1, size=12, bold=True)
        if cliente.documento_identificacion:
            run2 = p.add_run(cliente.documento_identificacion)
            set_run_font(run2, size=12)
        else:
            run2 = p.add_run('________________________')
            set_run_font(run2, size=12)
            run2.font.underline = True
        run3 = p.add_run('  INE / Pasaporte / Otro: ')
        set_run_font(run3, size=12)
        # El tipo de identificación siempre va con línea porque no lo tenemos en el sistema
        run4 = p.add_run('________')
        set_run_font(run4, size=12)
        run4.font.underline = True
        
        # Acompañantes
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run('Acompañantes:')
        set_run_font(run1, size=12, bold=True)
        
        # Procesar acompañantes: quitar saltos de línea y separar con comas
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(4)
        run1 = p.add_run('• Nombre(s) y edad(es): ')
        set_run_font(run1, size=12, bold=True)
        if venta.pasajeros:
            # Limpiar el texto: quitar saltos de línea y espacios extra, separar con comas
            pasajeros_texto = venta.pasajeros.replace('\n', ', ').replace('\r', ', ')
            # Limpiar espacios múltiples y comas múltiples
            import re
            pasajeros_texto = re.sub(r'\s+', ' ', pasajeros_texto)
            pasajeros_texto = re.sub(r',\s*,', ',', pasajeros_texto)
            pasajeros_texto = pasajeros_texto.strip().rstrip(',')
            run2 = p.add_run(pasajeros_texto)
            set_run_font(run2, size=12)
        else:
            run2 = p.add_run('________________________')
            set_run_font(run2, size=12)
            run2.font.underline = True
        
        # Sección 2: DATOS DEL SERVICIO TURÍSTICO CONTRATADO
        add_paragraph(doc, '2. DATOS DEL SERVICIO TURÍSTICO CONTRATADO', size=14, bold=True, color=MOVUMS_BLUE, space_before=15, space_after=8)
        
        campos_servicio = [
            ('Nombre del Paquete:', ''),
            ('Destino(s):', ''),
            ('Fecha de inicio:', format_date(venta.fecha_inicio_viaje) if venta.fecha_inicio_viaje else '//2025'),
            ('Fecha de término:', format_date(venta.fecha_fin_viaje) if venta.fecha_fin_viaje else '//2025'),
            ('Número total de viajeros:', ''),
        ]
        
        for label, value in campos_servicio:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run1 = p.add_run(f'{label} ')
            set_run_font(run1, size=12, bold=True)
            if value:
                run2 = p.add_run(value)
                set_run_font(run2, size=12)
            else:
                run2 = p.add_run('________________________')
                set_run_font(run2, size=12)
                run2.font.underline = True
        
        # Sección 3: TRANSPORTE
        add_paragraph(doc, '3. TRANSPORTE', size=14, bold=True, color=MOVUMS_BLUE, space_before=15, space_after=8)
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run('Tipo de transporte: ')
        set_run_font(run1, size=12, bold=True)
        run2 = p.add_run('________________________')
        set_run_font(run2, size=12)
        run2.font.underline = True
        run3 = p.add_run('  Avión / Autobús / Transporte terrestre privado / Otro')
        set_run_font(run3, size=12)
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run('Proveedor: ')
        set_run_font(run1, size=12, bold=True)
        if venta.proveedor:
            run2 = p.add_run(venta.proveedor.nombre)
            set_run_font(run2, size=12)
        else:
            run2 = p.add_run('________________________')
            set_run_font(run2, size=12)
            run2.font.underline = True
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run('Número de viaje o clave de reservación: ')
        set_run_font(run1, size=12, bold=True)
        run2 = p.add_run('________________________')
        set_run_font(run2, size=12)
        run2.font.underline = True
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run('Incluye equipaje?: ')
        set_run_font(run1, size=12, bold=True)
        run2 = p.add_run('________________________')
        set_run_font(run2, size=12)
        run2.font.underline = True
        run3 = p.add_run('  Sí / No / Detallar: ')
        set_run_font(run3, size=12)
        run4 = p.add_run('________________________')
        set_run_font(run4, size=12)
        run4.font.underline = True
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('(La AGENCIA no garantiza horarios o cambios operados por el transportista.)')
        set_run_font(run, size=11, color=RGBColor(102, 102, 102))
        run.italic = True
        
        # Sección 4: HOSPEDAJE
        add_paragraph(doc, '4. HOSPEDAJE', size=14, bold=True, color=MOVUMS_BLUE, space_before=15, space_after=8)
        
        campos_hospedaje = [
            ('Hotel:', ''),
            ('Categoría:', ''),
            ('Dirección:', ''),
            ('Noches incluidas:', ''),
            ('Tipo de habitación:', ''),
        ]
        
        for label, value in campos_hospedaje:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run1 = p.add_run(f'{label} ')
            set_run_font(run1, size=12, bold=True)
            run2 = p.add_run('________________________')
            set_run_font(run2, size=12)
            run2.font.underline = True
        
        # Plan de alimentos
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run('Plan de alimentos:')
        set_run_font(run1, size=12, bold=True)
        
        planes = ['Sólo hospedaje', 'Desayuno', 'Media pensión', 'Todo incluido', 'Otro:']
        for plan in planes:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(4)
            run1 = p.add_run('○ ')
            set_run_font(run1, size=12)
            run2 = p.add_run(plan)
            set_run_font(run2, size=12)
            if plan == 'Otro:':
                run3 = p.add_run(' ________________________')
                set_run_font(run3, size=12)
                run3.font.underline = True
        
        # Sección 5: SERVICIOS ADICIONALES INCLUIDOS
        add_paragraph(doc, '5. SERVICIOS ADICIONALES INCLUIDOS', size=14, bold=True, color=MOVUMS_BLUE, space_before=15, space_after=8)
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('(Marca los que apliquen)')
        set_run_font(run, size=11, color=RGBColor(102, 102, 102))
        run.italic = True
        
        servicios_adicionales = [
            'Traslado aeropuerto-hotel-aeropuerto',
            ('Tours o excursiones (describir):', True),
            'Coordinador o guía turístico',
            'Entradas o accesos',
            'Seguro de viajero nacional (si aplica)',
            ('Otros servicios incluidos:', True),
        ]
        
        for servicio in servicios_adicionales:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(4)
            run1 = p.add_run('☐ ')
            set_run_font(run1, size=12)
            if isinstance(servicio, tuple):
                run2 = p.add_run(servicio[0])
                set_run_font(run2, size=12)
                run3 = p.add_run(' ________________________')
                set_run_font(run3, size=12)
                run3.font.underline = True
            else:
                run2 = p.add_run(servicio)
                set_run_font(run2, size=12)
        
        # Sección 6: SERVICIOS NO INCLUIDOS
        doc.add_page_break()
        add_paragraph(doc, '6. SERVICIOS NO INCLUIDOS', size=14, bold=True, color=MOVUMS_BLUE, space_before=15, space_after=8)
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('(Para evitar reclamaciones)')
        set_run_font(run, size=11, color=RGBColor(102, 102, 102))
        run.italic = True
        
        servicios_no_incluidos = [
            'Impuestos locales, resort fees o cuotas gubernamentales',
            'Propinas',
            'Servicios no listados como incluidos',
            'Actividades opcionales',
            'Gastos personales',
            'Sobrepeso de equipaje',
            'Comidas no contempladas',
        ]
        
        for servicio in servicios_no_incluidos:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(f'• {servicio}')
            set_run_font(run, size=12)
        
        # Sección 7: PRECIO Y CONDICIONES ECONÓMICAS
        add_paragraph(doc, '7. PRECIO Y CONDICIONES ECONÓMICAS', size=14, bold=True, color=MOVUMS_BLUE, space_before=15, space_after=8)
        
        # Precio total
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run('Precio total del paquete: ')
        set_run_font(run1, size=12, bold=True)
        if venta.tipo_viaje == 'INT':
            if venta.costo_total_con_modificacion_usd:
                run2 = p.add_run(f'USD ${format_currency(venta.costo_total_con_modificacion_usd)}')
                set_run_font(run2, size=12)
            else:
                run2 = p.add_run('USD $________________________')
                set_run_font(run2, size=12)
                run2.font.underline = True
        else:
            if venta.costo_total_con_modificacion:
                run2 = p.add_run(f'${format_currency(venta.costo_total_con_modificacion)} MXN')
                set_run_font(run2, size=12)
            else:
                run2 = p.add_run('$________________________ MXN')
                set_run_font(run2, size=12)
                run2.font.underline = True
        
        # Anticipo recibido
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run('Anticipo recibido: ')
        set_run_font(run1, size=12, bold=True)
        if venta.tipo_viaje == 'INT':
            if venta.cantidad_apertura_usd:
                run2 = p.add_run(f'USD ${format_currency(venta.cantidad_apertura_usd)}')
                set_run_font(run2, size=12)
            else:
                run2 = p.add_run('USD $________________________ MXN')
                set_run_font(run2, size=12)
                run2.font.underline = True
        else:
            if venta.cantidad_apertura:
                run2 = p.add_run(f'${format_currency(venta.cantidad_apertura)} MXN')
                set_run_font(run2, size=12)
            else:
                run2 = p.add_run('$________________________ MXN')
                set_run_font(run2, size=12)
                run2.font.underline = True
        
        # Saldo pendiente
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run('Saldo pendiente: ')
        set_run_font(run1, size=12, bold=True)
        if venta.tipo_viaje == 'INT':
            if venta.saldo_restante_usd:
                run2 = p.add_run(f'USD ${format_currency(venta.saldo_restante_usd)}')
                set_run_font(run2, size=12)
            else:
                run2 = p.add_run('USD $________________________ MXN')
                set_run_font(run2, size=12)
                run2.font.underline = True
        else:
            if saldo_pendiente:
                run2 = p.add_run(f'${format_currency(saldo_pendiente)} MXN')
                set_run_font(run2, size=12)
            else:
                run2 = p.add_run('$________________________ MXN')
                set_run_font(run2, size=12)
                run2.font.underline = True
        
        # Fecha límite de pago
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run('Fecha límite de pago total: ')
        set_run_font(run1, size=12, bold=True)
        run2 = p.add_run('________________________')
        set_run_font(run2, size=12)
        run2.font.underline = True
        if venta.fecha_vencimiento_pago:
            run3 = p.add_run(f' {format_date(venta.fecha_vencimiento_pago)}')
            set_run_font(run3, size=12)
        else:
            run3 = p.add_run(' //2025')
            set_run_font(run3, size=12)
        
        # Desglose de pagos
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run('Desglose de pagos (si aplica):')
        set_run_font(run1, size=12, bold=True)
        
        for i in range(1, 4):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(1.0)
            p.paragraph_format.space_after = Pt(4)
            run1 = p.add_run(f'Parcialidad {i}: $________________________  Fecha: ________________________ //2025')
            set_run_font(run1, size=12)
            # Subrayar espacios en blanco
            for run in p.runs:
                if '________________' in run.text:
                    run.font.underline = True
        
        # Sección 8: DOCUMENTACIÓN ENTREGADA AL CLIENTE
        doc.add_page_break()
        add_paragraph(doc, '8. DOCUMENTACIÓN ENTREGADA AL CLIENTE', size=14, bold=True, color=MOVUMS_BLUE, space_before=15, space_after=8)
        
        documentacion = [
            'Contrato firmado',
            'Copia de esta caratula',
            'Itinerario preliminar',
            'Políticas del proveedor',
            'Comprobantes de pago',
            'Claves de reservación',
            'Información de contacto para emergencias',
        ]
        
        for doc_item in documentacion:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(f'☐ {doc_item}')
            set_run_font(run, size=12)
        
        # Sección 9: DECLARACIÓN DEL CLIENTE
        add_paragraph(doc, '9. DECLARACIÓN DEL CLIENTE', size=14, bold=True, color=MOVUMS_BLUE, space_before=15, space_after=8)
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run('EL CLIENTE declara que:')
        set_run_font(run1, size=12, bold=True)
        
        declaraciones = [
            'Ha revisado y entendido toda la información contenida en este Anexo.',
            'Proporcionó datos veraces y completos.',
            'Acepta las condiciones del servicio, políticas de proveedores y cláusulas del contrato.',
        ]
        
        for decl in declaraciones:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(f'• {decl}')
            set_run_font(run, size=12)
        
        # Sección 10: TÉRMINOS Y CONDICIONES
        doc.add_page_break()
        add_paragraph(doc, '10. TÉRMINOS Y CONDICIONES', size=14, bold=True, color=MOVUMS_BLUE, space_before=15, space_after=8)
        
        # Título del contrato
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('CONTRATO DE MEDIACIÓN PARA LA PRESTACIÓN DE SERVICIOS TURÍSTICOS')
        set_run_font(run, size=12, bold=True)
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run('QUE CELEBRAN POR UNA PARTE LA AGENCIA DE VIAJES "GRUPO IMVED, S.A. DE C.V." ACTUANDO EN USO DE SU NOMBRE COMERCIAL MOVUMS THE TRAVEL STORE, EN ADELANTE DENOMINADA COMO "LA AGENCIA", Y POR LA OTRA EL (LA) C')
        set_run_font(run, size=12)
        run2 = p.add_run('________________________')
        set_run_font(run2, size=12)
        run2.font.underline = True
        run3 = p.add_run(' A QUIEN EN LO SUCESIVO SE LE DENOMINARÁ "EL CLIENTE", AL TENOR DE LAS SIGUIENTES DEFINICIONES, DECLARACIONES Y CLÁUSULAS:')
        set_run_font(run3, size=12)
        
        # GLOSARIO
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('GLOSARIO')
        set_run_font(run, size=12, bold=True)
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run('Para efectos del presente contrato, se entiende por:')
        set_run_font(run, size=12)
        
        glosario_items = [
            'Agencia: Es el proveedor de servicios turísticos que intermedia, contrata u ofrece servicios o productos turístico nacionales, previo pago de un precio cierto y determinado.',
            'Cliente: Consumidor que contrata los servicios turísticos nacionales mediante el pago de un precio cierto y determinado.',
            'Paquete turístico: Integración de uno o más servicios turísticos en un solo producto, ofrecidos al Cliente y detallado en el Anexo del presente contrato.',
            'Servicio turístico: Prestación de carácter comercial en transporte nacional, hospedaje, alimentación, excursiones u otros servicios relacionados, detallados en el Anexo del presente contrato.',
            'Caratula: Documento que detalla servicios, fechas, precios y condiciones del servicio turístico contratado.',
        ]
        
        for item in glosario_items:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(4)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run1 = p.add_run('◆ ')
            set_run_font(run1, size=12)
            run2 = p.add_run(item)
            set_run_font(run2, size=12)
        
        # DECLARACIONES
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('DECLARACIONES')
        set_run_font(run, size=12, bold=True)
        
        # Declaración I - LA AGENCIA
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('I. Declara LA AGENCIA:')
        set_run_font(run, size=12, bold=True)
        
        declaraciones_agencia = [
            'Ser una persona moral legalmente constituida conforme a las leyes mexicanas.',
            'Ser un Prestador de Servicios Turísticos con Razón Social: GRUPO IMVED, S.A. de C.V.',
            'Ser la única propietaria de la marca MOVUMS THE TRAVEL STORE',
            'RFC GIM190722FS7 y domicilio ubicado en: Plaza Mora, Juárez Sur, 321, interior 18, Colonia Centro, Texcoco, Estado de México, C.P. 56100.',
            'Teléfono, correo electrónico y horario de atención al público: 59 59319954, 5951255279 ventas@movums.com, lunes a sábado de 09:00 a 18:00 horas.',
            'Contar con infraestructura, personal capacitado y experiencia suficiente para la prestación de los servicios turísticos contratados.',
            'Haber informado previamente al Cliente sobre los precios, tarifas, condiciones, características y costo total del servicio turístico contratado.',
        ]
        
        for item in declaraciones_agencia:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(4)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run1 = p.add_run('◆ ')
            set_run_font(run1, size=12)
            run2 = p.add_run(item)
            set_run_font(run2, size=12)
        
        # Declaración II - EL CLIENTE
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('II. Declara EL CLIENTE:')
        set_run_font(run, size=12, bold=True)
        
        declaraciones_cliente = [
            'Ser persona física/moral con capacidad legal y económica para obligarse en términos del presente contrato.',
            'En caso de persona moral: Ser una persona moral legalmente constituida conforme a las leyes mexicanas, conforme lo acredita con copia del instrumento número ____________, de fecha ___________, otorgado ante la Fe del Notario Público Número ____, de ____________, y que el(la) C. _________ __________________ en este acto interviene en su carácter de Representante Legal, calidad que acredita con copia del instrumento número _______, de fecha _________, otorgada ante la Fe del Notario Público número _______ del _________, facultad y calidad que no le han sido revocadas, modificadas o limitadas a la fecha de firma del presente contrato.',
            'Encontrarse inscrito en el Registro Federal de Contribuyentes con la clave que ha manifestado.',
            'Haber recibido previamente de LA AGENCIA información útil, precisa, veraz y detallada sobre los servicios objeto del presente contrato.',
            'Proporciona su nombre, domicilio, número telefónico y correo electrónico, tal y como lo ha señalado en la caratula de prestación de servicios, acreditando los mismos con copia de los documentos idóneos para tal efecto.',
        ]
        
        for item in declaraciones_cliente:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(4)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run1 = p.add_run('◆ ')
            set_run_font(run1, size=12)
            # Procesar líneas en blanco en el texto
            item_text = item
            parts = item_text.split('____________')
            for i, part in enumerate(parts):
                if part:
                    run2 = p.add_run(part)
                    set_run_font(run2, size=12)
                if i < len(parts) - 1:
                    run_blank = p.add_run('____________')
                    set_run_font(run_blank, size=12)
                    run_blank.font.underline = True
        
        # CLÁUSULAS - 3 saltos de línea para bajar de página
        add_paragraph(doc, '', space_after=0)
        add_paragraph(doc, '', space_after=0)
        add_paragraph(doc, '', space_after=0)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('CLÁUSULAS')
        set_run_font(run, size=12, bold=True)
        
        # Cláusulas
        clausulas = [
            ('PRIMERA. CONSENTIMIENTO.', 'Las partes manifiestan su voluntad de celebrar el presente contrato, cuya naturaleza jurídica es la mediación para la prestación de servicios turísticos.'),
            ('SEGUNDA. OBJETO.', 'LA AGENCIA intermediará, contratará u ofrecerá servicios turísticos detallados en la CARATULA, previo pago del Cliente de un precio cierto y determinado.'),
            ('TERCERA. PRECIO, FORMA Y LUGAR DE PAGO.', 'Las partes manifiestan su conformidad en que el precio total a pagar por EL CLIENTE como contraprestación del Servicio turístico, es la cantidad que por cada concepto se indica en la CARATULA de este Contrato. El importe señalado en la CARATULA, contempla todas las cantidades y conceptos referentes al Servicio turístico, por lo que LA AGENCIA se obliga a respetar en todo momento dicho costo sin poder cobrar otra cantidad o condicionar la prestación del Servicio turístico contratado a la adquisición de otro servicio no requerido por El cliente, salvo que El cliente autorice de manera escrita algún otro cobro no estipulado en el presente Contrato. EL CLIENTE efectuará el pago pactado por el Servicio turístico señalado en la caratula del presente Contrato en los términos y condiciones acordadas pudiendo ser:\na) Al contado: en efectivo, con tarjeta de débito, tarjeta de crédito, transferencia bancaria, y/o cheque en el domicilio de la agencia en moneda nacional, sin menoscabo de poderlo hacer en moneda extranjera al tipo de cambio publicado en el Diario Oficial de la Federación al día en que el pago se efectúe.\nb) A plazos: El cliente podrá, previo acuerdo con La agencia a pagar en parcialidades, para lo cual, La agencia deberá de entregar a El CLIENTE la información por escrito de las fechas, así como los montos parciales a pagar.\nc) En caso de que El cliente realice el pago con cheque y no se cubra el pago por causas imputables al librador, La agencia tendrá el derecho de realizar el cobro adicional del 20% (veinte por ciento) del valor del documento, por concepto de daños y perjuicios, en caso de que el cheque sea devuelto por causas imputables al librador, conforme al artículo 193 de la Ley General del Títulos y Operaciones de Crédito.'),
            ('QUINTA. OBLIGACIONES DE LA AGENCIA.', 'LA AGENCIA SE OBLIGA A:\nA) Cumplir lo pactado en el contrato.\nB) Entregar a EL CLIENTE copia del contrato y constancias de reservación.\nC) Proporcionar a EL CLIENTE boletos, claves de reservación y documentos de viaje.\nD) Auxiliar a EL CLIENTE en emergencias y gestionar indemnizaciones relacionadas con el servicio contratado\nE) Solicitar los Servicios turísticos que se especifican en la caratula de este Contrato por cuenta de EL CLIENTE de acuerdo a la disponibilidad de los mismos, a contratarlos fungiendo como intermediario entre éste y las personas encargadas de proporcionar directamente el Servicio turístico.\nF) Coadyuvar a EL CLIENTE para reclamar ante el prestador del servicio final, las indemnizaciones que correspondan.\nG) Respetar la Ley Federal de Protección al Consumidor y la NOM-010-TUR-2001.'),
            ('SEXTA. OBLIGACIONES DE EL CLIENTE:', 'Cumplir con lo establecido en el presente contrato:\nA) Proporcionar previo a la prestación del servicio los datos generales veraces y documentos requeridos para los servicios contratados (como pueden ser de manera enunciativa más no limitativa, el nombre, edad, identificación, comprobante de domicilio, pasaporte, visas, vacunas, constancia de situación fiscal, número telefónico, correo electrónico). Proporcionará sus propios datos y documentos de su persona así como el de las personas que lo acompañen.\nB) Realizar pagos a la AGENCIA conforme a lo pactado en el presente contrato.\nC) Respetar reglamentos de prestadores finales.\nD) Notificar por lo menos con CINCO DÍAS HÁBILES y por escrito a LA AGENCIA cualquier cambio o cancelación una vez aceptado el servicio.'),
            ('SÉPTIMA. VIGENCIA.', 'El contrato estará vigente mientras se presten los servicios y se cumplan las obligaciones de pago, tiempo en que el presente Contrato surtirá todos sus efectos legales.'),
            ('OCTAVA. CASO FORTUITO Y FUERZA MAYOR.', 'Se entiende por caso fortuito o fuerza mayor aquellos hechos o acontecimientos ajenos a la voluntad de las partes, que sean imprevisibles, irresistibles, insuperables y que no provengan de negligencia, dolo o falta de cuidado de alguna de ellas. No se considerarán caso fortuito o fuerza mayor las enfermedades personales de EL CLIENTE o de sus acompañantes. EL CLIENTE reconoce que la AGENCIA no será responsable por errores, omisiones, falta de entrega de documentos, información incompleta o inexacta, ni por cualquier otra actuación u omisión atribuible al propio CLIENTE que afecte la reservación, emisión de boletos, acceso a servicios turísticos, cambios, cancelaciones o cualquier trámite derivado del presente contrato. Cuando el servicio turístico no pueda prestarse total o parcialmente por caso fortuito o fuerza mayor, la AGENCIA reembolsará a EL CLIENTE las cantidades que, conforme a las políticas de los prestadores finales (aerolíneas, hoteles, operadores, etc.), sean efectivamente recuperables y devueltas a la AGENCIA. EL CLIENTE tendrá derecho a recibir el reembolso correspondiente únicamente respecto de los importes efectivamente recuperados. En caso de que el servicio turístico se haya prestado de manera parcial, EL CLIENTE tendrá derecho a un reembolso proporcional exclusivamente respecto de los servicios no utilizados, conforme a lo que determine el proveedor correspondiente.'),
            ('NOVENA. CAMBIOS DE ORDEN DE LOS SERVICIOS CON AUTORIZACIÓN DE EL CLIENTE.', 'La agencia podrá modificar el orden de los Servicios turísticos indicados en el presente Contrato, para un mejor desarrollo de los mismos o por las causas que así lo justifiquen, siempre y cuando respete la cantidad y calidad de los Servicios turísticos que se hayan contratado. Este será con la autorización por escrito de EL CLIENTE, sea cual fuese la causa. El cliente no podrá hacer cambios de fechas, rutas, ni servicios, sin previa autorización de La agencia, en caso de que dichos cambios tengan un costo, éste será indicado en al CARATULA del presente Contrato. EL CLIENTE reconoce que, una vez firmado el presente contrato y realizado el anticipo, pago parcial o total, los pagos efectuados no son cancelables ni reembolsables, en virtud de que la AGENCIA realiza de manera inmediata gestiones, reservaciones y pagos a terceros proveedores, los cuales se rigen por políticas propias de cancelación y reembolso que no dependen de la AGENCIA. EL CLIENTE acepta que cualquier solicitud de cambio, corrección o modificación respecto a fechas, nombres, itinerarios, servicios contratados o cualquier otro aspecto, estará sujeta a la disponibilidad de los proveedores, así como al pago de cargos adicionales o penalidades, conforme a las políticas vigentes de dichos proveedores.'),
            ('DÉCIMA. CANCELACIÓN.', 'EL CLIENTE reconoce que, una vez firmado el presente Contrato y realizado el anticipo, pago parcial o total, los pagos no son cancelables ni reembolsables, debido a que la AGENCIA realiza de manera inmediata pagos, reservaciones y gestiones con terceros proveedores, cuyas políticas no permiten cancelaciones ni devoluciones. Cualquier solicitud de cancelación o modificación deberá realizarse por escrito, pero no dará derecho a devolución, salvo que algún proveedor permita recuperar total o parcialmente los montos pagados, caso en el cual la AGENCIA entregará al CLIENTE únicamente las cantidades efectivamente devueltas por dicho proveedor. Las modificaciones estarán sujetas a disponibilidad y podrán generar cargos adicionales conforme a las políticas de los prestadores finales. La presente cláusula aplica únicamente a solicitudes voluntarias de cancelación formuladas por EL CLIENTE. Lo anterior es independiente de las consecuencias aplicables por rescisión por incumplimiento, reguladas en las cláusulas siguientes.'),
            ('DÉCIMA PRIMERA. VUELOS.', 'EL CLIENTE reconoce que los servicios aéreos incluidos en el paquete vacacional son operados exclusivamente por la aerolínea correspondiente, por lo que Movums The Travel Store no es responsable por cambios de itinerario, demoras, reprogramaciones, sobreventas, cancelaciones, modificaciones operativas o cualquier otra decisión adoptada por la aerolínea, toda vez que dichos actos son ajenos al control de la AGENCIA. EL CLIENTE acepta que toda compensación, reembolso, cambio o beneficio derivado de acciones de la aerolínea está sujeto exclusivamente a las políticas y procedimientos de dicha aerolínea, y que la AGENCIA actuará únicamente como intermediaria en la gestión correspondiente.'),
            ('DÉCIMA SEGUNDA. RESCISIÓN.', 'Procede si alguna parte incumple lo pactado o si el servicio no corresponde a lo solicitado. En caso de rescisión del presente Contrato, la parte que incumpla deberá de pagar lo correspondiente a la pena convencional. La AGENCIA podrá dar por terminado el presente contrato cuando EL CLIENTE no realice los depósitos o pagos en las fechas pactadas. En este supuesto, la AGENCIA notificará al CLIENTE mediante los medios de contacto proporcionados, y dicha terminación se considerará efectiva desde la fecha del incumplimiento. El CLIENTE reconoce que la falta de pago oportuno constituye un incumplimiento del contrato y acepta que los anticipos podrán aplicarse a cargos, penalidades o gastos ya generados conforme a las políticas de proveedores y prestadores de servicios turísticos. La rescisión no será considerada como una cancelación voluntaria, sino como una consecuencia jurídica del incumplimiento de cualquiera de las partes.'),
            ('DÉCIMA TERCERA. PENA CONVENCIONAL.', 'La parte incumplida pagará el 20% (veinte por ciento) del precio total del servicio, sin incluir IVA.'),
            ('DÉCIMA CUARTA. RESERVACIONES Y PAGOS.', 'La aceptación y formalización del presente contrato se considerará efectiva una vez que EL CLIENTE envíe el contrato debidamente firmado y efectúe el anticipo, pago parcial o total, mismo que no es reembolsable, en virtud de que Movums The Travel Store realiza gestiones inmediatas con terceros proveedores para asegurar la disponibilidad de los servicios solicitados.'),
            ('DÉCIMA QUINTA. JURISDICCIÓN.', 'Las partes se someten a PROFECO y, en su caso, a tribunales competentes de Texcoco, Estado de México.'),
        ]
        
        for idx, (titulo, contenido) in enumerate(clausulas):
            # Salto de página antes de QUINTA. OBLIGACIONES DE LA AGENCIA
            if 'QUINTA. OBLIGACIONES DE LA AGENCIA' in titulo:
                doc.add_page_break()
            
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run1 = p.add_run(titulo)
            set_run_font(run1, size=12, bold=True)
            
            # Procesar el contenido que puede tener saltos de línea y subsecciones
            contenido_lines = contenido.split('\n')
            for i, line in enumerate(contenido_lines):
                if i == 0 and line.strip():
                    # Primera línea va en el mismo párrafo
                    run2 = p.add_run(' ' + line)
                    set_run_font(run2, size=12)
                elif line.strip():
                    # Líneas subsecuentes en nuevos párrafos
                    p2 = doc.add_paragraph()
                    p2.paragraph_format.left_indent = Inches(0.5)
                    p2.paragraph_format.space_after = Pt(3)
                    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    run = p2.add_run(line)
                    set_run_font(run, size=12)
        
        # Sección 11: FIRMAS (combinada con las firmas del contrato)
        add_paragraph(doc, '11. FIRMAS', size=14, bold=True, color=MOVUMS_BLUE, space_before=15, space_after=10)
        
        # Firma Cliente
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('CLIENTE')
        set_run_font(run, size=12, bold=True)
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(15)
        run1 = p.add_run('Nombre y firma: ')
        set_run_font(run1, size=12, bold=True)
        run2 = p.add_run('___________________________________________')
        set_run_font(run2, size=12)
        run2.font.underline = True
        
        # Firma Agencia
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run('AGENCIA – Movums The Travel Store')
        set_run_font(run, size=12, bold=True)
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run('Nombre y firma del representante: ')
        set_run_font(run1, size=12, bold=True)
        run2 = p.add_run('____________________________')
        set_run_font(run2, size=12)
        run2.font.underline = True
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run1 = p.add_run('Fecha: ')
        set_run_font(run1, size=12, bold=True)
        run2 = p.add_run('//2025')
        set_run_font(run2, size=12)
        
        # Preparar respuesta HTTP
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        nombre_cliente_safe = venta.cliente.nombre_completo_display.replace(' ', '_').replace('/', '_')
        filename = f"Contrato_Venta_{venta.pk}_{nombre_cliente_safe}.docx"
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        response['Content-Length'] = str(len(buffer.getvalue()))
        response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response['Pragma'] = 'no-cache'
        response['Expires'] = '0'
        
        buffer.close()
        return response


class ContratoHospedajePDFView(LoginRequiredMixin, DetailView):
    """
    Vista para generar el Contrato de Hospedaje en formato DOCX.
    Formato específico para ventas de hospedaje basado en las imágenes proporcionadas.
    """
    model = VentaViaje
    
    def get(self, request, *args, **kwargs):
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.oxml.ns import qn
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.section import WD_SECTION
        except ImportError:
            return HttpResponse("Error: python-docx no está instalado. Ejecuta: pip install python-docx", status=500)

        self.object = self.get_object() 
        venta = self.object
        cliente = venta.cliente
        
        # Verificar que sea hospedaje
        if not venta.servicios_seleccionados or 'HOS' not in venta.servicios_seleccionados:
            return HttpResponse("Error: Esta venta no es de hospedaje.", status=400)
        
        # Obtener datos de la cotización si existe
        hotel_nombre = ''
        habitacion = ''
        plan_alimentos = ''
        destino = ''
        
        if venta.cotizacion_origen:
            cotizacion = venta.cotizacion_origen
            propuestas = cotizacion.propuestas if isinstance(cotizacion.propuestas, dict) else {}
            
            # Obtener información del hotel desde propuestas
            if propuestas.get('hoteles'):
                hoteles = propuestas.get('hoteles', [])
                if isinstance(hoteles, list) and len(hoteles) > 0:
                    hotel_data = hoteles[0] if isinstance(hoteles[0], dict) else {}
                    hotel_nombre = hotel_data.get('hotel', '') or hotel_data.get('nombre', '')
                    habitacion = hotel_data.get('habitacion', '') or hotel_data.get('tipo_habitacion', '')
                    plan_alimentos = hotel_data.get('plan_alimentos', '') or hotel_data.get('regimen', '')
            
            # Obtener destino desde cotización
            destino = cotizacion.destino or ''
        
        # Si no hay destino en cotización, usar servicios_detalle_desde_logistica o un valor por defecto
        if not destino:
            destino = venta.servicios_detalle_desde_logistica or 'HOSPEDAJE'
        
        # Calcular valores financieros
        from decimal import Decimal
        precio_total = venta.costo_total_con_modificacion or Decimal('0.00')
        anticipo = venta.cantidad_apertura or Decimal('0.00')
        saldo_pendiente = max(Decimal('0.00'), precio_total - venta.total_pagado)
        fecha_limite_pago = venta.fecha_vencimiento_pago
        
        # Convertir montos a texto (la función retorna "tres mil pesos 00/100 M.N.")
        anticipo_texto = numero_a_texto(float(anticipo))
        # Ajustar formato: cambiar "M.N." por "MXN" y ajustar formato según imagen
        anticipo_texto = anticipo_texto.replace('M.N.', 'MXN')
        
        # Usar plantilla con membrete si existe
        template_path = os.path.join(settings.BASE_DIR, 'static', 'docx', 'membrete.docx')
        if os.path.exists(template_path):
            doc = Document(template_path)
        else:
            doc = Document()
        
        # Configurar fuente predeterminada
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        
        # Colores
        MOVUMS_BLUE = RGBColor(0, 74, 142)  # #004a8e
        TEXT_COLOR = RGBColor(47, 47, 47)  # #2f2f2f
        
        def set_run_font(run, size=12, bold=False, color=TEXT_COLOR):
            run.font.name = 'Arial'
            run.font.size = Pt(size)
            run.bold = bold
            run.font.color.rgb = color
        
        def format_date(value):
            if not value:
                return ''
            try:
                if isinstance(value, date):
                    # Formato: DD MES YYYY (ej: 02 ABRIL 2026)
                    meses = {
                        1: 'ENERO', 2: 'FEBRERO', 3: 'MARZO', 4: 'ABRIL',
                        5: 'MAYO', 6: 'JUNIO', 7: 'JULIO', 8: 'AGOSTO',
                        9: 'SEPTIEMBRE', 10: 'OCTUBRE', 11: 'NOVIEMBRE', 12: 'DICIEMBRE'
                    }
                    return f"{value.day:02d} {meses[value.month]} {value.year}"
                return str(value)
            except:
                return ''
        
        def format_currency(value):
            if value in (None, '', 0):
                return '0.00'
            try:
                number = Decimal(str(value).replace(',', ''))
            except (ValueError, TypeError) as e:
                logger.warning(f"Error al formatear moneda: {e}")
                return str(value)
            return f"{number:,.2f}"
        
        # ============================================
        # PÁGINA 1: CONTRATO PRINCIPAL
        # ============================================
        # El membrete ya está en la plantilla, no necesitamos agregarlo manualmente
        
        # Título principal
        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_titulo.paragraph_format.space_before = Pt(0)
        p_titulo.paragraph_format.space_after = Pt(12)
        run_titulo = p_titulo.add_run('CONTRATO DE SERVICIOS TURÍSTICOS')
        set_run_font(run_titulo, size=18, bold=True, color=MOVUMS_BLUE)
        
        # Fecha y Destino (en la misma línea)
        p_fecha_destino = doc.add_paragraph()
        p_fecha_destino.paragraph_format.space_after = Pt(12)
        run_fecha_label = p_fecha_destino.add_run('Fecha: ')
        set_run_font(run_fecha_label, size=12, bold=True)
        run_fecha_val = p_fecha_destino.add_run(format_date(date.today()))
        set_run_font(run_fecha_val, size=12)
        
        # Destino en línea separada
        p_destino = doc.add_paragraph()
        p_destino.paragraph_format.space_after = Pt(12)
        run_destino_label = p_destino.add_run('Destino: ')
        set_run_font(run_destino_label, size=12, bold=True)
        run_destino_val = p_destino.add_run(destino.upper() if destino else 'HOSPEDAJE')
        set_run_font(run_destino_val, size=12, bold=True)
        
        # Información del cliente y monto recibido
        p_cliente = doc.add_paragraph()
        p_cliente.paragraph_format.space_after = Pt(6)
        run_texto1 = p_cliente.add_run('Movums The travel Store, con domicilio Plaza Mora, Juárez Sur 321 Local 18 CP. 56100 Texcoco Estado de México, recibió de: ')
        set_run_font(run_texto1, size=12)
        run_cliente = p_cliente.add_run(cliente.nombre_completo_display.upper())
        set_run_font(run_cliente, size=12, bold=True)
        run_cliente.font.underline = True
        run_texto2 = p_cliente.add_run(' la cantidad de:')
        set_run_font(run_texto2, size=12)
        
        # Monto recibido
        p_monto = doc.add_paragraph()
        p_monto.paragraph_format.space_after = Pt(12)
        run_monto_num = p_monto.add_run(f'${format_currency(anticipo)}')
        set_run_font(run_monto_num, size=12, bold=True)
        run_monto_num.font.underline = True
        run_monto_texto = p_monto.add_run(f' ({anticipo_texto}).')
        set_run_font(run_monto_texto, size=12)
        
        # Agregar espacio antes de los detalles del hospedaje
        p_espacio = doc.add_paragraph()
        p_espacio.paragraph_format.space_after = Pt(6)
        
        # Detalles del hospedaje
        p_detalles = doc.add_paragraph()
        p_detalles.paragraph_format.space_after = Pt(6)
        
        # FECHA DE IDA
        run_ida_label = p_detalles.add_run('FECHA DE IDA: ')
        set_run_font(run_ida_label, size=12, bold=True)
        run_ida_val = p_detalles.add_run(format_date(venta.fecha_inicio_viaje) if venta.fecha_inicio_viaje else '')
        set_run_font(run_ida_val, size=12, bold=True)
        
        p_detalles2 = doc.add_paragraph()
        p_detalles2.paragraph_format.space_after = Pt(6)
        run_regreso_label = p_detalles2.add_run('FECHA DE REGRESO: ')
        set_run_font(run_regreso_label, size=12, bold=True)
        run_regreso_val = p_detalles2.add_run(format_date(venta.fecha_fin_viaje) if venta.fecha_fin_viaje else '')
        set_run_font(run_regreso_val, size=12, bold=True)
        
        # Pasajeros - obtener desde cotización si existe
        adultos = 0
        menores = 0
        if venta.cotizacion_origen:
            cotizacion = venta.cotizacion_origen
            adultos = cotizacion.adultos or 0
            menores = cotizacion.menores or 0
        
        # Si no hay cotización, intentar contar desde pasajeros
        if adultos == 0 and venta.pasajeros:
            # Contar líneas no vacías como aproximación
            pasajeros_lista = [p.strip() for p in venta.pasajeros.split('\n') if p.strip()]
            adultos = len(pasajeros_lista) if pasajeros_lista else 1
        
        p_pasajeros = doc.add_paragraph()
        p_pasajeros.paragraph_format.space_after = Pt(6)
        edades_menores_texto = ''
        if venta.edades_menores:
            edades_menores_texto = f' ({venta.edades_menores})'
        elif menores > 0:
            edades_menores_texto = ' (Edades: )'  # Dejar espacio para completar manualmente
        
        run_pasajeros_label = p_pasajeros.add_run('PASAJEROS: ')
        set_run_font(run_pasajeros_label, size=12, bold=True)
        if menores > 0:
            run_pasajeros_val = p_pasajeros.add_run(f'{adultos} adultos + {menores} Menor{edades_menores_texto}')
        else:
            run_pasajeros_val = p_pasajeros.add_run(f'{adultos} adultos')
        set_run_font(run_pasajeros_val, size=12, bold=True)
        
        # Hotel
        p_hotel = doc.add_paragraph()
        p_hotel.paragraph_format.space_after = Pt(6)
        run_hotel_label = p_hotel.add_run('Hotel: ')
        set_run_font(run_hotel_label, size=12, bold=True)
        run_hotel_val = p_hotel.add_run(hotel_nombre.upper() if hotel_nombre else '')
        set_run_font(run_hotel_val, size=12, bold=True)
        
        # Habitación
        p_habitacion = doc.add_paragraph()
        p_habitacion.paragraph_format.space_after = Pt(6)
        run_habitacion_label = p_habitacion.add_run('Habitación: ')
        set_run_font(run_habitacion_label, size=12, bold=True)
        run_habitacion_val = p_habitacion.add_run(habitacion.upper() if habitacion else '')
        set_run_font(run_habitacion_val, size=12, bold=True)
        
        # Plan de Alimentos
        p_plan = doc.add_paragraph()
        p_plan.paragraph_format.space_after = Pt(12)
        run_plan_label = p_plan.add_run('Plan de Alimentos: ')
        set_run_font(run_plan_label, size=12, bold=True)
        if plan_alimentos:
            run_plan_val = p_plan.add_run(f'/ {plan_alimentos.upper()} /')
        else:
            run_plan_val = p_plan.add_run('/ SIN ALIMENTOS / SIN BEBIDAS /')
        set_run_font(run_plan_val, size=12, bold=True)
        
        # Agregar espacio antes de la sección económica
        p_espacio2 = doc.add_paragraph()
        p_espacio2.paragraph_format.space_after = Pt(6)
        
        # PRECIO Y CONDICIONES ECONÓMICAS
        p_seccion = doc.add_paragraph()
        p_seccion.paragraph_format.space_before = Pt(12)
        p_seccion.paragraph_format.space_after = Pt(8)
        run_seccion = p_seccion.add_run('PRECIO Y CONDICIONES ECONÓMICAS')
        set_run_font(run_seccion, size=12, bold=True)
        
        # Precio total
        p_precio = doc.add_paragraph()
        p_precio.paragraph_format.space_after = Pt(6)
        run_precio_label = p_precio.add_run('Precio total del paquete: ')
        set_run_font(run_precio_label, size=12, bold=True)
        run_precio_val = p_precio.add_run(f'${format_currency(precio_total)} MXN')
        set_run_font(run_precio_val, size=12, bold=True)
        run_precio_val.font.underline = True
        
        # Anticipo recibido
        p_anticipo = doc.add_paragraph()
        p_anticipo.paragraph_format.space_after = Pt(6)
        run_anticipo_label = p_anticipo.add_run('Anticipo recibido: ')
        set_run_font(run_anticipo_label, size=12, bold=True)
        run_anticipo_val = p_anticipo.add_run(f'${format_currency(anticipo)} MXN')
        set_run_font(run_anticipo_val, size=12, bold=True)
        run_anticipo_val.font.underline = True
        
        # Saldo pendiente
        p_saldo = doc.add_paragraph()
        p_saldo.paragraph_format.space_after = Pt(6)
        run_saldo_label = p_saldo.add_run('Saldo pendiente: ')
        set_run_font(run_saldo_label, size=12, bold=True)
        run_saldo_val = p_saldo.add_run(f'${format_currency(saldo_pendiente)} MXN')
        set_run_font(run_saldo_val, size=12, bold=True)
        run_saldo_val.font.underline = True
        
        # Fecha límite de pago
        p_fecha_limite = doc.add_paragraph()
        p_fecha_limite.paragraph_format.space_after = Pt(12)
        run_fecha_limite_label = p_fecha_limite.add_run('Fecha límite de pago total: ')
        set_run_font(run_fecha_limite_label, size=12, bold=True)
        if fecha_limite_pago:
            # Formato especial: DD/MES/YYYY (ej: 06/MARZO/2026)
            meses_cortos = {
                1: 'ENERO', 2: 'FEBRERO', 3: 'MARZO', 4: 'ABRIL',
                5: 'MAYO', 6: 'JUNIO', 7: 'JULIO', 8: 'AGOSTO',
                9: 'SEPTIEMBRE', 10: 'OCTUBRE', 11: 'NOVIEMBRE', 12: 'DICIEMBRE'
            }
            fecha_texto = f"{fecha_limite_pago.day:02d}/{meses_cortos[fecha_limite_pago.month]}/{fecha_limite_pago.year}"
            run_fecha_limite_val = p_fecha_limite.add_run(fecha_texto)
        else:
            run_fecha_limite_val = p_fecha_limite.add_run('')
        set_run_font(run_fecha_limite_val, size=12, bold=True)
        run_fecha_limite_val.font.underline = True
        
        # El footer ya está en el membrete de la plantilla, no necesitamos agregarlo manualmente
        
        # ============================================
        # PÁGINA 2: ANEXO DE TÉRMINOS Y CONDICIONES
        # ============================================
        
        doc.add_page_break()
        
        # Título del anexo (si es necesario, aunque en la imagen no aparece explícitamente)
        
        # EL CLIENTE declara que:
        p_declara = doc.add_paragraph()
        p_declara.paragraph_format.space_before = Pt(12)
        p_declara.paragraph_format.space_after = Pt(8)
        run_declara = p_declara.add_run('EL CLIENTE declara que:')
        set_run_font(run_declara, size=12, bold=True)
        
        declaraciones = [
            'Ha revisado y entendido toda la información contenida en este Anexo.',
            'Proporcionó datos veraces y completos.',
            'Acepta las condiciones del servicio, políticas de proveedores y cláusulas del contrato.',
        ]
        
        for decl in declaraciones:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(4)
            run_bullet = p.add_run('• ')
            set_run_font(run_bullet, size=12)
            run_text = p.add_run(decl)
            set_run_font(run_text, size=12)
        
        # NOTA
        p_nota = doc.add_paragraph()
        p_nota.paragraph_format.space_before = Pt(12)
        p_nota.paragraph_format.space_after = Pt(8)
        run_nota_label = p_nota.add_run('NOTA: Movums The Travel Store, ')
        set_run_font(run_nota_label, size=12, bold=True)
        run_nota_text = p_nota.add_run('se reserva el derecho de cancelar el contrato sin previo aviso en caso de que no se reciban los depósitos en las fechas estipuladas.')
        set_run_font(run_nota_text, size=12)
        
        # CANCELACIONES
        p_cancelaciones = doc.add_paragraph()
        p_cancelaciones.paragraph_format.space_before = Pt(12)
        p_cancelaciones.paragraph_format.space_after = Pt(8)
        run_cancelaciones = p_cancelaciones.add_run('CANCELACIONES:')
        set_run_font(run_cancelaciones, size=12, bold=True)
        
        textos_cancelacion = [
            'Entre la firma del contrato y pago de anticipo, parcial o total no se reembolsará ningún pago',
            'No es cancelable, ni reembolsable.',
            'Cualquier modificación puede ocasionar cargo extra; y están sujetos a disponibilidad'
        ]
        
        for texto in textos_cancelacion:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(texto)
            set_run_font(run, size=12)
        
        # FECHA DE VALIDEZ DEL CONTRATO
        p_validez = doc.add_paragraph()
        p_validez.paragraph_format.space_before = Pt(12)
        p_validez.paragraph_format.space_after = Pt(8)
        run_validez = p_validez.add_run('FECHA DE VALIDEZ DEL CONTRATO:')
        set_run_font(run_validez, size=12, bold=True)
        
        texto_validez = (
            'Las condiciones y precios antes mencionados serán mantenidos a la '
            'fecha límite de pago, a esta fecha el "CLIENTE" deberá haber cubierto el pago total del paquete o servicio turístico contratado. '
            'La aceptación de este contrato será efectiva una vez que el "CLIENTE" envié el contrato debidamente firmado y con el '
            'anticipo, pago parcial o pago total; no reembolsable; para que Movums The travel store, proceda con la reservación del servicio contratado'
        )
        
        p_texto_validez = doc.add_paragraph()
        p_texto_validez.paragraph_format.space_after = Pt(12)
        p_texto_validez.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Resaltar palabras clave
        partes = texto_validez.split('fecha límite de pago')
        if len(partes) > 1:
            run1 = p_texto_validez.add_run(partes[0])
            set_run_font(run1, size=12)
            run2 = p_texto_validez.add_run('fecha límite de pago')
            set_run_font(run2, size=12, bold=True)
            resto = 'anticipo, pago parcial o pago total'
            if resto in partes[1]:
                partes2 = partes[1].split(resto)
                run3 = p_texto_validez.add_run(partes2[0])
                set_run_font(run3, size=12)
                run4 = p_texto_validez.add_run(resto)
                set_run_font(run4, size=12, bold=True)
                if len(partes2) > 1:
                    run5 = p_texto_validez.add_run(partes2[1])
                    set_run_font(run5, size=12)
            else:
                run3 = p_texto_validez.add_run(partes[1])
                set_run_font(run3, size=12)
        else:
            run = p_texto_validez.add_run(texto_validez)
            set_run_font(run, size=12)
        
        # FIRMAS
        p_firmas = doc.add_paragraph()
        p_firmas.paragraph_format.space_before = Pt(20)
        p_firmas.paragraph_format.space_after = Pt(8)
        run_firmas = p_firmas.add_run('FIRMAS')
        set_run_font(run_firmas, size=12, bold=True)
        
        # CLIENTE
        p_cliente_firma = doc.add_paragraph()
        p_cliente_firma.paragraph_format.space_before = Pt(12)
        p_cliente_firma.paragraph_format.space_after = Pt(6)
        run_cliente_firma = p_cliente_firma.add_run('CLIENTE')
        set_run_font(run_cliente_firma, size=12, bold=True)
        
        p_nombre_firma = doc.add_paragraph()
        p_nombre_firma.paragraph_format.space_after = Pt(6)
        run_nombre_label = p_nombre_firma.add_run('Nombre y firma: ')
        set_run_font(run_nombre_label, size=12)
        run_linea = p_nombre_firma.add_run('_' * 50)
        set_run_font(run_linea, size=12)
        
        # AGENCIA
        p_agencia_firma = doc.add_paragraph()
        p_agencia_firma.paragraph_format.space_before = Pt(12)
        p_agencia_firma.paragraph_format.space_after = Pt(6)
        run_agencia_firma = p_agencia_firma.add_run('AGENCIA - Movums The Travel Store')
        set_run_font(run_agencia_firma, size=12, bold=True)
        
        p_representante = doc.add_paragraph()
        p_representante.paragraph_format.space_after = Pt(6)
        run_representante_label = p_representante.add_run('Nombre y firma del representante: ')
        set_run_font(run_representante_label, size=12)
        run_linea2 = p_representante.add_run('_' * 50)
        set_run_font(run_linea2, size=12)
        
        # El footer ya está en el membrete de la plantilla, no necesitamos agregarlo manualmente
        
        # Preparar respuesta HTTP
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        nombre_cliente_safe = cliente.nombre_completo_display.replace(' ', '_').replace('/', '_')
        filename = f"Contrato_Hospedaje_{venta.pk}_{nombre_cliente_safe}.docx"
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        response['Content-Length'] = str(len(buffer.getvalue()))
        response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response['Pragma'] = 'no-cache'
        response['Expires'] = '0'
        
        buffer.close()
        return response


class ContratoPaqueteNacionalPDFView(LoginRequiredMixin, DetailView):
    """
    Vista para generar contrato específico de paquetes nacionales en formato DOCX.
    Reemplaza el contrato genérico para ventas de tipo PAQ y tipo_viaje NAC.
    """
    model = VentaViaje
    
    def get(self, request, *args, **kwargs):
        self.object = self.get_object() 
        venta = self.object
        cliente = venta.cliente
        
        # Verificar que sea paquete nacional
        # Verificar si es PAQ directamente o si tiene cotización de tipo 'paquete'
        es_paquete = False
        if venta.servicios_seleccionados and 'PAQ' in venta.servicios_seleccionados:
            es_paquete = True
        elif venta.cotizacion_origen:
            # Si no tiene PAQ pero tiene cotización, verificar el tipo de cotización
            cotizacion = venta.cotizacion_origen
            if isinstance(cotizacion.propuestas, str):
                try:
                    import json
                    propuestas = json.loads(cotizacion.propuestas)
                except (ValueError, TypeError, json.JSONDecodeError) as e:
                    logger.warning(f"Error al decodificar propuestas de cotización: {e}")
                    propuestas = {}
            elif isinstance(cotizacion.propuestas, dict):
                propuestas = cotizacion.propuestas
            else:
                propuestas = {}
            
            tipo_cot = propuestas.get('tipo', '')
            if tipo_cot == 'paquete':
                es_paquete = True
        
        if not es_paquete:
            return HttpResponse("Error: Esta venta no es de paquete.", status=400)
        
        if venta.tipo_viaje not in ('NAC', 'INT_MXN'):
            return HttpResponse("Error: Esta venta no es nacional.", status=400)
        
        # Obtener datos de la cotización si existe
        origen = ''
        destino = ''
        vuelo_aerolinea = ''
        vuelo_salida = ''
        vuelo_regreso = ''
        vuelo_incluye = ''
        hotel_nombre = ''
        habitacion = ''
        plan_alimentos = ''
        tours = []
        # Traslado: usar "Opción de Traslado" del formulario (servicios_detalle), ej. "Privado"
        traslado_info = ''
        if venta.servicios_detalle:
            for linea in venta.servicios_detalle.split('\n'):
                linea = linea.strip()
                if linea.startswith('Traslado - ') and ' - Opción: ' in linea:
                    partes = linea.split(' - Opción: ', 1)
                    if len(partes) == 2:
                        traslado_info = partes[1].strip()
                    break
        if not traslado_info:
            # Fallback: nombre del proveedor desde LogisticaServicio TRA
            for s in venta.servicios_logisticos.filter(codigo_servicio='TRA'):
                opc = (s.opcion_proveedor or '').strip()
                if opc:
                    traslado_info = opc
                    break
        # Adicionales: usar opcion_proveedor de LogisticaServicio TOU (casilla "Opcion de tour y actividades")
        adicionales_info = []
        for s in venta.servicios_logisticos.filter(codigo_servicio='TOU').order_by('orden', 'pk'):
            opc = (s.opcion_proveedor or '').strip()
            if opc:
                adicionales_info.append(opc)
        adicionales_info = ' / '.join(adicionales_info) if adicionales_info else ''
        
        # DEBUG: Log para verificar datos
        import logging
        logger = logging.getLogger(__name__)
        
        if venta.cotizacion_origen:
            cotizacion = venta.cotizacion_origen
            
            # Convertir propuestas a dict si es string JSON
            if isinstance(cotizacion.propuestas, str):
                try:
                    import json
                    propuestas = json.loads(cotizacion.propuestas)
                except (ValueError, TypeError, json.JSONDecodeError) as e:
                    logger.warning(f"Error al decodificar propuestas de cotización: {e}")
                    propuestas = {}
            elif isinstance(cotizacion.propuestas, dict):
                propuestas = cotizacion.propuestas
            else:
                propuestas = {}
            
            logger.info(f"DEBUG: Propuestas tipo: {type(propuestas)}, contenido: {propuestas}")
            
            # Origen y destino desde cotización
            origen = (cotizacion.origen or '').strip()
            destino = (cotizacion.destino or '').strip()
            
            # Obtener información del paquete
            paquete = propuestas.get('paquete', {})
            logger.info(f"DEBUG: Paquete encontrado: {bool(paquete)}, tipo: {type(paquete)}")
            
            if paquete:
                # Información del vuelo - PRIMERO intentar desde paquete
                vuelo = paquete.get('vuelo', {})
                logger.info(f"DEBUG: Vuelo desde paquete: {bool(vuelo)}, tipo: {type(vuelo)}, contenido: {vuelo}")
                
                if vuelo and isinstance(vuelo, dict):
                    vuelo_aerolinea = (vuelo.get('aerolinea', '') or '').strip()
                    vuelo_salida = (vuelo.get('salida', '') or '').strip()
                    vuelo_regreso = (vuelo.get('regreso', '') or '').strip()
                    vuelo_incluye = (vuelo.get('incluye', '') or '').strip()
                    logger.info(f"DEBUG: Vuelo extraído desde paquete - aerolinea: {vuelo_aerolinea}, salida: {vuelo_salida}, regreso: {vuelo_regreso}, incluye: {vuelo_incluye}")
                
                # FALLBACK: Si no hay vuelo en paquete o está vacío, buscar en propuestas['vuelos']
                if not vuelo_aerolinea and propuestas.get('vuelos'):
                    vuelos_lista = propuestas.get('vuelos', [])
                    logger.info(f"DEBUG: Buscando vuelo en propuestas['vuelos']: {bool(vuelos_lista)}, tipo: {type(vuelos_lista)}")
                    
                    if isinstance(vuelos_lista, list) and len(vuelos_lista) > 0:
                        # Tomar el primer vuelo de la lista
                        vuelo_fallback = vuelos_lista[0] if isinstance(vuelos_lista[0], dict) else {}
                        logger.info(f"DEBUG: Vuelo fallback encontrado: {vuelo_fallback}")
                        
                        if vuelo_fallback:
                            # Si no tenemos aerolínea, intentar obtenerla del vuelo fallback
                            if not vuelo_aerolinea:
                                vuelo_aerolinea = (vuelo_fallback.get('aerolinea', '') or vuelo_fallback.get('aerolinea_nombre', '') or '').strip()
                            if not vuelo_salida:
                                vuelo_salida = (vuelo_fallback.get('salida', '') or vuelo_fallback.get('fecha_salida', '') or vuelo_fallback.get('ida', '') or '').strip()
                            if not vuelo_regreso:
                                vuelo_regreso = (vuelo_fallback.get('regreso', '') or vuelo_fallback.get('fecha_regreso', '') or vuelo_fallback.get('vuelta', '') or '').strip()
                            if not vuelo_incluye:
                                vuelo_incluye = (vuelo_fallback.get('incluye', '') or vuelo_fallback.get('equipaje', '') or '').strip()
                            logger.info(f"DEBUG: Vuelo extraído desde fallback - aerolinea: {vuelo_aerolinea}, salida: {vuelo_salida}, regreso: {vuelo_regreso}, incluye: {vuelo_incluye}")
                    elif isinstance(vuelos_lista, dict):
                        # Si es un diccionario único, usarlo directamente
                        if not vuelo_aerolinea:
                            vuelo_aerolinea = (vuelos_lista.get('aerolinea', '') or vuelos_lista.get('aerolinea_nombre', '') or '').strip()
                        if not vuelo_salida:
                            vuelo_salida = (vuelos_lista.get('salida', '') or vuelos_lista.get('fecha_salida', '') or vuelos_lista.get('ida', '') or '').strip()
                        if not vuelo_regreso:
                            vuelo_regreso = (vuelos_lista.get('regreso', '') or vuelos_lista.get('fecha_regreso', '') or vuelos_lista.get('vuelta', '') or '').strip()
                        if not vuelo_incluye:
                            vuelo_incluye = (vuelos_lista.get('incluye', '') or vuelos_lista.get('equipaje', '') or '').strip()
                
                # Información del hotel - PRIMERO intentar desde paquete
                hotel = paquete.get('hotel', {})
                logger.info(f"DEBUG: Hotel desde paquete: {bool(hotel)}, tipo: {type(hotel)}, contenido: {hotel}")
                
                if hotel and isinstance(hotel, dict):
                    hotel_nombre = (hotel.get('nombre', '') or '').strip()
                    habitacion = (hotel.get('habitacion', '') or '').strip()
                    plan_alimentos = (hotel.get('plan', '') or '').strip()
                    logger.info(f"DEBUG: Hotel extraído desde paquete - nombre: {hotel_nombre}, habitacion: {habitacion}, plan: {plan_alimentos}")
            
            # FALLBACK: Si no hay hotel en paquete o está vacío, buscar en propuestas['hoteles']
            if (not hotel_nombre or not habitacion or not plan_alimentos) and propuestas.get('hoteles'):
                hoteles_lista = propuestas.get('hoteles', [])
                logger.info(f"DEBUG: Buscando hotel en propuestas['hoteles']: {bool(hoteles_lista)}, tipo: {type(hoteles_lista)}")
                
                if isinstance(hoteles_lista, list) and len(hoteles_lista) > 0:
                    # Tomar el primer hotel de la lista
                    hotel_fallback = hoteles_lista[0] if isinstance(hoteles_lista[0], dict) else {}
                    logger.info(f"DEBUG: Hotel fallback encontrado: {hotel_fallback}")
                    
                    if hotel_fallback:
                        # Si no tenemos nombre, intentar obtenerlo del hotel fallback
                        if not hotel_nombre:
                            hotel_nombre = (hotel_fallback.get('nombre', '') or hotel_fallback.get('hotel', '') or '').strip()
                        if not habitacion:
                            habitacion = (hotel_fallback.get('habitacion', '') or hotel_fallback.get('tipo_habitacion', '') or '').strip()
                        if not plan_alimentos:
                            plan_alimentos = (hotel_fallback.get('plan', '') or hotel_fallback.get('plan_alimentos', '') or hotel_fallback.get('alimentos', '') or '').strip()
                        logger.info(f"DEBUG: Hotel extraído desde fallback - nombre: {hotel_nombre}, habitacion: {habitacion}, plan: {plan_alimentos}")
                elif isinstance(hoteles_lista, dict):
                    # Si es un diccionario único, usarlo directamente
                    if not hotel_nombre:
                        hotel_nombre = (hoteles_lista.get('nombre', '') or hoteles_lista.get('hotel', '') or '').strip()
                    if not habitacion:
                        habitacion = (hoteles_lista.get('habitacion', '') or hoteles_lista.get('tipo_habitacion', '') or '').strip()
                    if not plan_alimentos:
                        plan_alimentos = (hoteles_lista.get('plan', '') or hoteles_lista.get('plan_alimentos', '') or hoteles_lista.get('alimentos', '') or '').strip()
                
                # Tours/Adicionales - PRIMERO intentar desde paquete
                tours_raw = paquete.get('tours', [])
                logger.info(f"DEBUG: Tours desde paquete: {bool(tours_raw)}, tipo: {type(tours_raw)}, contenido: {tours_raw}")
                
                if tours_raw:
                    if isinstance(tours_raw, list):
                        tours = tours_raw
                    elif isinstance(tours_raw, dict):
                        # Si es un diccionario, convertirlo a lista
                        tours = [tours_raw]
                    else:
                        tours = []
                else:
                    tours = []
            
            # FALLBACK: Si no hay tours en paquete, buscar en propuestas['tours']
            if (not tours or len(tours) == 0) and propuestas.get('tours'):
                tours_fallback = propuestas.get('tours', [])
                logger.info(f"DEBUG: Buscando tours en propuestas['tours']: {bool(tours_fallback)}, tipo: {type(tours_fallback)}")
                
                if isinstance(tours_fallback, list) and len(tours_fallback) > 0:
                    tours = tours_fallback
                    logger.info(f"DEBUG: Tours encontrados desde fallback: {len(tours)}")
                elif isinstance(tours_fallback, dict):
                    # Si es un diccionario único, convertirlo a lista
                    tours = [tours_fallback]
                    logger.info(f"DEBUG: Tours encontrados desde fallback (dict único)")
            
            # Traslado y adicionales se obtienen desde servicios_logisticos (opcion_proveedor), no desde propuestas
        
        # Si no hay origen/destino en cotización, usar valores por defecto
        if not origen:
            origen = 'AIFA (FELIPE ANGELES)'  # Valor por defecto según imagen
        if not destino:
            destino = (venta.servicios_detalle_desde_logistica or 'PAQUETE NACIONAL').strip()
        
        # Calcular valores financieros
        from decimal import Decimal
        precio_total = venta.costo_total_con_modificacion or Decimal('0.00')
        anticipo = venta.cantidad_apertura or Decimal('0.00')
        saldo_pendiente = max(Decimal('0.00'), precio_total - venta.total_pagado)
        fecha_limite_pago = venta.fecha_vencimiento_pago
        
        # Convertir montos a texto
        anticipo_texto = numero_a_texto(float(anticipo))
        anticipo_texto = anticipo_texto.replace('M.N.', 'MXN')
        
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.oxml.ns import qn
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.section import WD_SECTION
        except ImportError:
            return HttpResponse("Error: python-docx no está instalado. Ejecuta: pip install python-docx", status=500)
        
        # Usar plantilla con membrete si existe
        template_path = os.path.join(settings.BASE_DIR, 'static', 'docx', 'membrete.docx')
        if os.path.exists(template_path):
            doc = Document(template_path)
        else:
            doc = Document()
        
        # Configurar fuente predeterminada
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        
        # Colores
        MOVUMS_BLUE = RGBColor(0, 74, 142)  # #004a8e
        TEXT_COLOR = RGBColor(47, 47, 47)  # #2f2f2f
        
        def set_run_font(run, size=12, bold=False, color=TEXT_COLOR):
            run.font.name = 'Arial'
            run.font.size = Pt(size)
            run.bold = bold
            run.font.color.rgb = color
        
        def format_date(value):
            if not value:
                return ''
            try:
                if isinstance(value, date):
                    # Formato: DD MES YYYY (ej: 15 DICIEMBRE 2025)
                    meses = {
                        1: 'ENERO', 2: 'FEBRERO', 3: 'MARZO', 4: 'ABRIL',
                        5: 'MAYO', 6: 'JUNIO', 7: 'JULIO', 8: 'AGOSTO',
                        9: 'SEPTIEMBRE', 10: 'OCTUBRE', 11: 'NOVIEMBRE', 12: 'DICIEMBRE'
                    }
                    return f"{value.day:02d} {meses[value.month]} {value.year}"
                return str(value)
            except:
                return ''
        
        def format_currency(value):
            """Formatea un valor Decimal como moneda"""
            if not value:
                return '0.00'
            return f"{value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
        
        # ============================================
        # PÁGINA 1: CONTRATO PRINCIPAL (tamaño 10)
        # ============================================
        
        # Título principal
        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_titulo.paragraph_format.space_before = Pt(0)
        p_titulo.paragraph_format.space_after = Pt(12)
        run_titulo = p_titulo.add_run('CONTRATO DE SERVICIOS TURÍSTICOS')
        set_run_font(run_titulo, size=10, bold=True, color=MOVUMS_BLUE)
        
        # Fecha (alineada a la derecha)
        p_fecha = doc.add_paragraph()
        p_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_fecha.paragraph_format.space_after = Pt(12)
        run_fecha_label = p_fecha.add_run('Fecha: ')
        set_run_font(run_fecha_label, size=10, bold=True)
        run_fecha_val = p_fecha.add_run(format_date(date.today()))
        set_run_font(run_fecha_val, size=10)
        
        # Origen / Destino
        p_origen_destino = doc.add_paragraph()
        p_origen_destino.paragraph_format.space_after = Pt(12)
        run_origen_label = p_origen_destino.add_run('Origen / Destino: ')
        set_run_font(run_origen_label, size=10, bold=True)
        run_origen_val = p_origen_destino.add_run(origen.upper() if origen else '')
        set_run_font(run_origen_val, size=10)
        run_separador = p_origen_destino.add_run(' / ')
        set_run_font(run_separador, size=10)
        run_destino_val = p_origen_destino.add_run(destino.upper() if destino else '')
        set_run_font(run_destino_val, size=10, bold=True)
        
        # Información del cliente y monto recibido
        p_cliente = doc.add_paragraph()
        p_cliente.paragraph_format.space_after = Pt(6)
        run_texto1 = p_cliente.add_run('Movums The travel Store, con domicilio Plaza Mora, Juárez Sur 321 Local 18 CP. 56100 Texcoco Estado de México, recibió de: ')
        set_run_font(run_texto1, size=10)
        run_cliente = p_cliente.add_run(cliente.nombre_completo_display.upper())
        set_run_font(run_cliente, size=10, bold=True)
        run_cliente.font.underline = True
        run_texto2 = p_cliente.add_run(' la cantidad de:')
        set_run_font(run_texto2, size=10)
        
        # Monto recibido
        p_monto = doc.add_paragraph()
        p_monto.paragraph_format.space_after = Pt(6)
        run_monto_num = p_monto.add_run(f'${format_currency(anticipo)}')
        set_run_font(run_monto_num, size=10, bold=True)
        run_monto_num.font.underline = True
        run_monto_texto = p_monto.add_run(f' ({anticipo_texto}).')
        set_run_font(run_monto_texto, size=10)
        
        # Agregar espacio
        p_espacio = doc.add_paragraph()
        p_espacio.paragraph_format.space_after = Pt(6)
        
        # Detalles del paquete - Labels sin negritas, valores en negritas, espaciado compacto
        # FECHA DE IDA
        p_ida = doc.add_paragraph()
        p_ida.paragraph_format.space_after = Pt(2)  # Espaciado compacto
        run_ida_label = p_ida.add_run('FECHA DE IDA: ')
        set_run_font(run_ida_label, size=10, bold=False)  # Label sin negritas
        run_ida_val = p_ida.add_run(format_date(venta.fecha_inicio_viaje) if venta.fecha_inicio_viaje else '')
        set_run_font(run_ida_val, size=10, bold=True)  # Valor en negritas
        
        # FECHA DE REGRESO
        p_regreso = doc.add_paragraph()
        p_regreso.paragraph_format.space_after = Pt(2)  # Espaciado compacto
        run_regreso_label = p_regreso.add_run('FECHA DE REGRESO: ')
        set_run_font(run_regreso_label, size=10, bold=False)  # Label sin negritas
        run_regreso_val = p_regreso.add_run(format_date(venta.fecha_fin_viaje) if venta.fecha_fin_viaje else '')
        set_run_font(run_regreso_val, size=10, bold=True)  # Valor en negritas
        
        # PASAJEROS
        adultos = 0
        menores = 0
        if venta.cotizacion_origen:
            cotizacion = venta.cotizacion_origen
            adultos = cotizacion.adultos or 0
            menores = cotizacion.menores or 0
        
        if adultos == 0 and venta.pasajeros:
            pasajeros_lista = [p.strip() for p in venta.pasajeros.split('\n') if p.strip()]
            adultos = len(pasajeros_lista) if pasajeros_lista else 1
        
        p_pasajeros = doc.add_paragraph()
        p_pasajeros.paragraph_format.space_after = Pt(2)  # Espaciado compacto
        run_pasajeros_label = p_pasajeros.add_run('PASAJEROS: ')
        set_run_font(run_pasajeros_label, size=10, bold=False)  # Label sin negritas
        if menores > 0:
            run_pasajeros_val = p_pasajeros.add_run(f'{adultos} adultos + {menores} Menor')
        else:
            run_pasajeros_val = p_pasajeros.add_run(f'{adultos} adultos')
        set_run_font(run_pasajeros_val, size=10, bold=True)  # Valor en negritas
        
        # ACOMPAÑANTES: Pasajeros (Nombres Completos para Contrato) separados por coma
        pasajeros_contrato = (venta.pasajeros or '').strip()
        if pasajeros_contrato:
            # Normalizar: quitar saltos de línea y separar por coma
            lineas = [n.strip() for n in pasajeros_contrato.replace('\r\n', '\n').replace('\r', '\n').split('\n') if n.strip()]
            acompanantes_texto = ', '.join(lineas)
        else:
            acompanantes_texto = ''
        p_acompanantes = doc.add_paragraph()
        p_acompanantes.paragraph_format.space_after = Pt(2)  # Espaciado compacto
        run_acompanantes_label = p_acompanantes.add_run('ACOMPAÑANTES: ')
        set_run_font(run_acompanantes_label, size=10, bold=False)  # Label sin negritas
        run_acompanantes_val = p_acompanantes.add_run(acompanantes_texto)
        set_run_font(run_acompanantes_val, size=10, bold=True)  # Valor en negritas
        
        # Hotel
        p_hotel = doc.add_paragraph()
        p_hotel.paragraph_format.space_after = Pt(2)  # Espaciado compacto
        run_hotel_label = p_hotel.add_run('Hotel: ')
        set_run_font(run_hotel_label, size=10, bold=False)  # Label sin negritas
        run_hotel_val = p_hotel.add_run(hotel_nombre.upper() if hotel_nombre else '')
        set_run_font(run_hotel_val, size=10, bold=True)  # Valor en negritas
        
        # Habitación
        p_habitacion = doc.add_paragraph()
        p_habitacion.paragraph_format.space_after = Pt(2)  # Espaciado compacto
        run_habitacion_label = p_habitacion.add_run('Habitación: ')
        set_run_font(run_habitacion_label, size=10, bold=False)  # Label sin negritas
        run_habitacion_val = p_habitacion.add_run(habitacion.upper() if habitacion else '')
        set_run_font(run_habitacion_val, size=10, bold=True)  # Valor en negritas
        
        # Plan de Alimentos
        p_plan = doc.add_paragraph()
        p_plan.paragraph_format.space_after = Pt(2)  # Espaciado compacto
        run_plan_label = p_plan.add_run('Plan de Alimentos: ')
        set_run_font(run_plan_label, size=10, bold=False)  # Label sin negritas
        run_plan_val = p_plan.add_run(plan_alimentos.upper() if plan_alimentos else '')
        set_run_font(run_plan_val, size=10, bold=True)  # Valor en negritas
        
        # Vuelos - SIEMPRE mostrar este campo
        p_vuelos = doc.add_paragraph()
        p_vuelos.paragraph_format.space_after = Pt(2)  # Espaciado compacto
        run_vuelos_label = p_vuelos.add_run('Vuelos: ')
        set_run_font(run_vuelos_label, size=10, bold=False)  # Label sin negritas
        vuelo_aerolinea_display = vuelo_aerolinea.upper() if vuelo_aerolinea else ''
        run_vuelos_val = p_vuelos.add_run(vuelo_aerolinea_display)
        set_run_font(run_vuelos_val, size=10, bold=True)  # Valor en negritas
        
        # IDA - SIEMPRE mostrar este campo
        p_ida_vuelo = doc.add_paragraph()
        p_ida_vuelo.paragraph_format.space_after = Pt(2)  # Espaciado compacto
        run_ida_vuelo_label = p_ida_vuelo.add_run('IDA : ')
        set_run_font(run_ida_vuelo_label, size=10, bold=False)  # Label sin negritas
        vuelo_salida_display = vuelo_salida if vuelo_salida else ''
        run_ida_vuelo_val = p_ida_vuelo.add_run(vuelo_salida_display)
        set_run_font(run_ida_vuelo_val, size=10, bold=True)  # Valor en negritas
        
        # REGRESO - SIEMPRE mostrar este campo
        p_regreso_vuelo = doc.add_paragraph()
        p_regreso_vuelo.paragraph_format.space_after = Pt(2)  # Espaciado compacto
        run_regreso_vuelo_label = p_regreso_vuelo.add_run('REGRESO : ')
        set_run_font(run_regreso_vuelo_label, size=10, bold=False)  # Label sin negritas
        vuelo_regreso_display = vuelo_regreso if vuelo_regreso else ''
        run_regreso_vuelo_val = p_regreso_vuelo.add_run(vuelo_regreso_display)
        set_run_font(run_regreso_vuelo_val, size=10, bold=True)  # Valor en negritas
        
        # EQUIPAJE (desde campo "incluye") - SIEMPRE mostrar este campo
        p_equipaje = doc.add_paragraph()
        p_equipaje.paragraph_format.space_after = Pt(2)  # Espaciado compacto
        run_equipaje_label = p_equipaje.add_run('EQUIPAJE: ')
        set_run_font(run_equipaje_label, size=10, bold=False)  # Label sin negritas
        vuelo_incluye_display = vuelo_incluye.upper() if vuelo_incluye else ''
        run_equipaje_val = p_equipaje.add_run(vuelo_incluye_display)
        set_run_font(run_equipaje_val, size=10, bold=True)  # Valor en negritas
        
        # Traslado - SIEMPRE mostrar este campo (usar información extraída si existe)
        p_traslado = doc.add_paragraph()
        p_traslado.paragraph_format.space_after = Pt(2)  # Espaciado compacto
        run_traslado_label = p_traslado.add_run('Traslado: ')
        set_run_font(run_traslado_label, size=10, bold=False)  # Label sin negritas
        traslado_display = traslado_info.upper() if traslado_info else ''
        run_traslado_val = p_traslado.add_run(traslado_display)
        set_run_font(run_traslado_val, size=10, bold=True)  # Valor en negritas
        
        # ADICIONALES: usar opcion_proveedor de LogisticaServicio TOU (casilla "Opcion de tour y actividades")
        p_adicionales = doc.add_paragraph()
        p_adicionales.paragraph_format.space_after = Pt(0)  # Sin espacio después para compactar
        run_adicionales_label = p_adicionales.add_run('ADICIONALES: ')
        set_run_font(run_adicionales_label, size=10, bold=False)  # Label sin negritas
        if adicionales_info:
            run_adicionales_val = p_adicionales.add_run(adicionales_info.upper())
            set_run_font(run_adicionales_val, size=10, bold=True)  # Valor en negritas
        
        # PRECIO Y CONDICIONES ECONÓMICAS - Sin espacio extra antes
        p_seccion = doc.add_paragraph()
        p_seccion.paragraph_format.space_before = Pt(0)  # Sin espacio antes
        p_seccion.paragraph_format.space_after = Pt(0)  # Sin espacio después para compactar
        run_seccion = p_seccion.add_run('PRECIO Y CONDICIONES ECONÓMICAS')
        set_run_font(run_seccion, size=10, bold=True)
        
        # Precio total del paquete
        p_precio = doc.add_paragraph()
        p_precio.paragraph_format.space_after = Pt(2)  # Espaciado compacto
        run_precio_label = p_precio.add_run('•Precio total del paquete: ')
        set_run_font(run_precio_label, size=10, bold=False)  # Label sin negritas
        run_precio_dollar = p_precio.add_run('$')
        set_run_font(run_precio_dollar, size=10, bold=False)
        run_precio_val = p_precio.add_run(f'{format_currency(precio_total)} MXN')
        set_run_font(run_precio_val, size=10, bold=True)  # Valor en negritas
        run_precio_val.font.underline = True
        
        # Anticipo recibido
        p_anticipo = doc.add_paragraph()
        p_anticipo.paragraph_format.space_after = Pt(2)  # Espaciado compacto
        run_anticipo_label = p_anticipo.add_run('•Anticipo recibido: ')
        set_run_font(run_anticipo_label, size=10, bold=False)  # Label sin negritas
        run_anticipo_dollar = p_anticipo.add_run('$')
        set_run_font(run_anticipo_dollar, size=10, bold=False)
        run_anticipo_val = p_anticipo.add_run(f'{format_currency(anticipo)} MXN')
        set_run_font(run_anticipo_val, size=10, bold=True)  # Valor en negritas
        run_anticipo_val.font.underline = True
        
        # Saldo pendiente
        p_saldo = doc.add_paragraph()
        p_saldo.paragraph_format.space_after = Pt(2)  # Espaciado compacto
        run_saldo_label = p_saldo.add_run('•Saldo pendiente: ')
        set_run_font(run_saldo_label, size=10, bold=False)  # Label sin negritas
        run_saldo_dollar = p_saldo.add_run('$')
        set_run_font(run_saldo_dollar, size=10, bold=False)
        run_saldo_val = p_saldo.add_run(f'{format_currency(saldo_pendiente)} MXN')
        set_run_font(run_saldo_val, size=10, bold=True)  # Valor en negritas
        run_saldo_val.font.underline = True
        
        # Fecha límite de pago total
        p_fecha_limite = doc.add_paragraph()
        p_fecha_limite.paragraph_format.space_after = Pt(2)  # Espaciado compacto
        run_fecha_limite_label = p_fecha_limite.add_run('•Fecha límite de pago total: ')
        set_run_font(run_fecha_limite_label, size=10, bold=False)  # Label sin negritas
        if fecha_limite_pago:
            meses_cortos = {
                1: 'ENERO', 2: 'FEBRERO', 3: 'MARZO', 4: 'ABRIL',
                5: 'MAYO', 6: 'JUNIO', 7: 'JULIO', 8: 'AGOSTO',
                9: 'SEPTIEMBRE', 10: 'OCTUBRE', 11: 'NOVIEMBRE', 12: 'DICIEMBRE'
            }
            fecha_texto = f"{fecha_limite_pago.day:02d}/{meses_cortos[fecha_limite_pago.month]}/{fecha_limite_pago.year}"
            run_fecha_limite_val = p_fecha_limite.add_run(fecha_texto)
        else:
            run_fecha_limite_val = p_fecha_limite.add_run('//2025')
        set_run_font(run_fecha_limite_val, size=10, bold=True)
        run_fecha_limite_val.font.underline = True
        
        # El footer ya está en el membrete de la plantilla
        
        # ============================================
        # PÁGINA 2: ANEXO DE TÉRMINOS Y CONDICIONES (tamaño 10)
        # ============================================
        
        # Usar page_break_before en lugar de add_page_break() para evitar página vacía
        # EL CLIENTE declara que: - Segunda página con fuente 10.5pt
        p_declara = doc.add_paragraph()
        p_declara.paragraph_format.page_break_before = True  # Forzar salto de página antes de este párrafo
        p_declara.paragraph_format.space_before = Pt(0)  # Sin espacio antes
        p_declara.paragraph_format.space_after = Pt(8)
        run_declara = p_declara.add_run('EL CLIENTE declara que:')
        set_run_font(run_declara, size=10, bold=True)
        
        declaraciones = [
            'Ha revisado y entendido toda la información contenida en este Anexo.',
            'Proporcionó datos veraces y completos.',
            'Acepta las condiciones del servicio, políticas de proveedores y cláusulas del contrato.',
        ]
        
        for decl in declaraciones:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(4)
            run_bullet = p.add_run('•')
            set_run_font(run_bullet, size=10)
            run_text = p.add_run(decl)
            set_run_font(run_text, size=10)
        
        # NOTA
        p_nota = doc.add_paragraph()
        p_nota.paragraph_format.space_before = Pt(12)
        p_nota.paragraph_format.space_after = Pt(8)
        run_nota_label = p_nota.add_run('NOTA: ')
        set_run_font(run_nota_label, size=10, bold=True)
        run_nota_text = p_nota.add_run('Movums The Travel Store, se reserva el derecho de cancelar sin previo aviso este contrato, si los depósitos no son recibidos en las fechas pactadas con el "CLIENTE" anteriormente estipuladas.')
        set_run_font(run_nota_text, size=10)
        
        # CANCELACIONES
        p_cancelaciones = doc.add_paragraph()
        p_cancelaciones.paragraph_format.space_before = Pt(12)
        p_cancelaciones.paragraph_format.space_after = Pt(8)
        run_cancelaciones = p_cancelaciones.add_run('CANCELACIONES:')
        set_run_font(run_cancelaciones, size=10, bold=True)
        
        textos_cancelacion = [
            'Entre la firma del contrato y pago de anticipo, parcial o total no se reembolsará ningún pago',
            'No es cancelable, ni reembolsable.',
            'Cualquier modificación puede ocasionar cargo extra; y están sujetos a disponibilidad'
        ]
        
        for texto in textos_cancelacion:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(texto)
            set_run_font(run, size=10)
        
        # VUELOS
        p_vuelos_seccion = doc.add_paragraph()
        p_vuelos_seccion.paragraph_format.space_before = Pt(12)
        p_vuelos_seccion.paragraph_format.space_after = Pt(8)
        run_vuelos_seccion = p_vuelos_seccion.add_run('VUELOS:')
        set_run_font(run_vuelos_seccion, size=10, bold=True)
        
        texto_vuelos = 'Movums The Travel Store, no se hace responsable por cambios y/o cancelaciones de la aerolínea contratada para el paquete vacacional, en este caso Movums The Travel Store ofrecera las alternativas que nos brinde directamente la aerolínea.'
        p_vuelos_texto = doc.add_paragraph()
        p_vuelos_texto.paragraph_format.space_after = Pt(8)
        run_vuelos_texto = p_vuelos_texto.add_run(texto_vuelos)
        set_run_font(run_vuelos_texto, size=10)
        
        # FECHA DE VALIDEZ DEL CONTRATO
        p_validez = doc.add_paragraph()
        p_validez.paragraph_format.space_before = Pt(12)
        p_validez.paragraph_format.space_after = Pt(8)
        run_validez = p_validez.add_run('FECHA DE VALIDEZ DEL CONTRATO:')
        set_run_font(run_validez, size=10, bold=True)
        
        texto_validez = (
            '•Las condiciones y precios antes mencionados serán mantenidos a la fecha límite de pago, a esta fecha el "CLIENTE" deberá haber cubierto el pago total del paquete o servicio turístico contratado. '
            'La aceptación de este contrato será efectiva una vez que el "CLIENTE" envié el contrato debidamente firmado y con el anticipo, pago parcial o pago total; no reembolsable; para que Movums The travel store, proceda con la reservación del servicio contratado'
        )
        
        p_texto_validez = doc.add_paragraph()
        p_texto_validez.paragraph_format.space_after = Pt(12)
        p_texto_validez.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Resaltar palabras clave
        partes = texto_validez.split('fecha límite de pago')
        if len(partes) > 1:
            run1 = p_texto_validez.add_run(partes[0])
            set_run_font(run1, size=10)
            run2 = p_texto_validez.add_run('fecha límite de pago')
            set_run_font(run2, size=10, bold=True)
            resto = 'anticipo, pago parcial o pago total'
            if resto in partes[1]:
                partes2 = partes[1].split(resto)
                run3 = p_texto_validez.add_run(partes2[0])
                set_run_font(run3, size=10)
                run4 = p_texto_validez.add_run(resto)
                set_run_font(run4, size=10, bold=True)
                if len(partes2) > 1:
                    run5 = p_texto_validez.add_run(partes2[1])
                    set_run_font(run5, size=10)
            else:
                run3 = p_texto_validez.add_run(partes[1])
                set_run_font(run3, size=10)
        else:
            run = p_texto_validez.add_run(texto_validez)
            set_run_font(run, size=10)
        
        # FIRMAS
        p_firmas = doc.add_paragraph()
        p_firmas.paragraph_format.space_before = Pt(12)
        p_firmas.paragraph_format.space_after = Pt(8)
        run_firmas = p_firmas.add_run('FIRMAS:')
        set_run_font(run_firmas, size=10, bold=True)
        
        # CLIENTE
        p_cliente_firma = doc.add_paragraph()
        p_cliente_firma.paragraph_format.space_before = Pt(12)
        p_cliente_firma.paragraph_format.space_after = Pt(6)
        run_cliente_firma = p_cliente_firma.add_run('CLIENTE:')
        set_run_font(run_cliente_firma, size=10, bold=True)
        
        p_nombre_firma = doc.add_paragraph()
        p_nombre_firma.paragraph_format.space_after = Pt(6)
        run_nombre_firma_label = p_nombre_firma.add_run('Nombre y firma: ')
        set_run_font(run_nombre_firma_label, size=10)
        run_linea1 = p_nombre_firma.add_run('_' * 50)
        set_run_font(run_linea1, size=10)
        
        # AGENCIA
        p_agencia_firma = doc.add_paragraph()
        p_agencia_firma.paragraph_format.space_before = Pt(12)
        p_agencia_firma.paragraph_format.space_after = Pt(6)
        run_agencia_firma = p_agencia_firma.add_run('AGENCIA – Movums The Travel Store')
        set_run_font(run_agencia_firma, size=10, bold=True)
        
        p_representante = doc.add_paragraph()
        p_representante.paragraph_format.space_after = Pt(6)
        run_representante_label = p_representante.add_run('Nombre y firma del representante: ')
        set_run_font(run_representante_label, size=10)
        run_linea2 = p_representante.add_run('_' * 50)
        set_run_font(run_linea2, size=10)
        
        # ANEXO - CONTRATO DE MEDIACIÓN (comienza en página 3, todo tamaño 7)
        nombre_cliente_anexo = (cliente.nombre_completo_display or '').strip() if cliente else ''
        p_anexo_titulo = doc.add_paragraph()
        p_anexo_titulo.paragraph_format.page_break_before = True  # Forzar que texto tamaño 7 comience en página 3
        p_anexo_titulo.paragraph_format.space_before = Pt(0)
        p_anexo_titulo.paragraph_format.space_after = Pt(8)
        p_anexo_titulo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_anexo_1 = p_anexo_titulo.add_run('CONTRATO DE MEDIACIÓN PARA LA PRESTACIÓN DE SERVICIOS TURÍSTICOS, QUE CELEBRAN POR UNA PARTE LA AGENCIA DE VIAJES "GRUPO IMVED, S.A. DE C.V." ACTUANDO EN USO DE SU NOMBRE COMERCIAL MOVUMS THE TRAVEL STORE, EN ADELANTE DENOMINADA COMO "LA AGENCIA", Y POR LA OTRA EL/LA C. ')
        set_run_font(run_anexo_1, size=7, bold=True)
        run_anexo_cliente = p_anexo_titulo.add_run(nombre_cliente_anexo)
        set_run_font(run_anexo_cliente, size=7, bold=True)
        run_anexo_cliente.font.underline = True
        run_anexo_2 = p_anexo_titulo.add_run(' A QUIEN EN LO SUCESIVO SE LE DENOMINARÁ "EL CLIENTE", AL TENOR DE LAS SIGUIENTES DEFINICIONES, DECLARACIONES Y CLÁUSULAS:')
        set_run_font(run_anexo_2, size=7, bold=True)
        
        # GLOSARIO (tamaño 7, justificado)
        p_glosario = doc.add_paragraph()
        p_glosario.paragraph_format.space_before = Pt(12)
        p_glosario.paragraph_format.space_after = Pt(6)
        p_glosario.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_glosario = p_glosario.add_run('GLOSARIO')
        set_run_font(run_glosario, size=7, bold=True)
        
        p_glosario_intro = doc.add_paragraph()
        p_glosario_intro.paragraph_format.space_after = Pt(6)
        p_glosario_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_glosario_intro = p_glosario_intro.add_run('Para efectos del presente contrato, se entiende por:')
        set_run_font(run_glosario_intro, size=7)
        
        p_agencia_def = doc.add_paragraph()
        p_agencia_def.paragraph_format.space_after = Pt(4)
        p_agencia_def.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_agencia_def_label = p_agencia_def.add_run('Agencia: ')
        set_run_font(run_agencia_def_label, size=7, bold=True)
        run_agencia_def_text = p_agencia_def.add_run('Es el proveedor de servicios turísticos que intermedia, contrata u ofrece servicios o productos turístico nacionales, previo pago de un precio cierto y determinado.')
        set_run_font(run_agencia_def_text, size=7)
        
        p_cliente_def = doc.add_paragraph()
        p_cliente_def.paragraph_format.space_after = Pt(4)
        p_cliente_def.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_cliente_def_label = p_cliente_def.add_run('Cliente: ')
        set_run_font(run_cliente_def_label, size=7, bold=True)
        run_cliente_def_text = p_cliente_def.add_run('Consumidor que contrata los servicios turísticos nacionales mediante el pago de un precio cierto y determinado.')
        set_run_font(run_cliente_def_text, size=7)
        
        p_paquete_def = doc.add_paragraph()
        p_paquete_def.paragraph_format.space_after = Pt(4)
        p_paquete_def.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_paquete_def_label = p_paquete_def.add_run('Paquete turístico: ')
        set_run_font(run_paquete_def_label, size=7, bold=True)
        run_paquete_def_text = p_paquete_def.add_run('Integración de uno o más servicios turísticos en un solo producto, ofrecidos al Cliente y detallado en el Anexo del presente contrato.')
        set_run_font(run_paquete_def_text, size=7)
        
        p_servicio_def = doc.add_paragraph()
        p_servicio_def.paragraph_format.space_after = Pt(4)
        p_servicio_def.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_servicio_def_label = p_servicio_def.add_run('Servicio turístico: ')
        set_run_font(run_servicio_def_label, size=7, bold=True)
        run_servicio_def_text = p_servicio_def.add_run('Prestación de carácter comercial en transporte nacional, hospedaje, alimentación, excursiones u otros servicios relacionados,  detallados en el Anexo del presente contrato.')
        set_run_font(run_servicio_def_text, size=7)
        
        p_caratula_def = doc.add_paragraph()
        p_caratula_def.paragraph_format.space_after = Pt(12)
        p_caratula_def.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_caratula_def_label = p_caratula_def.add_run('Caratula: ')
        set_run_font(run_caratula_def_label, size=7, bold=True)
        run_caratula_def_text = p_caratula_def.add_run('Documento que detalla servicios, fechas, precios y condiciones del servicio turístico contratado.')
        set_run_font(run_caratula_def_text, size=7)
        
        # DECLARACIONES (tamaño 7, justificado)
        p_declaraciones = doc.add_paragraph()
        p_declaraciones.paragraph_format.space_before = Pt(12)
        p_declaraciones.paragraph_format.space_after = Pt(6)
        p_declaraciones.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_declaraciones = p_declaraciones.add_run('DECLARACIONES')
        set_run_font(run_declaraciones, size=7, bold=True)
        
        p_declara_agencia = doc.add_paragraph()
        p_declara_agencia.paragraph_format.space_before = Pt(8)
        p_declara_agencia.paragraph_format.space_after = Pt(6)
        p_declara_agencia.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_declara_agencia = p_declara_agencia.add_run('Declara LA AGENCIA:')
        set_run_font(run_declara_agencia, size=7, bold=True)
        
        declaraciones_agencia = [
            'Ser una persona moral legalmente constituida conforme a las leyes mexicanas.',
            'Ser un Prestador de Servicios Turísticos con Razón Social: GRUPO IMVED, S.A. de C.V.',
            'Ser la única propietaria de la marca MOVUMS THE TRAVEL STORE',
            'RFC GIM190722FS7  y domicilio ubicado en:  Plaza Mora, Juárez Sur, 321, interior 18, Colonia Centro, Texcoco, Estado de México, C.P. 56100.',
            ' Teléfono, correo electrónico y horario de atención al público: 59 59319954, 5951255279  ventas@movums.com, lunes a viernes de 10:00 a 19:00 horas. Y sábados de 11:00 a 15:00 hrs.',
            'Contar con infraestructura, personal capacitado y experiencia suficiente para la prestación de los servicios turísticos contratados.',
            'Haber informado previamente al Cliente sobre los precios, tarifas, condiciones, características y costo total del servicio turístico contratado.'
        ]
        
        for texto in declaraciones_agencia:
            p_decl = doc.add_paragraph()
            p_decl.paragraph_format.space_after = Pt(4)
            p_decl.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run_decl = p_decl.add_run(texto)
            set_run_font(run_decl, size=7)
        
        p_declara_cliente = doc.add_paragraph()
        p_declara_cliente.paragraph_format.space_before = Pt(12)
        p_declara_cliente.paragraph_format.space_after = Pt(6)
        p_declara_cliente.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_declara_cliente = p_declara_cliente.add_run('II. Declara EL CLIENTE:')
        set_run_font(run_declara_cliente, size=7, bold=True)
        
        declaraciones_cliente = [
            'Ser persona física/moral con capacidad legal y económica para obligarse en términos del presente contrato.',
            'En caso de persona moral: Ser una persona moral legalmente constituida conforme a las leyes mexicanas, conforme lo acredita con copia del instrumento número ____________, de fecha ___________, otorgado ante la Fe del Notario Público Número ____, de ____________, y que el(la) C. _________ __________________ en este acto interviene en su carácter de Representante Legal, calidad que acredita con copia del instrumento número _______, de fecha _________, otorgada ante la Fe del Notario Público número _______ del _________, facultad y calidad  que no le han sido revocadas, modificadas o limitadas a la fecha de firma del presente contrato.',
            'Encontrarse inscrito en el Registro Federal de Contribuyentes con la clave que ha manifestado.',
            'Haber recibido previamente de LA AGENCIA  información útil, precisa, veraz y detallada sobre los servicios objeto del presente contrato.',
            'Proporciona su nombre, domicilio, número telefónico y correo electrónico, tal y como lo ha señalado en la caratula de prestación de servicios, acreditando los mismos con copia de los documentos idóneos para tal efecto.'
        ]
        
        for texto in declaraciones_cliente:
            p_decl = doc.add_paragraph()
            p_decl.paragraph_format.space_after = Pt(4)
            p_decl.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run_decl = p_decl.add_run(texto)
            set_run_font(run_decl, size=7)
        
        # CLÁUSULAS
        p_clausulas = doc.add_paragraph()
        p_clausulas.paragraph_format.space_before = Pt(12)
        p_clausulas.paragraph_format.space_after = Pt(6)
        p_clausulas.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_clausulas = p_clausulas.add_run('CLÁUSULAS')
        set_run_font(run_clausulas, size=7, bold=True)
        
        # PRIMERA
        p_primera = doc.add_paragraph()
        p_primera.paragraph_format.space_before = Pt(8)
        p_primera.paragraph_format.space_after = Pt(4)
        p_primera.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_primera = p_primera.add_run('PRIMERA. CONSENTIMIENTO. ')
        set_run_font(run_primera, size=7, bold=True)
        run_primera_text = p_primera.add_run('Las partes manifiestan su voluntad de celebrar el presente contrato, cuya naturaleza jurídica es la mediación para la prestación de servicios turísticos.')
        set_run_font(run_primera_text, size=7)
        
        # SEGUNDA
        p_segunda = doc.add_paragraph()
        p_segunda.paragraph_format.space_before = Pt(8)
        p_segunda.paragraph_format.space_after = Pt(4)
        p_segunda.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_segunda = p_segunda.add_run('SEGUNDA. OBJETO. ')
        set_run_font(run_segunda, size=7, bold=True)
        run_segunda_text = p_segunda.add_run('LA AGENCIA intermediará, contratará u ofrecerá servicios turísticos detallados en la CARATULA, previo pago del Cliente de un precio cierto y determinado.')
        set_run_font(run_segunda_text, size=7)
        
        # TERCERA
        p_tercera = doc.add_paragraph()
        p_tercera.paragraph_format.space_before = Pt(8)
        p_tercera.paragraph_format.space_after = Pt(4)
        p_tercera.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_tercera = p_tercera.add_run('TERCERA. PRECIO, FORMA Y LUGAR DE PAGO. ')
        set_run_font(run_tercera, size=7, bold=True)
        run_tercera_text = p_tercera.add_run('Las partes manifiestan su conformidad en que el precio total a pagar por EL CLIENTE como contraprestación del Servicio turístico, es la cantidad que por cada concepto se indica en la CARATULA de este Contrato. El importe señalado en la CARATULA, contempla todas las cantidades y conceptos referentes al Servicio turístico, por lo que LA AGENCIA se obliga a respetar en todo momento dicho costo sin poder cobrar otra cantidad o condicionar la prestación del Servicio turístico contratado a la adquisición de otro servicio no requerido por El cliente, salvo que El cliente autorice de manera escrita algún otro cobro no estipulado en el presente Contrato. EL CLIENTE efectuará el pago pactado por el Servicio turístico señalado en la caratula del presente Contrato en los términos y condiciones acordadas pudiendo ser: ')
        set_run_font(run_tercera_text, size=7)
        
        p_tercera_contado = doc.add_paragraph()
        p_tercera_contado.paragraph_format.space_after = Pt(4)
        p_tercera_contado.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_tercera_contado = p_tercera_contado.add_run('Al contado: en efectivo, con tarjeta de débito, tarjeta de crédito, transferencia bancaria, y/o cheque en el domicilio de la agencia en moneda nacional, sin menoscabo de poderlo hacer en moneda extranjera al tipo de cambio publicado en el Diario Oficial de la Federación al día en que el pago se efectúe.')
        set_run_font(run_tercera_contado, size=7)
        
        p_tercera_plazos = doc.add_paragraph()
        p_tercera_plazos.paragraph_format.space_after = Pt(4)
        p_tercera_plazos.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_tercera_plazos = p_tercera_plazos.add_run('A plazos: El cliente podrá, previo acuerdo con La agencia a pagar en parcialidades, para lo cual, La agencia deberá de entregar a El CLIENTE la información por escrito de las fechas, así como los montos parciales a pagar.')
        set_run_font(run_tercera_plazos, size=7)
        
        p_tercera_cheque = doc.add_paragraph()
        p_tercera_cheque.paragraph_format.space_after = Pt(8)
        p_tercera_cheque.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_tercera_cheque = p_tercera_cheque.add_run('En caso de que El cliente realice el pago con cheque y no se cubra el pago por causas imputables al librador, La agencia tendrá el derecho de realizar el cobro adicional del 20% (veinte por ciento) del valor del documento, por concepto de daños y perjuicios, en caso de que el cheque sea devuelto por causas imputables al librador, conforme al artículo 193 de la Ley General del Títulos y Operaciones de Crédito.')
        set_run_font(run_tercera_cheque, size=7)
        
        # QUINTA
        p_quinta = doc.add_paragraph()
        p_quinta.paragraph_format.space_before = Pt(8)
        p_quinta.paragraph_format.space_after = Pt(4)
        p_quinta.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_quinta = p_quinta.add_run('QUINTA. OBLIGACIONES DE LA AGENCIA. ')
        set_run_font(run_quinta, size=7, bold=True)
        run_quinta_text = p_quinta.add_run('LA AGENCIA SE OBLIGA A:')
        set_run_font(run_quinta_text, size=7, bold=True)
        
        obligaciones_agencia = [
            'Cumplir lo pactado en el contrato.',
            'Entregar  a EL CLIENTE copia del contrato y constancias de reservación.',
            'Proporcionar a EL CLIENTE boletos, claves de reservación y documentos de viaje.',
            'Auxiliar a EL CLIENTE en emergencias  y gestionar indemnizaciones relacionadas con el servicio contratado',
            'Solicitar los Servicios turísticos que se especifican en la caratula de este Contrato por cuenta de EL CLIENTE de acuerdo a la disponibilidad de los mismos, a contratarlos fungiendo como intermediario entre éste y las personas encargadas de proporcionar directamente el Servicio turístico.',
            'Coadyuvar a EL CLIENTE para reclamar ante el prestador del servicio final, las indemnizaciones que correspondan.',
            'Respetar la Ley Federal de Protección al Consumidor y la NOM-010-TUR-2001.'
        ]
        
        for texto in obligaciones_agencia:
            p_oblig = doc.add_paragraph()
            p_oblig.paragraph_format.space_after = Pt(4)
            p_oblig.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run_oblig = p_oblig.add_run(texto)
            set_run_font(run_oblig, size=7)
        
        # SEXTA
        p_sexta = doc.add_paragraph()
        p_sexta.paragraph_format.space_before = Pt(8)
        p_sexta.paragraph_format.space_after = Pt(4)
        p_sexta.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_sexta = p_sexta.add_run('SEXTA. OBLIGACIONES DE EL CLIENTE: ')
        set_run_font(run_sexta, size=7, bold=True)
        run_sexta_text = p_sexta.add_run('Cumplir con lo establecido en el presente contrato:')
        set_run_font(run_sexta_text, size=7)
        
        obligaciones_cliente = [
            'Proporcionar previo a la prestación del servicio los datos generales veraces y documentos requeridos para los servicios contratados (como pueden ser de manera enunciativa más no limitativa, el nombre, edad, identificación, comprobante de domicilio, pasaporte, visas, vacunas, constancia de situación fiscal, número telefónico, correo electrónico). Proporcionará sus propios datos y documentos de su persona así como el de las personas que lo acompañen.',
            'Realizar pagos a la AGENCIA conforme a lo pactado en el presente contrato.',
            'Respetar reglamentos de prestadores finales.',
            'Notificar por lo menos con  20 DÍAS HÁBILES y por escrito a LA AGENCIA cualquier cambio  una vez aceptado el servicio.'
        ]
        
        for texto in obligaciones_cliente:
            p_oblig = doc.add_paragraph()
            p_oblig.paragraph_format.space_after = Pt(4)
            p_oblig.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run_oblig = p_oblig.add_run(texto)
            set_run_font(run_oblig, size=7)
        
        # SÉPTIMA
        p_septima = doc.add_paragraph()
        p_septima.paragraph_format.space_before = Pt(8)
        p_septima.paragraph_format.space_after = Pt(4)
        p_septima.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_septima = p_septima.add_run('SÉPTIMA. VIGENCIA. ')
        set_run_font(run_septima, size=7, bold=True)
        run_septima_text = p_septima.add_run('El contrato estará vigente mientras se presten los servicios y se cumplan las obligaciones de pago, tiempo en que el presente Contrato surtirá todos sus efectos legales.')
        set_run_font(run_septima_text, size=7)
        
        # OCTAVA
        p_octava = doc.add_paragraph()
        p_octava.paragraph_format.space_before = Pt(8)
        p_octava.paragraph_format.space_after = Pt(4)
        p_octava.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_octava = p_octava.add_run('OCTAVA. CASO FORTUITO Y FUERZA MAYOR. ')
        set_run_font(run_octava, size=7, bold=True)
        run_octava_text = p_octava.add_run('Se entiende por caso fortuito o fuerza mayor aquellos hechos o acontecimientos ajenos a la voluntad de las partes, que sean imprevisibles, irresistibles, insuperables y que no provengan de negligencia, dolo o falta de cuidado de alguna de ellas. No se considerarán caso fortuito o fuerza mayor las enfermedades personales de EL CLIENTE o de sus acompañantes. EL CLIENTE reconoce que la AGENCIA no será responsable por errores, omisiones, falta de entrega de documentos, información incompleta o inexacta, ni por cualquier otra actuación u omisión atribuible al propio CLIENTE que afecte la reservación, emisión de boletos, acceso a servicios turísticos, cambios, cancelaciones o cualquier trámite derivado del presente contrato. Cuando el servicio turístico no pueda prestarse total o parcialmente por caso fortuito o fuerza mayor, la AGENCIA reembolsará a EL CLIENTE las cantidades que, conforme a las políticas de los prestadores finales (aerolíneas, hoteles, operadores, etc.), sean efectivamente recuperables y devueltas a la AGENCIA. EL CLIENTE tendrá derecho a recibir el reembolso correspondiente únicamente respecto de los importes efectivamente recuperados. En caso de que el servicio turístico se haya prestado de manera parcial, EL CLIENTE tendrá derecho a un reembolso proporcional exclusivamente respecto de los servicios no utilizados, conforme a lo que determine el proveedor correspondiente.')
        set_run_font(run_octava_text, size=7)
        
        # NOVENA
        p_novena = doc.add_paragraph()
        p_novena.paragraph_format.space_before = Pt(8)
        p_novena.paragraph_format.space_after = Pt(4)
        p_novena.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_novena = p_novena.add_run('NOVENA . CAMBIOS DE ORDEN DE LOS SERVICIOS CON AUTORIZACIÓN DE EL CLIENTE. ')
        set_run_font(run_novena, size=7, bold=True)
        run_novena_text = p_novena.add_run('La agencia podrá modificar el orden de los Servicios turísticos indicados en el presente Contrato, para un mejor desarrollo de los mismos o por las causas que así lo justifiquen, siempre y cuando respete la cantidad y calidad de los Servicios turísticos que se hayan contratado. Este será con la autorización por escrito de EL CLIENTE, sea cual fuese la causa. El cliente no podrá hacer cambios de fechas, rutas, ni servicios, sin previa autorización de La agencia, en caso de que dichos cambios tengan un costo, éste será indicado en al CARATULA del presente Contrato. EL CLIENTE reconoce que, una vez firmado el presente contrato y realizado el anticipo, pago parcial o total, los pagos efectuados no son cancelables ni reembolsables, en virtud de que la AGENCIA realiza de manera inmediata gestiones, reservaciones y pagos a terceros proveedores, los cuales se rigen por políticas propias de cancelación y reembolso que no dependen de la AGENCIA. EL CLIENTE acepta que cualquier solicitud de cambio, corrección o modificación respecto a fechas, nombres, itinerarios, servicios contratados o cualquier otro aspecto, estará sujeta a la disponibilidad de los proveedores, así como al pago de cargos adicionales o penalidades, conforme a las políticas vigentes de dichos proveedores.')
        set_run_font(run_novena_text, size=7)
        
        # DÉCIMA
        p_decima = doc.add_paragraph()
        p_decima.paragraph_format.space_before = Pt(8)
        p_decima.paragraph_format.space_after = Pt(4)
        p_decima.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_decima = p_decima.add_run('DÉCIMA. CANCELACIÓN. ')
        set_run_font(run_decima, size=7, bold=True)
        run_decima_text = p_decima.add_run('EL CLIENTE reconoce que, una vez firmado el presente Contrato y realizado el anticipo, pago parcial o total, los pagos no son cancelables ni reembolsables, debido a que la AGENCIA realiza de manera inmediata pagos, reservaciones y gestiones con terceros proveedores, cuyas políticas no permiten cancelaciones ni devoluciones. Cualquier solicitud de cancelación o modificación deberá realizarse por escrito, pero no dará derecho a devolución, salvo que algún proveedor permita recuperar total o parcialmente los montos pagados, caso en el cual la AGENCIA entregará al CLIENTE únicamente las cantidades efectivamente devueltas por dicho proveedor. Las modificaciones estarán sujetas a disponibilidad y podrán generar cargos adicionales conforme a las políticas de los prestadores finales. La presente cláusula aplica únicamente a solicitudes voluntarias de cancelación formuladas por EL CLIENTE. Lo anterior es independiente de las consecuencias aplicables por rescisión por incumplimiento, reguladas en las cláusulas siguientes.')
        set_run_font(run_decima_text, size=7)
        
        # DÉCIMA PRIMERA
        p_decima_primera = doc.add_paragraph()
        p_decima_primera.paragraph_format.space_before = Pt(8)
        p_decima_primera.paragraph_format.space_after = Pt(4)
        p_decima_primera.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_decima_primera = p_decima_primera.add_run('DÉCIMA PRIMERA. VUELOS. ')
        set_run_font(run_decima_primera, size=7, bold=True)
        run_decima_primera_text = p_decima_primera.add_run('EL CLIENTE reconoce que los servicios aéreos incluidos en el paquete vacacional son operados exclusivamente por la aerolínea correspondiente, por lo que Movums The Travel Store no es responsable por cambios de itinerario, demoras, reprogramaciones, sobreventas, cancelaciones, modificaciones operativas o cualquier otra decisión adoptada por la aerolínea, toda vez que dichos actos son ajenos al control de la AGENCIA.EL CLIENTE acepta que toda compensación, reembolso, cambio o beneficio derivado de acciones de la aerolínea está sujeto exclusivamente a las políticas y procedimientos de dicha aerolínea, y que la AGENCIA actuará únicamente como intermediaria en la gestión correspondiente.')
        set_run_font(run_decima_primera_text, size=7)
        
        # DÉCIMA SEGUNDA
        p_decima_segunda = doc.add_paragraph()
        p_decima_segunda.paragraph_format.space_before = Pt(8)
        p_decima_segunda.paragraph_format.space_after = Pt(4)
        p_decima_segunda.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_decima_segunda = p_decima_segunda.add_run('DÉCIMA SEGUNDA. RESCISIÓN. ')
        set_run_font(run_decima_segunda, size=7, bold=True)
        run_decima_segunda_text = p_decima_segunda.add_run('Procede si alguna parte incumple lo pactado o si el servicio no corresponde a lo solicitado. En caso de rescisión del presente Contrato, la parte que incumpla deberá de pagar lo correspondiente a la pena convencional. La AGENCIA podrá dar por terminado el presente contrato cuando EL CLIENTE no realice los depósitos o pagos en las fechas pactadas. En este supuesto, la AGENCIA notificará al CLIENTE mediante los medios de contacto proporcionados, y dicha terminación se considerará efectiva desde la fecha del incumplimiento. El CLIENTE reconoce que la falta de pago oportuno constituye un incumplimiento del contrato y acepta que los anticipos podrán aplicarse a cargos, penalidades o gastos ya generados conforme a las políticas de proveedores y prestadores de servicios turísticos. La rescisión no será considerada como una cancelación voluntaria, sino como una consecuencia jurídica del incumplimiento de cualquiera de las partes.')
        set_run_font(run_decima_segunda_text, size=7)
        
        # DÉCIMA TERCERA
        p_decima_tercera = doc.add_paragraph()
        p_decima_tercera.paragraph_format.space_before = Pt(8)
        p_decima_tercera.paragraph_format.space_after = Pt(4)
        p_decima_tercera.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_decima_tercera = p_decima_tercera.add_run('DÉCIMA TERCERA. PENA CONVENCIONAL. ')
        set_run_font(run_decima_tercera, size=7, bold=True)
        run_decima_tercera_text = p_decima_tercera.add_run('La parte incumplida pagará la penalidad indicada por el OPERADOR')
        set_run_font(run_decima_tercera_text, size=7)
        
        # DÉCIMA CUARTA
        p_decima_cuarta = doc.add_paragraph()
        p_decima_cuarta.paragraph_format.space_before = Pt(8)
        p_decima_cuarta.paragraph_format.space_after = Pt(4)
        p_decima_cuarta.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_decima_cuarta = p_decima_cuarta.add_run('DÉCIMA CUARTA.  RESERVACIONES Y PAGOS. ')
        set_run_font(run_decima_cuarta, size=7, bold=True)
        run_decima_cuarta_text = p_decima_cuarta.add_run('La aceptación y formalización del presente contrato se considerará efectiva una vez que EL CLIENTE envíe el contrato debidamente firmado y efectúe el anticipo, pago parcial o total, mismo que no es reembolsable, en virtud de que Movums The Travel Store realiza gestiones inmediatas con terceros proveedores para asegurar la disponibilidad de los servicios solicitados.')
        set_run_font(run_decima_cuarta_text, size=7)
        
        # DÉCIMA QUINTA + firmas - Sin salto de página para que suban al espacio en blanco de la página anterior
        p_decima_quinta = doc.add_paragraph()
        p_decima_quinta.paragraph_format.space_before = Pt(8)
        p_decima_quinta.paragraph_format.space_after = Pt(12)  # 1 salto de línea más después de "Estado de México."
        p_decima_quinta.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_decima_quinta = p_decima_quinta.add_run('DÉCIMA QUINTA. JURISDICCIÓN. ')
        set_run_font(run_decima_quinta, size=7, bold=True)
        run_decima_quinta_text = p_decima_quinta.add_run('Las partes se someten a PROFECO y, en su caso, a tribunales competentes de Texcoco, Estado de México.')
        set_run_font(run_decima_quinta_text, size=7)
        
        # FIRMAS DEL ANEXO - Centradas con líneas arriba de los textos
        # Línea y texto para LA AGENCIA (centrado)
        p_linea_agencia = doc.add_paragraph()
        p_linea_agencia.paragraph_format.space_before = Pt(12)
        p_linea_agencia.paragraph_format.space_after = Pt(2)
        p_linea_agencia.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_linea_agencia = p_linea_agencia.add_run('_' * 50)
        set_run_font(run_linea_agencia, size=7)
        
        p_agencia_label = doc.add_paragraph()
        p_agencia_label.paragraph_format.space_after = Pt(18)  # 3 saltos de línea de separación
        p_agencia_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_agencia_label = p_agencia_label.add_run('LA AGENCIA')
        set_run_font(run_agencia_label, size=7, bold=True)
        
        # Línea y texto para EL CLIENTE (centrado)
        p_linea_cliente = doc.add_paragraph()
        p_linea_cliente.paragraph_format.space_before = Pt(0)  # Ya hay espacio del anterior
        p_linea_cliente.paragraph_format.space_after = Pt(2)
        p_linea_cliente.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_linea_cliente = p_linea_cliente.add_run('_' * 50)
        set_run_font(run_linea_cliente, size=7)
        
        p_cliente_label = doc.add_paragraph()
        p_cliente_label.paragraph_format.space_after = Pt(6)
        p_cliente_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cliente_label = p_cliente_label.add_run('EL CLIENTE  (NOMBRE COMPLETO Y FIRMA)')
        set_run_font(run_cliente_label, size=7, bold=True)
        
        # El footer ya está en el membrete de la plantilla
        
        # Preparar respuesta HTTP
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        nombre_cliente_safe = cliente.nombre_completo_display.replace(' ', '_').replace('/', '_')
        filename = f"Contrato_Paquete_Nacional_{venta.pk}_{nombre_cliente_safe}.docx"
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        response['Content-Length'] = str(len(buffer.getvalue()))
        response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response['Pragma'] = 'no-cache'
        response['Expires'] = '0'
        
        buffer.close()
        return response


class ContratoPaqueteInternacionalPDFView(LoginRequiredMixin, DetailView):
    """
    Vista para generar contrato de servicios turísticos internacionales en formato DOCX.
    Aplica para TODAS las ventas internacionales (PAQ, VUE+HOS, solo HOS, otros).
    Formato basado en docs/Formato de contrato internacional.docx.
    """
    model = VentaViaje

    def get(self, request, *args, **kwargs):
        self.object = self.get_object()
        venta = self.object
        cliente = venta.cliente

        if venta.tipo_viaje != 'INT':
            return HttpResponse("Error: Esta venta no es internacional.", status=400)
        
        # ── Datos de cotización ──
        destino = ''
        nombre_paquete = ''
        if venta.cotizacion_origen:
            cotizacion = venta.cotizacion_origen
            if isinstance(cotizacion.propuestas, str):
                try:
                    import json
                    propuestas = json.loads(cotizacion.propuestas)
                except (ValueError, TypeError, json.JSONDecodeError):
                    propuestas = {}
            elif isinstance(cotizacion.propuestas, dict):
                propuestas = cotizacion.propuestas
            else:
                propuestas = {}
            destino = (cotizacion.destino or '').strip()
            paquete = propuestas.get('paquete', {})
            if paquete:
                nombre_paquete = (paquete.get('programa', '') or paquete.get('nombre', '') or '').strip()

        if not destino:
            destino = getattr(venta, 'destino', '') or ''
        if not nombre_paquete:
            nombre_paquete = venta.servicios_detalle_desde_logistica or ''

        # ── Pasajeros ──
        pasajeros_texto = ''
        if venta.pasajeros:
            lineas = [p.strip() for p in venta.pasajeros.split('\n') if p.strip()]
            if not lineas:
                lineas = [p.strip() for p in venta.pasajeros.split(',') if p.strip()]
            pasajeros_texto = ', '.join(lineas)
        num_viajeros = len([p.strip() for p in venta.pasajeros.split('\n') if p.strip()]) if venta.pasajeros else 1

        # ── Finanzas USD ──
        from decimal import Decimal
        precio_total_usd = venta.costo_total_con_modificacion_usd or venta.total_usd or Decimal('0.00')
        anticipo_usd = venta._cantidad_apertura_usd_resuelto()
        saldo_pendiente_usd = venta.saldo_restante_usd
        fecha_limite_pago = venta.fecha_vencimiento_pago
        fecha_salida = venta.fecha_inicio_viaje
        fecha_regreso = venta.fecha_fin_viaje

        def _fmt_date(val):
            if not val:
                return '_______________'
            if isinstance(val, datetime.date):
                return val.strftime('%d/%m/%Y')
            return str(val)

        try:
            from docx import Document
            from docx.shared import Pt, Emu, RGBColor
            from docx.oxml.ns import qn
            from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
        except ImportError:
            return HttpResponse("Error: python-docx no está instalado.", status=500)

        AZUL_MOVUMS = RGBColor(0, 0x4A, 0x8E)

        template_path = os.path.join(settings.BASE_DIR, 'static', 'docx', 'membrete.docx')
        doc = Document(template_path) if os.path.exists(template_path) else Document()

        # ── Márgenes (idénticos al formato original) ──
        for section in doc.sections:
            section.top_margin = Emu(2160270)
            section.bottom_margin = Emu(1440180)
            section.left_margin = Emu(1080135)
            section.right_margin = Emu(1080135)

        # ── Estilo base Calibri ──
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)
        rpr = style._element.get_or_add_rPr()
        rpr_fonts = rpr.find(qn('w:rFonts'))
        if rpr_fonts is None:
            from docx.oxml import OxmlElement
            rpr_fonts = OxmlElement('w:rFonts')
            rpr.append(rpr_fonts)
        rpr_fonts.set(qn('w:ascii'), 'Calibri')
        rpr_fonts.set(qn('w:hAnsi'), 'Calibri')
        rpr_fonts.set(qn('w:eastAsia'), 'Calibri')

        def _run(paragraph, text, bold=False, size=10, color=None):
            r = paragraph.add_run(text)
            r.font.name = 'Calibri'
            r.font.size = Pt(size)
            r.bold = bold
            if color:
                r.font.color.rgb = color
            return r

        # =====================================================================
        #  PARTE 1 — CARÁTULA  (Calibri 10 pt, labels bold, valores normal)
        # =====================================================================

        # Título
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, 'CONTRATO DE SERVICIOS TURÍSTICOS', bold=True, size=12, color=AZUL_MOVUMS)
        doc.add_paragraph()

        # --- Datos generales del cliente ---
        p = doc.add_paragraph()
        _run(p, 'DATOS GENERALES DEL CLIENTE', bold=True, color=AZUL_MOVUMS)

        p = doc.add_paragraph()
        _run(p, 'Nombre completo: ', bold=True)
        _run(p, cliente.nombre_completo_display.upper())

        p = doc.add_paragraph()
        _run(p, 'Teléfono: ', bold=True)
        _run(p, str(cliente.telefono or ''))

        p = doc.add_paragraph()
        _run(p, 'Correo electrónico: ', bold=True)
        _run(p, str(cliente.email or ''))

        p = doc.add_paragraph()
        _run(p, 'Identificación oficial: ', bold=True)
        _run(p, '_______________  INE / Pasaporte / Otro: ___')

        p = doc.add_paragraph()
        _run(p, 'Acompañantes:', bold=True)

        p = doc.add_paragraph()
        _run(p, '• Nombre(s) y edad(es): ', bold=True)
        _run(p, pasajeros_texto or '_______________')

        # --- Datos del servicio turístico contratado ---
        doc.add_paragraph()
        p = doc.add_paragraph()
        _run(p, 'DATOS DEL SERVICIO TURÍSTICO CONTRATADO', bold=True, color=AZUL_MOVUMS)
        doc.add_paragraph()

        p = doc.add_paragraph()
        _run(p, 'Nombre del Paquete: ', bold=True)
        _run(p, nombre_paquete or '_______________')

        p = doc.add_paragraph()
        _run(p, 'Destino(s): ', bold=True)
        _run(p, destino or '_______________')

        p = doc.add_paragraph()
        _run(p, 'Fecha de inicio: ', bold=True)
        _run(p, _fmt_date(fecha_salida))

        p = doc.add_paragraph()
        _run(p, 'Fecha de término: ', bold=True)
        _run(p, _fmt_date(fecha_regreso))

        p = doc.add_paragraph()
        _run(p, 'Número total de viajeros: ', bold=True)
        _run(p, f'{num_viajeros} adultos')

        # --- Servicios incluidos (bold 9pt) ---
        doc.add_paragraph()
        p = doc.add_paragraph()
        _run(p, 'SERVICIOS INCLUIDOS', bold=True, color=AZUL_MOVUMS)
        doc.add_paragraph()

        for txt in [
            'Vuelos Indicados en el itinerario',
            'Traslados aeropuerto-hotel-aeropuerto',
            'Tours/excursiones: Incluidos en el itinerario a excepción de los que se mencionan como OPCIONALES (que son con costo extra)',
            'Coordinador o guia en destino',
            'Entradas,accesos y actividades programadas MENCIONADAS EN EL ITINERARIO A EXCEPCION DE OPCIONALES',
        ]:
            p = doc.add_paragraph()
            _run(p, txt, bold=True, size=9)

        # --- Servicios no incluidos ---
        doc.add_paragraph()
        p = doc.add_paragraph()
        _run(p, 'SERVICIOS NO INCLUIDOS', bold=True, color=AZUL_MOVUMS)
        doc.add_paragraph()

        p = doc.add_paragraph()
        _run(p, 'ESPECIFICADOS EN EL ITINERARIO', bold=True, size=9)

        for txt in [
            '• Impuestos locales, resort fees o cuotas gubernamentales',
            '• Propinas',
            '• Servicios no listados como incluidos',
            '• Actividades opcionales',
            '• Gastos personales',
            '• Sobrepeso de equipaje',
            '• Comidas no contempladas',
        ]:
            p = doc.add_paragraph()
            _run(p, txt, size=9)

        # Salto de línea como en el documento original (P[29])
        doc.add_paragraph()

        # --- Documentación migratoria ---
        p = doc.add_paragraph()
        _run(p, 'DOCUMENTACION MIGRATORIO Y/O SANITARIA (OBLIGATORIA)', bold=True, size=9)

        p = doc.add_paragraph()
        _run(p, 'EL CLIENTE reconoce y acepta:', bold=True, size=9)

        p = doc.add_paragraph()
        _run(p, 'Que es su responsabilidad verificar requitos de visa, ETA, autorizaciones electornicas, pasaporte y/o vacunas del destino y escalas.', bold=True, size=9)

        p = doc.add_paragraph()
        _run(p, 'Que la autorización de ingreso depende exclusivamente de autoridades migratorias del país destino.', bold=True, size=9)

        # --- Documentos requeridos ---
        doc.add_paragraph()
        p = doc.add_paragraph()
        _run(p, 'DOCUMENTOS REQUERIDOS PARA ESTE VIAJE:', bold=True, color=AZUL_MOVUMS)
        doc.add_paragraph()

        p = doc.add_paragraph()
        _run(p, 'Pasaporte vigente', bold=True, size=9)

        p = doc.add_paragraph()
        _run(p, '☐ Visa/ Autorización electronica (especificar): ', size=9)
        _run(p, '________________________________________', size=9)

        # --- Servicios adicionales ---
        p = doc.add_paragraph()
        _run(p, 'SERVICIOS ADICIONALES INCLUIDOS', bold=True, size=9)

        p = doc.add_paragraph()
        _run(p, '(Marca los que apliquen)', size=9)

        p = doc.add_paragraph()
        _run(p, '☐ Seguro de viajero (si aplica)', size=9)

        p = doc.add_paragraph()
        _run(p, '☐ Otros servicios incluidos:', size=9)
        _run(p, ' ________________________', size=9)

        # --- Precio y condiciones económicas ---
        doc.add_paragraph()
        p = doc.add_paragraph()
        _run(p, 'PRECIO Y CONDICIONES ECONÓMICAS', bold=True, color=AZUL_MOVUMS)
        doc.add_paragraph()

        p = doc.add_paragraph()
        _run(p, 'Precio total del paquete: ', bold=True)
        _run(p, f'USD ${precio_total_usd:,.2f}')

        p = doc.add_paragraph()
        _run(p, 'Anticipo recibido: ', bold=True)
        _run(p, f'USD ${anticipo_usd:,.2f}')

        p = doc.add_paragraph()
        _run(p, 'Saldo pendiente: ', bold=True)
        _run(p, f'USD ${saldo_pendiente_usd:,.2f}')

        p = doc.add_paragraph()
        _run(p, 'Fecha límite de pago total: ', bold=True)
        _run(p, _fmt_date(fecha_limite_pago))

        # Salto de línea como en el documento original (P[46])
        doc.add_paragraph()

        # --- Documentación entregada al cliente ---
        doc.add_paragraph()
        p = doc.add_paragraph()
        _run(p, 'DOCUMENTACIÓN ENTREGADA AL CLIENTE', bold=True, color=AZUL_MOVUMS)
        doc.add_paragraph()

        for txt in [
            '☐ Contrato firmado',
            '☐ Copia de esta caratula',
            '☐ Itinerario preliminar',
            '☐ Políticas del proveedor',
            '☐ Comprobantes de pago',
            '☐ Claves de reservación',
            '☐ Información de contacto para emergencias',
        ]:
            p = doc.add_paragraph()
            _run(p, txt, size=9)

        # --- Declaración del cliente ---
        doc.add_paragraph()
        p = doc.add_paragraph()
        _run(p, 'DECLARACIÓN DEL CLIENTE', bold=True, color=AZUL_MOVUMS)
        doc.add_paragraph()

        p = doc.add_paragraph()
        _run(p, 'EL CLIENTE declara que:', bold=True)

        for txt in [
            '• Ha revisado y entendido toda la información contenida en este Anexo.',
            '• Proporcionó datos veraces y completos.',
            '• Acepta las condiciones del servicio, políticas de proveedores y cláusulas del contrato.',
        ]:
            p = doc.add_paragraph()
            _run(p, txt, size=9)

        # =====================================================================
        #  PARTE 2 — CONTRATO LEGAL  (Calibri 7 pt, justify)
        #  Inicia en página nueva
        # =====================================================================

        LEGAL_SIZE = 7
        JU = WD_ALIGN_PARAGRAPH.JUSTIFY

        nombre_cliente_upper = cliente.nombre_completo_display.upper()

        # Salto de página antes del contrato legal
        p_break = doc.add_paragraph()
        run_pb = p_break.add_run()
        run_pb.add_break(WD_BREAK.PAGE)

        p = doc.add_paragraph()
        p.alignment = JU
        _run(p, 'CONTRATO DE MEDIACIÓN PARA LA PRESTACIÓN DE SERVICIOS TURÍSTICOS QUE CELEBRAN POR UNA PARTE LA AGENCIA DE VIAJES "GRUPO IMVED, S.A. DE C.V." ACTUANDO EN USO DE SU NOMBRE COMERCIAL MOVUMS THE TRAVEL STORE, EN ADELANTE DENOMINADA COMO "LA AGENCIA", Y POR LA OTRA EL (LA) C_ ', bold=True, size=LEGAL_SIZE)
        _run(p, f' {nombre_cliente_upper} ', size=LEGAL_SIZE)
        _run(p, ' A QUIEN EN LO SUCESIVO SE LE DENOMINARÁ "EL CLIENTE", AL TENOR DE LAS SIGUIENTES DEFINICIONES, DECLARACIONES Y CLÁUSULAS:', bold=True, size=LEGAL_SIZE)

        # Glosario
        p = doc.add_paragraph(); p.alignment = JU; _run(p, 'GLOSARIO', bold=True, size=LEGAL_SIZE)
        p = doc.add_paragraph(); p.alignment = JU; _run(p, 'Para efectos del presente contrato, se entiende por:', size=LEGAL_SIZE)

        glosario = [
            '◆ Agencia: Es el proveedor de servicios turísticos que intermedia, contrata u ofrece servicios o productos turístico nacionales, previo pago de un precio cierto y determinado.',
            '◆ Cliente: Consumidor que contrata los servicios turísticos nacionales mediante el pago de un precio cierto y determinado.',
            '◆ Paquete turístico: Integración de uno o más servicios turísticos en un solo producto, ofrecidos al Cliente y detallado en el Anexo del presente contrato.',
            '◆ Servicio turístico: Prestación de carácter comercial en transporte nacional, hospedaje, alimentación, excursiones u otros servicios relacionados, detallados en el Anexo del presente contrato.',
            '◆ Caratula: Documento que detalla servicios, fechas, precios y condiciones del servicio turístico contratado.',
        ]
        for txt in glosario:
            p = doc.add_paragraph(); p.alignment = JU; _run(p, txt, size=LEGAL_SIZE)

        # Declaraciones
        p = doc.add_paragraph(); p.alignment = JU; _run(p, 'DECLARACIONES', bold=True, size=LEGAL_SIZE)
        p = doc.add_paragraph(); p.alignment = JU; _run(p, 'I. Declara LA AGENCIA:', bold=True, size=LEGAL_SIZE)

        decl_agencia = [
            '◆ Ser una persona moral legalmente constituida conforme a las leyes mexicanas.',
            '◆ Ser un Prestador de Servicios Turísticos con Razón Social: GRUPO IMVED, S.A. de C.V.',
            '◆ Ser la única propietaria de la marca MOVUMS THE TRAVEL STORE',
            '◆ RFC GIM190722FS7 y domicilio ubicado en: Plaza Mora, Juárez Sur, 321, interior 18, Colonia Centro, Texcoco, Estado de México, C.P. 56100.',
            '◆ Teléfono, correo electrónico y horario de atención al público: 59 59319954, 5951255279 ventas@movums.com, lunes a sábado de 09:00 a 18:00 horas.',
            '◆ Contar con infraestructura, personal capacitado y experiencia suficiente para la prestación de los servicios turísticos contratados.',
            '◆ Haber informado previamente al Cliente sobre los precios, tarifas, condiciones, características y costo total del servicio turístico contratado.',
        ]
        for txt in decl_agencia:
            p = doc.add_paragraph(); p.alignment = JU; _run(p, txt, size=LEGAL_SIZE)

        doc.add_paragraph()
        p = doc.add_paragraph(); p.alignment = JU; _run(p, 'II. Declara EL CLIENTE:', bold=True, size=LEGAL_SIZE)

        decl_cliente = [
            '◆ Ser persona física/moral con capacidad legal y económica para obligarse en términos del presente contrato.',
            '◆ En caso de persona moral: Ser una persona moral legalmente constituida conforme a las leyes mexicanas, conforme lo acredita con copia del instrumento número ____________, de fecha ___________, otorgado ante la Fe del Notario Público Número ____, de ____________, y que el(la) C. _________ __________________ en este acto interviene en su carácter de Representante Legal, calidad que acredita con copia del instrumento número _______, de fecha _________, otorgada ante la Fe del Notario Público número _______ del _________, facultad y calidad que no le han sido revocadas, modificadas o limitadas a la fecha de firma del presente contrato.',
            '◆ Encontrarse inscrito en el Registro Federal de Contribuyentes con la clave que ha manifestado.',
            '◆ Haber recibido previamente de LA AGENCIA información útil, precisa, veraz y detallada sobre los servicios objeto del presente contrato.',
            '◆ Proporciona su nombre, domicilio, número telefónico y correo electrónico, tal y como lo ha señalado en la caratula de prestación de servicios, acreditando los mismos con copia de los documentos idóneos para tal efecto.',
        ]
        for txt in decl_cliente:
            p = doc.add_paragraph(); p.alignment = JU; _run(p, txt, size=LEGAL_SIZE)

        doc.add_paragraph()

        # Cláusulas
        p = doc.add_paragraph(); p.alignment = JU; _run(p, 'CLÁUSULAS', bold=True, size=LEGAL_SIZE)

        clausulas_simple = [
            ('PRIMERA. CONSENTIMIENTO.', ' Las partes manifiestan su voluntad de celebrar el presente contrato, cuya naturaleza jurídica es la mediación para la prestación de servicios turísticos.'),
            ('SEGUNDA. OBJETO.', ' LA AGENCIA intermediará, contratará u ofrecerá servicios turísticos detallados en la CARATULA, previo pago del Cliente de un precio cierto y determinado.'),
        ]
        for titulo, cuerpo in clausulas_simple:
            p = doc.add_paragraph(); p.alignment = JU
            _run(p, titulo, bold=True, size=LEGAL_SIZE)
            _run(p, cuerpo, size=LEGAL_SIZE)

        # TERCERA (larga, con sub-incisos)
        p = doc.add_paragraph(); p.alignment = JU
        _run(p, 'TERCERA. PRECIO, FORMA Y LUGAR DE PAGO.', bold=True, size=LEGAL_SIZE)
        _run(p, ' Las partes manifiestan su conformidad en que el precio total a pagar por EL CLIENTE como contraprestación del Servicio turístico, es la cantidad que por cada concepto se indica en la CARATULA de este Contrato. El importe señalado en la CARATULA, contempla todas las cantidades y conceptos referentes al Servicio turístico, por lo que LA AGENCIA se obliga a respetar en todo momento dicho costo sin poder cobrar otra cantidad o condicionar la prestación del Servicio turístico contratado a la adquisición de otro servicio no requerido por El cliente, salvo que El cliente autorice de manera escrita algún otro cobro no estipulado en el presente Contrato. EL CLIENTE efectuará el pago pactado por el Servicio turístico señalado en la caratula del presente Contrato en los términos y condiciones acordadas pudiendo ser:', size=LEGAL_SIZE)

        tercera_incisos = [
            'a) Al contado: en efectivo, con tarjeta de débito, tarjeta de crédito, transferencia bancaria, y/o cheque en el domicilio de la agencia en moneda nacional, sin menoscabo de poderlo hacer en moneda extranjera al tipo de cambio publicado en el Diario Oficial de la Federación al día en que el pago se efectúe.',
            'b) A plazos: El cliente podrá, previo acuerdo con La agencia a pagar en parcialidades, para lo cual, La agencia deberá de entregar a El CLIENTE la información por escrito de las fechas, así como los montos parciales a pagar.',
            'c) En caso de que El cliente realice el pago con cheque y no se cubra el pago por causas imputables al librador, La agencia tendrá el derecho de realizar el cobro adicional del 20% (veinte por ciento) del valor del documento, por concepto de daños y perjuicios, en caso de que el cheque sea devuelto por causas imputables al librador, conforme al artículo 193 de la Ley General del Títulos y Operaciones de Crédito.',
        ]
        for inciso in tercera_incisos:
            p = doc.add_paragraph(); p.alignment = JU; _run(p, inciso, size=LEGAL_SIZE)

        doc.add_paragraph()

        # QUINTA - Obligaciones de la Agencia
        p = doc.add_paragraph(); p.alignment = JU; _run(p, 'QUINTA. OBLIGACIONES DE LA AGENCIA. LA AGENCIA SE OBLIGA A:', bold=True, size=LEGAL_SIZE)
        doc.add_paragraph()

        quinta_incisos = [
            'A) Cumplir lo pactado en el contrato.',
            'B) Entregar a EL CLIENTE copia del contrato y constancias de reservación.',
            'C) Proporcionar a EL CLIENTE boletos, claves de reservación y documentos de viaje.',
            'D) Auxiliar a EL CLIENTE en emergencias y gestionar indemnizaciones relacionadas con el servicio contratado',
            'E) Solicitar los Servicios turísticos que se especifican en la caratula de este Contrato por cuenta de EL CLIENTE de acuerdo a la disponibilidad de los mismos, a contratarlos fungiendo como intermediario entre éste y las personas encargadas de proporcionar directamente el Servicio turístico.',
            'F) Coadyuvar a EL CLIENTE para reclamar ante el prestador del servicio final, las indemnizaciones que correspondan.',
            'G) Respetar la Ley Federal de Protección al Consumidor y la NOM-010-TUR-2001.',
        ]
        for inciso in quinta_incisos:
            p = doc.add_paragraph(); p.alignment = JU; _run(p, inciso, size=LEGAL_SIZE)

        # SEXTA - Obligaciones del Cliente
        p = doc.add_paragraph(); p.alignment = JU
        _run(p, 'SEXTA. OBLIGACIONES DE EL CLIENTE:', bold=True, size=LEGAL_SIZE)
        _run(p, ' Cumplir con lo establecido en el presente contrato:', size=LEGAL_SIZE)

        sexta_incisos = [
            'A) Proporcionar previo a la prestación del servicio los datos generales veraces y documentos requeridos para los servicios contratados (como pueden ser de manera enunciativa más no limitativa, el nombre, edad, identificación, comprobante de domicilio, pasaporte, visas, vacunas, constancia de situación fiscal, número telefónico, correo electrónico). Proporcionará sus propios datos y documentos de su persona así como el de las personas que lo acompañen.',
            'B) Realizar pagos a la AGENCIA conforme a lo pactado en el presente contrato.',
            'C) Respetar reglamentos de prestadores finales.',
            'D) Notificar por lo menos con CINCO DÍAS HÁBILES y por escrito a LA AGENCIA cualquier cambio o cancelación una vez aceptado el servicio.',
        ]
        for inciso in sexta_incisos:
            p = doc.add_paragraph(); p.alignment = JU; _run(p, inciso, size=LEGAL_SIZE)

        # Cláusulas SÉPTIMA a DÉCIMA QUINTA
        clausulas_restantes = [
            ('SÉPTIMA. VIGENCIA.', ' El contrato estará vigente mientras se presten los servicios y se cumplan las obligaciones de pago, tiempo en que el presente Contrato surtirá todos sus efectos legales.'),
            ('OCTAVA. CASO FORTUITO Y FUERZA MAYOR.', ' Se entiende por caso fortuito o fuerza mayor aquellos hechos o acontecimientos ajenos a la voluntad de las partes, que sean imprevisibles, irresistibles, insuperables y que no provengan de negligencia, dolo o falta de cuidado de alguna de ellas. No se considerarán caso fortuito o fuerza mayor las enfermedades personales de EL CLIENTE o de sus acompañantes. EL CLIENTE reconoce que la AGENCIA no será responsable por errores, omisiones, falta de entrega de documentos, información incompleta o inexacta, ni por cualquier otra actuación u omisión atribuible al propio CLIENTE que afecte la reservación, emisión de boletos, acceso a servicios turísticos, cambios, cancelaciones o cualquier trámite derivado del presente contrato. Cuando el servicio turístico no pueda prestarse total o parcialmente por caso fortuito o fuerza mayor, la AGENCIA reembolsará a EL CLIENTE las cantidades que, conforme a las políticas de los prestadores finales (aerolíneas, hoteles, operadores, etc.), sean efectivamente recuperables y devueltas a la AGENCIA. EL CLIENTE tendrá derecho a recibir el reembolso correspondiente únicamente respecto de los importes efectivamente recuperados. En caso de que el servicio turístico se haya prestado de manera parcial, EL CLIENTE tendrá derecho a un reembolso proporcional exclusivamente respecto de los servicios no utilizados, conforme a lo que determine el proveedor correspondiente.'),
            ('NOVENA. CAMBIOS DE ORDEN DE LOS SERVICIOS CON AUTORIZACIÓN DE EL CLIENTE.', ' La agencia podrá modificar el orden de los Servicios turísticos indicados en el presente Contrato, para un mejor desarrollo de los mismos o por las causas que así lo justifiquen, siempre y cuando respete la cantidad y calidad de los Servicios turísticos que se hayan contratado. Este será con la autorización por escrito de EL CLIENTE, sea cual fuese la causa. El cliente no podrá hacer cambios de fechas, rutas, ni servicios, sin previa autorización de La agencia, en caso de que dichos cambios tengan un costo, éste será indicado en al CARATULA del presente Contrato. EL CLIENTE reconoce que, una vez firmado el presente contrato y realizado el anticipo, pago parcial o total, los pagos efectuados no son cancelables ni reembolsables, en virtud de que la AGENCIA realiza de manera inmediata gestiones, reservaciones y pagos a terceros proveedores, los cuales se rigen por políticas propias de cancelación y reembolso que no dependen de la AGENCIA. EL CLIENTE acepta que cualquier solicitud de cambio, corrección o modificación respecto a fechas, nombres, itinerarios, servicios contratados o cualquier otro aspecto, estará sujeta a la disponibilidad de los proveedores, así como al pago de cargos adicionales o penalidades, conforme a las políticas vigentes de dichos proveedores.'),
            ('DÉCIMA. CANCELACIÓN.', ' EL CLIENTE reconoce que, una vez firmado el presente Contrato y realizado el anticipo, pago parcial o total, los pagos no son cancelables ni reembolsables, debido a que la AGENCIA realiza de manera inmediata pagos, reservaciones y gestiones con terceros proveedores, cuyas políticas no permiten cancelaciones ni devoluciones. Cualquier solicitud de cancelación o modificación deberá realizarse por escrito, pero no dará derecho a devolución, salvo que algún proveedor permita recuperar total o parcialmente los montos pagados, caso en el cual la AGENCIA entregará al CLIENTE únicamente las cantidades efectivamente devueltas por dicho proveedor. Las modificaciones estarán sujetas a disponibilidad y podrán generar cargos adicionales conforme a las políticas de los prestadores finales. La presente cláusula aplica únicamente a solicitudes voluntarias de cancelación formuladas por EL CLIENTE. Lo anterior es independiente de las consecuencias aplicables por rescisión por incumplimiento, reguladas en las cláusulas siguientes.'),
            ('DÉCIMA PRIMERA. VUELOS.', ' EL CLIENTE reconoce que los servicios aéreos incluidos en el paquete vacacional son operados exclusivamente por la aerolínea correspondiente, por lo que Movums The Travel Store no es responsable por cambios de itinerario, demoras, reprogramaciones, sobreventas, cancelaciones, modificaciones operativas o cualquier otra decisión adoptada por la aerolínea, toda vez que dichos actos son ajenos al control de la AGENCIA. EL CLIENTE acepta que toda compensación, reembolso, cambio o beneficio derivado de acciones de la aerolínea está sujeto exclusivamente a las políticas y procedimientos de dicha aerolínea, y que la AGENCIA actuará únicamente como intermediaria en la gestión correspondiente.'),
            ('DÉCIMA SEGUNDA. RESCISIÓN.', ' Procede si alguna parte incumple lo pactado o si el servicio no corresponde a lo solicitado. En caso de rescisión del presente Contrato, la parte que incumpla deberá de pagar lo correspondiente a la pena convencional. La AGENCIA podrá dar por terminado el presente contrato cuando EL CLIENTE no realice los depósitos o pagos en las fechas pactadas. En este supuesto, la AGENCIA notificará al CLIENTE mediante los medios de contacto proporcionados, y dicha terminación se considerará efectiva desde la fecha del incumplimiento. El CLIENTE reconoce que la falta de pago oportuno constituye un incumplimiento del contrato y acepta que los anticipos podrán aplicarse a cargos, penalidades o gastos ya generados conforme a las políticas de proveedores y prestadores de servicios turísticos. La rescisión no será considerada como una cancelación voluntaria, sino como una consecuencia jurídica del incumplimiento de cualquiera de las partes.'),
            ('DÉCIMA TERCERA. PENA CONVENCIONAL.', ' La parte incumplida pagará el 20% (veinte por ciento) del precio total del servicio, sin incluir IVA.'),
            ('DÉCIMA CUARTA. RESERVACIONES Y PAGOS.', ' La aceptación y formalización del presente contrato se considerará efectiva una vez que EL CLIENTE envíe el contrato debidamente firmado y efectúe el anticipo, pago parcial o total, mismo que no es reembolsable, en virtud de que Movums The Travel Store realiza gestiones inmediatas con terceros proveedores para asegurar la disponibilidad de los servicios solicitados.'),
            ('DÉCIMA QUINTA. JURISDICCIÓN.', ' Las partes se someten a PROFECO y, en su caso, a tribunales competentes de Texcoco, Estado de México.'),
        ]
        for titulo, cuerpo in clausulas_restantes:
            p = doc.add_paragraph(); p.alignment = JU
            _run(p, titulo, bold=True, size=LEGAL_SIZE)
            _run(p, cuerpo, size=LEGAL_SIZE)

        # ── Firmas (parte legal) ──
        p = doc.add_paragraph(); p.alignment = JU; _run(p, 'FIRMAS', bold=True, size=LEGAL_SIZE)
        p = doc.add_paragraph(); p.alignment = JU; _run(p, 'CLIENTE', bold=True, size=LEGAL_SIZE)
        p = doc.add_paragraph(); p.alignment = JU
        _run(p, 'Nombre y firma: ', bold=True, size=LEGAL_SIZE)
        _run(p, '___________________________________________', size=LEGAL_SIZE)

        p = doc.add_paragraph(); p.alignment = JU; _run(p, 'AGENCIA – Movums The Travel Store', bold=True, size=LEGAL_SIZE)
        p = doc.add_paragraph(); p.alignment = JU
        _run(p, 'Nombre y firma del representante: ', bold=True, size=LEGAL_SIZE)
        _run(p, '____________________________', size=LEGAL_SIZE)

        fecha_creacion = venta.fecha_creacion if hasattr(venta, 'fecha_creacion') and venta.fecha_creacion else timezone.now().date()
        p = doc.add_paragraph(); p.alignment = JU
        _run(p, 'Fecha: ', bold=True, size=LEGAL_SIZE)
        _run(p, _fmt_date(fecha_creacion), size=LEGAL_SIZE)

        # =====================================================================
        #  Respuesta HTTP
        # =====================================================================
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        nombre_cliente_safe = cliente.nombre_completo_display.replace(' ', '_').replace('/', '_')
        filename = f"Contrato_Internacional_{venta.pk}_{nombre_cliente_safe}.docx"

        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        response['Content-Length'] = str(len(buffer.getvalue()))
        response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response['Pragma'] = 'no-cache'
        response['Expires'] = '0'

        buffer.close()
        return response

# ------------------- 11. ELIMINACIÓN DE VENTA -------------------

class VentaViajeDeleteView(LoginRequiredMixin, UserPassesTestMixin, DeleteView):
    """
    Vista para eliminar una VentaViaje.
    Solo accesible para roles 'JEFE' y solo cuando la venta está cerrada, liquidada o cancelada.
    """
    model = VentaViaje
    template_name = 'ventas/venta_confirm_delete.html' 
    success_url = reverse_lazy('lista_ventas')

    def test_func(self):
        # Solo permite eliminar a JEFE
        user_rol = perm.get_user_role(self.request.user, self.request)
        if user_rol != 'JEFE':
            return False
        
        # Verificar que la venta esté cerrada, liquidada o cancelada
        venta = self.get_object()
        
        # Una venta está cerrada/liquidada si:
        # 1. estado_confirmacion == 'COMPLETADO' (pagada completamente)
        # 2. O esta_pagada == True (saldo <= 0)
        esta_liquidada = venta.estado_confirmacion == 'COMPLETADO' or venta.esta_pagada
        
        # Una venta está cancelada si:
        # estado == 'CANCELADA'
        esta_cancelada = venta.estado == 'CANCELADA'
        
        # Solo se puede eliminar si está liquidada o cancelada
        return esta_liquidada or esta_cancelada

    def handle_no_permission(self):
        # Redirige al detalle de la venta si no tiene permiso o la venta no está en estado válido
        venta = self.get_object()
        user_rol = perm.get_user_role(self.request.user, self.request)
        
        if user_rol != 'JEFE':
            messages.error(self.request, "Solo el JEFE puede eliminar ventas.")
        else:
            messages.error(self.request, "Solo se pueden eliminar ventas que estén cerradas, liquidadas o canceladas.")
        
        return redirect('detalle_venta', pk=venta.pk, slug=venta.slug_safe)
    
# ------------------- 12. SISTEMA ROBUSTO DE COMISIONES (MOSTRADOR) -------------------

class ComisionesMensualesView(LoginRequiredMixin, TemplateView):
    """
    Vista para calcular y mostrar comisiones mensuales de Asesores de Mostrador.
    - JEFE/CONTADOR: Ve todas las comisiones
    - VENDEDOR: Solo ve sus propias comisiones
    """
    template_name = 'ventas/comisiones_mensuales.html'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        user = self.request.user
        user_rol = perm.get_user_role(user)
        
        # Obtener mes y año del request (por defecto mes actual)
        mes_actual = timezone.now().month
        anio_actual = timezone.now().year
        
        mes_filtro = safe_int(self.request.GET.get('mes'), default=mes_actual)
        anio_filtro = safe_int(self.request.GET.get('anio'), default=anio_actual)
        
        # Determinar qué vendedores mostrar
        if user_rol == 'VENDEDOR':
            vendedores_a_mostrar = User.objects.filter(pk=user.pk)
        elif perm.has_full_access(user) or perm.is_contador(user):
            # Solo mostrar vendedores de tipo MOSTRADOR
            vendedores_a_mostrar = User.objects.filter(
                perfil__rol='VENDEDOR',
                perfil__tipo_vendedor='MOSTRADOR'
            ).order_by('username')
        else:
            vendedores_a_mostrar = User.objects.none()
        
        # Obtener o calcular comisiones mensuales
        comisiones_mensuales = []
        for vendedor in vendedores_a_mostrar:
            # Verificar si ya existe el cálculo para este mes
            comision_mensual = ComisionMensual.objects.filter(
                vendedor=vendedor,
                mes=mes_filtro,
                anio=anio_filtro,
                tipo_vendedor='MOSTRADOR'
            ).first()
            
            # Si no existe, calcular automáticamente
            if not comision_mensual:
                try:
                    from ventas.services.comisiones import calcular_comisiones_mensuales_mostrador
                    comision_mensual = calcular_comisiones_mensuales_mostrador(
                        vendedor, mes_filtro, anio_filtro
                    )
                except Exception as e:
                    logger.error(f"Error al calcular comisiones para {vendedor.username}: {e}", exc_info=True)
                    continue
            
            # Obtener comisiones de ventas para el detalle
            comisiones_ventas = ComisionVenta.objects.filter(
                vendedor=vendedor,
                mes=mes_filtro,
                anio=anio_filtro
            ).select_related('venta', 'venta__cliente').order_by('-venta__fecha_creacion')
            
            comisiones_mensuales.append({
                'vendedor': vendedor,
                'comision_mensual': comision_mensual,
                'comisiones_ventas': comisiones_ventas,
                'es_usuario_actual': (vendedor.pk == user.pk),
            })
        
        context['comisiones_mensuales'] = comisiones_mensuales
        context['mes_filtro'] = mes_filtro
        context['anio_filtro'] = anio_filtro
        context['user_rol'] = user_rol
        
        # Generar lista de meses y años para el selector
        meses = [(i, datetime.datetime(2000, i, 1).strftime('%B')) for i in range(1, 13)]
        anios = list(range(anio_actual - 2, anio_actual + 1))
        
        context['meses'] = meses
        context['anios'] = anios
        
        return context
    
    def post(self, request, *args, **kwargs):
        """Recalcular comisiones para un mes específico"""
        from ventas.validators import safe_int
        mes = safe_int(request.POST.get('mes'), default=timezone.now().month)
        anio = safe_int(request.POST.get('anio'), default=timezone.now().year)
        
        user = request.user
        user_rol = perm.get_user_role(user)
        
        # Determinar qué vendedores recalcular
        if user_rol == 'VENDEDOR':
            vendedores = [user]
        elif perm.has_full_access(user) or perm.is_contador(user):
            vendedores = User.objects.filter(
                perfil__rol='VENDEDOR',
                perfil__tipo_vendedor='MOSTRADOR'
            )
        else:
            messages.error(request, "No tienes permiso para recalcular comisiones.")
            return redirect('comisiones_mensuales')
        
        # Recalcular comisiones
        from ventas.services.comisiones import calcular_comisiones_mensuales_mostrador
        
        recalculos = 0
        for vendedor in vendedores:
            try:
                calcular_comisiones_mensuales_mostrador(vendedor, mes, anio)
                recalculos += 1
            except Exception as e:
                logger.error(f"Error al recalcular comisiones para {vendedor.username}: {e}", exc_info=True)
        
        messages.success(request, f"Comisiones recalculadas para {recalculos} vendedor(es).")
        return redirect(f"{reverse('comisiones_mensuales')}?mes={mes}&anio={anio}")


class DetalleComisionesMensualesView(LoginRequiredMixin, UserPassesTestMixin, DetailView):
    """
    Vista de detalle de comisiones mensuales de un vendedor específico.
    Muestra el desglose por venta.
    """
    model = User
    template_name = 'ventas/detalle_comisiones_mensuales.html'
    context_object_name = 'vendedor'
    
    def test_func(self):
        """Verificar permisos"""
        user = self.request.user
        user_rol = perm.get_user_role(user)
        vendedor = self.get_object()
        
        # El vendedor puede ver sus propias comisiones, JEFE/CONTADOR pueden ver todas
        if user_rol == 'VENDEDOR':
            return vendedor.pk == user.pk
        elif perm.has_full_access(user) or perm.is_contador(user):
            return True
        return False
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        vendedor = self.get_object()
        
        # Obtener mes y año del request
        mes_actual = timezone.now().month
        anio_actual = timezone.now().year
        
        mes_filtro = safe_int(self.request.GET.get('mes'), default=mes_actual)
        anio_filtro = safe_int(self.request.GET.get('anio'), default=anio_actual)
        
        # Obtener comisión mensual
        comision_mensual = ComisionMensual.objects.filter(
            vendedor=vendedor,
            mes=mes_filtro,
            anio=anio_filtro,
            tipo_vendedor='MOSTRADOR'
        ).first()
        
        # Si no existe, calcular automáticamente
        if not comision_mensual:
            try:
                from ventas.services.comisiones import calcular_comisiones_mensuales_mostrador
                comision_mensual = calcular_comisiones_mensuales_mostrador(
                    vendedor, mes_filtro, anio_filtro
                )
            except Exception as e:
                logger.error(f"Error al calcular comisiones para {vendedor.username}: {e}", exc_info=True)
        
        # Obtener comisiones de ventas
        comisiones_ventas = ComisionVenta.objects.filter(
            vendedor=vendedor,
            mes=mes_filtro,
            anio=anio_filtro
        ).select_related('venta', 'venta__cliente').order_by('-venta__fecha_creacion')
        
        context['comision_mensual'] = comision_mensual
        context['comisiones_ventas'] = comisiones_ventas
        context['mes_filtro'] = mes_filtro
        context['anio_filtro'] = anio_filtro
        
        # Generar lista de meses y años
        meses = [(i, datetime.datetime(2000, i, 1).strftime('%B')) for i in range(1, 13)]
        anios = list(range(anio_actual - 2, anio_actual + 1))
        
        context['meses'] = meses
        context['anios'] = anios
        
        return context


class ExportarComisionesMensualesExcelView(LoginRequiredMixin, View):
    """
    Exporta las comisiones mensuales a Excel.
    """
    def get(self, request, pk):
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        except ImportError:
            messages.error(request, "Error: openpyxl no está instalado. Instala con: pip install openpyxl")
            return redirect('comisiones_mensuales')
        
        vendedor = get_object_or_404(User, pk=pk)
        
        # Verificar permisos
        user = request.user
        user_rol = perm.get_user_role(user)
        
        if user_rol == 'VENDEDOR' and vendedor.pk != user.pk:
            messages.error(request, "No tienes permiso para exportar estas comisiones.")
            return redirect('comisiones_mensuales')
        elif user_rol not in ['JEFE', 'CONTADOR', 'VENDEDOR']:
            messages.error(request, "No tienes permiso para exportar comisiones.")
            return redirect('comisiones_mensuales')
        
        # Obtener mes y año del request
        mes_actual = timezone.now().month
        anio_actual = timezone.now().year
        
        mes_filtro = safe_int(request.GET.get('mes'), default=mes_actual)
        anio_filtro = safe_int(request.GET.get('anio'), default=anio_actual)
        
        # Obtener comisión mensual
        comision_mensual = ComisionMensual.objects.filter(
            vendedor=vendedor,
            mes=mes_filtro,
            anio=anio_filtro,
            tipo_vendedor='MOSTRADOR'
        ).first()
        
        if not comision_mensual:
            messages.warning(request, "No hay comisiones calculadas para este período.")
            return redirect('detalle_comisiones_mensuales', pk=vendedor.pk)
        
        # Obtener comisiones de ventas
        comisiones_ventas = ComisionVenta.objects.filter(
            vendedor=vendedor,
            mes=mes_filtro,
            anio=anio_filtro
        ).select_related('venta', 'venta__cliente').order_by('venta__fecha_creacion')
        
        # Crear workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Comisiones Mensuales"
        
        # Estilos
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=12)
        title_font = Font(bold=True, size=14)
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Título
        ws['A1'] = f"Reporte de Comisiones Mensuales - {vendedor.get_full_name() or vendedor.username}"
        ws['A1'].font = title_font
        ws.merge_cells('A1:F1')
        
        # Información del período
        ws['A3'] = f"Período: {datetime.datetime(anio_filtro, mes_filtro, 1).strftime('%B %Y')}"
        ws['A3'].font = Font(bold=True)
        
        # Resumen mensual
        row = 5
        ws[f'A{row}'] = "RESUMEN MENSUAL"
        ws[f'A{row}'].font = Font(bold=True, size=12)
        ws.merge_cells(f'A{row}:B{row}')
        
        row += 1
        ws[f'A{row}'] = "Total de Ventas del Mes:"
        ws[f'B{row}'] = f"${comision_mensual.total_ventas_mes:,.2f}"
        ws[f'A{row}'].font = Font(bold=True)
        
        row += 1
        ws[f'A{row}'] = "Porcentaje de Comisión:"
        ws[f'B{row}'] = f"{comision_mensual.porcentaje_comision:.2f}%"
        ws[f'A{row}'].font = Font(bold=True)
        
        row += 1
        ws[f'A{row}'] = "Bono Extra (1% sobre $500,000+):"
        ws[f'B{row}'] = f"${comision_mensual.bono_extra:,.2f}"
        ws[f'A{row}'].font = Font(bold=True)
        
        row += 1
        ws[f'A{row}'] = "Comisión Total Pagada:"
        ws[f'B{row}'] = f"${comision_mensual.comision_total_pagada:,.2f}"
        ws[f'A{row}'].font = Font(bold=True)
        ws[f'B{row}'].fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
        
        row += 1
        ws[f'A{row}'] = "Comisión Total Pendiente:"
        ws[f'B{row}'] = f"${comision_mensual.comision_total_pendiente:,.2f}"
        ws[f'A{row}'].font = Font(bold=True)
        ws[f'B{row}'].fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
        
        row += 1
        ws[f'A{row}'] = "COMISIÓN TOTAL:"
        ws[f'B{row}'] = f"${comision_mensual.comision_total:,.2f}"
        ws[f'A{row}'].font = Font(bold=True, size=12)
        ws[f'B{row}'].font = Font(bold=True, size=12)
        ws[f'B{row}'].fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
        
        # Detalle por venta
        row += 3
        ws[f'A{row}'] = "DETALLE POR VENTA"
        ws[f'A{row}'].font = Font(bold=True, size=12)
        ws.merge_cells(f'A{row}:H{row}')
        
        row += 1
        headers = ['Venta #', 'Cliente', 'Tipo', 'Monto Base', 'Porcentaje', 'Comisión Calculada', 'Estado Pago', 'Comisión Pagada', 'Comisión Pendiente']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border
        
        # Datos de ventas
        for comision_venta in comisiones_ventas:
            row += 1
            venta = comision_venta.venta
            
            ws.cell(row=row, column=1, value=venta.pk)
            ws.cell(row=row, column=2, value=venta.cliente.nombre_completo_display)
            ws.cell(row=row, column=3, value=comision_venta.get_tipo_venta_display())
            ws.cell(row=row, column=4, value=float(comision_venta.monto_base_comision))
            ws.cell(row=row, column=5, value=f"{comision_venta.porcentaje_aplicado:.2f}%")
            ws.cell(row=row, column=6, value=float(comision_venta.comision_calculada))
            ws.cell(row=row, column=7, value=comision_venta.get_estado_pago_venta_display())
            ws.cell(row=row, column=8, value=float(comision_venta.comision_pagada))
            ws.cell(row=row, column=9, value=float(comision_venta.comision_pendiente))
            
            # Aplicar formato
            for col in range(1, 10):
                cell = ws.cell(row=row, column=col)
                cell.border = border
                if col in [4, 6, 8, 9]:  # Columnas numéricas
                    cell.number_format = '#,##0.00'
                    cell.alignment = Alignment(horizontal='right')
        
        # Ajustar ancho de columnas
        ws.column_dimensions['A'].width = 12
        ws.column_dimensions['B'].width = 30
        ws.column_dimensions['C'].width = 15
        ws.column_dimensions['D'].width = 15
        ws.column_dimensions['E'].width = 12
        ws.column_dimensions['F'].width = 18
        ws.column_dimensions['G'].width = 15
        ws.column_dimensions['H'].width = 18
        ws.column_dimensions['I'].width = 18
        
        # Preparar respuesta
        nombre_vendedor_safe = (vendedor.get_full_name() or vendedor.username).replace(' ', '_')
        nombre_mes = datetime.datetime(anio_filtro, mes_filtro, 1).strftime('%B_%Y')
        filename = f"comisiones_{nombre_vendedor_safe}_{nombre_mes}.xlsx"
        
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        wb.save(response)
        return response


class ExportarComisionesMensualesTodosExcelView(LoginRequiredMixin, View):
    """
    Exporta las comisiones mensuales de TODOS los vendedores a Excel.
    Solo disponible para JEFE y CONTADOR.
    """
    def get(self, request):
        # Determinar desde dónde se llamó para redirigir correctamente
        referer = request.META.get('HTTP_REFERER', '')
        viene_de_reporte_comisiones = 'reporte_comisiones' in referer or 'comisiones/' in referer
        
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        except ImportError:
            messages.error(request, "Error: openpyxl no está instalado. Instala con: pip install openpyxl")
            if viene_de_reporte_comisiones:
                return redirect('reporte_comisiones')
            return redirect('comisiones_mensuales')
        
        # Verificar permisos
        user = request.user
        user_rol = perm.get_user_role(user)
        
        if user_rol not in ['JEFE', 'CONTADOR']:
            messages.error(request, "No tienes permiso para exportar comisiones de todos los vendedores.")
            if viene_de_reporte_comisiones:
                return redirect('reporte_comisiones')
            return redirect('comisiones_mensuales')
        
        # Obtener mes y año del request
        mes_actual = timezone.now().month
        anio_actual = timezone.now().year
        
        mes_filtro = safe_int(request.GET.get('mes'), default=mes_actual)
        anio_filtro = safe_int(request.GET.get('anio'), default=anio_actual)
        
        logger.info(f"Exportando comisiones para mes={mes_filtro}, año={anio_filtro}")
        
        # Obtener todos los vendedores de tipo MOSTRADOR
        vendedores = User.objects.filter(
            perfil__rol='VENDEDOR',
            perfil__tipo_vendedor='MOSTRADOR'
        ).order_by('username')
        
        logger.info(f"Vendedores encontrados: {vendedores.count()}")
        
        if not vendedores.exists():
            messages.warning(request, "No hay vendedores de tipo MOSTRADOR.")
            if viene_de_reporte_comisiones:
                return redirect('reporte_comisiones')
            return redirect('comisiones_mensuales')
        
        # Crear workbook
        wb = Workbook()
        
        # ===== HOJA 1: RESUMEN GENERAL =====
        ws_resumen = wb.active
        ws_resumen.title = "Resumen General"
        
        # Estilos
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=12)
        title_font = Font(bold=True, size=14)
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Título
        ws_resumen['A1'] = f"Reporte de Comisiones Mensuales - Todos los Vendedores"
        ws_resumen['A1'].font = title_font
        ws_resumen.merge_cells('A1:J1')
        
        # Información del período
        ws_resumen['A3'] = f"Período: {datetime.datetime(anio_filtro, mes_filtro, 1).strftime('%B %Y')}"
        ws_resumen['A3'].font = Font(bold=True)
        
        # Encabezados de resumen
        row = 5
        headers_resumen = [
            'Vendedor', 'Total Ventas', 'Porcentaje', 'Bono Extra', 
            'Comisión Pagada', 'Comisión Pendiente', 'Comisión Total'
        ]
        for col, header in enumerate(headers_resumen, 1):
            cell = ws_resumen.cell(row=row, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border
        
        # Datos de resumen por vendedor
        total_general_ventas = Decimal('0.00')
        total_general_pagada = Decimal('0.00')
        total_general_pendiente = Decimal('0.00')
        total_general_comision = Decimal('0.00')
        total_general_bono = Decimal('0.00')
        
        vendedores_con_datos = 0
        
        for vendedor in vendedores:
            # Obtener o calcular comisión mensual
            comision_mensual = ComisionMensual.objects.filter(
                vendedor=vendedor,
                mes=mes_filtro,
                anio=anio_filtro,
                tipo_vendedor='MOSTRADOR'
            ).first()
            
            if not comision_mensual:
                try:
                    from ventas.services.comisiones import calcular_comisiones_mensuales_mostrador
                    comision_mensual = calcular_comisiones_mensuales_mostrador(
                        vendedor, mes_filtro, anio_filtro
                    )
                except Exception as e:
                    logger.error(f"Error al calcular comisiones para {vendedor.username}: {e}", exc_info=True)
                    continue
            
            row += 1
            nombre_vendedor = vendedor.get_full_name() or vendedor.username
            
            ws_resumen.cell(row=row, column=1, value=nombre_vendedor)
            ws_resumen.cell(row=row, column=2, value=float(comision_mensual.total_ventas_mes))
            ws_resumen.cell(row=row, column=3, value=f"{comision_mensual.porcentaje_comision:.2f}%")
            ws_resumen.cell(row=row, column=4, value=float(comision_mensual.bono_extra))
            ws_resumen.cell(row=row, column=5, value=float(comision_mensual.comision_total_pagada))
            ws_resumen.cell(row=row, column=6, value=float(comision_mensual.comision_total_pendiente))
            ws_resumen.cell(row=row, column=7, value=float(comision_mensual.comision_total))
            
            # Aplicar formato
            for col in range(1, 8):
                cell = ws_resumen.cell(row=row, column=col)
                cell.border = border
                if col in [2, 4, 5, 6, 7]:  # Columnas numéricas
                    cell.number_format = '#,##0.00'
                    cell.alignment = Alignment(horizontal='right')
                else:
                    cell.alignment = Alignment(horizontal='left')
            
            # Colorear comisión pagada y pendiente
            ws_resumen.cell(row=row, column=5).fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
            ws_resumen.cell(row=row, column=6).fill = PatternFill(start_color="FFEB9C", end_color="FFEB9C", fill_type="solid")
            ws_resumen.cell(row=row, column=7).fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
            
            # Acumular totales
            total_general_ventas += comision_mensual.total_ventas_mes
            total_general_pagada += comision_mensual.comision_total_pagada
            total_general_pendiente += comision_mensual.comision_total_pendiente
            total_general_comision += comision_mensual.comision_total
            total_general_bono += comision_mensual.bono_extra
            vendedores_con_datos += 1
        
        # Verificar si hay datos
        logger.info(f"Vendedores con datos: {vendedores_con_datos}")
        if vendedores_con_datos == 0:
            messages.warning(request, f"No hay comisiones calculadas para {datetime.datetime(anio_filtro, mes_filtro, 1).strftime('%B %Y')}. Por favor, recalcula las comisiones primero.")
            if viene_de_reporte_comisiones:
                return redirect('reporte_comisiones')
            return redirect('comisiones_mensuales')
        
        # Fila de totales
        row += 1
        ws_resumen.cell(row=row, column=1, value="TOTAL GENERAL").font = Font(bold=True, size=12)
        ws_resumen.cell(row=row, column=2, value=float(total_general_ventas)).font = Font(bold=True)
        ws_resumen.cell(row=row, column=3, value="-").font = Font(bold=True)
        ws_resumen.cell(row=row, column=4, value=float(total_general_bono)).font = Font(bold=True)
        ws_resumen.cell(row=row, column=5, value=float(total_general_pagada)).font = Font(bold=True)
        ws_resumen.cell(row=row, column=6, value=float(total_general_pendiente)).font = Font(bold=True)
        ws_resumen.cell(row=row, column=7, value=float(total_general_comision)).font = Font(bold=True, size=12)
        
        # Aplicar formato a totales
        for col in range(1, 8):
            cell = ws_resumen.cell(row=row, column=col)
            cell.border = border
            cell.fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
            if col in [2, 4, 5, 6, 7]:
                cell.number_format = '#,##0.00'
                cell.alignment = Alignment(horizontal='right')
        
        # Ajustar ancho de columnas (Resumen)
        ws_resumen.column_dimensions['A'].width = 30
        ws_resumen.column_dimensions['B'].width = 18
        ws_resumen.column_dimensions['C'].width = 12
        ws_resumen.column_dimensions['D'].width = 15
        ws_resumen.column_dimensions['E'].width = 18
        ws_resumen.column_dimensions['F'].width = 18
        ws_resumen.column_dimensions['G'].width = 18
        
        # ===== HOJA 2: DETALLE POR VENTA (TODOS LOS VENDEDORES) =====
        ws_detalle = wb.create_sheet("Detalle por Venta")
        
        # Título
        ws_detalle['A1'] = f"Detalle de Comisiones por Venta - Todos los Vendedores"
        ws_detalle['A1'].font = title_font
        ws_detalle.merge_cells('A1:J1')
        
        # Información del período
        ws_detalle['A3'] = f"Período: {datetime.datetime(anio_filtro, mes_filtro, 1).strftime('%B %Y')}"
        ws_detalle['A3'].font = Font(bold=True)
        
        # Encabezados de detalle
        row = 5
        headers_detalle = [
            'Vendedor', 'Venta #', 'Cliente', 'Tipo', 'Monto Base', 
            'Porcentaje', 'Comisión Calculada', 'Estado Pago', 'Comisión Pagada', 'Comisión Pendiente'
        ]
        for col, header in enumerate(headers_detalle, 1):
            cell = ws_detalle.cell(row=row, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border
        
        # Obtener todas las comisiones de ventas ordenadas por vendedor y fecha
        todas_comisiones_ventas = ComisionVenta.objects.filter(
            vendedor__in=vendedores,
            mes=mes_filtro,
            anio=anio_filtro
        ).select_related('vendedor', 'venta', 'venta__cliente').order_by('vendedor__username', 'venta__fecha_creacion')
        
        # Datos de detalle
        for comision_venta in todas_comisiones_ventas:
            row += 1
            venta = comision_venta.venta
            nombre_vendedor = comision_venta.vendedor.get_full_name() or comision_venta.vendedor.username
            
            ws_detalle.cell(row=row, column=1, value=nombre_vendedor)
            ws_detalle.cell(row=row, column=2, value=venta.pk)
            ws_detalle.cell(row=row, column=3, value=venta.cliente.nombre_completo_display)
            ws_detalle.cell(row=row, column=4, value=comision_venta.get_tipo_venta_display())
            ws_detalle.cell(row=row, column=5, value=float(comision_venta.monto_base_comision))
            ws_detalle.cell(row=row, column=6, value=f"{comision_venta.porcentaje_aplicado:.2f}%")
            ws_detalle.cell(row=row, column=7, value=float(comision_venta.comision_calculada))
            ws_detalle.cell(row=row, column=8, value=comision_venta.get_estado_pago_venta_display())
            ws_detalle.cell(row=row, column=9, value=float(comision_venta.comision_pagada))
            ws_detalle.cell(row=row, column=10, value=float(comision_venta.comision_pendiente))
            
            # Aplicar formato
            for col in range(1, 11):
                cell = ws_detalle.cell(row=row, column=col)
                cell.border = border
                if col in [5, 7, 9, 10]:  # Columnas numéricas
                    cell.number_format = '#,##0.00'
                    cell.alignment = Alignment(horizontal='right')
                else:
                    cell.alignment = Alignment(horizontal='left')
        
        # Totales por vendedor en detalle
        row += 2
        ws_detalle.cell(row=row, column=1, value="TOTALES POR VENDEDOR").font = Font(bold=True, size=12)
        ws_detalle.merge_cells(f'A{row}:D{row}')
        
        for vendedor in vendedores:
            comisiones_vendedor = todas_comisiones_ventas.filter(vendedor=vendedor)
            if comisiones_vendedor.exists():
                row += 1
                nombre_vendedor = vendedor.get_full_name() or vendedor.username
                total_pagada_v = sum(c.comision_pagada for c in comisiones_vendedor)
                total_pendiente_v = sum(c.comision_pendiente for c in comisiones_vendedor)
                total_v = total_pagada_v + total_pendiente_v
                
                ws_detalle.cell(row=row, column=1, value=nombre_vendedor).font = Font(bold=True)
                ws_detalle.cell(row=row, column=7, value=float(total_pagada_v + total_pendiente_v)).font = Font(bold=True)
                ws_detalle.cell(row=row, column=9, value=float(total_pagada_v)).font = Font(bold=True)
                ws_detalle.cell(row=row, column=10, value=float(total_pendiente_v)).font = Font(bold=True)
                
                for col in [7, 9, 10]:
                    cell = ws_detalle.cell(row=row, column=col)
                    cell.number_format = '#,##0.00'
                    cell.border = border
        
        # Ajustar ancho de columnas (Detalle)
        ws_detalle.column_dimensions['A'].width = 25
        ws_detalle.column_dimensions['B'].width = 10
        ws_detalle.column_dimensions['C'].width = 30
        ws_detalle.column_dimensions['D'].width = 15
        ws_detalle.column_dimensions['E'].width = 15
        ws_detalle.column_dimensions['F'].width = 12
        ws_detalle.column_dimensions['G'].width = 18
        ws_detalle.column_dimensions['H'].width = 15
        ws_detalle.column_dimensions['I'].width = 18
        ws_detalle.column_dimensions['J'].width = 18
        
        # Preparar respuesta
        try:
            nombre_mes = datetime.datetime(anio_filtro, mes_filtro, 1).strftime('%B_%Y')
            filename = f"comisiones_todos_vendedores_{nombre_mes}.xlsx"
            
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            
            wb.save(response)
            logger.info(f"Excel generado exitosamente: {filename}")
            return response
        except Exception as e:
            logger.error(f"Error al generar Excel de comisiones: {e}", exc_info=True)
            messages.error(request, f"Error al generar el archivo Excel: {str(e)}")
            if viene_de_reporte_comisiones:
                return redirect('reporte_comisiones')
            return redirect('comisiones_mensuales')

# ------------------- 12. REPORTE DE COMISIONES POR VENDEDOR (LEGACY) -------------------

def calcular_comision_por_tipo(total_ventas, tipo_vendedor):
    """
    Calcula la comisión según el tipo de vendedor y el total de ventas.
    
    Sistema de comisiones:
    - MOSTRADOR/ISLA (Ventas Internas): Escalonado
      * $0 - $99,999: 1%
      * $100,000 - $199,999: 2%
      * $200,000 - $299,999: 3%
      * $300,000 - $399,999: 4%
      * $400,000 - $500,000: 5%
      * Más de $500,000: 5% (tope máximo)
    
    - CAMPO (Ventas de Campo): Fijo 4%
    
    Args:
        total_ventas: Decimal - Total de ventas pagadas del vendedor
        tipo_vendedor: str - 'MOSTRADOR', 'ISLA', 'CAMPO', o 'OFICINA'/'CALLE' (legacy)
    
    Returns:
        tuple: (porcentaje_comision, monto_comision)
            porcentaje_comision: Decimal (ej: 0.03 para 3%)
            monto_comision: Decimal (monto calculado)
    """
    # Mapeo: CAMPO usa comisión fija, el resto usa escalonado
    if tipo_vendedor == 'CAMPO' or tipo_vendedor == 'CALLE':
        # Ejecutivos de Ventas de Campo: 4% fijo siempre
        porcentaje = Decimal('0.04')
        return porcentaje, total_ventas * porcentaje
    
    elif tipo_vendedor in ['MOSTRADOR', 'ISLA', 'OFICINA']:
        # Ejecutivos de Ventas Internas (Mostrador/Isla): Sistema escalonado
        if total_ventas < Decimal('100000'):
            porcentaje = Decimal('0.01')  # 1%
        elif total_ventas < Decimal('200000'):
            porcentaje = Decimal('0.02')  # 2%
        elif total_ventas < Decimal('300000'):
            porcentaje = Decimal('0.03')  # 3%
        elif total_ventas < Decimal('400000'):
            porcentaje = Decimal('0.04')  # 4%
        else:  # >= $400,000 (hasta $500,000 y más, máximo 5%)
            porcentaje = Decimal('0.05')  # 5%
        
        return porcentaje, total_ventas * porcentaje
    
    # Fallback: Por defecto 4% si no se identifica (nunca debería llegar aquí)
    porcentaje = Decimal('0.04')
    return porcentaje, total_ventas * porcentaje


class ComisionesVendedoresView(LoginRequiredMixin, TemplateView):
    """
    Muestra un reporte de sueldo fijo y comisiones calculadas.
    - JEFE: Ve a todos los vendedores de la agencia.
    - VENDEDOR: Solo ve su propia información.
    """
    template_name = 'ventas/comisiones_vendedores.html'
    
    # SUELDO BASE POR DEFECTO (si no tiene ejecutivo asociado)
    SUELDO_BASE = Decimal('10000.00') 

    def get_queryset_base(self):
        """Prepara el queryset base de ventas con el total pagado."""
        # OPTIMIZACIÓN N+1: select_related, prefetch y anotación de abonos confirmados
        return VentaViaje.objects.select_related(
            'cliente', 'vendedor', 'vendedor__perfil', 'proveedor'
        ).prefetch_related(
            Prefetch(
                'abonos',
                queryset=AbonoPago.objects.filter(Q(confirmado=True) | Q(forma_pago='EFE'))
            )
        ).annotate(
            # Anotar total de abonos confirmados para cálculos eficientes
            total_abonos_confirmados=Coalesce(
                Sum('abonos__monto', filter=Q(abonos__confirmado=True) | Q(abonos__forma_pago='EFE')),
                Value(Decimal('0.00')),
                output_field=ModelDecimalField()
            )
        )

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        user_rol = perm.get_user_role(self.request.user, self.request)
        vendedores_query = User.objects.filter(perfil__rol='VENDEDOR').order_by('username')
        
        # 1. Determinar qué usuarios ver
        if user_rol == 'VENDEDOR':
            # Si es vendedor, solo se ve a sí mismo
            vendedores_a_mostrar = User.objects.filter(pk=self.request.user.pk)
        elif perm.has_full_access(self.request.user, self.request) or perm.is_contador(self.request.user, self.request):
            # Si es jefe o contador, ve a todos los vendedores (solo lectura para contador)
            vendedores_a_mostrar = vendedores_query
        else:
            # Otros roles no deberían acceder
            vendedores_a_mostrar = User.objects.none() 

        # 2. Obtener todas las ventas del queryset base
        # Nota: No podemos filtrar directamente por total_pagado (es una propiedad de Python)
        # así que obtenemos todas las ventas y luego filtramos en Python
        ventas_base = self.get_queryset_base()

        # Obtener mes y año del filtro ANTES del bucle
        mes_actual = timezone.now().month
        anio_actual = timezone.now().year
        
        # Manejar parámetros vacíos o inválidos
        mes_param = self.request.GET.get('mes', '').strip()
        anio_param = self.request.GET.get('anio', '').strip()
        
        mes_filtro = safe_int(mes_param, default=mes_actual)
        anio_filtro = safe_int(anio_param, default=anio_actual)
        
        # Rango del periodo seleccionado (ventas generadas en este mes/año)
        from datetime import date as date_type
        fecha_inicio = date_type(anio_filtro, mes_filtro, 1)
        if mes_filtro == 12:
            fecha_fin = date_type(anio_filtro + 1, 1, 1)
        else:
            fecha_fin = date_type(anio_filtro, mes_filtro + 1, 1)
        
        lista_comisiones = []
        
        for vendedor in vendedores_a_mostrar:
            # Ventas del vendedor generadas en el periodo (fecha_creacion en mes/año)
            ventas_periodo = ventas_base.filter(
                vendedor=vendedor,
                fecha_creacion__gte=fecha_inicio,
                fecha_creacion__lt=fecha_fin
            )
            
            # Base de comisión: total de ventas generadas en el periodo (no solo pagadas).
            # INT: convertir USD a MXN con tipo_cambio; NAC: costo_venta_final.
            total_ventas_periodo = Decimal('0.00')
            for venta in ventas_periodo:
                if getattr(venta, 'tipo_viaje', 'NAC') == 'INT':
                    total_usd = getattr(venta, 'costo_venta_final_usd', None) or (getattr(venta, 'total_usd', None) if hasattr(venta, 'total_usd') else None)
                    tc = getattr(venta, 'tipo_cambio', None)
                    if total_usd and tc and Decimal(str(tc)) > 0:
                        total_ventas_periodo += (Decimal(str(total_usd)) * Decimal(str(tc))).quantize(Decimal('0.01'))
                else:
                    total_ventas_periodo += (venta.costo_venta_final or Decimal('0.00'))
            
            ejecutivo = getattr(vendedor, 'ejecutivo_asociado', None)

            # Obtener tipo de vendedor (por defecto MOSTRADOR si no tiene ejecutivo asociado)
            tipo_vendedor = ejecutivo.tipo_vendedor if ejecutivo else 'MOSTRADOR'
            
            # Sueldo base (ISLA SÍ tiene sueldo base)
            sueldo_base = ejecutivo.sueldo_base if ejecutivo and ejecutivo.sueldo_base else self.SUELDO_BASE

            # Base de comisión = ventas generadas en el periodo (para tabla y cálculo de %)
            total_ventas_pagadas = total_ventas_periodo
            
            # Para ISLA, las comisiones son 100% manuales (no hay cálculo automático)
            porcentaje_comision = None
            if tipo_vendedor == 'ISLA':
                comision_mensual = ComisionMensual.objects.filter(
                    vendedor=vendedor,
                    mes=mes_filtro,
                    anio=anio_filtro,
                    tipo_vendedor='ISLA'
                ).first()
                
                if comision_mensual and comision_mensual.porcentaje_ajustado_manual:
                    # Usar porcentaje ajustado manualmente por el JEFE
                    porcentaje_comision = comision_mensual.porcentaje_ajustado_manual / Decimal('100')
                    comision_ganada = total_ventas_pagadas * porcentaje_comision
                else:
                    # Si no hay porcentaje ajustado, comisión = 0 (debe ser asignada manualmente)
                    porcentaje_comision = Decimal('0.00')
                    comision_ganada = Decimal('0.00')
            else:
                # Para otros tipos (MOSTRADOR, CAMPO), calcular normalmente
                porcentaje_comision, comision_ganada = calcular_comision_por_tipo(
                    total_ventas_pagadas, 
                    tipo_vendedor
                )
            
            ingreso_total = sueldo_base + comision_ganada

            lista_comisiones.append({
                'vendedor': vendedor,
                'sueldo_base': sueldo_base,
                'comision_porcentaje': porcentaje_comision * 100,  # Para mostrar como porcentaje
                'total_ventas_pagadas': total_ventas_pagadas,
                'comision_ganada': comision_ganada,
                'ingreso_total_estimado': ingreso_total,
                'es_usuario_actual': (vendedor.pk == self.request.user.pk),
                'ejecutivo': ejecutivo,
                'tipo_vendedor': tipo_vendedor,
            })

        context['lista_comisiones'] = lista_comisiones
        context['titulo_reporte'] = "Reporte de Comisiones de Ventas"
        context['user_rol'] = user_rol
        context['mes_filtro'] = mes_filtro
        context['anio_filtro'] = anio_filtro
        
        # Generar fecha_desde para mostrar en el template
        try:
            fecha_desde = date_type(context['anio_filtro'], context['mes_filtro'], 1)
            context['fecha_desde'] = fecha_desde
        except (ValueError, TypeError):
            fecha_desde = date_type(anio_actual, mes_actual, 1)
            context['fecha_desde'] = fecha_desde
        
        return context

    def _generar_username_unico(self, nombre_base):
        base_slug = slugify(nombre_base).replace('-', '')
        base_slug = base_slug or 'ejecutivo'
        base_slug = base_slug[:20]
        username = base_slug
        contador = 1
        while User.objects.filter(username=username).exists():
            sufijo = str(contador)
            username = f"{base_slug[:20-len(sufijo)]}{sufijo}"
            contador += 1
        return username

    def _split_nombre(self, nombre):
        partes = (nombre or '').strip().split()
        if not partes:
            return '', ''
        first = partes[0]
        last = ' '.join(partes[1:]) if len(partes) > 1 else ''
        return first, last

    def _crear_o_actualizar_usuario(self, ejecutivo, tipo_usuario='VENDEDOR', forzar_password=False):
        password_plano = None
        first_name, last_name = self._split_nombre(ejecutivo.nombre_completo)

        if not ejecutivo.usuario:
            username = self._generar_username_unico(ejecutivo.nombre_completo)
            password_plano = secrets.token_urlsafe(10)
            user = User.objects.create_user(
                username=username,
                password=password_plano,
                email=ejecutivo.email or '',
                first_name=first_name,
                last_name=last_name
            )
            ejecutivo.usuario = user
            # SEGURIDAD: No se almacena la contraseña. Usar sistema de reseteo si es necesario.
            ejecutivo.save(update_fields=['usuario'])
        else:
            user = ejecutivo.usuario
            user.email = ejecutivo.email or ''
            user.first_name = first_name
            user.last_name = last_name
            if forzar_password:
                password_plano = secrets.token_urlsafe(10)
                user.set_password(password_plano)
            user.save()
            # SEGURIDAD: No se almacena la contraseña. Usar sistema de reseteo si es necesario.

        # Asegurar el rol seleccionado (VENDEDOR o CONTADOR)
        perfil = getattr(user, 'perfil', None)
        if perfil:
            # Asegurar que el rol sea válido
            rol_final = tipo_usuario if tipo_usuario in ['VENDEDOR', 'CONTADOR'] else 'VENDEDOR'
            if perfil.rol != rol_final:
                perfil.rol = rol_final
            perfil.save(update_fields=['rol'])

        return user, password_plano



class GestionRolesView(LoginRequiredMixin, usuarios_mixins.ManageRolesRequiredMixin, TemplateView):
    """
    Vista para gestionar roles y usuarios del sistema.
    Accesible para JEFE, Director General y Director Administrativo.
    """
    template_name = 'ventas/gestion_roles.html'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['ejecutivos'] = Ejecutivo.objects.all().select_related('usuario', 'usuario__perfil', 'oficina')
        context['ejecutivo_form'] = kwargs.get('ejecutivo_form') or EjecutivoForm()
        context['mostrar_modal_ejecutivo'] = kwargs.get('mostrar_modal_ejecutivo', False)
        context['oficinas'] = Oficina.objects.all().order_by('nombre')
        context['oficina_form'] = kwargs.get('oficina_form') or OficinaForm()
        context['mostrar_modal_oficina'] = kwargs.get('mostrar_modal_oficina', False)
        return context
    
    def _generar_username_unico(self, nombre_base):
        """Genera un username único basado en el nombre."""
        base_slug = slugify(nombre_base).replace('-', '')
        base_slug = base_slug or 'ejecutivo'
        base_slug = base_slug[:20]
        username = base_slug
        contador = 1
        while User.objects.filter(username=username).exists():
            sufijo = str(contador)
            username = f"{base_slug[:20-len(sufijo)]}{sufijo}"
            contador += 1
        return username
    
    def _split_nombre(self, nombre):
        """Divide el nombre completo en first_name y last_name."""
        partes = (nombre or '').strip().split()
        if not partes:
            return '', ''
        first = partes[0]
        last = ' '.join(partes[1:]) if len(partes) > 1 else ''
        return first, last
    
    def _crear_o_actualizar_usuario(self, ejecutivo, tipo_usuario='VENDEDOR', forzar_password=False, nueva_password=None):
        """Crea o actualiza el usuario asociado al ejecutivo."""
        password_plano = None
        first_name, last_name = self._split_nombre(ejecutivo.nombre_completo)
        
        try:
            if not ejecutivo.usuario:
                # Si no hay usuario, crear uno nuevo
                # Verificar que haya email para crear el usuario (email es necesario para crear credenciales)
                if not ejecutivo.email:
                    logger.error(f"No se puede crear usuario para ejecutivo {ejecutivo.pk} sin email")
                    raise ValueError("El correo electrónico es obligatorio para crear el usuario.")
                
                # Verificar que el email no esté en uso por otro usuario
                if User.objects.filter(email__iexact=ejecutivo.email).exists():
                    logger.error(f"El email {ejecutivo.email} ya está en uso por otro usuario")
                    raise ValueError(f"El correo electrónico '{ejecutivo.email}' ya está registrado por otro usuario en el sistema.")
                
                username = self._generar_username_unico(ejecutivo.nombre_completo)
                
                # Verificar que el username no exista (aunque _generar_username_unico debería generar uno único)
                if User.objects.filter(username=username).exists():
                    # Si el username ya existe, generar uno nuevo con contador
                    base_username = username
                    contador = 1
                    while User.objects.filter(username=username).exists():
                        username = f"{base_username[:20-len(str(contador))]}{contador}"
                        contador += 1
                
                password_plano = secrets.token_urlsafe(10)
                try:
                    user = User.objects.create_user(
                        username=username,
                        password=password_plano,
                        email=ejecutivo.email,
                        first_name=first_name,
                        last_name=last_name
                    )
                except Exception as e:
                    error_msg = str(e)
                    # Capturar errores específicos de base de datos
                    if "unique constraint" in error_msg.lower() or "duplicate" in error_msg.lower():
                        if "username" in error_msg.lower():
                            raise ValueError("El nombre de usuario generado ya existe. Intenta nuevamente.")
                        elif "email" in error_msg.lower():
                            raise ValueError(f"El correo electrónico '{ejecutivo.email}' ya está registrado por otro usuario en el sistema.")
                    logger.error(f"Error al crear usuario para ejecutivo {ejecutivo.pk}: {error_msg}", exc_info=True)
                    raise ValueError(f"Error al crear el usuario: {error_msg}")
                
                # Refrescar el usuario desde la base de datos para asegurar que la señal se haya ejecutado
                user.refresh_from_db()
                ejecutivo.usuario = user
                # SEGURIDAD: No se almacena la contraseña. Usar sistema de reseteo si es necesario.
                ejecutivo.save(update_fields=['usuario'])
            else:
                # Si ya existe usuario, actualizarlo
                user = ejecutivo.usuario
                if not user:
                    logger.error(f"Ejecutivo {ejecutivo.pk} tiene referencia a usuario pero el usuario es None")
                    return None, None
                
                # Actualizar email solo si se proporciona y es diferente
                if ejecutivo.email and user.email != ejecutivo.email:
                    user.email = ejecutivo.email
                
                user.first_name = first_name
                user.last_name = last_name
                if forzar_password:
                    # Si se proporciona una nueva contraseña, usarla; si no, generar una aleatoria
                    password_plano = nueva_password if nueva_password else secrets.token_urlsafe(10)
                    user.set_password(password_plano)
                
                try:
                    user.save()
                except Exception as e:
                    logger.error(f"Error al guardar usuario {user.pk}: {str(e)}", exc_info=True)
                    raise
                
                # SEGURIDAD: No se almacena la contraseña. Usar sistema de reseteo si es necesario.
            
            # Verificar que user no sea None antes de continuar
            if not user:
                logger.error(f"Error: user es None después de crear/actualizar para ejecutivo {ejecutivo.pk}")
                return None, None
            
            # Asegurar que el Perfil exista (puede no haberse creado por la señal si hubo algún problema)
            try:
                perfil = user.perfil
            except Perfil.DoesNotExist:
                # Si no existe, crearlo manualmente
                perfil = Perfil.objects.create(user=user, rol='VENDEDOR')
            
            # Asegurar el rol seleccionado
            roles_validos = ['JEFE', 'DIRECTOR_GENERAL', 'DIRECTOR_VENTAS', 'DIRECTOR_ADMINISTRATIVO', 'GERENTE', 'CONTADOR', 'VENDEDOR']
            rol_final = tipo_usuario if tipo_usuario in roles_validos else 'VENDEDOR'
            if perfil.rol != rol_final:
                perfil.rol = rol_final
                perfil.save(update_fields=['rol'])
            
            return user, password_plano
        except Exception as e:
            error_msg = str(e)
            logger.error(f"Error al crear/actualizar usuario para ejecutivo {ejecutivo.pk}: {error_msg}", exc_info=True)
            # Re-raise la excepción para que el código que llama pueda manejarla mejor
            raise
    
    def post(self, request, *args, **kwargs):
        """Maneja las acciones de crear, editar y eliminar ejecutivos y oficinas."""
        # Verificar si es una acción de oficina
        if 'oficina_action' in request.POST:
            return self._handle_oficina_action(request)
        
        # Acción de ejecutivo
        action = request.POST.get('ejecutivo_action', 'crear')
        ejecutivo_id = request.POST.get('ejecutivo_id')
        
        if action == 'eliminar':
            if not ejecutivo_id:
                messages.error(request, "No se pudo identificar al ejecutivo a eliminar.")
                return redirect('gestion_roles')
            ejecutivo = get_object_or_404(Ejecutivo, pk=ejecutivo_id)
            usuario_rel = ejecutivo.usuario
            nombre = ejecutivo.nombre_completo
            ejecutivo.delete()
            if usuario_rel:
                usuario_rel.delete()
            messages.success(request, f"Ejecutivo '{nombre}' eliminado correctamente.")
            return redirect('gestion_roles')
        
        instance = None
        if action == 'editar':
            if not ejecutivo_id:
                messages.error(request, "No se pudo identificar al ejecutivo a editar.")
                return redirect('gestion_roles')
            instance = get_object_or_404(Ejecutivo, pk=ejecutivo_id)
        
        form = EjecutivoForm(request.POST, request.FILES, instance=instance)
        if form.is_valid():
            try:
                # INICIO DEL BLOQUE DE SEGURIDAD - Transacción atómica
                with transaction.atomic():
                    ejecutivo = form.save()  # Ahora esto es provisional dentro de la transacción
                    tipo_usuario = form.cleaned_data.get('tipo_usuario', 'VENDEDOR')
                    
                    # Manejar cambio de contraseña
                    nueva_contrasena = request.POST.get('nueva_contrasena', '').strip()
                    regenerar_password = False
                    if action == 'editar':
                        # En edición, solo regenerar si se proporciona una nueva contraseña
                        regenerar_password = bool(nueva_contrasena)
                    else:
                        # En creación, siempre generar nueva contraseña
                        regenerar_password = True
                    
                    # Intentar crear usuario
                    user, password = self._crear_o_actualizar_usuario(
                        ejecutivo,
                        tipo_usuario=tipo_usuario,
                        forzar_password=regenerar_password,
                        nueva_password=nueva_contrasena if nueva_contrasena else None
                    )
                    
                    # Verificar que el usuario se creó correctamente
                    # Si no se creó, esto forzará el rollback automático del ejecutivo guardado arriba
                    if not user:
                        error_msg = "Error al crear/actualizar el usuario. Por favor, verifica que el correo electrónico sea válido y único."
                        raise ValueError(error_msg)
                # FIN DEL BLOQUE - Si llegamos aquí, todo se confirma (commit)
                
                # Refrescar el ejecutivo desde la base de datos para asegurar que las relaciones estén cargadas
                ejecutivo.refresh_from_db()
                
                if password:
                    # TODO: Implementar envío de contraseña por email en lugar de mostrarla en UI
                    # La contraseña generada es: {password}
                    # Por seguridad, NO se muestra en la interfaz de usuario
                    if action == 'editar':
                        messages.success(
                            request,
                            f"Ejecutivo '{ejecutivo.nombre_completo}' actualizado correctamente. "
                            f"Se ha generado una nueva contraseña. Por favor, comuníquela al usuario de forma segura."
                        )
                    else:
                        messages.success(
                            request,
                            f"Ejecutivo '{ejecutivo.nombre_completo}' agregado correctamente. "
                            f"Se ha generado una contraseña temporal. Por favor, comuníquela al usuario de forma segura."
                        )
                    # Log para administrador (no visible al usuario)
                    logger.info(f"Contraseña generada para ejecutivo {ejecutivo.nombre_completo} (ID: {ejecutivo.pk})")
                else:
                    messages.success(request, f"Ejecutivo '{ejecutivo.nombre_completo}' actualizado correctamente.")
                return redirect('gestion_roles')
                
            except Exception as e:
                # Si ocurre CUALQUIER error dentro del 'with transaction.atomic', 
                # Django deshace automáticamente el form.save() inicial.
                # La BD queda limpia como si nada hubiera pasado.
                
                error_msg = str(e)
                # Simplificar el mensaje si no es un mensaje específico ya formateado
                if "Error al guardar el ejecutivo" not in error_msg and not any(keyword in error_msg.lower() for keyword in ["correo", "email", "usuario", "username"]):
                    error_msg = f"Error al guardar el ejecutivo: {str(e)}"
                
                messages.error(request, error_msg)
                import traceback
                logger.error(f"Error al crear ejecutivo: {traceback.format_exc()}")
                
                # Mantener el formulario con errores para que se muestren
                # El formulario original tiene los datos del POST, así que se mantendrán
                context = self.get_context_data()
                context['ejecutivo_form'] = form  # El formulario ya tiene los datos del POST
                context['mostrar_modal_ejecutivo'] = True
                # Pasar el ejecutivo_id para que el modal se abra en modo edición (si es edición)
                if action == 'editar' and ejecutivo_id:
                    context['ejecutivo_id_editar'] = ejecutivo_id
                return self.render_to_response(context)
        else:
            # Si el formulario no es válido, mostrar los errores
            error_messages = []
            for field, errors in form.errors.items():
                for error in errors:
                    error_messages.append(f"{field}: {error}")
            if error_messages:
                messages.error(request, "Por favor, corrige los siguientes errores: " + " | ".join(error_messages))
        
        # Al renderizar el formulario con errores, asegurarse de que se pasa con los datos del POST
        # para que el usuario no tenga que volver a llenar los campos
        context = self.get_context_data(ejecutivo_form=form, mostrar_modal_ejecutivo=True)
        # Si estamos creando un ejecutivo que ya se guardó pero falló el usuario, pasar su ID para editarlo
        if action == 'crear' and 'ejecutivo' in locals() and ejecutivo.pk:
            context['ejecutivo_id_editar'] = ejecutivo.pk
        elif action == 'editar' and ejecutivo_id:
            context['ejecutivo_id_editar'] = ejecutivo_id
        return self.render_to_response(context)
    
    def _handle_oficina_action(self, request):
        """Maneja las acciones de crear, editar, deshabilitar y habilitar oficinas."""
        action = request.POST.get('oficina_action', 'crear')
        oficina_id = request.POST.get('oficina_id')
        
        if action == 'deshabilitar':
            if not oficina_id:
                messages.error(request, "No se pudo identificar la oficina a deshabilitar.")
                return redirect('gestion_roles')
            oficina = get_object_or_404(Oficina, pk=oficina_id)
            nombre = oficina.nombre
            oficina.activa = False
            oficina.save()
            messages.warning(request, f"Oficina '{nombre}' deshabilitada correctamente.")
            return redirect('gestion_roles')
        
        if action == 'habilitar':
            if not oficina_id:
                messages.error(request, "No se pudo identificar la oficina a habilitar.")
                return redirect('gestion_roles')
            oficina = get_object_or_404(Oficina, pk=oficina_id)
            nombre = oficina.nombre
            oficina.activa = True
            oficina.save()
            messages.success(request, f"Oficina '{nombre}' habilitada correctamente.")
            return redirect('gestion_roles')
        
        instance = None
        if action == 'editar':
            if not oficina_id:
                messages.error(request, "No se pudo identificar la oficina a editar.")
                return redirect('gestion_roles')
            instance = get_object_or_404(Oficina, pk=oficina_id)
        
        form = OficinaForm(request.POST, instance=instance)
        if form.is_valid():
            try:
                oficina = form.save()
                if action == 'editar':
                    messages.success(request, f"Oficina '{oficina.nombre}' actualizada correctamente.")
                else:
                    messages.success(request, f"Oficina '{oficina.nombre}' creada correctamente.")
                return redirect('gestion_roles')
            except Exception as e:
                messages.error(request, f"Error al guardar la oficina: {str(e)}")
        else:
            messages.error(request, "Por favor, corrige los errores en el formulario.")
            # Retornar con el formulario con errores
            context = self.get_context_data()
            context['oficina_form'] = form
            context['mostrar_modal_oficina'] = True
            return self.render_to_response(context)
        
        return redirect('gestion_roles')


class EjecutivoDetailView(LoginRequiredMixin, UserPassesTestMixin, DetailView):
    """
    Vista para ver los detalles completos de un ejecutivo.
    Solo accesible para JEFE.
    """
    model = Ejecutivo
    template_name = 'ventas/ejecutivo_detail.html'
    context_object_name = 'ejecutivo'
    
    def test_func(self):
        """Solo JEFE puede ver detalles de ejecutivos."""
        user_rol = perm.get_user_role(self.request.user, self.request)
        return perm.has_full_access(self.request.user)
    
    def handle_no_permission(self):
        messages.error(self.request, "No tienes permiso para ver los detalles del ejecutivo. Solo el JEFE puede acceder.")
        return redirect('gestion_roles')
    
    def post(self, request, *args, **kwargs):
        """Maneja las acciones desde la vista de detalle: eliminar, cambiar contraseña, cambiar sueldo, desactivar/activar."""
        ejecutivo = self.get_object()
        action = request.POST.get('action')
        
        if action == 'cambiar_contrasena':
            nueva_contrasena = request.POST.get('nueva_contrasena', '').strip()
            if nueva_contrasena and ejecutivo.usuario:
                ejecutivo.usuario.set_password(nueva_contrasena)
                ejecutivo.usuario.save()
                # SEGURIDAD: No se muestra la contraseña en la UI
                messages.success(request, "Contraseña actualizada correctamente. El usuario debe usar el sistema de recuperación de contraseña si la olvida.")
                return redirect(reverse('ejecutivo_detail', kwargs={'pk': ejecutivo.pk}) + '?contrasena_actualizada=1')
            else:
                messages.error(request, "No se pudo actualizar la contraseña. Verifica que el ejecutivo tenga un usuario asociado.")
            return redirect('ejecutivo_detail', pk=ejecutivo.pk)
        
        elif action == 'desactivar' or action == 'activar':
            # Desactivar o activar el usuario
            if not ejecutivo.usuario:
                messages.error(request, "El ejecutivo no tiene un usuario asociado.")
                return redirect('ejecutivo_detail', pk=ejecutivo.pk)
            
            if action == 'desactivar':
                ejecutivo.usuario.is_active = False
                ejecutivo.usuario.save()
                messages.warning(request, f"Usuario '{ejecutivo.usuario.username}' desactivado correctamente. No podrá iniciar sesión hasta que sea reactivado.")
            else:  # activar
                ejecutivo.usuario.is_active = True
                ejecutivo.usuario.save()
                messages.success(request, f"Usuario '{ejecutivo.usuario.username}' activado correctamente. Ya puede iniciar sesión.")
            
            return redirect('ejecutivo_detail', pk=ejecutivo.pk)
        
        elif action == 'eliminar' or not action:
            # Eliminación del ejecutivo
            usuario_rel = ejecutivo.usuario
            nombre = ejecutivo.nombre_completo
            ejecutivo.delete()
            if usuario_rel:
                usuario_rel.delete()
            messages.success(request, f"Ejecutivo '{nombre}' eliminado correctamente.")
            return redirect('gestion_roles')
        
        return redirect('ejecutivo_detail', pk=ejecutivo.pk)


class ProveedorListCreateView(LoginRequiredMixin, usuarios_mixins.ManageSuppliersRequiredMixin, TemplateView):
    """
    Gestiona el catálogo de proveedores. Accesible para JEFE, Director General y Director Administrativo.
    """
    template_name = 'ventas/proveedores.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        proveedores = Proveedor.objects.all()
        proveedores_por_servicio = {clave: [] for clave, _ in Proveedor.SERVICIO_CHOICES}

        # ✅ Agrupar proveedores por cada servicio que ofrecen (pueden estar en múltiples grupos)
        for proveedor in proveedores:
            if proveedor.servicios:
                servicios_list = [s.strip() for s in proveedor.servicios.split(',') if s.strip()]
                for servicio_codigo in servicios_list:
                    if servicio_codigo in proveedores_por_servicio:
                        proveedores_por_servicio[servicio_codigo].append(proveedor)
            # Si no tiene servicios definidos, no se agrupa

        grupos_proveedores = [
            {
                'clave': clave,
                'label': label,
                'proveedores': sorted(
                    proveedores_por_servicio.get(clave, []),
                    key=lambda p: p.nombre.lower()
                ),
            }
            for clave, label in Proveedor.SERVICIO_CHOICES
        ]

        context['grupos_proveedores'] = grupos_proveedores
        context['total_proveedores'] = len(proveedores)  # Total único de proveedores
        context['form'] = kwargs.get('form') or ProveedorForm()
        context['servicio_choices'] = Proveedor.SERVICIO_CHOICES
        return context

    def post(self, request, *args, **kwargs):
        form = ProveedorForm(request.POST)
        if form.is_valid():
            proveedor = form.save()
            messages.success(request, f"Proveedor '{proveedor.nombre}' agregado correctamente.")
            return redirect('proveedores')
        
        # Si el formulario no es válido, mostrar errores
        messages.error(request, "Por favor, corrige los errores en el formulario.")
        context = self.get_context_data(form=form)
        context['form'] = form
        return self.render_to_response(context)

# ------------------- VISTAS PARA NOTIFICACIONES -------------------

class ConfirmarPagoView(LoginRequiredMixin, UserPassesTestMixin, View):
    """Vista para que el CONTADOR confirme pagos pendientes."""
    
    def test_func(self):
        """Solo CONTADOR puede confirmar pagos."""
        return perm.is_contador(self.request.user, self.request)
    
    def handle_no_permission(self):
        messages.error(self.request, "No tienes permiso para confirmar pagos.")
        return redirect('dashboard')
    
    def post(self, request, notificacion_id):
        from django.utils import timezone
        
        try:
            notificacion = Notificacion.objects.get(
                pk=notificacion_id,
                usuario=request.user,
                tipo='PAGO_PENDIENTE',
                confirmado=False
            )
            
            venta = notificacion.venta
            abono = notificacion.abono
            
            # Confirmar el abono (si existe): requiere comprobante subido
            if abono:
                if abono.forma_pago in ['TRN', 'TAR', 'DEP'] and not abono.comprobante_subido:
                    messages.error(request, "No se puede confirmar este abono: debe estar subido el comprobante.")
                    return redirect(reverse('pagos_por_confirmar'))
                abono.confirmado = True
                abono.confirmado_por = request.user
                abono.confirmado_en = timezone.now()
                abono.save(update_fields=['confirmado', 'confirmado_por', 'confirmado_en'])
            
            # Actualizar estado de la venta
            # Refrescar la venta desde la BD para obtener los datos actualizados
            venta.refresh_from_db()
            
            if abono:
                # Verificar si la venta está completada (liquidada)
                # Ahora total_pagado incluirá el abono confirmado
                # IMPORTANTE: Usar costo_total_con_modificacion, no solo costo_venta_final
                nuevo_total = venta.total_pagado
                venta_liquidada = nuevo_total >= venta.costo_total_con_modificacion
                if venta_liquidada:
                    venta.estado_confirmacion = 'COMPLETADO'
                    # Crear notificación de liquidación para el VENDEDOR
                    if venta.vendedor:
                        mensaje_liquidacion = f"¡Venta #{venta.pk} completamente liquidada! - Cliente: {venta.cliente.nombre_completo_display} - Total: ${venta.costo_venta_final:,.2f}"
                        Notificacion.objects.create(
                            usuario=venta.vendedor,
                            tipo='LIQUIDACION',
                            mensaje=mensaje_liquidacion,
                            venta=venta,
                            confirmado=False
                        )
                else:
                    venta.estado_confirmacion = 'PENDIENTE'
            else:
                # Si no hay abono, es una apertura - confirmarla solo si hay comprobante subido
                if venta.modo_pago_apertura in ['TRN', 'TAR', 'DEP'] and not venta.comprobante_apertura_subido:
                    messages.error(request, "No se puede confirmar la apertura: debe estar subido el comprobante de apertura.")
                    return redirect(reverse('pagos_por_confirmar'))
                venta.apertura_confirmada = True
                venta.estado_confirmacion = 'COMPLETADO'
                # No hacer refresh_from_db() aquí: borraría apertura_confirmada=True en memoria
                # y al guardar se persistiría False. total_pagado usa venta.apertura_confirmada.
                venta_liquidada = venta.total_pagado >= venta.costo_total_con_modificacion
                if venta_liquidada:
                    venta.estado_confirmacion = 'COMPLETADO'
                    if venta.vendedor:
                        mensaje_liquidacion = f"¡Venta #{venta.pk} completamente liquidada! - Cliente: {venta.cliente.nombre_completo_display} - Total: ${venta.costo_venta_final:,.2f}"
                        Notificacion.objects.create(
                            usuario=venta.vendedor,
                            tipo='LIQUIDACION',
                            mensaje=mensaje_liquidacion,
                            venta=venta,
                            confirmado=False
                        )
                else:
                    venta.estado_confirmacion = 'PENDIENTE'
            
            update_fields = ['estado_confirmacion']
            if not abono:
                update_fields.append('apertura_confirmada')
            venta.save(update_fields=update_fields)
            
            # Actualizar notificación del CONTADOR: cambiar de PAGO_PENDIENTE a PAGO_CONFIRMADO
            notificacion.tipo = 'PAGO_CONFIRMADO'
            mensaje_contador_actualizado = f"✅ Pago confirmado: ${abono.monto if abono else venta.cantidad_apertura:,.2f} - Venta #{venta.pk}"
            notificacion.mensaje = mensaje_contador_actualizado
            notificacion.confirmado = True
            notificacion.confirmado_por = request.user
            notificacion.confirmado_en = timezone.now()
            notificacion.marcar_como_vista()
            notificacion.save(update_fields=['tipo', 'mensaje', 'confirmado', 'confirmado_por', 'confirmado_en', 'vista', 'fecha_vista'])
            
            # Crear notificación para el VENDEDOR (solo si NO es JEFE para evitar duplicados)
            if venta.vendedor:
                # Verificar si el vendedor es JEFE
                vendedor_es_jefe = venta.vendedor.perfil.rol == 'JEFE' if hasattr(venta.vendedor, 'perfil') else False
                
                if not vendedor_es_jefe:
                    # Solo crear notificación para vendedor si NO es JEFE
                    mensaje_vendedor = f"✅ Tu pago ha sido confirmado por el contador. Puedes proceder con la venta #{venta.pk}"
                    Notificacion.objects.create(
                        usuario=venta.vendedor,
                        tipo='PAGO_CONFIRMADO',
                        mensaje=mensaje_vendedor,
                        venta=venta,
                        abono=abono,
                        confirmado=True
                    )
            
            # Actualizar notificaciones existentes del JEFE en lugar de crear nuevas
            jefes = User.objects.filter(perfil__rol='JEFE')
            mensaje_jefe_actualizado = f"✅ Pago confirmado por el contador: ${abono.monto if abono else venta.cantidad_apertura:,.2f} - Venta #{venta.pk}"
            
            for jefe in jefes:
                # Buscar la notificación pendiente existente
                notificacion_jefe = Notificacion.objects.filter(
                    usuario=jefe,
                    venta=venta,
                    tipo='PAGO_PENDIENTE',
                    confirmado=False
                ).first()
                
                if notificacion_jefe and abono:
                    # Si existe una notificación pendiente del mismo abono, actualizarla
                    notificacion_jefe_abono = Notificacion.objects.filter(
                        usuario=jefe,
                        venta=venta,
                        abono=abono,
                        tipo='PAGO_PENDIENTE',
                        confirmado=False
                    ).first()
                    
                    if notificacion_jefe_abono:
                        # Actualizar la notificación existente
                        notificacion_jefe_abono.tipo = 'PAGO_CONFIRMADO'
                        notificacion_jefe_abono.mensaje = mensaje_jefe_actualizado
                        notificacion_jefe_abono.confirmado = True
                        notificacion_jefe_abono.confirmado_por = request.user
                        notificacion_jefe_abono.confirmado_en = timezone.now()
                        notificacion_jefe_abono.save(update_fields=['tipo', 'mensaje', 'confirmado', 'confirmado_por', 'confirmado_en'])
                    elif notificacion_jefe:
                        # Si no hay una específica del abono, actualizar la genérica
                        notificacion_jefe.tipo = 'PAGO_CONFIRMADO'
                        notificacion_jefe.mensaje = mensaje_jefe_actualizado
                        notificacion_jefe.confirmado = True
                        notificacion_jefe.confirmado_por = request.user
                        notificacion_jefe.confirmado_en = timezone.now()
                        notificacion_jefe.save(update_fields=['tipo', 'mensaje', 'confirmado', 'confirmado_por', 'confirmado_en'])
                    else:
                        # Si no existe ninguna, crear una nueva (caso especial)
                        Notificacion.objects.create(
                            usuario=jefe,
                            tipo='PAGO_CONFIRMADO',
                            mensaje=mensaje_jefe_actualizado,
                            venta=venta,
                            abono=abono,
                            confirmado=True
                        )
                elif notificacion_jefe:
                    # Si existe una notificación pero no hay abono específico (apertura)
                    notificacion_jefe.tipo = 'PAGO_CONFIRMADO'
                    notificacion_jefe.mensaje = mensaje_jefe_actualizado
                    notificacion_jefe.confirmado = True
                    notificacion_jefe.confirmado_por = request.user
                    notificacion_jefe.confirmado_en = timezone.now()
                    notificacion_jefe.save(update_fields=['tipo', 'mensaje', 'confirmado', 'confirmado_por', 'confirmado_en'])
                else:
                    # Si no existe ninguna notificación previa, crear una nueva (caso raro)
                    Notificacion.objects.create(
                        usuario=jefe,
                        tipo='PAGO_CONFIRMADO',
                        mensaje=mensaje_jefe_actualizado,
                        venta=venta,
                        abono=abono,
                        confirmado=True
                    )
            
            messages.success(request, "Pago confirmado exitosamente. ✅")
            # Redirigir al detalle de la venta en la pestaña de abonos para que el contador vea el pago confirmado
            return redirect(reverse('detalle_venta', kwargs={'pk': venta.pk, 'slug': venta.slug_safe}) + '?tab=abonos')
            
        except Notificacion.DoesNotExist:
            messages.error(request, "No se encontró la notificación o ya fue confirmada.")
            return redirect('dashboard')
        except Exception as e:
            messages.error(request, f"Error al confirmar el pago: {str(e)}")
            # Intentar redirigir a la venta si está disponible en el contexto
            try:
                if 'venta' in locals() and venta:
                    return redirect('detalle_venta', pk=venta.pk, slug=venta.slug_safe)
            except:
                pass
            return redirect('dashboard')


class ConfirmarAbonoView(LoginRequiredMixin, UserPassesTestMixin, View):
    """Vista para que el CONTADOR confirme abonos pendientes directamente desde el detalle de venta."""
    
    def test_func(self):
        """Solo CONTADOR puede confirmar abonos."""
        return perm.is_contador(self.request.user, self.request)
    
    def handle_no_permission(self):
        messages.error(self.request, "No tienes permiso para confirmar abonos.")
        return redirect('dashboard')
    
    def post(self, request, abono_id):
        from django.utils import timezone
        
        try:
            abono = get_object_or_404(AbonoPago, pk=abono_id)
            
            # Verificar que el abono está pendiente de confirmación
            if abono.confirmado:
                messages.warning(request, "Este abono ya fue confirmado anteriormente.")
                return redirect(reverse('detalle_venta', kwargs={'pk': abono.venta.pk, 'slug': abono.venta.slug_safe}) + '?tab=abonos')
            
            # Verificar que el abono es de tipo Transferencia/Tarjeta/Depósito (requiere confirmación)
            if abono.forma_pago not in ['TRN', 'TAR', 'DEP']:
                messages.warning(request, "Solo se pueden confirmar abonos de Transferencia, Tarjeta o Depósito.")
                return redirect(reverse('detalle_venta', kwargs={'pk': abono.venta.pk, 'slug': abono.venta.slug_safe}) + '?tab=abonos')
            
            # Requiere comprobante subido para que el contador pueda confirmar
            if not abono.comprobante_subido:
                messages.error(request, "No se puede confirmar este abono: debe estar subido el comprobante.")
                return redirect(reverse('detalle_venta', kwargs={'pk': abono.venta.pk, 'slug': abono.venta.slug_safe}) + '?tab=abonos')
            
            venta = abono.venta
            
            # Confirmar el abono
            abono.confirmado = True
            abono.confirmado_por = request.user
            abono.confirmado_en = timezone.now()
            abono.save(update_fields=['confirmado', 'confirmado_por', 'confirmado_en'])
            
            # Actualizar estado de la venta
            venta.refresh_from_db()
            nuevo_total = venta.total_pagado
            if nuevo_total >= venta.costo_venta_final:
                venta.estado_confirmacion = 'COMPLETADO'
            else:
                venta.estado_confirmacion = 'PENDIENTE'
            venta.save(update_fields=['estado_confirmacion'])
            
            # Actualizar SOLO las notificaciones relacionadas con ESTE abono específico del CONTADOR
            # IMPORTANTE: Solo actualizar notificaciones que están relacionadas con este abono específico
            notificaciones_contador = Notificacion.objects.filter(
                usuario=request.user,
                abono=abono,
                tipo='PAGO_PENDIENTE',
                confirmado=False
            )
            mensaje_contador_actualizado = f"✅ Abono confirmado: ${abono.monto:,.2f} ({abono.get_forma_pago_display()}) - Venta #{venta.pk}"
            for notificacion in notificaciones_contador:
                notificacion.tipo = 'PAGO_CONFIRMADO'
                notificacion.mensaje = mensaje_contador_actualizado
                notificacion.confirmado = True
                notificacion.confirmado_por = request.user
                notificacion.confirmado_en = timezone.now()
                notificacion.marcar_como_vista()
                notificacion.save(update_fields=['tipo', 'mensaje', 'confirmado', 'confirmado_por', 'confirmado_en', 'vista', 'fecha_vista'])
            
            # Crear notificación para el VENDEDOR (solo si NO es JEFE para evitar duplicados)
            if venta.vendedor:
                # Verificar si el vendedor es JEFE
                vendedor_es_jefe = venta.vendedor.perfil.rol == 'JEFE' if hasattr(venta.vendedor, 'perfil') else False
                
                if not vendedor_es_jefe:
                    # Solo crear notificación para vendedor si NO es JEFE (INT: mostrar USD)
                    if venta.tipo_viaje == 'INT' and abono.monto_usd_para_display is not None:
                        monto_texto = f"USD ${abono.monto_usd_para_display:,.2f}"
                    else:
                        monto_texto = f"${abono.monto:,.2f}"
                    mensaje_vendedor = f"✅ Tu abono de {monto_texto} ha sido confirmado por el contador. Venta #{venta.pk}"
                    Notificacion.objects.create(
                        usuario=venta.vendedor,
                        tipo='PAGO_CONFIRMADO',
                        mensaje=mensaje_vendedor,
                        venta=venta,
                        abono=abono,
                        confirmado=True
                    )
            
            # Actualizar notificaciones existentes del JEFE en lugar de crear nuevas (INT: mostrar USD)
            jefes = User.objects.filter(perfil__rol='JEFE')
            if venta.tipo_viaje == 'INT' and abono.monto_usd_para_display is not None:
                monto_texto_jefe = f"USD ${abono.monto_usd_para_display:,.2f}"
            else:
                monto_texto_jefe = f"${abono.monto:,.2f}"
            mensaje_jefe_actualizado = f"✅ Abono confirmado por el contador: {monto_texto_jefe} ({abono.get_forma_pago_display()}) - Venta #{venta.pk}"
            
            for jefe in jefes:
                # Buscar la notificación pendiente existente del mismo abono
                notificacion_jefe_abono = Notificacion.objects.filter(
                    usuario=jefe,
                    venta=venta,
                    abono=abono,
                    tipo='PAGO_PENDIENTE',
                    confirmado=False
                ).first()
                
                if notificacion_jefe_abono:
                    # Actualizar la notificación existente
                    notificacion_jefe_abono.tipo = 'PAGO_CONFIRMADO'
                    notificacion_jefe_abono.mensaje = mensaje_jefe_actualizado
                    notificacion_jefe_abono.confirmado = True
                    notificacion_jefe_abono.confirmado_por = request.user
                    notificacion_jefe_abono.confirmado_en = timezone.now()
                    notificacion_jefe_abono.save(update_fields=['tipo', 'mensaje', 'confirmado', 'confirmado_por', 'confirmado_en'])
                else:
                    # Si no existe, crear una nueva (caso raro donde no se creó la inicial)
                    Notificacion.objects.create(
                        usuario=jefe,
                        tipo='PAGO_CONFIRMADO',
                        mensaje=mensaje_jefe_actualizado,
                        venta=venta,
                        abono=abono,
                        confirmado=True
                    )
            
            if venta.tipo_viaje == 'INT' and abono.monto_usd_para_display is not None:
                messages.success(request, f"✅ Abono de USD ${abono.monto_usd_para_display:,.2f} confirmado exitosamente.")
            else:
                messages.success(request, f"✅ Abono de ${abono.monto:,.2f} confirmado exitosamente.")
            return redirect(reverse('detalle_venta', kwargs={'pk': venta.pk, 'slug': venta.slug_safe}) + '?tab=abonos')
            
        except AbonoPago.DoesNotExist:
            messages.error(request, "El abono no existe.")
            return redirect('dashboard')
        except Exception as e:
            messages.error(request, f"Error al confirmar el abono: {str(e)}")
            return redirect('dashboard')


class MarcarNotificacionVistaView(LoginRequiredMixin, View):
    """Vista AJAX para marcar una notificación como vista."""
    
    def post(self, request, pk):
        import logging
        logger = logging.getLogger(__name__)
        
        try:
            logger.info(f"Intentando marcar notificación {pk} como vista para usuario {request.user.username}")
            
            # Obtener la notificación sin restricción de usuario
            # Cualquier usuario autenticado puede marcar cualquier notificación como vista
            notificacion = Notificacion.objects.get(pk=pk)
            logger.info(f"Notificación encontrada: {notificacion.pk}, usuario original: {notificacion.usuario.username}, vista: {notificacion.vista}")
            
            # Marcar como vista
            notificacion.marcar_como_vista()
            logger.info(f"Notificación {pk} marcada como vista exitosamente por usuario {request.user.username}")
            
            return JsonResponse({'success': True, 'message': 'Notificación marcada como vista'})
        except Notificacion.DoesNotExist:
            logger.warning(f"Notificación {pk} no existe en la base de datos")
            return JsonResponse({
                'success': False, 
                'message': 'Notificación no encontrada. Puede que ya haya sido eliminada.'
            }, status=404)
        except Exception as e:
            logger.error(f"Error al marcar notificación {pk} como vista: {str(e)}", exc_info=True)
            return JsonResponse({'success': False, 'message': f'Error del servidor: {str(e)}'}, status=500)


class EliminarNotificacionView(LoginRequiredMixin, View):
    """Vista AJAX para eliminar una notificación."""
    
    def post(self, request, pk):
        try:
            notificacion = Notificacion.objects.get(pk=pk, usuario=request.user)
            notificacion.delete()
            return JsonResponse({'success': True, 'message': 'Notificación eliminada'})
        except Notificacion.DoesNotExist:
            return JsonResponse({'success': False, 'message': 'Notificación no encontrada'}, status=404)
        except Exception as e:
            return JsonResponse({'success': False, 'message': str(e)}, status=500)


class IncrementarCotizacionClienteView(LoginRequiredMixin, View):
    """Incrementa el contador de cotizaciones generadas por un cliente asociado a una venta."""

    def post(self, request, slug, pk, *args, **kwargs):
        try:
            venta = get_object_or_404(VentaViaje, slug=slug, pk=pk)
            
            # SEGURIDAD: Validar acceso a la venta (prevenir IDOR)
            if not self._user_can_access_venta(request.user, venta):
                return JsonResponse({
                    'success': False, 
                    'message': 'No tienes permiso para modificar esta venta.'
                }, status=403)
            
            cliente = venta.cliente
            Cliente.objects.filter(pk=cliente.pk).update(
                cotizaciones_generadas=F('cotizaciones_generadas') + 1
            )
            cliente.refresh_from_db(fields=['cotizaciones_generadas'])
            return JsonResponse({'success': True, 'total': cliente.cotizaciones_generadas})
        except VentaViaje.DoesNotExist:
            return JsonResponse({'success': False, 'message': 'Venta no encontrada'}, status=404)
        except Exception as exc:
            logger.exception("Error incrementando cotizaciones para la venta %s", pk)
            return JsonResponse({'success': False, 'message': str(exc)}, status=500)

    def _user_can_access_venta(self, user, venta):
        """
        SEGURIDAD: Verifica si el usuario tiene permiso para acceder/modificar la venta.
        - JEFE y ADMIN pueden acceder a cualquier venta
        - VENDEDOR solo puede acceder a sus propias ventas
        """
        if not hasattr(user, 'perfil'):
            return False
        
        rol = user.perfil.rol
        
        # JEFE y ADMIN tienen acceso total
        if rol in ['JEFE', 'ADMIN']:
            return True
        
        # VENDEDOR: solo sus propias ventas
        if rol == 'VENDEDOR' and hasattr(user.perfil, 'ejecutivo'):
            return venta.vendedor == user.perfil.ejecutivo
        
        return False


class GenerarCotizacionDocxView(LoginRequiredMixin, View):
    """Genera una cotización en formato DOCX usando la plantilla con membrete."""

    def post(self, request, *args, **kwargs):
        try:
            payload = json.loads(request.body.decode('utf-8'))
        except json.JSONDecodeError:
            return JsonResponse({'success': False, 'message': 'Formato JSON inválido.'}, status=400)

        cotizacion = payload.get('cotizacion') or {}
        general = cotizacion.get('general') or {}
        tipo = cotizacion.get('tipo')

        if not tipo:
            return JsonResponse({'success': False, 'message': 'El tipo de cotización es obligatorio.'}, status=400)

        venta = None
        cliente_nombre = general.get('cliente') or ''
        venta_id = payload.get('venta_id')

        if venta_id:
            venta = VentaViaje.objects.filter(pk=venta_id).select_related('cliente').first()
            if venta:
                cliente_nombre = cliente_nombre or venta.cliente.nombre_completo_display
                Cliente.objects.filter(pk=venta.cliente.pk).update(
                    cotizaciones_generadas=F('cotizaciones_generadas') + 1
                )

        general['cliente'] = cliente_nombre

        template_path = os.path.join(settings.BASE_DIR, 'static', 'docx', 'membrete.docx')
        if not os.path.exists(template_path):
            return JsonResponse({'success': False, 'message': 'La plantilla DOCX no fue encontrada en static/docx/membrete.docx.'}, status=500)

        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.oxml.ns import qn
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return JsonResponse({'success': False, 'message': 'python-docx no está instalado en el entorno.'}, status=500)

        def format_date(value):
            if not value:
                return '-'
            try:
                if isinstance(value, datetime.date):
                    parsed = value
                else:
                    parsed = datetime.date.fromisoformat(str(value))
                return parsed.strftime('%d/%m/%Y')
            except Exception:
                return str(value)

        def format_currency(value):
            if value in (None, '', 0):
                return '0.00'
            try:
                number = Decimal(str(value).replace(',', ''))
            except Exception:
                return str(value)
            return f"{number:,.2f}"

        doc = Document(template_path)

        section = doc.sections[0]

        MOVUMS_BLUE = RGBColor(15, 92, 192)
        MOVUMS_LIGHT_BLUE = RGBColor(92, 141, 214)
        TEXT_COLOR = RGBColor(20, 20, 20)

        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        def set_run_font(run, size=12, bold=False, color=TEXT_COLOR):
            run.font.name = 'Arial'
            run.font.size = Pt(size)
            run.bold = bold
            run.font.color.rgb = color

        def add_paragraph(doc_obj, text='', size=12, bold=False, color=TEXT_COLOR, space_before=0, space_after=0):
            paragraph = doc_obj.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(space_before)
            paragraph.paragraph_format.space_after = Pt(space_after)
            run = paragraph.add_run(text)
            set_run_font(run, size=size, bold=bold, color=color)
            return paragraph

        # Respetar el membretado existente de la plantilla desplazando el contenido hacia abajo
        # sin párrafos extra: respetamos el margen del membrete pero iniciamos lo antes posible

        fecha_paragraph = doc.add_paragraph()
        fecha_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fecha_paragraph.paragraph_format.space_after = Pt(6)
        fecha_run = fecha_paragraph.add_run(f"Fecha de Cotización: {format_date(general.get('fechaCotizacion'))}")
        set_run_font(fecha_run, size=14, bold=True, color=MOVUMS_BLUE)

        info_table = doc.add_table(rows=4, cols=3)
        info_data = [
            ("Origen / Destino", general.get('origen') or '-', general.get('destino') or '-'),
            ("Inicio / Fin", format_date(general.get('fechaInicio')), format_date(general.get('fechaFin'))),
            ("Pasajeros", general.get('pasajeros') or '1', f"{general.get('adultos') or 0} Adultos / {general.get('menores') or 0} Menores"),
            ("Viaje", f"{general.get('dias') or '-'} días", f"{general.get('noches') or '-'} noches"),
        ]
        for row_idx, (label, v1, v2) in enumerate(info_data):
            row = info_table.rows[row_idx].cells
            label_run = row[0].paragraphs[0].add_run(label)
            set_run_font(label_run, size=14, bold=True, color=MOVUMS_BLUE)
            val1_run = row[1].paragraphs[0].add_run(v1)
            set_run_font(val1_run, size=12)
            val2_run = row[2].paragraphs[0].add_run(v2)
            set_run_font(val2_run, size=12)

        add_paragraph(doc, "", space_after=6)
        
        # Salto de línea entre la información general y el contenido de la cotización
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

        # Funciones helper para el nuevo formato (similar a confirmaciones)
        MOVUMS_BLUE_CORP = RGBColor(0, 74, 142)  # Color corporativo #004a8e
        
        def agregar_subtitulo_con_vineta(texto):
            """Agrega un subtítulo con viñeta azul."""
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            bullet_run = p.add_run('• ')
            set_run_font(bullet_run, size=14, bold=True, color=MOVUMS_BLUE_CORP)
            texto_run = p.add_run(texto)
            set_run_font(texto_run, size=14, bold=True, color=MOVUMS_BLUE_CORP)
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(2)
            return p
        
        def agregar_info_line(etiqueta, valor):
            """Agrega una línea de información."""
            if not valor:
                return
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            label_run = p.add_run(f'{etiqueta}: ')
            set_run_font(label_run, size=12, bold=True)
            value_run = p.add_run(str(valor))
            set_run_font(value_run, size=12)
            return p
        
        def agregar_info_inline(*pares_etiqueta_valor, separador=' | '):
            """Agrega múltiples campos en una sola línea."""
            if not pares_etiqueta_valor:
                return
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            for idx, (etiqueta, valor) in enumerate(pares_etiqueta_valor):
                if not valor:
                    continue
                if idx > 0:
                    sep_run = p.add_run(separador)
                    set_run_font(sep_run, size=12)
                label_run = p.add_run(f'{etiqueta}: ')
                set_run_font(label_run, size=12, bold=True)
                value_run = p.add_run(str(valor))
                set_run_font(value_run, size=12)
            return p
        
        def agregar_salto_entre_secciones():
            """Agrega un salto de línea entre secciones."""
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)
            return spacer
        
        def agregar_titulo_principal(texto):
            """Agrega un título principal (tamaño 18)."""
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(texto)
            set_run_font(run, size=18, bold=True, color=MOVUMS_BLUE_CORP)
            return p

        def render_vuelos(vuelos):
            """Renderiza cotización de vuelos con el nuevo formato."""
            for vuelo in vuelos:
                agregar_titulo_principal("VUELO")
                
                agregar_subtitulo_con_vineta('Información del Vuelo')
                agregar_info_inline(
                    ('Aerolínea', vuelo.get('aerolinea') or '-'),
                    ('Salida', vuelo.get('salida') or '-'),
                    ('Regreso', vuelo.get('regreso') or '-')
                )
                agregar_info_line('Incluye', vuelo.get('incluye') or '-')
                
                # Forma de pago si está presente
                if vuelo.get('forma_pago'):
                    agregar_info_line('Forma de Pago', vuelo.get('forma_pago'))
                
                # Total sin viñeta, tamaño 18 y subrayado (sin salto antes)
                total_p = doc.add_paragraph()
                total_p.paragraph_format.space_before = Pt(0)
                total_p.paragraph_format.space_after = Pt(6)
                total_run = total_p.add_run(f"Total MXN {format_currency(vuelo.get('total'))} Pesos")
                set_run_font(total_run, size=18, bold=True, color=MOVUMS_BLUE_CORP)
                total_run.font.underline = True
                
                agregar_salto_entre_secciones()

        def render_hospedaje(hoteles):
            """Renderiza cotización de hospedaje con el nuevo formato."""
            for hotel in hoteles:
                agregar_titulo_principal("HOSPEDAJE")
                
                agregar_subtitulo_con_vineta('Información del Alojamiento')
                
                # Crear tabla de 2 columnas para la información (como muestra la imagen)
                info_table = doc.add_table(rows=4, cols=2)
                info_table.autofit = False
                
                # Configurar ancho de columnas (50% cada una)
                for col in info_table.columns:
                    for cell in col.cells:
                        cell.width = Inches(3.25)
                
                # Columna izquierda: Nombre, Habitación, Dirección
                # Nombre (fila 0, columna 0)
                nombre_cell = info_table.rows[0].cells[0]
                nombre_label = nombre_cell.paragraphs[0].add_run('Nombre: ')
                set_run_font(nombre_label, size=12, bold=True)
                nombre_val = nombre_cell.paragraphs[0].add_run(hotel.get('nombre') or 'Hotel propuesto')
                set_run_font(nombre_val, size=12)
                
                # Habitación (fila 1, columna 0)
                habitacion_cell = info_table.rows[1].cells[0]
                habitacion_label = habitacion_cell.paragraphs[0].add_run('Habitación: ')
                set_run_font(habitacion_label, size=12, bold=True)
                habitacion_val = habitacion_cell.paragraphs[0].add_run(hotel.get('habitacion') or '-')
                set_run_font(habitacion_val, size=12)
                
                # Dirección (fila 2, columna 0)
                direccion_cell = info_table.rows[2].cells[0]
                direccion_label = direccion_cell.paragraphs[0].add_run('Dirección: ')
                set_run_font(direccion_label, size=12, bold=True)
                direccion_val = direccion_cell.paragraphs[0].add_run(hotel.get('direccion') or '-')
                set_run_font(direccion_val, size=12)
                
                # Columna derecha: Plan de Alimentos (fila 0, columna 1)
                plan_cell = info_table.rows[0].cells[1]
                plan_label = plan_cell.paragraphs[0].add_run('Plan de Alimentos: ')
                set_run_font(plan_label, size=12, bold=True)
                plan_val = plan_cell.paragraphs[0].add_run(hotel.get('plan') or '-')
                set_run_font(plan_val, size=12)
                
                # Forma de Pago (en la segunda columna, segunda fila) si está presente
                if hotel.get('forma_pago'):
                    forma_pago_cell = info_table.rows[1].cells[1]
                    forma_pago_label = forma_pago_cell.paragraphs[0].add_run('Forma de Pago: ')
                    set_run_font(forma_pago_label, size=12, bold=True)
                    forma_pago_val = forma_pago_cell.paragraphs[0].add_run(hotel.get('forma_pago'))
                    set_run_font(forma_pago_val, size=12)
                
                # Total sin viñeta, tamaño 18 y subrayado (sin salto antes)
                total_p = doc.add_paragraph()
                total_p.paragraph_format.space_before = Pt(0)
                total_p.paragraph_format.space_after = Pt(6)
                total_run = total_p.add_run(f"Total MXN {format_currency(hotel.get('total'))} Pesos")
                set_run_font(total_run, size=18, bold=True, color=MOVUMS_BLUE_CORP)
                total_run.font.underline = True
                
                agregar_salto_entre_secciones()

        def render_tours(tours_data):
            """Renderiza cotización de tours con el nuevo formato."""
            agregar_titulo_principal("TOUR")
            
            agregar_subtitulo_con_vineta('Información del Tour')
            agregar_info_line('Número de Reserva', tours_data.get('numero_reserva') or '-')
            agregar_info_line('Nombre del Tour', tours_data.get('nombre') or '-')
            
            if tours_data.get('especificaciones'):
                agregar_salto_entre_secciones()
                agregar_subtitulo_con_vineta('Especificaciones')
                especificaciones = tours_data.get('especificaciones', '').strip()
                if especificaciones:
                    lineas = especificaciones.split('\n')
                    for linea in lineas:
                        if linea.strip():
                            p = doc.add_paragraph()
                            p.paragraph_format.space_after = Pt(4)
                            run = p.add_run(linea.strip())
                            set_run_font(run, size=12)
            
            if tours_data.get('forma_pago'):
                agregar_salto_entre_secciones()
                agregar_subtitulo_con_vineta('Forma de Pago')
                agregar_info_line('Forma de Pago', tours_data.get('forma_pago'))
            
            agregar_salto_entre_secciones()

        def render_paquete(paquete):
            """Renderiza cotización de paquete con el nuevo formato."""
            vuelo = paquete.get('vuelo') or {}
            hotel = paquete.get('hotel') or {}
            
            agregar_titulo_principal("PAQUETE")
            
            agregar_subtitulo_con_vineta('Vuelo')
            agregar_info_inline(
                ('Aerolínea', vuelo.get('aerolinea') or '-'),
                ('Salida', vuelo.get('salida') or '-'),
                ('Regreso', vuelo.get('regreso') or '-')
            )
            agregar_info_line('Incluye', vuelo.get('incluye') or '-')
            
            # Total y Forma de pago del vuelo si están presentes
            if vuelo.get('total'):
                agregar_info_line('Total MXN', format_currency(vuelo.get('total')))
            if vuelo.get('forma_pago'):
                agregar_info_line('Forma de Pago', vuelo.get('forma_pago'))
            
            agregar_salto_entre_secciones()
            
            agregar_subtitulo_con_vineta('Hospedaje')
            
            # Crear tabla de 2 columnas para la información del hospedaje
            # Aumentar filas para incluir todos los campos
            num_filas = 5  # Nombre, Habitación, Dirección, Plan, Notas
            hospedaje_table = doc.add_table(rows=num_filas, cols=2)
            hospedaje_table.autofit = False
            
            # Configurar ancho de columnas (50% cada una)
            for col in hospedaje_table.columns:
                for cell in col.cells:
                    cell.width = Inches(3.25)
            
            # Nombre
            nombre_cell = hospedaje_table.rows[0].cells[0]
            nombre_label = nombre_cell.paragraphs[0].add_run('Nombre: ')
            set_run_font(nombre_label, size=12, bold=True)
            nombre_val = nombre_cell.paragraphs[0].add_run(hotel.get('nombre') or 'Hotel incluido')
            set_run_font(nombre_val, size=12)
            
            # Habitación
            habitacion_cell = hospedaje_table.rows[1].cells[0]
            habitacion_label = habitacion_cell.paragraphs[0].add_run('Habitación: ')
            set_run_font(habitacion_label, size=12, bold=True)
            habitacion_val = habitacion_cell.paragraphs[0].add_run(hotel.get('habitacion') or '-')
            set_run_font(habitacion_val, size=12)
            
            # Dirección
            direccion_cell = hospedaje_table.rows[2].cells[0]
            direccion_label = direccion_cell.paragraphs[0].add_run('Dirección: ')
            set_run_font(direccion_label, size=12, bold=True)
            direccion_val = direccion_cell.paragraphs[0].add_run(hotel.get('direccion') or '-')
            set_run_font(direccion_val, size=12)
            
            # Plan de alimentos
            plan_cell = hospedaje_table.rows[3].cells[0]
            plan_label = plan_cell.paragraphs[0].add_run('Plan de alimentos: ')
            set_run_font(plan_label, size=12, bold=True)
            plan_val = plan_cell.paragraphs[0].add_run(hotel.get('plan') or '-')
            set_run_font(plan_val, size=12)
            
            # Notas
            notas_cell = hospedaje_table.rows[4].cells[0]
            notas_label = notas_cell.paragraphs[0].add_run('Notas: ')
            set_run_font(notas_label, size=12, bold=True)
            notas_val = notas_cell.paragraphs[0].add_run(hotel.get('notas') or '-')
            set_run_font(notas_val, size=12)
            
            # Total y Forma de pago del hotel si están presentes
            if hotel.get('total'):
                agregar_info_line('Total MXN', format_currency(hotel.get('total')))
            if hotel.get('forma_pago'):
                agregar_info_line('Forma de Pago', hotel.get('forma_pago'))
            
            # Forma de pago del paquete si está presente
            if paquete.get('forma_pago'):
                agregar_info_line('Forma de Pago del Paquete', paquete.get('forma_pago'))
            
            # Total sin viñeta, tamaño 18 y subrayado (sin salto antes)
            total_p = doc.add_paragraph()
            total_p.paragraph_format.space_before = Pt(0)
            total_p.paragraph_format.space_after = Pt(6)
            total_run = total_p.add_run(f"Total MXN {format_currency(paquete.get('total'))} Pesos")
            set_run_font(total_run, size=18, bold=True, color=MOVUMS_BLUE_CORP)
            total_run.font.underline = True
            
            agregar_salto_entre_secciones()
            
            agregar_subtitulo_con_vineta('Términos y Condiciones')
            terms = [
                "Los boletos de avión no son reembolsables.",
                "Una vez emitido el boleto no puede ser asignado a otra persona o aerolínea.",
                "Los cambios pueden generar cargos extra y están sujetos a disponibilidad y políticas de cada aerolínea.",
                "Para vuelos nacionales presentarse 2 horas antes; para internacionales 3 horas antes.",
                "Las tarifas están sujetas a cambios y disponibilidad mientras no se reserve.",
            ]
            for term in terms:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                bullet_run = p.add_run('• ')
                set_run_font(bullet_run, size=12)
                term_run = p.add_run(term)
                set_run_font(term_run, size=12)
            
            agregar_salto_entre_secciones()

        if tipo == 'vuelos':
            render_vuelos(cotizacion.get('vuelos') or [])
        elif tipo == 'hospedaje':
            render_hospedaje(cotizacion.get('hoteles') or [])
        elif tipo == 'tours':
            render_tours(cotizacion.get('tours') or {})
        elif tipo == 'paquete':
            render_paquete(cotizacion.get('paquete') or {})

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename_cliente = slugify(general.get('cliente') or 'movums')
        filename = f"cotizacion_{filename_cliente}_{timezone.now().strftime('%Y%m%d%H%M%S')}.docx"
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response

# ------------------- PLANTILLAS DE CONFIRMACIÓN -------------------

class ListarConfirmacionesView(LoginRequiredMixin, DetailView):
    """
    Vista para listar todas las plantillas de confirmación de una venta.
    """
    model = VentaViaje
    template_name = 'ventas/listar_confirmaciones.html'
    context_object_name = 'venta'
    
    def get_object(self, queryset=None):
        pk = self.kwargs.get('pk')
        slug = self.kwargs.get('slug')
        return get_object_or_404(VentaViaje, pk=pk, slug=slug)
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        venta = context['venta']
        context['plantillas'] = PlantillaConfirmacion.objects.filter(venta=venta).order_by('tipo', '-fecha_creacion')
        
        # Obtener los servicios contratados para filtrar qué plantillas mostrar
        servicios_contratados = []
        if venta.servicios_seleccionados:
            servicios_contratados = [s.strip() for s in venta.servicios_seleccionados.split(',')]
        
        # Determinar qué plantillas mostrar según los servicios contratados
        plantillas_disponibles = {
            'mostrar_vuelo_unico': 'VUE' in servicios_contratados,
            'mostrar_vuelo_redondo': 'VUE' in servicios_contratados,
            'mostrar_hospedaje': 'HOS' in servicios_contratados,
            'mostrar_traslado': 'TRA' in servicios_contratados,
            'mostrar_generica': True,  # La genérica siempre está disponible
        }
        
        context.update(plantillas_disponibles)
        return context


class CrearPlantillaConfirmacionView(LoginRequiredMixin, View):
    """
    Vista base para crear una plantilla de confirmación.
    Las vistas específicas heredan de esta y definen el tipo.
    """
    tipo_plantilla = None
    template_name = None
    
    def dispatch(self, request, *args, **kwargs):
        if not self.tipo_plantilla or not self.template_name:
            raise ValueError("Las vistas hijas deben definir 'tipo_plantilla' y 'template_name'")
        return super().dispatch(request, *args, **kwargs)
    
    def get(self, request, pk, slug):
        venta = get_object_or_404(VentaViaje, pk=pk, slug=slug)
        # Si ya existe una plantilla de este tipo, la editamos en lugar de crear una nueva
        plantilla = PlantillaConfirmacion.objects.filter(venta=venta, tipo=self.tipo_plantilla).first()
        
        # Datos por defecto según el tipo
        datos_default = self.get_datos_default()
        
        if plantilla:
            # Si existe, usamos sus datos existentes, sino los defaults
            datos = plantilla.datos if plantilla.datos else datos_default
        else:
            datos = datos_default
        
        # Asegurar que escalas existe y es una lista
        if 'escalas' not in datos:
            datos['escalas'] = []
        elif not isinstance(datos['escalas'], list):
            datos['escalas'] = []
        
        # Convertir escalas a JSON para el template
        escalas_json = json.dumps(datos.get('escalas', []))
        
        # Para hospedaje, construir la URL completa de la imagen si existe
        if self.tipo_plantilla == 'HOSPEDAJE' and datos.get('imagen_hospedaje_url'):
            from django.conf import settings
            # Si la URL ya es absoluta, usarla tal cual, sino construirla
            if datos['imagen_hospedaje_url'].startswith('http'):
                datos['imagen_hospedaje_url'] = datos['imagen_hospedaje_url']
            else:
                # Asegurar que la URL empiece con /media/
                if not datos['imagen_hospedaje_url'].startswith('/'):
                    datos['imagen_hospedaje_url'] = f"/{datos['imagen_hospedaje_url']}"
        
        # Para traslados, preparar la lista de traslados
        traslados_json = "[]"
        if self.tipo_plantilla == 'TRASLADO':
            if 'traslados' in datos and isinstance(datos['traslados'], list) and len(datos['traslados']) > 0:
                traslados_json = json.dumps(datos['traslados'])
            else:
                # Si no hay traslados, crear uno con datos por defecto
                traslados_json = json.dumps([self.get_datos_default().get('TRASLADO', {})])
        
        # Para vuelo redondo, preparar escalas de ida y regreso
        escalas_ida_json = "[]"
        escalas_regreso_json = "[]"
        if self.tipo_plantilla == 'VUELO_REDONDO':
            if 'escalas_ida' not in datos:
                datos['escalas_ida'] = []
            if 'escalas_regreso' not in datos:
                datos['escalas_regreso'] = []
            escalas_ida_json = json.dumps(datos.get('escalas_ida', []))
            escalas_regreso_json = json.dumps(datos.get('escalas_regreso', []))
        
        context = {
            'venta': venta,
            'tipo_plantilla': self.tipo_plantilla,
            'datos': datos,
            'plantilla': plantilla,
            'escalas_json': escalas_json,
            'traslados_json': traslados_json,
            'escalas_ida_json': escalas_ida_json,
            'escalas_regreso_json': escalas_regreso_json,
        }
        return render(request, self.template_name, context)
    
    def post(self, request, pk, slug):
        venta = get_object_or_404(VentaViaje, pk=pk, slug=slug)
        
        # Recopilar todos los datos del POST
        datos = {}
        escalas = []
        
        # Procesar escalas si existen (formato: escalas[0][ciudad], escalas[0][aeropuerto], etc.)
        import re
        escalas_dict = {}
        
        for key, value in request.POST.items():
            if key.startswith('escalas[') and not key.startswith('escalas_ida[') and not key.startswith('escalas_regreso['):
                # Extraer el índice y el campo de la escala (para vuelo único)
                match = re.match(r'escalas\[(\d+)\]\[(\w+)\]', key)
                if match:
                    escala_index = int(match.group(1))
                    campo = match.group(2)
                    if escala_index not in escalas_dict:
                        escalas_dict[escala_index] = {}
                    escalas_dict[escala_index][campo] = value
            elif key.startswith('escalas_ida[') or key.startswith('escalas_regreso[') or key.startswith('traslados['):
                # Estos se procesan por separado más abajo
                continue
            elif key not in ['csrfmiddlewaretoken']:
                datos[key] = value
        
        # Convertir escalas_dict a lista ordenada
        if escalas_dict:
            for i in sorted(escalas_dict.keys()):
                escalas.append(escalas_dict[i])
            datos['escalas'] = escalas
        elif self.tipo_plantilla == 'VUELO_UNICO':
            datos['escalas'] = []
        
        # Procesar escalas de ida y regreso para vuelo redondo
        if self.tipo_plantilla == 'VUELO_REDONDO':
            # Escalas de Ida
            escalas_ida_dict = {}
            for key, value in request.POST.items():
                if key.startswith('escalas_ida['):
                    match = re.match(r'escalas_ida\[(\d+)\]\[(.*?)\]', key)
                    if match:
                        idx = int(match.group(1))
                        campo = match.group(2)
                        if idx not in escalas_ida_dict:
                            escalas_ida_dict[idx] = {}
                        escalas_ida_dict[idx][campo] = value
            
            if escalas_ida_dict:
                escalas_ida = []
                for i in sorted(escalas_ida_dict.keys()):
                    escalas_ida.append(escalas_ida_dict[i])
                datos['escalas_ida'] = escalas_ida
            else:
                datos['escalas_ida'] = []
            
            # Escalas de Regreso
            escalas_regreso_dict = {}
            for key, value in request.POST.items():
                if key.startswith('escalas_regreso['):
                    match = re.match(r'escalas_regreso\[(\d+)\]\[(.*?)\]', key)
                    if match:
                        idx = int(match.group(1))
                        campo = match.group(2)
                        if idx not in escalas_regreso_dict:
                            escalas_regreso_dict[idx] = {}
                        escalas_regreso_dict[idx][campo] = value
            
            if escalas_regreso_dict:
                escalas_regreso = []
                for i in sorted(escalas_regreso_dict.keys()):
                    escalas_regreso.append(escalas_regreso_dict[i])
                datos['escalas_regreso'] = escalas_regreso
            else:
                datos['escalas_regreso'] = []
        
        # Procesar múltiples traslados si existe
        if self.tipo_plantilla == 'TRASLADO':
            from ventas.validators import safe_int
            traslados = []
            traslados_count = safe_int(request.POST.get('traslados_count'), default=0)
            traslados_dict = {}
            
            # Recopilar todos los campos de traslados
            for key, value in request.POST.items():
                if key.startswith('traslados['):
                    match = re.match(r'traslados\[(\d+)\]\[(\w+)\]', key)
                    if match:
                        idx = int(match.group(1))
                        campo = match.group(2)
                        if idx not in traslados_dict:
                            traslados_dict[idx] = {}
                        traslados_dict[idx][campo] = value
            
            # Convertir a lista ordenada
            if traslados_dict:
                for i in sorted(traslados_dict.keys()):
                    traslados.append(traslados_dict[i])
            
            if traslados:
                datos['traslados'] = traslados
            else:
                # Si no hay traslados en POST, crear uno vacío
                datos['traslados'] = []
        
        # Buscar si ya existe una plantilla de este tipo
        plantilla, created = PlantillaConfirmacion.objects.get_or_create(
            venta=venta,
            tipo=self.tipo_plantilla,
            defaults={
                'datos': datos,
                'creado_por': request.user
            }
        )
        
        # Manejar imagen en base64 para plantilla de hospedaje
        if self.tipo_plantilla == 'HOSPEDAJE' and datos.get('imagen_hospedaje_base64'):
            from django.core.files.storage import default_storage
            from django.core.files.base import ContentFile
            import uuid
            import base64
            
            try:
                # Decodificar base64
                base64_data = datos['imagen_hospedaje_base64']
                # Remover el prefijo data:image/...;base64, si existe
                extension = 'png'  # Por defecto
                if ',' in base64_data:
                    header, base64_data = base64_data.split(',', 1)
                    # Intentar obtener extensión del header MIME
                    if 'image/png' in header:
                        extension = 'png'
                    elif 'image/jpeg' in header or 'image/jpg' in header:
                        extension = 'jpg'
                    elif 'image/gif' in header:
                        extension = 'gif'
                    elif 'image/webp' in header:
                        extension = 'webp'
                
                imagen_data = base64.b64decode(base64_data)
                
                # Generar nombre único para la imagen
                nombre_archivo = f"hospedaje_{venta.pk}_{uuid.uuid4().hex[:8]}.{extension}"
                ruta = default_storage.save(
                    f'plantillas_confirmacion/{nombre_archivo}',
                    ContentFile(imagen_data)
                )
                # Guardar la URL
                url_completa = default_storage.url(ruta)
                datos['imagen_hospedaje_url'] = url_completa
                # Remover el base64 de los datos (no queremos guardarlo)
                datos.pop('imagen_hospedaje_base64', None)
            except Exception as e:
                # Si hay error, no guardar la imagen pero continuar
                logger.error(f"Error al procesar imagen base64: {e}", exc_info=True)
                datos.pop('imagen_hospedaje_base64', None)
        
        # Si no se subió nueva imagen pero ya existe una, preservarla
        if not created and self.tipo_plantilla == 'HOSPEDAJE':
            if 'imagen_hospedaje_url' not in datos and plantilla.datos.get('imagen_hospedaje_url'):
                datos['imagen_hospedaje_url'] = plantilla.datos.get('imagen_hospedaje_url')
        
        if not created:
            # Actualizar la existente
            plantilla.datos = datos
            if not plantilla.creado_por:
                plantilla.creado_por = request.user
            plantilla.save()
        
        messages.success(request, f"Plantilla {self.get_tipo_display()} guardada correctamente.")
        return redirect('listar_confirmaciones', pk=venta.pk, slug=venta.slug_safe)
    
    def get_datos_default(self):
        """Retorna los datos por defecto según el tipo de plantilla."""
        defaults = {
            'VUELO_UNICO': {
                'clave_reserva': '',
                'numero_vuelo': '',
                'aerolinea': '',
                'fecha': '',
                'hora_salida': '',
                'hora_llegada': '',
                'origen_terminal': '',
                'destino_terminal': '',
                'tipo_vuelo': '',
                'escalas': [],
                'equipaje': '',
                'pasajeros': '',
                'vuelo': '',
            },
            'VUELO_REDONDO': {
                'clave_reserva': '',
                'numero_vuelo_ida': '',
                'aerolinea_ida': '',
                'fecha_salida_ida': '',
                'hora_salida_ida': '',
                'hora_llegada_ida': '',
                'origen_ida': '',
                'destino_ida': '',
                'numero_vuelo_regreso': '',
                'aerolinea_regreso': '',
                'fecha_salida_regreso': '',
                'hora_salida_regreso': '',
                'hora_llegada_regreso': '',
                'origen_regreso': '',
                'destino_regreso': '',
                'pasajeros': '',
                'equipaje': '',
            },
            'HOSPEDAJE': {
                'nombre_alojamiento': '',
                'numero_referencia': '',
                'fecha_checkin': '',
                'fecha_checkout': '',
                'tipo_habitacion': '',
                'adultos': '1',
                'ninos': '0',
                'regimen': '',
                'viajero_principal': '',
                'observaciones': '',
                'imagen_hospedaje_url': '',
            },
            'TRASLADO': {
                'compania': '',
                'codigo_reserva': '',
                'tipo_servicio': '',
                'horario_inicio': '',
                'desde': '',
                'hasta': '',
                'adultos': '1',
                'ninos': '0',
                'informacion_adicional': '',
            },
            'GENERICA': {
                'titulo': '',
                'contenido': '',
            },
        }
        return defaults.get(self.tipo_plantilla, {})
    
    def get_tipo_display(self):
        """Retorna el nombre legible del tipo de plantilla."""
        return dict(PlantillaConfirmacion.TIPO_CHOICES).get(self.tipo_plantilla, self.tipo_plantilla)


# Vistas específicas para cada tipo de plantilla
class CrearVueloUnicoView(CrearPlantillaConfirmacionView):
    tipo_plantilla = 'VUELO_UNICO'
    template_name = 'ventas/plantillas/vuelo_unico.html'


class CrearVueloRedondoView(CrearPlantillaConfirmacionView):
    tipo_plantilla = 'VUELO_REDONDO'
    template_name = 'ventas/plantillas/vuelo_redondo.html'


class CrearHospedajeView(CrearPlantillaConfirmacionView):
    tipo_plantilla = 'HOSPEDAJE'
    template_name = 'ventas/plantillas/hospedaje.html'


class CrearTrasladoView(CrearPlantillaConfirmacionView):
    tipo_plantilla = 'TRASLADO'
    template_name = 'ventas/plantillas/traslado.html'


class CrearGenericaView(CrearPlantillaConfirmacionView):
    tipo_plantilla = 'GENERICA'
    template_name = 'ventas/plantillas/generica.html'


class EliminarPlantillaConfirmacionView(LoginRequiredMixin, View):
    """Vista para eliminar una plantilla de confirmación."""
    
    def post(self, request, pk, slug, plantilla_pk):
        venta = get_object_or_404(VentaViaje, pk=pk, slug=slug)
        plantilla = get_object_or_404(PlantillaConfirmacion, pk=plantilla_pk, venta=venta)
        
        tipo_nombre = plantilla.get_tipo_display()
        plantilla.delete()
        
        messages.success(request, f'Plantilla "{tipo_nombre}" eliminada correctamente.')
        return redirect('listar_confirmaciones', pk=pk, slug=slug)


class GenerarDocumentoConfirmacionView(LoginRequiredMixin, DetailView):
    def _asegurar_estilo_heading(self, doc, style_name, size_pt):
        """Garantiza que exista un estilo de encabezado con nombre dado."""
        from docx.enum.style import WD_STYLE_TYPE
        from docx.shared import Pt

        styles = doc.styles
        try:
            styles[style_name]
        except KeyError:
            normal = styles['Normal']
            nuevo = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            nuevo.base_style = normal
            font = nuevo.font
            font.name = 'Arial'
            font.size = Pt(size_pt)

    """
    Vista para generar un documento .docx combinando todas las plantillas de confirmación de una venta.
    """
    model = VentaViaje
    
    def _normalizar_texto(self, texto):
        """
        Normaliza un texto limpiando espacios y saltos de línea excesivos.
        Convierte múltiples saltos de línea en uno solo y elimina espacios extras.
        """
        if not texto:
            return texto
        
        import re
        texto = str(texto)
        # Reemplazar múltiples saltos de línea consecutivos por uno solo o espacio
        texto = re.sub(r'\n\s*\n+', '\n', texto)
        # Reemplazar múltiples espacios o tabs por uno solo
        texto = re.sub(r'[ \t]+', ' ', texto)
        # Eliminar saltos de línea solitarios entre palabras (convertirlos a espacio)
        texto = re.sub(r'(?<!\n)\n(?!\n)', ' ', texto)
        # Eliminar espacios al inicio y final
        texto = texto.strip()
        # Limpiar espacios múltiples nuevamente después de convertir saltos de línea
        texto = re.sub(r'[ \t]+', ' ', texto)
        return texto
    
    def _capitalizar_nombre_propio(self, texto):
        """
        Capitaliza nombres propios correctamente según reglas del español.
        Primera letra mayúscula, resto minúsculas.
        Maneja nombres compuestos y artículos/preposiciones.
        """
        if not texto:
            return texto
        
        import re
        texto = str(texto).strip()
        if not texto:
            return texto
        
        # Palabras que deben mantenerse en minúsculas cuando están en medio de un nombre
        # (artículos, preposiciones)
        palabras_minusculas = {
            'de', 'del', 'la', 'las', 'el', 'los', 'y', 'e', 'o', 'u',
            'a', 'al', 'con', 'en', 'por', 'para', 'sin', 'sobre', 'entre',
            'da', 'das', 'do', 'dos', 'von', 'van', 'le', 'les'
        }
        
        # Palabras que siempre deben ir en mayúscula (apellidos comunes, títulos)
        palabras_mayusculas = {'iii', 'ii', 'iv', 'sr', 'sra', 'srta', 'dr', 'dra', 'mtro', 'mtra'}
        
        # Convertir todo a minúsculas primero (preservando estructura)
        texto_lower = texto.lower()
        
        # Dividir por espacios para procesar palabra por palabra
        palabras = texto_lower.split()
        palabras_capitalizadas = []
        
        for i, palabra_orig in enumerate(palabras):
            # Limpiar la palabra de separadores al inicio/final pero preservar estructura
            palabra = palabra_orig.strip()
            
            if not palabra:
                continue
            
            # Manejar palabras con guiones o apóstrofes
            tiene_guion = '-' in palabra
            tiene_apostrofe = "'" in palabra
            
            if tiene_guion or tiene_apostrofe:
                # Dividir por guiones/apóstrofes y capitalizar cada parte
                partes = re.split(r'([\-\']+)', palabra)
                partes_cap = []
                for j, parte in enumerate(partes):
                    if parte in ['-', "'"]:
                        partes_cap.append(parte)
                    elif parte.strip():
                        if j == 0 or (j > 0 and partes[j-1] in ['-', "'"]):
                            # Capitalizar después de guion/apóstrofe
                            partes_cap.append(parte[0].upper() + parte[1:] if len(parte) > 1 else parte.upper())
                        else:
                            partes_cap.append(parte)
                palabra_final = ''.join(partes_cap)
            else:
                # Procesar palabra normal
                if i == 0:
                    # Primera palabra siempre capitalizada
                    palabra_final = palabra[0].upper() + palabra[1:] if len(palabra) > 1 else palabra.upper()
                elif palabra in palabras_mayusculas:
                    # Títulos y números romanos en mayúsculas
                    palabra_final = palabra.upper()
                elif palabra in palabras_minusculas:
                    # Artículos/preposiciones en minúsculas (excepto si es la última palabra)
                    palabra_final = palabra if i < len(palabras) - 1 else palabra[0].upper() + palabra[1:] if len(palabra) > 1 else palabra.upper()
                else:
                    # Capitalizar normalmente
                    palabra_final = palabra[0].upper() + palabra[1:] if len(palabra) > 1 else palabra.upper()
            
            palabras_capitalizadas.append(palabra_final)
        
        # Reconstruir el texto con espacios
        texto_final = ' '.join(palabras_capitalizadas)
        # Limpiar espacios múltiples
        texto_final = re.sub(r'[ \t]+', ' ', texto_final).strip()
        
        return texto_final
    
    def _normalizar_valor_campo(self, valor, es_nombre_propio=False, limpiar_saltos_linea=True):
        """
        Normaliza un valor de campo aplicando las transformaciones necesarias.
        
        Args:
            valor: El valor a normalizar
            es_nombre_propio: Si es True, aplica capitalización de nombre propio
            limpiar_saltos_linea: Si es True, limpia saltos de línea excesivos
        """
        if not valor:
            return valor
        
        texto = str(valor).strip()
        if not texto:
            return valor
        
        # Limpiar saltos de línea si es necesario
        if limpiar_saltos_linea:
            texto = self._normalizar_texto(texto)
        
        # Capitalizar si es nombre propio
        if es_nombre_propio:
            texto = self._capitalizar_nombre_propio(texto)
        
        return texto
    
    def get_object(self, queryset=None):
        pk = self.kwargs.get('pk')
        slug = self.kwargs.get('slug')
        return get_object_or_404(VentaViaje, pk=pk, slug=slug)
    
    def get(self, request, *args, **kwargs):
        """
        Genera PDF de confirmaciones usando WeasyPrint (igual que comprobantes de abonos).
        Formato profesional igual que cotizaciones.
        """
        if not WEASYPRINT_AVAILABLE:
            return HttpResponse("Error en la generación de PDF. Faltan dependencias (WeasyPrint).", status=503)
        
        try:
            venta = self.get_object()
        except Exception as e:
            logger.error(f"Error obteniendo venta: {e}")
            return HttpResponse(f"Error: No se pudo obtener la venta. {str(e)}", status=400)
        
        # Orden: 1) Vuelos (sencillo + redondo), 2) Hospedaje, 3) Traslado, 4) Genérica
        from django.db.models import Case, When, Value, IntegerField
        plantillas = PlantillaConfirmacion.objects.filter(venta=venta).annotate(
            orden_tipo=Case(
                When(tipo='VUELO_UNICO', then=Value(0)),
                When(tipo='VUELO_REDONDO', then=Value(1)),
                When(tipo='HOSPEDAJE', then=Value(2)),
                When(tipo='TRASLADO', then=Value(3)),
                When(tipo='GENERICA', then=Value(4)),
                default=Value(5),
                output_field=IntegerField(),
            )
        ).order_by('orden_tipo', '-fecha_creacion')
        
        if not plantillas.exists():
            messages.warning(request, "No hay plantillas de confirmación para generar el documento.")
            return redirect('listar_confirmaciones', pk=venta.pk, slug=venta.slug_safe)
        
        def format_date(value):
            """Formatea una fecha para mostrar en el documento."""
            if not value:
                return '-'
            try:
                if isinstance(value, datetime.date):
                    return value.strftime('%d/%m/%Y')
                elif isinstance(value, str):
                    parsed = datetime.date.fromisoformat(value)
                    return parsed.strftime('%d/%m/%Y')
                return str(value)
            except Exception:
                return str(value)
        
        # Preparar ruta del membrete para WeasyPrint
        membrete_path = os.path.join(settings.BASE_DIR, 'static', 'img', 'membrete_movums.jpg')
        membrete_url = None
        if os.path.exists(membrete_path):
            membrete_abs_path = os.path.abspath(membrete_path)
            if os.name == 'nt':
                membrete_url = f"file:///{membrete_abs_path.replace(os.sep, '/')}"
            else:
                membrete_url = f"file://{membrete_abs_path}"
        
        # Procesar cada plantilla y generar HTML
        plantillas_html = []
        for plantilla in plantillas:
            datos = plantilla.datos or {}
            tipo = plantilla.tipo
            
            html_plantilla = ""
            if tipo == 'VUELO_UNICO':
                html_plantilla = self._generar_html_vuelo_unico(datos, format_date)
            elif tipo == 'VUELO_REDONDO':
                html_plantilla = self._generar_html_vuelo_redondo(datos, format_date)
            elif tipo == 'HOSPEDAJE':
                html_plantilla = self._generar_html_hospedaje(datos, format_date, request)
            elif tipo == 'TRASLADO':
                traslados_list = datos.get('traslados', [])
                if traslados_list and isinstance(traslados_list, list):
                    # Cada traslado como plantilla separada: evita bug WeasyPrint con page-break-before que oculta la última tabla
                    for i, traslado in enumerate(traslados_list):
                        card_html = self._generar_html_traslado(traslado, format_date)
                        if i == 0:
                            card_html = f'<div class="traslado-primera-espacio">{card_html}</div>'
                        else:
                            card_html = f'<div class="traslado-tabla-grande">{card_html}</div>'
                        plantillas_html.append(card_html)
                else:
                    card_html = self._generar_html_traslado(datos, format_date)
                    html_plantilla = f'<div class="traslado-primera-espacio">{card_html}</div>'
                    plantillas_html.append(html_plantilla)
                continue  # ya se añadieron a plantillas_html
            elif tipo == 'GENERICA':
                html_plantilla = self._generar_html_generica(datos)
            
            if html_plantilla:
                plantillas_html.append(html_plantilla)
        
        # Contexto para la plantilla HTML
        context = {
            'venta': venta,
            'fecha_generacion': datetime.datetime.now(),
            'plantillas_html': plantillas_html,
            'membrete_url': membrete_url,
            'STATIC_URL': settings.STATIC_URL,
        }
        
        # Renderizar plantilla HTML
        html_string = render_to_string('ventas/confirmaciones_pdf.html', context, request=request)
        
        # Generar PDF con WeasyPrint
        static_dir = os.path.join(settings.BASE_DIR, 'static')
        static_dir_abs = os.path.abspath(static_dir)
        base_url = f"file://{static_dir_abs}/"
        
        html = HTML(string=html_string, base_url=base_url)
        pdf_file = html.write_pdf(stylesheets=[])
        
        # Preparar respuesta HTTP
        nombre_cliente_safe = venta.cliente.nombre_completo_display.replace(' ', '_').replace('/', '_')
        timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"Confirmaciones_Venta_{venta.pk}_{nombre_cliente_safe}_{timestamp}.pdf"
        
        response = HttpResponse(pdf_file, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response['Pragma'] = 'no-cache'
        response['Expires'] = '0'
        
        return response
    
    # ===== MÉTODOS DE GENERACIÓN HTML PARA PDF =====
    # Estos métodos generan HTML para convertir a PDF usando WeasyPrint
    
    def _generar_html_vuelo_unico(self, datos, format_date):
        """Genera HTML para plantilla de vuelo único (EXACTAMENTE igual que cotizaciones con cards)."""
        html_parts = []
        
        # Card principal con header (IGUAL QUE COTIZACIONES - seccion_vuelo.html)
        html_parts.append('<div class="card">')
        html_parts.append('<div class="card-header">')
        html_parts.append('<span class="icon">')
        html_parts.append('<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">')
        html_parts.append('<path d="M21 16v-2l-8-5V3.5c0-.83-.67-1.5-1.5-1.5S10 2.67 10 3.5V9l-8 5v2l8-2.5V19l-2 1.5V22l3.5-1 3.5 1v-1.5L13 19v-5.5l8 2.5z" fill="white" stroke="none"/>')
        html_parts.append('</svg>')
        html_parts.append('</span>')
        html_parts.append('<span>VUELO ÚNICO</span>')
        html_parts.append('</div>')
        
        # Tabla de datos (IGUAL QUE COTIZACIONES - data-table)
        html_parts.append('<table class="data-table">')
        
        if datos.get('clave_reserva'):
            html_parts.append('<tr>')
            html_parts.append('<td style="width: 30%;"><strong>Clave de Reserva:</strong></td>')
            html_parts.append(f'<td>{datos.get("clave_reserva")}</td>')
            html_parts.append('</tr>')
        
        if datos.get('aerolinea'):
            html_parts.append('<tr>')
            html_parts.append('<td><strong>Aerolínea:</strong></td>')
            html_parts.append(f'<td>{datos.get("aerolinea")}</td>')
            html_parts.append('</tr>')
        
        if datos.get('numero_vuelo'):
            html_parts.append('<tr>')
            html_parts.append('<td><strong>Vuelo:</strong></td>')
            html_parts.append(f'<td>{datos.get("numero_vuelo")}</td>')
            html_parts.append('</tr>')
        
        if datos.get('fecha'):
            html_parts.append('<tr>')
            html_parts.append('<td><strong>Fecha:</strong></td>')
            html_parts.append(f'<td>{datos.get("fecha")}</td>')
            html_parts.append('</tr>')
        
        salida_info = []
        if datos.get('hora_salida'):
            salida_info.append(datos.get('hora_salida'))
        if datos.get('origen_terminal'):
            salida_info.append(f"Desde: {datos.get('origen_terminal')}")
        if salida_info:
            html_parts.append('<tr>')
            html_parts.append('<td><strong>Salida:</strong></td>')
            html_parts.append(f'<td>{" | ".join(salida_info)}</td>')
            html_parts.append('</tr>')
        
        llegada_info = []
        if datos.get('hora_llegada'):
            llegada_info.append(datos.get('hora_llegada'))
        if datos.get('destino_terminal'):
            llegada_info.append(f"Hasta: {datos.get('destino_terminal')}")
        if llegada_info:
            html_parts.append('<tr>')
            html_parts.append('<td><strong>Llegada:</strong></td>')
            html_parts.append(f'<td>{" | ".join(llegada_info)}</td>')
            html_parts.append('</tr>')
        
        if datos.get('tipo_vuelo'):
            html_parts.append('<tr>')
            html_parts.append('<td><strong>Tipo de Vuelo:</strong></td>')
            html_parts.append(f'<td>{datos.get("tipo_vuelo")}</td>')
            html_parts.append('</tr>')
        
        pasajeros = datos.get('pasajeros', '')
        if pasajeros:
            pasajeros_texto = ', '.join([p.strip() for p in str(pasajeros).replace('\r\n', '\n').split('\n') if p.strip()])
            html_parts.append('<tr>')
            html_parts.append('<td><strong>Pasajeros:</strong></td>')
            html_parts.append(f'<td>{pasajeros_texto}</td>')
            html_parts.append('</tr>')
        
        if datos.get('equipaje'):
            html_parts.append('<tr>')
            html_parts.append('<td><strong>Equipaje:</strong></td>')
            html_parts.append(f'<td>{datos.get("equipaje")}</td>')
            html_parts.append('</tr>')
        
        if datos.get('informacion_adicional'):
            html_parts.append('<tr>')
            html_parts.append('<td><strong>Información Adicional:</strong></td>')
            html_parts.append(f'<td>{datos.get("informacion_adicional")}</td>')
            html_parts.append('</tr>')
        
        # Escalas si aplica (dentro de la misma tabla)
        if datos.get('tipo_vuelo') == 'Escalas' and datos.get('escalas'):
            escalas = datos.get('escalas', [])
            if escalas:
                for i, escala in enumerate(escalas, 1):
                    html_parts.append('<tr>')
                    html_parts.append(f'<td><strong>Escala {i}:</strong></td>')
                    escala_texto = f"{escala.get('ciudad', '')} - {escala.get('aeropuerto', '')}"
                    detalles = []
                    if escala.get('hora_llegada'):
                        detalles.append(f"Llegada: {escala.get('hora_llegada')}")
                    if escala.get('hora_salida'):
                        detalles.append(f"Salida: {escala.get('hora_salida')}")
                    if escala.get('numero_vuelo'):
                        detalles.append(f"Vuelo: {escala.get('numero_vuelo')}")
                    if escala.get('duracion'):
                        detalles.append(f"Duración: {escala.get('duracion')}")
                    if detalles:
                        escala_texto += f" ({' | '.join(detalles)})"
                    html_parts.append(f'<td>{escala_texto}</td>')
                    html_parts.append('</tr>')
        
        html_parts.append('</table>')
        html_parts.append('</div>')  # Cierre de card
        
        # Información Completa del Vuelo (campo "vuelo") - en una nueva card si existe
        if datos.get('vuelo'):
            html_parts.append('<div style="page-break-before: always;"></div>')
            html_parts.append('<div class="card">')
            html_parts.append('<div class="card-header">')
            html_parts.append('<span class="icon">')
            html_parts.append('<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">')
            html_parts.append('<path d="M21 16v-2l-8-5V3.5c0-.83-.67-1.5-1.5-1.5S10 2.67 10 3.5V9l-8 5v2l8-2.5V19l-2 1.5V22l3.5-1 3.5 1v-1.5L13 19v-5.5l8 2.5z" fill="white" stroke="none"/>')
            html_parts.append('</svg>')
            html_parts.append('</span>')
            html_parts.append('<span>INFORMACIÓN COMPLETA DEL VUELO</span>')
            html_parts.append('</div>')
            vuelo_info = datos.get('vuelo', '').replace('\n', '<br>')
            html_parts.append(f'<div style="padding: 12px 18px; font-size: 9pt; line-height: 1.5;">{vuelo_info}</div>')
            html_parts.append('</div>')
        
        return "".join(html_parts)
    
    def _valor_o_guion(self, valor):
        """Retorna el valor como string o '-' si está vacío."""
        if valor is None:
            return '-'
        s = str(valor).strip()
        return s if s else '-'

    def _generar_html_vuelo_redondo(self, datos, format_date):
        """Genera HTML para plantilla de vuelo redondo. Incluye ida, regreso, escalas y toda la información llenada."""
        html_parts = []
        datos = datos or {}
        # Asegurar que escalas sean listas (por si el JSON devuelve otro tipo)
        escalas_ida = datos.get('escalas_ida')
        if not isinstance(escalas_ida, list):
            escalas_ida = []
        escalas_regreso = datos.get('escalas_regreso')
        if not isinstance(escalas_regreso, list):
            escalas_regreso = []

        # ----- VUELO DE IDA -----
        html_parts.append('<div class="card">')
        html_parts.append('<div class="card-header">')
        html_parts.append('<span class="icon">')
        html_parts.append('<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">')
        html_parts.append('<path d="M21 16v-2l-8-5V3.5c0-.83-.67-1.5-1.5-1.5S10 2.67 10 3.5V9l-8 5v2l8-2.5V19l-2 1.5V22l3.5-1 3.5 1v-1.5L13 19v-5.5l8 2.5z" fill="white" stroke="none"/>')
        html_parts.append('</svg>')
        html_parts.append('</span>')
        html_parts.append('<span>VUELO DE IDA</span>')
        html_parts.append('</div>')
        html_parts.append('<table class="data-table">')

        html_parts.append('<tr><td style="width: 30%;"><strong>Clave de Reserva:</strong></td>')
        html_parts.append(f'<td>{self._valor_o_guion(datos.get("clave_reserva"))}</td></tr>')
        html_parts.append('<tr><td><strong>Aerolínea:</strong></td>')
        html_parts.append(f'<td>{self._valor_o_guion(datos.get("aerolinea_ida"))}</td></tr>')
        html_parts.append('<tr><td><strong>Vuelo:</strong></td>')
        html_parts.append(f'<td>{self._valor_o_guion(datos.get("numero_vuelo_ida"))}</td></tr>')

        salida_ida_info = []
        if datos.get('fecha_salida_ida'):
            salida_ida_info.append(format_date(datos.get('fecha_salida_ida')) if format_date else datos.get('fecha_salida_ida'))
        if datos.get('hora_salida_ida'):
            salida_ida_info.append(datos.get('hora_salida_ida'))
        if datos.get('origen_ida'):
            salida_ida_info.append(f"Desde: {datos.get('origen_ida')}")
        html_parts.append('<tr><td><strong>Salida:</strong></td>')
        html_parts.append(f'<td>{" | ".join(salida_ida_info) if salida_ida_info else self._valor_o_guion(None)}</td></tr>')

        llegada_ida_info = []
        if datos.get('hora_llegada_ida'):
            llegada_ida_info.append(datos.get('hora_llegada_ida'))
        if datos.get('destino_ida'):
            llegada_ida_info.append(f"Hasta: {datos.get('destino_ida')}")
        html_parts.append('<tr><td><strong>Llegada:</strong></td>')
        html_parts.append(f'<td>{" | ".join(llegada_ida_info) if llegada_ida_info else self._valor_o_guion(None)}</td></tr>')

        html_parts.append('<tr><td><strong>Tipo de Vuelo:</strong></td>')
        html_parts.append(f'<td>{self._valor_o_guion(datos.get("tipo_vuelo_ida"))}</td></tr>')

        for i, escala in enumerate(escalas_ida, 1):
            html_parts.append('<tr>')
            html_parts.append(f'<td><strong>Escala {i} (Ida):</strong></td>')
            ciudad = (escala.get('ciudad') or '').strip()
            aeropuerto = (escala.get('aeropuerto') or '').strip()
            escala_texto = f"{ciudad} - {aeropuerto}" if (ciudad or aeropuerto) else '-'
            detalles = []
            if escala.get('hora_llegada'):
                detalles.append(f"Llegada: {escala.get('hora_llegada')}")
            if escala.get('hora_salida'):
                detalles.append(f"Salida: {escala.get('hora_salida')}")
            if escala.get('numero_vuelo'):
                detalles.append(f"Vuelo: {escala.get('numero_vuelo')}")
            if escala.get('duracion'):
                detalles.append(f"Duración: {escala.get('duracion')}")
            if detalles:
                escala_texto += f" ({' | '.join(detalles)})"
            html_parts.append(f'<td>{escala_texto}</td>')
            html_parts.append('</tr>')

        pasajeros = datos.get('pasajeros', '')
        pasajeros_texto = ', '.join([p.strip() for p in str(pasajeros).replace('\r\n', '\n').split('\n') if p.strip()]) if pasajeros else ''
        html_parts.append('<tr><td><strong>Pasajeros:</strong></td>')
        html_parts.append(f'<td>{pasajeros_texto or self._valor_o_guion(None)}</td></tr>')
        html_parts.append('<tr><td><strong>Equipaje:</strong></td>')
        html_parts.append(f'<td>{self._valor_o_guion(datos.get("equipaje"))}</td></tr>')

        html_parts.append('</table>')
        html_parts.append('</div>')  # Cierre card Ida

        # ----- VUELO DE REGRESO (nueva página con espacio superior para no quedar en el borde) -----
        html_parts.append('<div class="vuelo-regreso-nueva-pagina">')
        html_parts.append('<div class="card" style="margin-top: 12px;">')
        html_parts.append('<div class="card-header">')
        html_parts.append('<span class="icon">')
        html_parts.append('<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">')
        html_parts.append('<path d="M21 16v-2l-8-5V3.5c0-.83-.67-1.5-1.5-1.5S10 2.67 10 3.5V9l-8 5v2l8-2.5V19l-2 1.5V22l3.5-1 3.5 1v-1.5L13 19v-5.5l8 2.5z" fill="white" stroke="none"/>')
        html_parts.append('</svg>')
        html_parts.append('</span>')
        html_parts.append('<span>VUELO DE REGRESO</span>')
        html_parts.append('</div>')
        html_parts.append('<table class="data-table">')

        html_parts.append('<tr><td style="width: 30%;"><strong>Clave de Reserva:</strong></td>')
        html_parts.append(f'<td>{self._valor_o_guion(datos.get("clave_reserva"))}</td></tr>')
        html_parts.append('<tr><td><strong>Aerolínea:</strong></td>')
        html_parts.append(f'<td>{self._valor_o_guion(datos.get("aerolinea_regreso"))}</td></tr>')
        html_parts.append('<tr><td><strong>Vuelo:</strong></td>')
        html_parts.append(f'<td>{self._valor_o_guion(datos.get("numero_vuelo_regreso"))}</td></tr>')

        salida_regreso_info = []
        if datos.get('fecha_salida_regreso'):
            salida_regreso_info.append(format_date(datos.get('fecha_salida_regreso')) if format_date else datos.get('fecha_salida_regreso'))
        if datos.get('hora_salida_regreso'):
            salida_regreso_info.append(datos.get('hora_salida_regreso'))
        if datos.get('origen_regreso'):
            salida_regreso_info.append(f"Desde: {datos.get('origen_regreso')}")
        html_parts.append('<tr><td><strong>Salida:</strong></td>')
        html_parts.append(f'<td>{" | ".join(salida_regreso_info) if salida_regreso_info else self._valor_o_guion(None)}</td></tr>')

        llegada_regreso_info = []
        if datos.get('hora_llegada_regreso'):
            llegada_regreso_info.append(datos.get('hora_llegada_regreso'))
        if datos.get('destino_regreso'):
            llegada_regreso_info.append(f"Hasta: {datos.get('destino_regreso')}")
        html_parts.append('<tr><td><strong>Llegada:</strong></td>')
        html_parts.append(f'<td>{" | ".join(llegada_regreso_info) if llegada_regreso_info else self._valor_o_guion(None)}</td></tr>')

        html_parts.append('<tr><td><strong>Tipo de Vuelo:</strong></td>')
        html_parts.append(f'<td>{self._valor_o_guion(datos.get("tipo_vuelo_regreso"))}</td></tr>')

        for i, escala in enumerate(escalas_regreso, 1):
            html_parts.append('<tr>')
            html_parts.append(f'<td><strong>Escala {i} (Regreso):</strong></td>')
            ciudad = (escala.get('ciudad') or '').strip()
            aeropuerto = (escala.get('aeropuerto') or '').strip()
            escala_texto = f"{ciudad} - {aeropuerto}" if (ciudad or aeropuerto) else '-'
            detalles = []
            if escala.get('hora_llegada'):
                detalles.append(f"Llegada: {escala.get('hora_llegada')}")
            if escala.get('hora_salida'):
                detalles.append(f"Salida: {escala.get('hora_salida')}")
            if escala.get('numero_vuelo'):
                detalles.append(f"Vuelo: {escala.get('numero_vuelo')}")
            if escala.get('duracion'):
                detalles.append(f"Duración: {escala.get('duracion')}")
            if detalles:
                escala_texto += f" ({' | '.join(detalles)})"
            html_parts.append(f'<td>{escala_texto}</td>')
            html_parts.append('</tr>')

        html_parts.append('<tr><td><strong>Pasajeros:</strong></td>')
        html_parts.append(f'<td>{pasajeros_texto or self._valor_o_guion(None)}</td></tr>')
        html_parts.append('<tr><td><strong>Equipaje:</strong></td>')
        html_parts.append(f'<td>{self._valor_o_guion(datos.get("equipaje"))}</td></tr>')
        html_parts.append('<tr><td><strong>Información Adicional:</strong></td>')
        html_parts.append(f'<td>{self._valor_o_guion(datos.get("informacion_adicional"))}</td></tr>')

        html_parts.append('</table>')
        html_parts.append('</div>')  # Cierre card Regreso
        html_parts.append('</div>')  # Cierre vuelo-regreso-nueva-pagina

        return "".join(html_parts)
    
    def _generar_html_hospedaje(self, datos, format_date, request):
        """Genera HTML para plantilla de hospedaje (mismo formato que cotizaciones con cards)."""
        html_parts = []
        
        # Card principal con header (IGUAL QUE COTIZACIONES)
        html_parts.append('<div class="card">')
        html_parts.append('<div class="card-header">')
        html_parts.append('<span class="icon">')
        html_parts.append('<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">')
        html_parts.append('<path d="M7 13c1.66 0 3-1.34 3-3S8.66 7 7 7s-3 1.34-3 3 1.34 3 3 3zm12-6h-8v7h8V7zM7 14c-2.33 0-7 1.17-7 3.5V19h14v-1.5c0-2.33-4.67-3.5-7-3.5zM19 14c-2.33 0-7 1.17-7 3.5V19h14v-1.5c0-2.33-4.67-3.5-7-3.5z" fill="white" stroke="none"/>')
        html_parts.append('</svg>')
        html_parts.append('</span>')
        html_parts.append('<span>HOSPEDAJE</span>')
        html_parts.append('</div>')
        
        # Tabla de información (IGUAL QUE COTIZACIONES - data-table)
        nombre_alojamiento = self._normalizar_valor_campo(datos.get('nombre_alojamiento', ''), limpiar_saltos_linea=True) or 'Hotel propuesto'
        referencia = self._normalizar_valor_campo(datos.get('numero_referencia', ''), limpiar_saltos_linea=True) or '-'
        viajero_principal = self._normalizar_valor_campo(datos.get('viajero_principal', ''), es_nombre_propio=True) or '-'
        tipo_habitacion = self._normalizar_valor_campo(datos.get('tipo_habitacion', ''), limpiar_saltos_linea=True) or '-'
        direccion = self._normalizar_valor_campo(datos.get('direccion', ''), limpiar_saltos_linea=True) if datos.get('direccion') else '-'
        plan_alimentos = self._normalizar_valor_campo(datos.get('plan_alimentos', ''), limpiar_saltos_linea=True) if datos.get('plan_alimentos') else (datos.get('regimen', '') if datos.get('regimen') else '-')
        
        html_parts.append('<table class="data-table">')
        html_parts.append('<tr>')
        html_parts.append(f'<td style="width: 30%;"><strong>Nombre:</strong></td>')
        html_parts.append(f'<td>{nombre_alojamiento}</td>')
        html_parts.append('</tr>')
        html_parts.append('<tr>')
        html_parts.append(f'<td><strong>Habitación:</strong></td>')
        html_parts.append(f'<td>{tipo_habitacion}</td>')
        html_parts.append('</tr>')
        html_parts.append('<tr>')
        html_parts.append(f'<td><strong>Dirección:</strong></td>')
        html_parts.append(f'<td>{direccion}</td>')
        html_parts.append('</tr>')
        html_parts.append('<tr>')
        html_parts.append(f'<td><strong>Plan de Alimentos:</strong></td>')
        html_parts.append(f'<td>{plan_alimentos}</td>')
        html_parts.append('</tr>')
        html_parts.append('<tr>')
        html_parts.append(f'<td><strong>Referencia:</strong></td>')
        html_parts.append(f'<td>{referencia}</td>')
        html_parts.append('</tr>')
        if viajero_principal != '-':
            html_parts.append('<tr>')
            html_parts.append(f'<td><strong>Viajero Principal:</strong></td>')
            html_parts.append(f'<td><u>{viajero_principal}</u></td>')
            html_parts.append('</tr>')
        hora_checkin = datos.get('hora_checkin', '')
        hora_checkout = datos.get('hora_checkout', '')
        fecha_checkin = datos.get('fecha_checkin', '')
        fecha_checkout = datos.get('fecha_checkout', '')
        if fecha_checkin or fecha_checkout:
            html_parts.append('<tr>')
            html_parts.append('<td><strong>Fechas:</strong></td>')
            checkin_str = f"Check-in: {fecha_checkin}" + (f" {hora_checkin}" if hora_checkin else "")
            checkout_str = f"Check-out: {fecha_checkout}" + (f" {hora_checkout}" if hora_checkout else "")
            html_parts.append(f'<td>{checkin_str} | {checkout_str}</td>')
            html_parts.append('</tr>')
        
        adultos = datos.get('adultos', '0')
        ninos = datos.get('ninos', '0')
        ocupacion_str = f"{adultos} Adulto(s)"
        if int(ninos) > 0:
            ocupacion_str += f", {ninos} Niño(s)"
        regimen = datos.get('regimen', '')
        if ocupacion_str or regimen:
            html_parts.append('<tr>')
            html_parts.append('<td><strong>Huéspedes:</strong></td>')
            huéspedes_info = [ocupacion_str] if ocupacion_str else []
            if regimen:
                huéspedes_info.append(f"Régimen: {regimen}")
            html_parts.append(f'<td>{" | ".join(huéspedes_info)}</td>')
            html_parts.append('</tr>')
        
        if datos.get('observaciones'):
            html_parts.append('<tr>')
            html_parts.append('<td><strong>Observaciones:</strong></td>')
            html_parts.append(f'<td>{datos.get("observaciones")}</td>')
            html_parts.append('</tr>')
        
        html_parts.append('</table>')
        html_parts.append('</div>')  # Cierre de card
        
        # Imagen de hospedaje (si existe) - fuera de la card
        if datos.get('imagen_hospedaje_url'):
            try:
                image_url = datos['imagen_hospedaje_url']
                if image_url.startswith('/media/'):
                    image_url = request.build_absolute_uri(image_url)
                elif not image_url.startswith('http'):
                    image_url = request.build_absolute_uri(f'/media/{image_url}')
                html_parts.append(f'<img src="{image_url}" class="hospedaje-imagen" alt="Imagen del hospedaje">')
            except Exception:
                pass
        
        return "".join(html_parts)
    
    def _generar_html_traslado(self, datos, format_date):
        """Genera HTML para plantilla de traslado (EXACTAMENTE igual que cotizaciones con cards)."""
        html_parts = []
        
        # Card principal con header (IGUAL QUE COTIZACIONES)
        html_parts.append('<div class="card">')
        html_parts.append('<div class="card-header">')
        html_parts.append('<span class="icon">')
        html_parts.append('<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">')
        html_parts.append('<path d="M18.92 2.01C18.72 1.42 18.16 1 17.5 1h-11c-.66 0-1.21.42-1.42 1.01L3 8v8c0 .55.45 1 1 1h1c.55 0 1-.45 1-1v-1h12v1c0 .55.45 1 1 1h1c.55 0 1-.45 1-1V8l-2.08-5.99zM6.5 12c-.83 0-1.5-.67-1.5-1.5S5.67 9 6.5 9 8 9.67 8 10.5 7.33 12 6.5 12zm11 0c-.83 0-1.5-.67-1.5-1.5S16.67 9 17.5 9 19 9.67 19 10.5 18.33 12 17.5 12zM5 7l1.5-4.5h11L19 7H5z" fill="white" stroke="none"/>')
        html_parts.append('</svg>')
        html_parts.append('</span>')
        html_parts.append('<span>TRASLADO</span>')
        html_parts.append('</div>')
        
        # Tabla de datos (IGUAL QUE COTIZACIONES - data-table)
        html_parts.append('<table class="data-table">')
        
        if datos.get('compania'):
            html_parts.append('<tr>')
            html_parts.append('<td style="width: 30%;"><strong>Compañía:</strong></td>')
            html_parts.append(f'<td>{datos.get("compania")}</td>')
            html_parts.append('</tr>')
        
        if datos.get('codigo_reserva'):
            html_parts.append('<tr>')
            html_parts.append('<td><strong>Código Reserva:</strong></td>')
            html_parts.append(f'<td>{datos.get("codigo_reserva")}</td>')
            html_parts.append('</tr>')
        
        if datos.get('horario_inicio'):
            html_parts.append('<tr>')
            html_parts.append('<td><strong>Horario de Inicio:</strong></td>')
            html_parts.append(f'<td>{datos.get("horario_inicio")}</td>')
            html_parts.append('</tr>')
        
        if datos.get('tipo_servicio'):
            html_parts.append('<tr>')
            html_parts.append('<td><strong>Tipo de Servicio:</strong></td>')
            html_parts.append(f'<td>{datos.get("tipo_servicio")}</td>')
            html_parts.append('</tr>')
        
        ruta_info = []
        if datos.get('desde'):
            ruta_info.append(f"Desde: {datos.get('desde')}")
        if datos.get('hasta'):
            ruta_info.append(f"Hasta: {datos.get('hasta')}")
        if ruta_info:
            html_parts.append('<tr>')
            html_parts.append('<td><strong>Ruta:</strong></td>')
            html_parts.append(f'<td>{" | ".join(ruta_info)}</td>')
            html_parts.append('</tr>')
        
        adultos = datos.get('adultos', '0')
        ninos = datos.get('ninos', '0')
        pasajeros_str = f"{adultos} Adulto(s)"
        if int(ninos) > 0:
            pasajeros_str += f", {ninos} Niño(s)"
        html_parts.append('<tr>')
        html_parts.append('<td><strong>Pasajeros:</strong></td>')
        html_parts.append(f'<td>{pasajeros_str}</td>')
        html_parts.append('</tr>')
        
        if datos.get('informacion_adicional'):
            html_parts.append('<tr>')
            html_parts.append('<td><strong>Información Adicional:</strong></td>')
            html_parts.append(f'<td>{datos.get("informacion_adicional")}</td>')
            html_parts.append('</tr>')
        
        html_parts.append('</table>')
        html_parts.append('</div>')  # Cierre de card
        
        return "".join(html_parts)
    
    def _generar_html_generica(self, datos):
        """Genera HTML para plantilla genérica (EXACTAMENTE igual que cotizaciones con cards)."""
        html_parts = []
        
        # Card principal con header (IGUAL QUE COTIZACIONES)
        html_parts.append('<div class="card">')
        html_parts.append('<div class="card-header">')
        html_parts.append('<span class="icon">')
        html_parts.append('<svg viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">')
        html_parts.append('<path d="M19 3H5c-1.1 0-2 .9-2 2v14c0 1.1.9 2 2 2h14c1.1 0 2-.9 2-2V5c0-1.1-.9-2-2-2zm-5 14H7v-2h7v2zm3-4H7v-2h10v2zm0-4H7V7h10v2z" fill="white" stroke="none"/>')
        html_parts.append('</svg>')
        html_parts.append('</span>')
        titulo = datos.get('titulo', 'INFORMACIÓN ADICIONAL')
        html_parts.append(f'<span>{titulo.upper()}</span>')
        html_parts.append('</div>')
        
        contenido = datos.get('contenido', '')
        if contenido:
            contenido_normalizado = self._normalizar_valor_campo(contenido, limpiar_saltos_linea=True)
            contenido_html = contenido_normalizado.replace('\n', '<br>')
            html_parts.append(f'<div style="padding: 12px 18px; font-size: 9pt; line-height: 1.5;">{contenido_html}</div>')
        else:
            html_parts.append('<div style="padding: 12px 18px; font-size: 9pt; line-height: 1.5;">-</div>')
        
        html_parts.append('</div>')  # Cierre de card
        return "".join(html_parts)
    
    # ===== MÉTODOS ANTIGUOS (deprecated - mantenidos por compatibilidad pero ya no se usan) =====
    
    def _procesar_vuelo_unico(self, doc, datos, agregar_titulo_principal, agregar_subtitulo_con_vineta,
                              agregar_info_inline, agregar_info_line, agregar_salto_entre_secciones,
                              set_run_font, MOVUMS_BLUE_CORP, Pt):
        """Procesa plantilla de vuelo único usando funciones helper locales (mismo formato que cotizaciones)."""
        from docx.shared import Inches
        
        agregar_titulo_principal("VUELO ÚNICO")
        agregar_subtitulo_con_vineta('Información de Reserva')
        agregar_info_inline(
            ('Clave de Reserva', datos.get('clave_reserva', '')),
            ('Aerolínea', datos.get('aerolinea', '')),
            ('Vuelo', datos.get('numero_vuelo', ''))
        )
        agregar_salto_entre_secciones()
        agregar_subtitulo_con_vineta('Detalles del Vuelo')
        agregar_info_inline(
            ('Fecha', datos.get('fecha', '')),
            ('Hora Salida', datos.get('hora_salida', '')),
            ('Hora Llegada', datos.get('hora_llegada', ''))
        )
        agregar_info_inline(
            ('Origen', datos.get('origen_terminal', '')),
            ('Destino', datos.get('destino_terminal', ''))
        )
        agregar_info_line('Tipo de Vuelo', datos.get('tipo_vuelo', ''))
        
        # Escalas si aplica
        if datos.get('tipo_vuelo') == 'Escalas' and datos.get('escalas'):
            escalas = datos.get('escalas', [])
            if escalas:
                agregar_salto_entre_secciones()
                agregar_subtitulo_con_vineta('Detalles de Escalas')
                for i, escala in enumerate(escalas, 1):
                    escala_p = doc.add_paragraph()
                    escala_p.paragraph_format.space_after = Pt(6)
                    escala_run = escala_p.add_run(f'Escala {i}: ')
                    set_run_font(escala_run, size=12, bold=True)
                    escala_val_run = escala_p.add_run(f"{escala.get('ciudad', '')} - {escala.get('aeropuerto', '')}")
                    set_run_font(escala_val_run, size=12)
                    
                    escala_info = doc.add_paragraph()
                    escala_info.paragraph_format.left_indent = Inches(0.3)
                    escala_info.paragraph_format.space_after = Pt(6)
                    detalles = []
                    if escala.get('hora_llegada'):
                        detalles.append(f"Hora Llegada: {escala.get('hora_llegada')}")
                    if escala.get('hora_salida'):
                        detalles.append(f"Hora Salida: {escala.get('hora_salida')}")
                    if escala.get('numero_vuelo'):
                        detalles.append(f"Vuelo: {escala.get('numero_vuelo')}")
                    if escala.get('duracion'):
                        detalles.append(f"Duración: {escala.get('duracion')}")
                    detalle_run = escala_info.add_run(' | '.join(detalles))
                    set_run_font(detalle_run, size=12)
        
        agregar_salto_entre_secciones()
        agregar_subtitulo_con_vineta('Información de Pasajeros')
        pasajeros = datos.get('pasajeros', '')
        if pasajeros:
            pasajeros_texto = ', '.join([p.strip() for p in str(pasajeros).replace('\r\n', '\n').split('\n') if p.strip()])
            agregar_info_line('Pasajeros', pasajeros_texto)
        agregar_info_line('Equipaje', datos.get('equipaje', ''))
        if datos.get('informacion_adicional'):
            agregar_info_line('Información Adicional', datos.get('informacion_adicional', ''))
        
        # Información Completa del Vuelo (campo "vuelo") - en hoja aparte si existe
        if datos.get('vuelo'):
            from docx.enum.text import WD_BREAK
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            agregar_subtitulo_con_vineta('Información Completa del Vuelo')
            vuelo_info = datos.get('vuelo', '')
            for linea in vuelo_info.split('\n'):
                if linea.strip():
                    p = doc.add_paragraph(linea.strip())
                    p.paragraph_format.space_after = Pt(4)
                    for run in p.runs:
                        set_run_font(run, size=12)
    
    def _procesar_vuelo_redondo(self, doc, datos, agregar_titulo_principal, agregar_subtitulo_con_vineta,
                                agregar_info_inline, agregar_info_line, agregar_salto_entre_secciones,
                                set_run_font, MOVUMS_BLUE_CORP, Pt, RGBColor):
        """Procesa plantilla de vuelo redondo usando funciones helper locales (mismo formato que cotizaciones)."""
        from docx.shared import Inches
        
        agregar_titulo_principal("VUELO REDONDO")
        agregar_subtitulo_con_vineta('Información de Reserva')
        agregar_info_line('Clave de Reserva', datos.get('clave_reserva', ''))
        agregar_salto_entre_secciones()
        agregar_subtitulo_con_vineta('Vuelo de Ida')
        agregar_info_inline(
            ('Aerolínea', datos.get('aerolinea_ida', '')),
            ('Vuelo', datos.get('numero_vuelo_ida', ''))
        )
        agregar_info_inline(
            ('Fecha Salida', datos.get('fecha_salida_ida', '')),
            ('Hora Salida', datos.get('hora_salida_ida', '')),
            ('Hora Llegada', datos.get('hora_llegada_ida', ''))
        )
        agregar_info_inline(
            ('Origen', datos.get('origen_ida', '')),
            ('Destino', datos.get('destino_ida', ''))
        )
        agregar_info_line('Tipo de Vuelo', datos.get('tipo_vuelo_ida', ''))
        
        # Escalas de Ida si aplica
        escalas_ida = datos.get('escalas_ida', [])
        if escalas_ida and isinstance(escalas_ida, list) and len(escalas_ida) > 0:
            agregar_salto_entre_secciones()
            p_titulo = doc.add_paragraph()
            p_titulo.paragraph_format.space_before = Pt(8)
            p_titulo.paragraph_format.space_after = Pt(4)
            titulo_run = p_titulo.add_run('Escalas del Vuelo de Ida:')
            set_run_font(titulo_run, size=12, bold=True, color=MOVUMS_BLUE_CORP)
            
            for i, escala in enumerate(escalas_ida, 1):
                escala_p = doc.add_paragraph()
                escala_p.paragraph_format.space_after = Pt(4)
                escala_p.paragraph_format.left_indent = Inches(0.2)
                escala_run = escala_p.add_run(f'Escala {i}: ')
                set_run_font(escala_run, size=12, bold=True)
                ciudad = escala.get('ciudad', '')
                aeropuerto = escala.get('aeropuerto', '')
                escala_val = escala_p.add_run(f"{ciudad} - {aeropuerto}")
                set_run_font(escala_val, size=12)
                
                detalle_p = doc.add_paragraph()
                detalle_p.paragraph_format.space_after = Pt(6)
                detalle_p.paragraph_format.left_indent = Inches(0.4)
                detalles = []
                if escala.get('hora_llegada'):
                    detalles.append(f"Llegada: {escala.get('hora_llegada')}")
                if escala.get('hora_salida'):
                    detalles.append(f"Salida: {escala.get('hora_salida')}")
                if escala.get('numero_vuelo'):
                    detalles.append(f"Vuelo: {escala.get('numero_vuelo')}")
                if escala.get('duracion'):
                    detalles.append(f"Duración: {escala.get('duracion')}")
                detalle_run = detalle_p.add_run(' | '.join(detalles))
                set_run_font(detalle_run, size=11)
        
        # Salto de página antes del Vuelo de Regreso
        doc.add_page_break()
        
        agregar_subtitulo_con_vineta('Vuelo de Regreso')
        agregar_info_inline(
            ('Aerolínea', datos.get('aerolinea_regreso', '')),
            ('Vuelo', datos.get('numero_vuelo_regreso', ''))
        )
        agregar_info_inline(
            ('Fecha Salida', datos.get('fecha_salida_regreso', '')),
            ('Hora Salida', datos.get('hora_salida_regreso', '')),
            ('Hora Llegada', datos.get('hora_llegada_regreso', ''))
        )
        agregar_info_inline(
            ('Origen', datos.get('origen_regreso', '')),
            ('Destino', datos.get('destino_regreso', ''))
        )
        agregar_info_line('Tipo de Vuelo', datos.get('tipo_vuelo_regreso', ''))
        
        # Escalas de Regreso si aplica
        escalas_regreso = datos.get('escalas_regreso', [])
        if escalas_regreso and isinstance(escalas_regreso, list) and len(escalas_regreso) > 0:
            agregar_salto_entre_secciones()
            p_titulo = doc.add_paragraph()
            p_titulo.paragraph_format.space_before = Pt(8)
            p_titulo.paragraph_format.space_after = Pt(4)
            titulo_run = p_titulo.add_run('Escalas del Vuelo de Regreso:')
            set_run_font(titulo_run, size=12, bold=True, color=MOVUMS_BLUE_CORP)
            
            for i, escala in enumerate(escalas_regreso, 1):
                escala_p = doc.add_paragraph()
                escala_p.paragraph_format.space_after = Pt(4)
                escala_p.paragraph_format.left_indent = Inches(0.2)
                escala_run = escala_p.add_run(f'Escala {i}: ')
                set_run_font(escala_run, size=12, bold=True)
                ciudad = escala.get('ciudad', '')
                aeropuerto = escala.get('aeropuerto', '')
                escala_val = escala_p.add_run(f"{ciudad} - {aeropuerto}")
                set_run_font(escala_val, size=12)
                
                detalle_p = doc.add_paragraph()
                detalle_p.paragraph_format.space_after = Pt(6)
                detalle_p.paragraph_format.left_indent = Inches(0.4)
                detalles = []
                if escala.get('hora_llegada'):
                    detalles.append(f"Llegada: {escala.get('hora_llegada')}")
                if escala.get('hora_salida'):
                    detalles.append(f"Salida: {escala.get('hora_salida')}")
                if escala.get('numero_vuelo'):
                    detalles.append(f"Vuelo: {escala.get('numero_vuelo')}")
                if escala.get('duracion'):
                    detalles.append(f"Duración: {escala.get('duracion')}")
                detalle_run = detalle_p.add_run(' | '.join(detalles))
                set_run_font(detalle_run, size=11)
        
        agregar_salto_entre_secciones()
        agregar_subtitulo_con_vineta('Información General')
        pasajeros = datos.get('pasajeros', '')
        if pasajeros:
            pasajeros_texto = ', '.join([p.strip() for p in str(pasajeros).replace('\r\n', '\n').split('\n') if p.strip()])
            agregar_info_line('Pasajeros', pasajeros_texto)
        agregar_info_line('Equipaje', datos.get('equipaje', ''))
        if datos.get('informacion_adicional'):
            agregar_info_line('Información Adicional', datos.get('informacion_adicional', ''))
    
    def _procesar_hospedaje(self, doc, datos, agregar_titulo_principal, agregar_subtitulo_con_vineta,
                           agregar_info_inline, agregar_info_line, agregar_salto_entre_secciones,
                           set_run_font, MOVUMS_BLUE_CORP, Pt, Inches, request):
        """Procesa plantilla de hospedaje usando funciones helper locales (mismo formato que cotizaciones)."""
        agregar_titulo_principal("HOSPEDAJE")
        agregar_subtitulo_con_vineta('Información del Alojamiento')
        
        # Crear tabla de 2 columnas como en cotizaciones (líneas 7026-7056)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.autofit = False
        for col in info_table.columns:
            for cell in col.cells:
                cell.width = Inches(3.25)
        
        # Normalizar valores
        nombre_alojamiento = self._normalizar_valor_campo(datos.get('nombre_alojamiento', ''), limpiar_saltos_linea=True)
        referencia = self._normalizar_valor_campo(datos.get('numero_referencia', ''), limpiar_saltos_linea=True)
        viajero_principal = self._normalizar_valor_campo(datos.get('viajero_principal', ''), es_nombre_propio=True)
        tipo_habitacion = self._normalizar_valor_campo(datos.get('tipo_habitacion', ''), limpiar_saltos_linea=True)
        direccion = self._normalizar_valor_campo(datos.get('direccion', ''), limpiar_saltos_linea=True) if datos.get('direccion') else '-'
        plan_alimentos = self._normalizar_valor_campo(datos.get('plan_alimentos', ''), limpiar_saltos_linea=True) if datos.get('plan_alimentos') else (datos.get('regimen', '') if datos.get('regimen') else '-')
        
        # Nombre (fila 0, columna 0)
        nombre_cell = info_table.rows[0].cells[0]
        nombre_label = nombre_cell.paragraphs[0].add_run('Nombre: ')
        set_run_font(nombre_label, size=12, bold=True)
        nombre_val = nombre_cell.paragraphs[0].add_run(nombre_alojamiento or 'Hotel propuesto')
        set_run_font(nombre_val, size=12)
        
        # Habitación (fila 1, columna 0)
        habitacion_cell = info_table.rows[1].cells[0]
        habitacion_label = habitacion_cell.paragraphs[0].add_run('Habitación: ')
        set_run_font(habitacion_label, size=12, bold=True)
        habitacion_val = habitacion_cell.paragraphs[0].add_run(tipo_habitacion or '-')
        set_run_font(habitacion_val, size=12)
        
        # Dirección (fila 2, columna 0)
        direccion_cell = info_table.rows[2].cells[0]
        direccion_label = direccion_cell.paragraphs[0].add_run('Dirección: ')
        set_run_font(direccion_label, size=12, bold=True)
        direccion_val = direccion_cell.paragraphs[0].add_run(direccion)
        set_run_font(direccion_val, size=12)
        
        # Viajero Principal (fila 3, columna 0)
        viajero_cell = info_table.rows[3].cells[0]
        viajero_label = viajero_cell.paragraphs[0].add_run('Viajero Principal: ')
        set_run_font(viajero_label, size=12, bold=True)
        viajero_val = viajero_cell.paragraphs[0].add_run(viajero_principal or '-')
        set_run_font(viajero_val, size=12)
        if viajero_principal:
            viajero_val.font.underline = True
        
        # Plan de Alimentos (fila 0, columna 1)
        plan_cell = info_table.rows[0].cells[1]
        plan_label = plan_cell.paragraphs[0].add_run('Plan de Alimentos: ')
        set_run_font(plan_label, size=12, bold=True)
        plan_val = plan_cell.paragraphs[0].add_run(plan_alimentos)
        set_run_font(plan_val, size=12)
        
        # Referencia (fila 1, columna 1)
        ref_cell = info_table.rows[1].cells[1]
        ref_label = ref_cell.paragraphs[0].add_run('Referencia: ')
        set_run_font(ref_label, size=12, bold=True)
        ref_val = ref_cell.paragraphs[0].add_run(referencia or '-')
        set_run_font(ref_val, size=12)
        
        agregar_salto_entre_secciones()
        agregar_subtitulo_con_vineta('Fechas y Estancia')
        hora_checkin = datos.get('hora_checkin', '')
        hora_checkout = datos.get('hora_checkout', '')
        fecha_checkin = datos.get('fecha_checkin', '')
        fecha_checkout = datos.get('fecha_checkout', '')
        checkin_str = f"{fecha_checkin}" + (f" {hora_checkin}" if hora_checkin else "")
        checkout_str = f"{fecha_checkout}" + (f" {hora_checkout}" if hora_checkout else "")
        agregar_info_inline(
            ('Check-in', checkin_str),
            ('Check-out', checkout_str)
        )
        
        agregar_salto_entre_secciones()
        agregar_subtitulo_con_vineta('Información de Huéspedes')
        adultos = datos.get('adultos', '0')
        ninos = datos.get('ninos', '0')
        ocupacion_str = f"{adultos} Adulto(s)"
        if int(ninos) > 0:
            ocupacion_str += f", {ninos} Niño(s)"
        agregar_info_inline(
            ('Ocupación', ocupacion_str),
            ('Régimen', datos.get('regimen', ''))
        )
        if datos.get('observaciones'):
            agregar_info_line('Observaciones', datos.get('observaciones', ''))
        
        # Imagen de hospedaje (si existe)
        if datos.get('imagen_hospedaje_url'):
            try:
                import requests
                from io import BytesIO
                image_url = datos['imagen_hospedaje_url']
                if image_url.startswith('/media/'):
                    image_url = request.build_absolute_uri(image_url)
                response = requests.get(image_url, timeout=5)
                response.raise_for_status()
                image_stream = BytesIO(response.content)
                doc.add_picture(image_stream, width=Inches(4.5))
            except Exception:
                pass
    
    def _procesar_traslado(self, doc, datos, agregar_titulo_principal, agregar_subtitulo_con_vineta,
                          agregar_info_inline, agregar_info_line, agregar_salto_entre_secciones,
                          set_run_font, MOVUMS_BLUE_CORP, Pt, format_date):
        """Procesa plantilla de traslado usando funciones helper locales (mismo formato que cotizaciones)."""
        agregar_titulo_principal("TRASLADO")
        agregar_subtitulo_con_vineta('Información de la Compañía')
        agregar_info_inline(
            ('Compañía', datos.get('compania', '')),
            ('Código Reserva', datos.get('codigo_reserva', ''))
        )
        agregar_salto_entre_secciones()
        agregar_subtitulo_con_vineta('Detalles del Traslado')
        agregar_info_line('Horario de Inicio de Viaje', datos.get('horario_inicio', ''))
        agregar_info_line('Tipo de Servicio', datos.get('tipo_servicio', ''))
        agregar_info_inline(
            ('Desde', datos.get('desde', '')),
            ('Hasta', datos.get('hasta', ''))
        )
        
        # Información de Pasajeros
        adultos = datos.get('adultos', '0')
        ninos = datos.get('ninos', '0')
        pasajeros_str = f"{adultos} Adulto(s)"
        if int(ninos) > 0:
            pasajeros_str += f", {ninos} Niño(s)"
        agregar_info_line('Pasajeros', pasajeros_str)
        
        if datos.get('informacion_adicional'):
            agregar_info_line('Información Adicional', datos.get('informacion_adicional', ''))
    
    def _procesar_generica(self, doc, datos, agregar_titulo_principal, agregar_subtitulo_con_vineta,
                          agregar_info_line, agregar_salto_entre_secciones, set_run_font, Pt):
        """Procesa plantilla genérica usando funciones helper locales (mismo formato que cotizaciones)."""
        titulo = datos.get('titulo', 'Información Adicional')
        if titulo:
            agregar_titulo_principal(titulo.upper())
        
        contenido = datos.get('contenido', '')
        if contenido:
            contenido_normalizado = self._normalizar_valor_campo(contenido, limpiar_saltos_linea=True)
            for linea in contenido_normalizado.split('\n'):
                if linea.strip():
                    p = doc.add_paragraph(linea.strip())
                    p.paragraph_format.space_after = Pt(4)
                    for run in p.runs:
                        set_run_font(run, size=12)
    
    # ===== MÉTODOS ANTIGUOS (deprecated - se mantienen por compatibilidad pero ya no se usan) =====
    
    def _agregar_info_line(self, doc, etiqueta, valor, mostrar_si_vacio=False, es_nombre_propio=False, separar_con_comas=False):
        """Helper para agregar una línea de información formateada (mismo formato que cotizaciones)."""
        from docx.shared import Pt, RGBColor
        
        if not valor and not mostrar_si_vacio:
            return
        
        # Usar la función helper definida en get()
        set_run_font = getattr(self, '_set_run_font', None)
        TEXT_COLOR = getattr(self, '_TEXT_COLOR', RGBColor(20, 20, 20))
        
        # Si no existe la función helper, definirla localmente
        if not set_run_font:
            def set_run_font(run, size=12, bold=False, color=TEXT_COLOR):
                run.font.name = 'Arial'
                run.font.size = Pt(size)
                run.bold = bold
                run.font.color.rgb = color
        
        # Si separar_con_comas, convertir saltos de línea en comas
        if separar_con_comas and valor:
            # Reemplazar saltos de línea por comas
            valor = ', '.join([v.strip() for v in str(valor).replace('\r\n', '\n').split('\n') if v.strip()])
        
        # Normalizar el valor
        valor_normalizado = self._normalizar_valor_campo(
            valor, 
            es_nombre_propio=es_nombre_propio, 
            limpiar_saltos_linea=True
        )
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)  # Espaciado entre líneas de información
        p.paragraph_format.line_spacing = 1.15  # Interlineado
        
        # Etiqueta en negrita
        label_run = p.add_run(f'{etiqueta}: ')
        set_run_font(label_run, size=12, bold=True)
        
        # Valor normalizado
        value_run = p.add_run(valor_normalizado if valor_normalizado else 'No especificado')
        set_run_font(value_run, size=12)
        
        return p
    
    def _agregar_info_inline(self, doc, *pares_etiqueta_valor, separador=' | ', es_nombre_propio=False):
        """Helper para agregar múltiples campos en una sola línea (mismo formato que cotizaciones)."""
        from docx.shared import Pt, RGBColor
        
        if not pares_etiqueta_valor:
            return
        
        # Usar la función helper definida en get()
        set_run_font = getattr(self, '_set_run_font', None)
        TEXT_COLOR = getattr(self, '_TEXT_COLOR', RGBColor(20, 20, 20))
        
        # Si no existe la función helper, definirla localmente
        if not set_run_font:
            def set_run_font(run, size=12, bold=False, color=TEXT_COLOR):
                run.font.name = 'Arial'
                run.font.size = Pt(size)
                run.bold = bold
                run.font.color.rgb = color
        
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        
        for idx, (etiqueta, valor) in enumerate(pares_etiqueta_valor):
            if not valor:
                continue
            
            # Normalizar el valor
            valor_normalizado = self._normalizar_valor_campo(
                valor,
                es_nombre_propio=es_nombre_propio,
                limpiar_saltos_linea=True
            )
            
            if idx > 0:
                sep_run = p.add_run(separador)
                set_run_font(sep_run, size=12)
            
            label_run = p.add_run(f'{etiqueta}: ')
            set_run_font(label_run, size=12, bold=True)
            
            value_run = p.add_run(str(valor_normalizado))
            set_run_font(value_run, size=12)
        
        return p
    
    def _agregar_subtitulo_con_vineta(self, doc, texto):
        """Agrega un subtítulo con viñeta azul (mismo formato que cotizaciones)."""
        from docx.shared import Pt, RGBColor
        
        # Usar el color corporativo definido en get() (con valor por defecto)
        MOVUMS_BLUE_CORP = getattr(self, '_MOVUMS_BLUE_CORP', None) or RGBColor(0, 74, 142)
        set_run_font = getattr(self, '_set_run_font', None)
        TEXT_COLOR = getattr(self, '_TEXT_COLOR', None) or RGBColor(20, 20, 20)
        
        # Si no existe la función helper, definirla localmente
        if not set_run_font:
            def set_run_font(run, size=12, bold=False, color=TEXT_COLOR):
                run.font.name = 'Arial'
                run.font.size = Pt(size)
                run.bold = bold
                run.font.color.rgb = color
        
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        
        # Agregar viñeta manualmente (carácter bullet) - tamaño 14 como cotizaciones
        bullet_run = p.add_run('• ')
        set_run_font(bullet_run, size=14, bold=True, color=MOVUMS_BLUE_CORP)
        
        # Agregar texto del subtítulo (tamaño 14 como cotizaciones)
        texto_run = p.add_run(texto)
        set_run_font(texto_run, size=14, bold=True, color=MOVUMS_BLUE_CORP)
        
        # Agregar espaciado después como en cotizaciones
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(2)
        
        return p
    
    def _agregar_salto_entre_secciones(self, doc):
        """Agrega un salto de línea entre secciones para separar títulos con viñeta."""
        from docx.shared import Pt
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)
        return spacer
    
    def _agregar_vuelo_unico(self, doc, datos):
        """Agrega contenido de vuelo único al documento con formato profesional (mismo formato que cotizaciones)."""
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Color corporativo (mismo que cotizaciones)
        MOVUMS_BLUE_CORP = RGBColor(0, 74, 142)
        
        # Información de Reserva (con viñeta)
        self._agregar_subtitulo_con_vineta(doc, 'Información de Reserva')
        
        # Agrupar información de reserva en líneas compactas
        self._agregar_info_inline(doc,
            ('Clave de Reserva', datos.get('clave_reserva', '')),
            ('Aerolínea', datos.get('aerolinea', '')),
            ('Vuelo', datos.get('numero_vuelo', ''))
        )
        
        # Salto de línea entre secciones
        self._agregar_salto_entre_secciones(doc)
        
        # Detalles del Vuelo (con viñeta)
        self._agregar_subtitulo_con_vineta(doc, 'Detalles del Vuelo')
        
        # Agrupar campos relacionados en la misma línea
        self._agregar_info_inline(doc, 
            ('Fecha', datos.get('fecha', '')),
            ('Hora Salida', datos.get('hora_salida', '')),
            ('Hora Llegada', datos.get('hora_llegada', ''))
        )
        self._agregar_info_inline(doc,
            ('Origen', datos.get('origen_terminal', '')),
            ('Destino', datos.get('destino_terminal', ''))
        )
        self._agregar_info_line(doc, 'Tipo de Vuelo', datos.get('tipo_vuelo', ''))
        
        # Escalas si aplica (compacto)
        if datos.get('tipo_vuelo') == 'Escalas' and datos.get('escalas'):
            self._asegurar_estilo_heading(doc, 'Heading 4', 12)
            escalas_titulo = doc.add_heading('Detalles de Escalas', level=4)
            escalas_titulo.paragraph_format.space_before = Pt(4)
            escalas_titulo.paragraph_format.space_after = Pt(6)
            escalas_titulo.runs[0].font.size = Pt(10)
            for i, escala in enumerate(datos['escalas'], 1):
                escala_p = doc.add_paragraph()
                escala_p.paragraph_format.space_after = Pt(6)
                escala_run = escala_p.add_run(f'Escala {i}: ')
                escala_run.font.name = 'Arial'
                escala_run.font.size = Pt(12)
                escala_run.bold = True
                escala_val_run = escala_p.add_run(f"{escala.get('ciudad', '')} - {escala.get('aeropuerto', '')}")
                escala_val_run.font.name = 'Arial'
                escala_val_run.font.size = Pt(12)
                
                escala_info = doc.add_paragraph()
                escala_info.paragraph_format.left_indent = Inches(0.3)
                escala_info.paragraph_format.space_after = Pt(6)
                for part in [
                    (f"Hora Llegada: {escala.get('hora_llegada', '')}", True),
                    (" | ", False),
                    (f"Hora Salida: {escala.get('hora_salida', '')}", True),
                    (" | ", False),
                    (f"Vuelo: {escala.get('numero_vuelo', '')}", True),
                    (" | ", False),
                    (f"Duración: {escala.get('duracion', '')}", False)
                ]:
                    run = escala_info.add_run(part[0])
                    run.font.name = 'Arial'
                    run.font.size = Pt(12)
                    if part[1]:
                        run.bold = True
        
        # Salto de línea entre secciones
        self._agregar_salto_entre_secciones(doc)
        
        # Información de Pasajeros (con viñeta)
        self._agregar_subtitulo_con_vineta(doc, 'Información de Pasajeros')
        
        self._agregar_info_line(doc, 'Pasajeros', datos.get('pasajeros', ''), separar_con_comas=True)
        self._agregar_info_line(doc, 'Equipaje', datos.get('equipaje', ''))
        
        if datos.get('informacion_adicional'):
            info_normalizada = self._normalizar_valor_campo(datos.get('informacion_adicional', ''), limpiar_saltos_linea=True)
            info_p = doc.add_paragraph()
            info_p.paragraph_format.space_after = Pt(6)
            info_label = info_p.add_run('Información Adicional: ')
            info_label.font.name = 'Arial'
            info_label.font.size = Pt(12)
            info_label.bold = True
            info_val = info_p.add_run(info_normalizada)
            info_val.font.name = 'Arial'
            info_val.font.size = Pt(12)
        
        # Información Completa del Vuelo (campo "vuelo") - en hoja aparte
        if datos.get('vuelo'):
            from docx.enum.text import WD_BREAK
            # Agregar salto de página
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            self._agregar_subtitulo_con_vineta(doc, 'Información Completa del Vuelo')
            vuelo_info = self._normalizar_valor_campo(datos.get('vuelo', ''), limpiar_saltos_linea=False)
            vuelo_p = doc.add_paragraph()
            vuelo_p.paragraph_format.space_after = Pt(6)
            vuelo_val = vuelo_p.add_run(vuelo_info)
            vuelo_val.font.name = 'Arial'
            vuelo_val.font.size = Pt(12)
    
    def _agregar_vuelo_redondo(self, doc, datos):
        """Agrega contenido de vuelo redondo al documento con formato compacto (mismo formato que cotizaciones)."""
        from docx.shared import Pt, Inches
        from docx.shared import RGBColor
        
        # Color corporativo (mismo que cotizaciones)
        MOVUMS_BLUE_CORP = RGBColor(0, 74, 142)
        
        # Información de Reserva (con viñeta)
        self._agregar_subtitulo_con_vineta(doc, 'Información de Reserva')
        
        self._agregar_info_line(doc, 'Clave de Reserva', datos.get('clave_reserva', ''))
        
        # Salto de línea entre secciones
        self._agregar_salto_entre_secciones(doc)
        
        # Vuelo de Ida (con viñeta)
        self._agregar_subtitulo_con_vineta(doc, 'Vuelo de Ida')
        
        # Agrupar campos relacionados en líneas compactas
        self._agregar_info_inline(doc,
            ('Aerolínea', datos.get('aerolinea_ida', '')),
            ('Vuelo', datos.get('numero_vuelo_ida', ''))
        )
        self._agregar_info_inline(doc,
            ('Fecha Salida', datos.get('fecha_salida_ida', '')),
            ('Hora Salida', datos.get('hora_salida_ida', '')),
            ('Hora Llegada', datos.get('hora_llegada_ida', ''))
        )
        self._agregar_info_inline(doc,
            ('Origen', datos.get('origen_ida', '')),
            ('Destino', datos.get('destino_ida', ''))
        )
        self._agregar_info_line(doc, 'Tipo de Vuelo', datos.get('tipo_vuelo_ida', ''))
        
        # Escalas de Ida si aplica
        escalas_ida = datos.get('escalas_ida', [])
        if escalas_ida and isinstance(escalas_ida, list) and len(escalas_ida) > 0:
                p_titulo = doc.add_paragraph()
                p_titulo.paragraph_format.space_before = Pt(8)
                p_titulo.paragraph_format.space_after = Pt(4)
                titulo_run = p_titulo.add_run('Escalas del Vuelo de Ida:')
                titulo_run.font.name = 'Arial'
                titulo_run.font.size = Pt(12)
                titulo_run.font.bold = True
                titulo_run.font.color.rgb = MOVUMS_BLUE_CORP
                
                for i, escala in enumerate(escalas_ida, 1):
                    escala_p = doc.add_paragraph()
                    escala_p.paragraph_format.space_after = Pt(4)
                    escala_p.paragraph_format.left_indent = Inches(0.2)
                    
                    escala_run = escala_p.add_run(f'Escala {i}: ')
                    escala_run.font.name = 'Arial'
                    escala_run.font.size = Pt(12)
                    escala_run.bold = True
                    
                    ciudad = escala.get('ciudad', '')
                    aeropuerto = escala.get('aeropuerto', '')
                    escala_val = escala_p.add_run(f"{ciudad} - {aeropuerto}")
                    escala_val.font.name = 'Arial'
                    escala_val.font.size = Pt(12)
                    
                    # Detalles de la escala
                    detalle_p = doc.add_paragraph()
                    detalle_p.paragraph_format.space_after = Pt(6)
                    detalle_p.paragraph_format.left_indent = Inches(0.4)
                    
                    detalles = []
                    if escala.get('hora_llegada'):
                        detalles.append(f"Llegada: {escala.get('hora_llegada')}")
                    if escala.get('hora_salida'):
                        detalles.append(f"Salida: {escala.get('hora_salida')}")
                    if escala.get('numero_vuelo'):
                        detalles.append(f"Vuelo: {escala.get('numero_vuelo')}")
                    if escala.get('duracion'):
                        detalles.append(f"Duración: {escala.get('duracion')}")
                    
                    detalle_run = detalle_p.add_run(' | '.join(detalles))
                    detalle_run.font.name = 'Arial'
                    detalle_run.font.size = Pt(11)
        
        # Salto de página antes del Vuelo de Regreso
        doc.add_page_break()
        
        # Vuelo de Regreso (con viñeta)
        self._agregar_subtitulo_con_vineta(doc, 'Vuelo de Regreso')
        
        # Agrupar campos relacionados en líneas compactas
        self._agregar_info_inline(doc,
            ('Aerolínea', datos.get('aerolinea_regreso', '')),
            ('Vuelo', datos.get('numero_vuelo_regreso', ''))
        )
        self._agregar_info_inline(doc,
            ('Fecha Salida', datos.get('fecha_salida_regreso', '')),
            ('Hora Salida', datos.get('hora_salida_regreso', '')),
            ('Hora Llegada', datos.get('hora_llegada_regreso', ''))
        )
        self._agregar_info_inline(doc,
            ('Origen', datos.get('origen_regreso', '')),
            ('Destino', datos.get('destino_regreso', ''))
        )
        self._agregar_info_line(doc, 'Tipo de Vuelo', datos.get('tipo_vuelo_regreso', ''))
        
        # Escalas de Regreso si aplica
        escalas_regreso = datos.get('escalas_regreso', [])
        if escalas_regreso and isinstance(escalas_regreso, list) and len(escalas_regreso) > 0:
                p_titulo = doc.add_paragraph()
                p_titulo.paragraph_format.space_before = Pt(8)
                p_titulo.paragraph_format.space_after = Pt(4)
                titulo_run = p_titulo.add_run('Escalas del Vuelo de Regreso:')
                titulo_run.font.name = 'Arial'
                titulo_run.font.size = Pt(12)
                titulo_run.font.bold = True
                titulo_run.font.color.rgb = MOVUMS_BLUE_CORP
                
                for i, escala in enumerate(escalas_regreso, 1):
                    escala_p = doc.add_paragraph()
                    escala_p.paragraph_format.space_after = Pt(4)
                    escala_p.paragraph_format.left_indent = Inches(0.2)
                    
                    escala_run = escala_p.add_run(f'Escala {i}: ')
                    escala_run.font.name = 'Arial'
                    escala_run.font.size = Pt(12)
                    escala_run.bold = True
                    
                    ciudad = escala.get('ciudad', '')
                    aeropuerto = escala.get('aeropuerto', '')
                    escala_val = escala_p.add_run(f"{ciudad} - {aeropuerto}")
                    escala_val.font.name = 'Arial'
                    escala_val.font.size = Pt(12)
                    
                    # Detalles de la escala
                    detalle_p = doc.add_paragraph()
                    detalle_p.paragraph_format.space_after = Pt(6)
                    detalle_p.paragraph_format.left_indent = Inches(0.4)
                    
                    detalles = []
                    if escala.get('hora_llegada'):
                        detalles.append(f"Llegada: {escala.get('hora_llegada')}")
                    if escala.get('hora_salida'):
                        detalles.append(f"Salida: {escala.get('hora_salida')}")
                    if escala.get('numero_vuelo'):
                        detalles.append(f"Vuelo: {escala.get('numero_vuelo')}")
                    if escala.get('duracion'):
                        detalles.append(f"Duración: {escala.get('duracion')}")
                    
                    detalle_run = detalle_p.add_run(' | '.join(detalles))
                    detalle_run.font.name = 'Arial'
                    detalle_run.font.size = Pt(11)
        
        # Salto de línea entre secciones
        self._agregar_salto_entre_secciones(doc)
        
        # Información General (con viñeta)
        self._agregar_subtitulo_con_vineta(doc, 'Información General')
        
        self._agregar_info_line(doc, 'Pasajeros', datos.get('pasajeros', ''), es_nombre_propio=True, separar_con_comas=True)
        self._agregar_info_line(doc, 'Equipaje', datos.get('equipaje', ''))
        
        if datos.get('informacion_adicional'):
            info_normalizada = self._normalizar_valor_campo(datos.get('informacion_adicional', ''), limpiar_saltos_linea=True)
            info_p = doc.add_paragraph()
            info_p.paragraph_format.space_after = Pt(6)
            info_label = info_p.add_run('Información Adicional: ')
            info_label.font.name = 'Arial'
            info_label.font.size = Pt(12)
            info_label.bold = True
            info_val = info_p.add_run(info_normalizada)
            info_val.font.name = 'Arial'
            info_val.font.size = Pt(12)
    
    def _agregar_hospedaje(self, doc, datos):
        """Agrega contenido de hospedaje al documento con formato compacto (mismo formato que cotizaciones)."""
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor
        
        # Color corporativo (mismo que cotizaciones)
        MOVUMS_BLUE_CORP = RGBColor(0, 74, 142)
        
        # Información del Alojamiento (con viñeta)
        self._agregar_subtitulo_con_vineta(doc, 'Información del Alojamiento')
        
        # Agrupar información del alojamiento (normalizada)
        nombre_alojamiento = self._normalizar_valor_campo(datos.get('nombre_alojamiento', ''), limpiar_saltos_linea=True)
        referencia = self._normalizar_valor_campo(datos.get('numero_referencia', ''), limpiar_saltos_linea=True)
        viajero_principal = self._normalizar_valor_campo(datos.get('viajero_principal', ''), es_nombre_propio=True)
        tipo_habitacion = self._normalizar_valor_campo(datos.get('tipo_habitacion', ''), limpiar_saltos_linea=True)
        
        # Crear línea con viajero principal subrayado
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        
        # Alojamiento
        aloj_run = p.add_run('Alojamiento: ')
        aloj_run.font.name = 'Arial'
        aloj_run.font.size = Pt(12)
        aloj_run.font.bold = True
        aloj_val_run = p.add_run(f'{nombre_alojamiento} | ')
        aloj_val_run.font.name = 'Arial'
        aloj_val_run.font.size = Pt(12)
        
        # Referencia
        ref_run = p.add_run('Referencia: ')
        ref_run.font.name = 'Arial'
        ref_run.font.size = Pt(12)
        ref_run.font.bold = True
        ref_val_run = p.add_run(f'{referencia} | ')
        ref_val_run.font.name = 'Arial'
        ref_val_run.font.size = Pt(12)
        
        # Viajero Principal (subrayado y normalizado)
        viaj_run = p.add_run('Viajero Principal: ')
        viaj_run.font.name = 'Arial'
        viaj_run.font.size = Pt(12)
        viaj_run.font.bold = True
        viaj_val_run = p.add_run(viajero_principal if viajero_principal else '')
        viaj_val_run.font.name = 'Arial'
        viaj_val_run.font.size = Pt(12)
        viaj_val_run.font.underline = True  # Subrayado
        viaj_sep_run = p.add_run(' | ')
        viaj_sep_run.font.name = 'Arial'
        viaj_sep_run.font.size = Pt(12)
        
        # Tipo Habitación
        tipo_run = p.add_run('Tipo Habitación: ')
        tipo_run.font.name = 'Arial'
        tipo_run.font.size = Pt(12)
        tipo_run.font.bold = True
        tipo_val_run = p.add_run(tipo_habitacion if tipo_habitacion else '')
        tipo_val_run.font.name = 'Arial'
        tipo_val_run.font.size = Pt(12)
        
        # Salto de línea entre secciones
        self._agregar_salto_entre_secciones(doc)
        
        # Fechas y Estancia (con viñeta)
        self._agregar_subtitulo_con_vineta(doc, 'Fechas y Estancia')
        
        # Agrupar fechas y horas en una sola línea
        hora_checkin = datos.get('hora_checkin', '')
        hora_checkout = datos.get('hora_checkout', '')
        fecha_checkin = datos.get('fecha_checkin', '')
        fecha_checkout = datos.get('fecha_checkout', '')
        
        checkin_str = f"{fecha_checkin}" + (f" {hora_checkin}" if hora_checkin else "")
        checkout_str = f"{fecha_checkout}" + (f" {hora_checkout}" if hora_checkout else "")
        
        self._agregar_info_inline(doc,
            ('Check-in', checkin_str),
            ('Check-out', checkout_str)
        )
        
        # Salto de línea entre secciones
        self._agregar_salto_entre_secciones(doc)
        
        # Información de Huéspedes (con viñeta)
        self._agregar_subtitulo_con_vineta(doc, 'Información de Huéspedes')
        
        adultos = datos.get('adultos', '0')
        ninos = datos.get('ninos', '0')
        ocupacion_str = f"{adultos} Adulto(s)"
        if int(ninos) > 0:
            ocupacion_str += f", {ninos} Niño(s)"
        
        # Agrupar ocupación y régimen en una línea
        self._agregar_info_inline(doc,
            ('Ocupación', ocupacion_str),
            ('Régimen', datos.get('regimen', ''))
        )
        
        if datos.get('observaciones'):
            obs_normalizada = self._normalizar_valor_campo(datos.get('observaciones', ''), limpiar_saltos_linea=True)
            obs_p = doc.add_paragraph()
            obs_p.paragraph_format.space_after = Pt(6)
            obs_label = obs_p.add_run('Observaciones: ')
            obs_label.font.name = 'Arial'
            obs_label.font.size = Pt(12)
            obs_label.bold = True
            obs_val = obs_p.add_run(obs_normalizada)
            obs_val.font.name = 'Arial'
            obs_val.font.size = Pt(12)
        
        # Imagen de hospedaje (si existe)
        if datos.get('imagen_hospedaje_url'):
            try:
                # Intentar importar requests (opcional, puede no estar instalado)
                # type: ignore para evitar advertencia del linter si requests no está instalado
                import requests  # type: ignore[reportMissingModuleSource]
                from io import BytesIO
                image_url = datos['imagen_hospedaje_url']
                if image_url.startswith('/media/'):
                    image_url = self.request.build_absolute_uri(image_url)
                response = requests.get(image_url, timeout=5)
                response.raise_for_status()
                image_stream = BytesIO(response.content)
                doc.add_picture(image_stream, width=Inches(4.5))
            except ImportError:
                # Si requests no está disponible, solo agregar una nota
                img_note = doc.add_paragraph()
                img_note.paragraph_format.space_after = Pt(3)
                img_note.add_run('Nota: ').bold = True
                img_note.add_run('Imagen de confirmación disponible en el sistema (requests no disponible).')
            except Exception as e:
                img_note = doc.add_paragraph()
                img_note.paragraph_format.space_after = Pt(3)
                img_note.add_run('Nota: ').bold = True
                img_note.add_run('Imagen de confirmación disponible en el sistema.')
    
    def _agregar_traslado(self, doc, datos):
        """Agrega contenido de traslado al documento con formato compacto (mismo formato que cotizaciones)."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor
        
        # Color corporativo (mismo que cotizaciones)
        MOVUMS_BLUE_CORP = RGBColor(0, 74, 142)
        
        # Información de la Compañía (con viñeta)
        self._agregar_subtitulo_con_vineta(doc, 'Información de la Compañía')
        
        # Agrupar información de compañía
        self._agregar_info_inline(doc,
            ('Compañía', datos.get('compania', '')),
            ('Código Reserva', datos.get('codigo_reserva', ''))
        )
        
        # Salto de línea entre secciones
        self._agregar_salto_entre_secciones(doc)
        
        # Detalles del Traslado (con viñeta)
        self._agregar_subtitulo_con_vineta(doc, 'Detalles del Traslado')
        
        self._agregar_info_line(doc, 'Horario de Inicio de Viaje', datos.get('horario_inicio', ''))
        self._agregar_info_line(doc, 'Tipo de Servicio', datos.get('tipo_servicio', ''))
        self._agregar_info_line(doc, 'Desde', datos.get('desde', ''))
        self._agregar_info_line(doc, 'Hasta', datos.get('hasta', ''))
        
        # Información de Pasajeros (compacto)
        adultos = datos.get('adultos', '0')
        ninos = datos.get('ninos', '0')
        pasajeros_p = doc.add_paragraph()
        pasajeros_p.paragraph_format.space_after = Pt(3)
        pasajeros_p.add_run('Pasajeros: ').bold = True
        pasajeros_p.add_run(f'{adultos} Adulto(s)')
        if int(ninos) > 0:
            pasajeros_p.add_run(f', {ninos} Niño(s)')
        
        if datos.get('informacion_adicional'):
            info_normalizada = self._normalizar_valor_campo(datos.get('informacion_adicional', ''), limpiar_saltos_linea=True)
            info_p = doc.add_paragraph()
            info_p.paragraph_format.space_after = Pt(6)
            info_label = info_p.add_run('Información Adicional: ')
            info_label.font.name = 'Arial'
            info_label.font.size = Pt(12)
            info_label.bold = True
            info_val = info_p.add_run(info_normalizada)
            info_val.font.name = 'Arial'
            info_val.font.size = Pt(12)
    
    def _agregar_generica(self, doc, datos):
        """Agrega contenido genérico al documento con formato compacto."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor
        
        titulo = datos.get('titulo', 'Información Adicional')
        if titulo:
            self._agregar_subtitulo_con_vineta(doc, titulo)
        
        contenido = datos.get('contenido', '')
        if contenido:
            # Normalizar el contenido (limpiar saltos de línea excesivos)
            contenido_normalizado = self._normalizar_valor_campo(contenido, limpiar_saltos_linea=True)
            # Dividir por líneas para mejor formato (compacto)
            for linea in contenido_normalizado.split('\n'):
                if linea.strip():
                    p = doc.add_paragraph(linea.strip())
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1.1
                    for run in p.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(12)


# ------------------- API: Previsualización de promociones -------------------
from django.views.decorators.http import require_POST
from django.contrib.auth.decorators import login_required
from ventas.services.promociones import PromocionesService


@login_required
@require_POST
def preview_promociones(request):
    """
    Devuelve las promociones aplicables en tiempo real para el formulario de venta.
    Espera: cliente_id, tipo_viaje, costo_venta_final (opcional), costo_modificacion (opcional).
    """
    try:
        cliente_id_str = request.POST.get('cliente_id', '').strip()
        tipo_viaje = request.POST.get('tipo_viaje', '').strip()
        costo_venta_final_str = request.POST.get('costo_venta_final', '0').strip()
        costo_mod_str = request.POST.get('costo_modificacion', '0').strip()
        
        logger.info(f"preview_promociones recibido - cliente_id: '{cliente_id_str}', tipo_viaje: '{tipo_viaje}', costo_venta_final: '{costo_venta_final_str}', costo_mod: '{costo_mod_str}'")
        
        if not cliente_id_str or not tipo_viaje:
            logger.warning(f"Parámetros faltantes - cliente_id: '{cliente_id_str}', tipo_viaje: '{tipo_viaje}'")
            return JsonResponse({'ok': True, 'promos': []})
        
        cliente_id = int(cliente_id_str)
        costo_venta_final = Decimal(costo_venta_final_str or '0')
        costo_mod = Decimal(costo_mod_str or '0')
        
    except (ValueError, InvalidOperation) as e:
        logger.error(f"Error al parsear parámetros: {str(e)}")
        return JsonResponse({'ok': False, 'error': f'Parámetros inválidos: {str(e)}'}, status=400)
    except Exception as e:
        logger.error(f"Error inesperado al procesar parámetros: {str(e)}")
        return JsonResponse({'ok': False, 'error': 'Error al procesar parámetros'}, status=400)

    try:
        cliente = Cliente.objects.get(pk=cliente_id)
        logger.info(f"Cliente encontrado: {cliente.nombre_completo_display} (ID: {cliente_id})")
    except Cliente.DoesNotExist:
        logger.warning(f"Cliente no encontrado con ID: {cliente_id}")
        return JsonResponse({'ok': True, 'promos': []})
    except Exception as e:
        logger.error(f"Error al buscar cliente: {str(e)}")
        return JsonResponse({'ok': False, 'error': 'Error al buscar cliente'}, status=500)

    total_base = costo_venta_final + costo_mod
    logger.info(f"Buscando promociones - cliente: {cliente_id}, tipo_viaje: {tipo_viaje}, total_base: {total_base}")
    
    promos = PromocionesService.obtener_promos_aplicables(
        cliente=cliente,
        tipo_viaje=tipo_viaje,
        total_base_mxn=total_base
    )
    
    logger.info(f"Promociones encontradas: {len(promos)}")
    for p in promos:
        logger.info(f"  - {p['promo'].nombre} (tipo: {p['promo'].tipo}, alcance: {p['promo'].alcance}, monto: {p.get('monto_descuento', 0)})")
    promos_serialized = []
    for p in promos:
        promo = p['promo']
        promos_serialized.append({
            'id': promo.id,
            'nombre': promo.nombre,
            'tipo': promo.tipo,
            'porcentaje': str(p.get('porcentaje') or '0'),
            'monto_descuento': str(p.get('monto_descuento') or '0'),
            'km_bono': str(p.get('km_bono') or '0'),
            'requiere_confirmacion': p.get('requiere_confirmacion', False),
            'condicion': promo.get_condicion_display(),
            'alcance': promo.alcance,  # Agregar alcance para validación frontend
        })

    return JsonResponse({'ok': True, 'promos': promos_serialized})


# ------------------- VISTAS DE DETALLE Y EXPORTACIÓN DE COMISIONES -------------------

class DetalleComisionesView(LoginRequiredMixin, TemplateView):
    """
    Vista detallada de comisiones de un vendedor específico.
    Muestra desglose por venta.
    """
    template_name = 'ventas/detalle_comisiones.html'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        vendedor_id = kwargs.get('pk')
        
        try:
            vendedor = User.objects.get(pk=vendedor_id, perfil__rol='VENDEDOR')
        except User.DoesNotExist:
            from django.http import Http404
            raise Http404("Vendedor no encontrado")
        
        # Verificar permisos: solo puede ver su propio detalle o ser JEFE/CONTADOR
        user_rol = perm.get_user_role(self.request.user, self.request)
        if user_rol == 'VENDEDOR' and self.request.user.pk != vendedor_id:
            from django.core.exceptions import PermissionDenied
            raise PermissionDenied("No tienes permiso para ver este detalle")
        
        # Obtener mes y año del request ANTES de filtrar ventas
        mes_actual = timezone.now().month
        anio_actual = timezone.now().year
        
        # Manejar parámetros vacíos o inválidos
        mes_param = self.request.GET.get('mes', '').strip()
        anio_param = self.request.GET.get('anio', '').strip()
        
        mes_filtro = safe_int(mes_param, default=mes_actual)
        anio_filtro = safe_int(anio_param, default=anio_actual)
        
        # Rango del periodo seleccionado (ventas generadas en este mes/año)
        from datetime import date as date_type
        fecha_inicio = date_type(anio_filtro, mes_filtro, 1)
        if mes_filtro == 12:
            fecha_fin = date_type(anio_filtro + 1, 1, 1)
        else:
            fecha_fin = date_type(anio_filtro, mes_filtro + 1, 1)
        
        # OPTIMIZACIÓN N+1: Obtener ventas del vendedor GENERADAS EN EL PERIODO con prefetch y anotación
        ventas_periodo = VentaViaje.objects.filter(
            vendedor=vendedor,
            fecha_creacion__gte=fecha_inicio,
            fecha_creacion__lt=fecha_fin
        ).select_related(
            'cliente', 'proveedor'
        ).prefetch_related(
            Prefetch(
                'abonos',
                queryset=AbonoPago.objects.filter(Q(confirmado=True) | Q(forma_pago='EFE'))
            )
        ).annotate(
            total_abonos_confirmados=Coalesce(
                Sum('abonos__monto', filter=Q(abonos__confirmado=True) | Q(abonos__forma_pago='EFE')),
                Value(Decimal('0.00')),
                output_field=ModelDecimalField()
            )
        ).order_by('-fecha_creacion')
        
        # Calcular información detallada por venta
        ejecutivo = getattr(vendedor, 'ejecutivo_asociado', None)
        tipo_vendedor = ejecutivo.tipo_vendedor if ejecutivo else 'MOSTRADOR'
        
        # Base de comisión: total de ventas generadas en el periodo (no solo pagadas)
        # INT: convertir USD a MXN con tipo_cambio; NAC: costo_venta_final
        total_ventas_periodo = Decimal('0.00')
        ventas_detalle = []
        
        for venta in ventas_periodo:
            # Calcular base de comisión según tipo de venta
            if getattr(venta, 'tipo_viaje', 'NAC') == 'INT':
                total_usd = getattr(venta, 'costo_venta_final_usd', None) or (getattr(venta, 'total_usd', None) if hasattr(venta, 'total_usd') else None)
                tc = getattr(venta, 'tipo_cambio', None)
                if total_usd and tc and Decimal(str(tc)) > 0:
                    base_comision = (Decimal(str(total_usd)) * Decimal(str(tc))).quantize(Decimal('0.01'))
                    moneda_base = 'USD'
                    monto_base_usd = Decimal(str(total_usd))
                else:
                    base_comision = Decimal('0.00')
                    moneda_base = 'USD'
                    monto_base_usd = Decimal('0.00')
            else:
                base_comision = venta.costo_venta_final or Decimal('0.00')
                moneda_base = 'MXN'
                monto_base_usd = None
            
            total_ventas_periodo += base_comision
            
            # Calcular estado de pago
            total_abonos = getattr(venta, 'total_abonos_confirmados', Decimal('0.00')) or Decimal('0.00')
            apertura = venta.cantidad_apertura or Decimal('0.00')
            total_pagado_calc = total_abonos + apertura
            costo_total = (venta.costo_venta_final or Decimal('0.00')) + (venta.costo_modificacion or Decimal('0.00'))
            esta_pagada = total_pagado_calc >= costo_total
            
            # Porcentaje de pago
            porcentaje_pago = Decimal('0.00')
            if costo_total > 0:
                porcentaje_pago = (total_pagado_calc / costo_total * 100).quantize(Decimal('0.01'))
            
            ventas_detalle.append({
                'venta': venta,
                'tipo_viaje': getattr(venta, 'tipo_viaje', 'NAC'),
                'base_comision': base_comision,
                'moneda_base': moneda_base,
                'monto_base_usd': monto_base_usd,
                'tipo_cambio': getattr(venta, 'tipo_cambio', None),
                'total_pagado': total_pagado_calc,
                'costo_total': costo_total,
                'esta_pagada': esta_pagada,
                'porcentaje_pago': porcentaje_pago,
                'fecha_creacion': venta.fecha_creacion,
            })
        
        # Usar total_ventas_periodo como base (no total_ventas_pagadas)
        total_ventas_pagadas = total_ventas_periodo  # Mantener nombre para compatibilidad con template
        
        # Verificar si existe ComisionMensual para ISLA con ajuste manual
        comision_mensual = None
        porcentaje_a_usar = None
        if tipo_vendedor == 'ISLA':
            comision_mensual = ComisionMensual.objects.filter(
                vendedor=vendedor,
                mes=mes_filtro,
                anio=anio_filtro,
                tipo_vendedor='ISLA'
            ).first()
            
            if comision_mensual and comision_mensual.porcentaje_ajustado_manual:
                # Usar porcentaje ajustado manualmente por el JEFE
                porcentaje_a_usar = comision_mensual.porcentaje_ajustado_manual / Decimal('100')
            else:
                # ISLA no tiene cálculo automático, debe ser asignado manualmente
                porcentaje_a_usar = Decimal('0.00')
        else:
            # Para otros tipos, calcular normalmente usando total_ventas_periodo (ventas generadas)
            porcentaje_a_usar, _ = calcular_comision_por_tipo(
                total_ventas_periodo, 
                tipo_vendedor
            )
        
        # Calcular comisión por venta y total
        # Para ventas pagadas: 100% de comisión; para pendientes: 30% pagada, 70% pendiente
        comision_total_pagada = Decimal('0.00')
        comision_total_pendiente = Decimal('0.00')
        
        for detalle in ventas_detalle:
            comision_venta = detalle['base_comision'] * porcentaje_a_usar
            if detalle['esta_pagada']:
                # Venta pagada al 100% = 100% de comisión
                detalle['comision_pagada'] = comision_venta
                detalle['comision_pendiente'] = Decimal('0.00')
                comision_total_pagada += comision_venta
            else:
                # Venta pendiente = 30% pagada, 70% pendiente
                detalle['comision_pagada'] = comision_venta * Decimal('0.30')
                detalle['comision_pendiente'] = comision_venta * Decimal('0.70')
                comision_total_pagada += detalle['comision_pagada']
                comision_total_pendiente += detalle['comision_pendiente']
            detalle['comision_total'] = comision_venta
        
        comision_total = comision_total_pagada + comision_total_pendiente
        
        # Sueldo base (ISLA SÍ tiene sueldo base)
        sueldo_base = ejecutivo.sueldo_base if ejecutivo and ejecutivo.sueldo_base else Decimal('10000.00')
        
        ingreso_total = sueldo_base + comision_total
        
        # Determinar si se puede ajustar (solo JEFE/CONTADOR para ISLA)
        puede_ajustar_comision = (
            tipo_vendedor == 'ISLA' and 
            (perm.has_full_access(self.request.user, self.request) or perm.is_contador(self.request.user, self.request))
        )
        
        context.update({
            'vendedor': vendedor,
            'ejecutivo': ejecutivo,
            'tipo_vendedor': tipo_vendedor,
            'ventas_detalle': ventas_detalle,
            'user_rol': user_rol,
            'sueldo_base': sueldo_base,
            'total_ventas_periodo': total_ventas_periodo,
            'total_ventas_pagadas': total_ventas_periodo,  # Mantener para compatibilidad
            'porcentaje_comision': porcentaje_a_usar * 100,
            'comision_total': comision_total,
            'comision_total_pagada': comision_total_pagada,
            'comision_total_pendiente': comision_total_pendiente,
            'ingreso_total': ingreso_total,
            'mes_filtro': mes_filtro,
            'anio_filtro': anio_filtro,
            'comision_mensual': comision_mensual,
            'puede_ajustar_comision': puede_ajustar_comision,
            'porcentaje_ajustado_manual': comision_mensual.porcentaje_ajustado_manual if comision_mensual and comision_mensual.porcentaje_ajustado_manual else None,
        })
        
        # Generar fecha_desde para mostrar en el template
        try:
            fecha_desde = date_type(anio_filtro, mes_filtro, 1)
            context['fecha_desde'] = fecha_desde
        except (ValueError, TypeError):
            fecha_desde = date_type(anio_actual, mes_actual, 1)
            context['fecha_desde'] = fecha_desde
        
        return context


class AjustarComisionIslaView(LoginRequiredMixin, UserPassesTestMixin, View):
    """
    Vista para ajustar manualmente el porcentaje de comisión para Asesores de Isla.
    Solo JEFE y CONTADOR pueden realizar ajustes.
    """
    
    def test_func(self):
        """Solo JEFE y CONTADOR pueden ajustar comisiones."""
        user_rol = perm.get_user_role(self.request.user, self.request)
        return user_rol in ['JEFE', 'CONTADOR']
    
    def handle_no_permission(self):
        messages.error(self.request, "No tienes permiso para ajustar comisiones.")
        return redirect('reporte_comisiones')
    
    def post(self, request, pk):
        """Guarda el ajuste manual del porcentaje de comisión."""
        vendedor = get_object_or_404(User, pk=pk, perfil__rol='VENDEDOR')
        
        # Verificar que el vendedor sea ISLA
        ejecutivo = getattr(vendedor, 'ejecutivo_asociado', None)
        tipo_vendedor = ejecutivo.tipo_vendedor if ejecutivo else 'MOSTRADOR'
        
        if tipo_vendedor != 'ISLA':
            messages.error(request, "El ajuste manual solo está disponible para Asesores de Isla.")
            return redirect(f"{reverse('detalle_comisiones', kwargs={'pk': vendedor.pk})}?mes={timezone.now().month}&anio={timezone.now().year}")
        
        # Obtener parámetros
        from ventas.validators import safe_int
        mes = safe_int(request.POST.get('mes'), default=timezone.now().month)
        anio = safe_int(request.POST.get('anio'), default=timezone.now().year)
        porcentaje_ajustado = request.POST.get('porcentaje_ajustado', '').strip()
        nota_ajuste = request.POST.get('nota_ajuste', '').strip()
        
        # Validar porcentaje
        try:
            porcentaje_decimal = Decimal(porcentaje_ajustado)
            if porcentaje_decimal < 0 or porcentaje_decimal > 100:
                messages.error(request, "El porcentaje debe estar entre 0 y 100.")
                return redirect(f"{reverse('detalle_comisiones', kwargs={'pk': vendedor.pk})}?mes={mes}&anio={anio}")
        except (ValueError, InvalidOperation):
            messages.error(request, "Porcentaje inválido.")
            return redirect(f"{reverse('detalle_comisiones', kwargs={'pk': vendedor.pk})}?mes={mes}&anio={anio}")
        
        # Obtener o crear ComisionMensual
        comision_mensual, created = ComisionMensual.objects.get_or_create(
            vendedor=vendedor,
            mes=mes,
            anio=anio,
            tipo_vendedor='ISLA',
            defaults={
                'total_ventas_mes': Decimal('0.00'),
                'porcentaje_comision': Decimal('0.00'),
                'comision_total': Decimal('0.00'),
            }
        )
        
        # Si no tiene total_ventas_mes, calcularlo
        if comision_mensual.total_ventas_mes == Decimal('0.00'):
            fecha_inicio = date(anio, mes, 1)
            if mes == 12:
                fecha_fin = date(anio + 1, 1, 1)
            else:
                fecha_fin = date(anio, mes + 1, 1)
            
            ventas_mes = VentaViaje.objects.filter(
                vendedor=vendedor,
                fecha_creacion__gte=fecha_inicio,
                fecha_creacion__lt=fecha_fin
            )
            total_ventas_mes = sum(venta.costo_venta_final for venta in ventas_mes) or Decimal('0.00')
            comision_mensual.total_ventas_mes = total_ventas_mes
        
        # Actualizar con el ajuste manual
        comision_mensual.porcentaje_ajustado_manual = porcentaje_decimal
        comision_mensual.ajustado_por = request.user
        comision_mensual.fecha_ajuste = timezone.now()
        comision_mensual.nota_ajuste = nota_ajuste
        
        # Recalcular comisión total con el nuevo porcentaje
        # Para ISLA, calcular comisiones por venta individual
        fecha_inicio = date(anio, mes, 1)
        if mes == 12:
            fecha_fin = date(anio + 1, 1, 1)
        else:
            fecha_fin = date(anio, mes + 1, 1)
        
        ventas_mes = VentaViaje.objects.filter(
            vendedor=vendedor,
            fecha_creacion__gte=fecha_inicio,
            fecha_creacion__lt=fecha_fin
        )
        
        # Recalcular comisiones de ventas con el nuevo porcentaje
        porcentaje_decimal_normalizado = porcentaje_decimal / Decimal('100')
        comision_total_calculada = Decimal('0.00')
        
        for venta in ventas_mes:
            # Calcular comisión para esta venta (INT: base en USD)
            total_pagado = venta.total_pagado
            costo_total = venta.costo_total_con_modificacion
            monto_base = (venta.costo_venta_final_usd or venta.total_usd) if venta.tipo_viaje == 'INT' else (venta.costo_venta_final or Decimal('0.00'))

            if total_pagado >= costo_total:
                # Venta pagada: 100% de comisión
                comision_venta = monto_base * porcentaje_decimal_normalizado
                comision_pagada = comision_venta
                comision_pendiente = Decimal('0.00')
                estado_pago = 'PAGADA'
            else:
                # Venta pendiente: 30% pagada, 70% pendiente
                comision_venta = monto_base * porcentaje_decimal_normalizado
                comision_pagada = comision_venta * Decimal('0.30')
                comision_pendiente = comision_venta * Decimal('0.70')
                estado_pago = 'PENDIENTE'

            # Actualizar o crear ComisionVenta
            tipo_venta = 'INTERNACIONAL' if venta.tipo_viaje == 'INT' else ('INTERNACIONAL MXN' if venta.tipo_viaje == 'INT_MXN' else 'NACIONAL')
            ComisionVenta.objects.update_or_create(
                venta=venta,
                mes=mes,
                anio=anio,
                defaults={
                    'vendedor': vendedor,
                    'tipo_venta': tipo_venta,
                    'monto_base_comision': monto_base,
                    'porcentaje_aplicado': porcentaje_decimal,
                    'comision_calculada': comision_venta,
                    'comision_pagada': comision_pagada,
                    'comision_pendiente': comision_pendiente,
                    'estado_pago_venta': estado_pago,
                }
            )
            
            comision_total_calculada += comision_venta
        
        # Actualizar ComisionMensual
        comision_mensual.porcentaje_comision = porcentaje_decimal  # Actualizar también el porcentaje principal
        comision_mensual.comision_total = comision_total_calculada
        comision_mensual.comision_total_pagada = sum(
            cv.comision_pagada for cv in ComisionVenta.objects.filter(
                vendedor=vendedor, mes=mes, anio=anio
            )
        ) or Decimal('0.00')
        comision_mensual.comision_total_pendiente = sum(
            cv.comision_pendiente for cv in ComisionVenta.objects.filter(
                vendedor=vendedor, mes=mes, anio=anio
            )
        ) or Decimal('0.00')
        comision_mensual.bono_extra = Decimal('0.00')  # ISLA no tiene bono extra
        comision_mensual.save()
        
        messages.success(
            request, 
            f"Porcentaje de comisión ajustado a {porcentaje_decimal}% para {vendedor.get_full_name() or vendedor.username}."
        )
        logger.info(
            f"✅ Comisión ajustada manualmente para {vendedor.username} ({mes}/{anio}): {porcentaje_decimal}% por {request.user.username}"
        )
        
        # Redirigir al reporte de comisiones con los filtros de mes y año
        return redirect(f"{reverse('reporte_comisiones')}?mes={mes}&anio={anio}")


class ExportarComisionesExcelView(LoginRequiredMixin, View):
    """
    Exporta las comisiones de un vendedor a Excel.
    """
    def get(self, request, pk):
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        except ImportError:
            messages.error(request, "La librería openpyxl no está instalada. Instálala con: pip install openpyxl")
            return redirect('reporte_comisiones')
        
        from django.http import HttpResponse
        
        try:
            vendedor = User.objects.get(pk=pk, perfil__rol='VENDEDOR')
        except User.DoesNotExist:
            from django.http import Http404
            raise Http404("Vendedor no encontrado")
        
        # Verificar permisos
        user_rol = request.user.perfil.rol if hasattr(request.user, 'perfil') else 'INVITADO'
        if user_rol == 'VENDEDOR' and request.user.pk != pk:
            from django.core.exceptions import PermissionDenied
            raise PermissionDenied("No tienes permiso para exportar este detalle")
        
        # OPTIMIZACIÓN N+1: Obtener ventas con prefetch y anotación
        ventas_base = VentaViaje.objects.filter(vendedor=vendedor).select_related(
            'cliente'
        ).prefetch_related(
            Prefetch(
                'abonos',
                queryset=AbonoPago.objects.filter(Q(confirmado=True) | Q(forma_pago='EFE'))
            )
        ).annotate(
            total_abonos_confirmados=Coalesce(
                Sum('abonos__monto', filter=Q(abonos__confirmado=True) | Q(abonos__forma_pago='EFE')),
                Value(Decimal('0.00')),
                output_field=ModelDecimalField()
            )
        )
        
        # OPTIMIZACIÓN N+1: Filtrar ventas pagadas usando anotación
        ventas_pagadas = []
        for venta in ventas_base:
            total_abonos = getattr(venta, 'total_abonos_confirmados', Decimal('0.00')) or Decimal('0.00')
            apertura = venta.cantidad_apertura or Decimal('0.00')
            total_pagado_calc = total_abonos + apertura
            costo_total = (venta.costo_venta_final or Decimal('0.00')) + (venta.costo_modificacion or Decimal('0.00'))
            if total_pagado_calc >= costo_total:
                ventas_pagadas.append(venta)
        
        # Calcular comisiones
        ejecutivo = getattr(vendedor, 'ejecutivo_asociado', None)
        tipo_vendedor = ejecutivo.tipo_vendedor if ejecutivo else 'MOSTRADOR'
        
        total_ventas_pagadas = sum(
            venta.costo_venta_final for venta in ventas_pagadas
        ) or Decimal('0.00')
        
        porcentaje_comision, comision_total = calcular_comision_por_tipo(
            total_ventas_pagadas, 
            tipo_vendedor
        )
        
        sueldo_base = ejecutivo.sueldo_base if ejecutivo and ejecutivo.sueldo_base else Decimal('10000.00')
        ingreso_total = sueldo_base + comision_total
        
        # Crear workbook — mismo formato que reporte financiero (Movums)
        from openpyxl.utils import get_column_letter
        from openpyxl.styles import Alignment

        wb = Workbook()
        ws = wb.active
        ws.title = "Comisiones"

        # Estilos Movums (igual que reporte financiero)
        header_fill = PatternFill(start_color="667EEA", end_color="667EEA", fill_type="solid")
        title_fill = PatternFill(start_color="5C0CD1", end_color="5C0CD1", fill_type="solid")
        row_fill_even = PatternFill(start_color="F5F3FF", end_color="F5F3FF", fill_type="solid")
        row_fill_odd = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=11)
        title_font = Font(bold=True, color="FFFFFF", size=14)
        border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )
        border_medium = Border(
            left=Side(style='medium'), right=Side(style='medium'),
            top=Side(style='medium'), bottom=Side(style='medium')
        )
        align_center = Alignment(horizontal='center', vertical='center', wrap_text=True)
        align_left = Alignment(horizontal='left', vertical='center', wrap_text=True)

        # Título (fila 1, estilo Movums)
        ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=6)
        titulo_cell = ws.cell(row=1, column=1, value=f"Reporte de Comisiones — {vendedor.get_full_name() or vendedor.username} · Movums")
        titulo_cell.fill = title_fill
        titulo_cell.font = title_font
        titulo_cell.alignment = Alignment(horizontal='center', vertical='center')
        titulo_cell.border = border_medium
        ws.row_dimensions[1].height = 28

        row = 3

        # Resumen general
        ws.cell(row=row, column=1, value="RESUMEN GENERAL").font = Font(bold=True, size=12)
        row += 1

        for c in range(1, 3):
            cell = ws.cell(row=row, column=c, value="Concepto" if c == 1 else "Monto")
            cell.fill = header_fill
            cell.font = header_font
            cell.border = border
            cell.alignment = align_center
        ws.row_dimensions[row].height = 22
        row += 1

        totales_fill = PatternFill(start_color="E8E4F8", end_color="E8E4F8", fill_type="solid")
        for label, valor in [
            ("Sueldo Base", float(sueldo_base)),
            ("Comisión", float(comision_total)),
            ("TOTAL", float(ingreso_total)),
        ]:
            ws.cell(row=row, column=1, value=label).fill = totales_fill
            ws.cell(row=row, column=1).border = border
            ws.cell(row=row, column=1).alignment = align_left
            if label == "TOTAL":
                ws.cell(row=row, column=1).font = Font(bold=True)
            c2 = ws.cell(row=row, column=2, value=valor)
            c2.number_format = '#,##0.00'
            c2.fill = totales_fill
            c2.border = border
            c2.alignment = align_center
            if label == "TOTAL":
                c2.font = Font(bold=True)
            row += 1
        row += 1

        # Detalle de ventas (título de sección)
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
        sec_cell = ws.cell(row=row, column=1, value="DETALLE DE VENTAS")
        sec_cell.font = Font(bold=True, color="FFFFFF", size=12)
        sec_cell.fill = header_fill
        sec_cell.alignment = align_center
        sec_cell.border = border_medium
        row += 1

        headers = ['Venta ID', 'Cliente', 'Total Venta', 'Pagado', 'Comisión']
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.border = border
            cell.alignment = align_center
        ws.row_dimensions[row].height = 22
        row += 1

        for idx, venta in enumerate(ventas_pagadas):
            comision_venta = venta.costo_venta_final * porcentaje_comision
            row_fill = row_fill_even if idx % 2 == 0 else row_fill_odd
            for c, val in enumerate([
                venta.pk,
                str(venta.cliente),
                float(venta.costo_venta_final),
                float(venta.total_pagado),
                float(comision_venta),
            ], 1):
                cell = ws.cell(row=row, column=c, value=val)
                cell.fill = row_fill
                cell.border = border
                cell.alignment = align_center if c >= 3 else align_left
                if c >= 3:
                    cell.number_format = '#,##0.00'
            row += 1

        # Anchos de columna fijos (evita MergedCell)
        ws.column_dimensions['A'].width = 12
        ws.column_dimensions['B'].width = 32
        ws.column_dimensions['C'].width = 14
        ws.column_dimensions['D'].width = 14
        ws.column_dimensions['E'].width = 14
        ws.column_dimensions['F'].width = 14
        ws.freeze_panes = 'A3'

        # Preparar respuesta
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        filename = f"comisiones_{vendedor.username}.xlsx"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        wb.save(response)
        return response


# ------------------- VISTAS PARA PROVEEDORES -------------------

class ProveedorUpdateView(LoginRequiredMixin, usuarios_mixins.ManageSuppliersRequiredMixin, UpdateView):
    """
    Vista para editar un proveedor. Accesible para JEFE, Director General y Director Administrativo.
    """
    model = Proveedor
    form_class = ProveedorForm
    template_name = 'ventas/proveedores.html'
    
    def get_success_url(self):
        messages.success(self.request, "Proveedor actualizado correctamente.")
        return reverse('proveedores')
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Incluir el contexto de la vista de lista
        proveedores = Proveedor.objects.all()
        proveedores_por_servicio = {clave: [] for clave, _ in Proveedor.SERVICIO_CHOICES}

        # ✅ Agrupar proveedores por cada servicio que ofrecen (pueden estar en múltiples grupos)
        for proveedor in proveedores:
            if proveedor.servicios:
                servicios_list = [s.strip() for s in proveedor.servicios.split(',') if s.strip()]
                for servicio_codigo in servicios_list:
                    if servicio_codigo in proveedores_por_servicio:
                        proveedores_por_servicio[servicio_codigo].append(proveedor)
        
        grupos_proveedores = [
            {
                'clave': clave,
                'label': label,
                'proveedores': sorted(
                    proveedores_por_servicio.get(clave, []),
                    key=lambda p: p.nombre.lower()
                ),
            }
            for clave, label in Proveedor.SERVICIO_CHOICES
        ]
        
        context['grupos_proveedores'] = grupos_proveedores
        context['total_proveedores'] = sum(len(grupo['proveedores']) for grupo in grupos_proveedores)
        context['servicio_choices'] = Proveedor.SERVICIO_CHOICES
        context['proveedor_editando'] = self.object
        context['form'] = self.get_form()
        return context


class ProveedorDeleteView(LoginRequiredMixin, usuarios_mixins.ManageSuppliersRequiredMixin, DeleteView):
    """
    Vista para eliminar un proveedor. Accesible para JEFE, Director General y Director Administrativo.
    """
    model = Proveedor
    template_name = 'ventas/proveedor_confirm_delete.html'
    
    def get_success_url(self):
        messages.success(self.request, "Proveedor eliminado correctamente.")
        return reverse('proveedores')


# ------------------- VISTAS PARA GESTIÓN DE COMPROBANTES Y PAGOS POR CONFIRMAR -------------------

class SubirComprobanteAbonoView(LoginRequiredMixin, View):
    """Vista para subir el comprobante de un abono."""
    
    def post(self, request, pk):
        abono = get_object_or_404(AbonoPago, pk=pk)
        
        # Verificar que el usuario tenga permiso (vendedor de la venta, JEFE o CONTADOR)
        user_rol = perm.get_user_role(request.user, request)
        puede_subir = (
            perm.has_full_access(request.user, request) or perm.is_contador(request.user, request) or
            (abono.venta.vendedor == request.user)
        )
        
        if not puede_subir:
            messages.error(request, "No tienes permiso para subir comprobantes.")
            return redirect(reverse('detalle_venta', kwargs={'pk': abono.venta.pk, 'slug': abono.venta.slug_safe}) + '?tab=abonos')
        
        # Verificar que el abono requiera comprobante (TRN, TAR, DEP)
        if abono.forma_pago not in ['TRN', 'TAR', 'DEP']:
            messages.error(request, "Este tipo de pago no requiere comprobante.")
            return redirect(reverse('detalle_venta', kwargs={'pk': abono.venta.pk, 'slug': abono.venta.slug_safe}) + '?tab=abonos')
        
        # Obtener la imagen del request
        imagen = request.FILES.get('comprobante_imagen')
        if not imagen:
            messages.error(request, "Debes seleccionar una imagen del comprobante.")
            return redirect(reverse('detalle_venta', kwargs={'pk': abono.venta.pk, 'slug': abono.venta.slug_safe}) + '?tab=abonos')
        
        # ✅ SEGURIDAD: Validar archivo antes de guardar
        from ventas.validators import validate_uploaded_file
        from django.core.exceptions import ValidationError
        try:
            validate_uploaded_file(imagen)
        except ValidationError as e:
            messages.error(request, f"Archivo inválido: {str(e)}")
            return redirect(reverse('detalle_venta', kwargs={'pk': abono.venta.pk, 'slug': abono.venta.slug_safe}) + '?tab=abonos')
        
        # Guardar el comprobante
        abono.comprobante_imagen = imagen
        abono.comprobante_subido = True
        abono.comprobante_subido_en = timezone.now()
        abono.comprobante_subido_por = request.user
        abono.save()
        
        # Crear notificaciones para CONTADOR
        contadores = User.objects.filter(perfil__rol='CONTADOR')
        forma_pago_display = dict(AbonoPago.FORMA_PAGO_CHOICES).get(abono.forma_pago, abono.forma_pago)
        es_usd = abono.venta.tipo_viaje == 'INT'
        monto_display = f"USD ${abono.monto_usd:,.2f}" if es_usd and abono.monto_usd else f"${abono.monto:,.2f}"
        mensaje_contador = f"{'⚠️ DÓLARES - ' if es_usd else ''}Abono pendiente de confirmación: {monto_display} ({forma_pago_display}) - Venta #{abono.venta.pk} - Cliente: {abono.venta.cliente.nombre_completo_display}"
        
        for contador in contadores:
            Notificacion.objects.filter(
                usuario=contador,
                venta=abono.venta,
                abono=abono,
                tipo='PAGO_PENDIENTE',
                confirmado=False
            ).delete()
            
            Notificacion.objects.create(
                usuario=contador,
                tipo='PAGO_PENDIENTE',
                mensaje=mensaje_contador,
                venta=abono.venta,
                abono=abono,
                confirmado=False
            )
        
        jefes = User.objects.filter(perfil__rol='JEFE')
        mensaje_jefe = f"{'⚠️ DÓLARES - ' if es_usd else ''}Abono pendiente de confirmación: {monto_display} ({forma_pago_display}) - Venta #{abono.venta.pk} - Cliente: {abono.venta.cliente.nombre_completo_display}"
        for jefe in jefes:
            # Eliminar notificaciones previas del mismo abono
            Notificacion.objects.filter(
                usuario=jefe,
                venta=abono.venta,
                abono=abono,
                tipo='PAGO_PENDIENTE',
                confirmado=False
            ).delete()
            
            Notificacion.objects.create(
                usuario=jefe,
                tipo='PAGO_PENDIENTE',
                mensaje=mensaje_jefe,
                venta=abono.venta,
                abono=abono,
                confirmado=False
            )
        
        messages.success(request, "Comprobante subido exitosamente. El contador ha sido notificado.")
        return redirect(reverse('detalle_venta', kwargs={'pk': abono.venta.pk, 'slug': abono.venta.slug_safe}) + '?tab=abonos')


class SubirComprobanteAperturaView(LoginRequiredMixin, View):
    """Vista para subir el comprobante de un pago de apertura."""
    
    def post(self, request, pk):
        venta = get_object_or_404(VentaViaje, pk=pk)
        
        # Verificar que el usuario tenga permiso
        user_rol = perm.get_user_role(request.user, request)
        puede_subir = (
            perm.has_full_access(request.user, request) or perm.is_contador(request.user, request) or
            (venta.vendedor == request.user)
        )
        
        if not puede_subir:
            messages.error(request, "No tienes permiso para subir comprobantes.")
            return redirect(reverse('detalle_venta', kwargs={'pk': venta.pk, 'slug': venta.slug_safe}) + '?tab=abonos')
        
        # Verificar que la venta tenga apertura y requiera comprobante (no aplica para crédito)
        if venta.modo_pago_apertura == 'CRE':
            messages.error(request, "El crédito no requiere comprobante. El contador validará el crédito directamente.")
            return redirect(reverse('detalle_venta', kwargs={'pk': venta.pk, 'slug': venta.slug_safe}) + '?tab=abonos')
        
        # Para ventas internacionales puede haber solo cantidad_apertura_usd
        tiene_apertura = (venta.cantidad_apertura and venta.cantidad_apertura > 0) or (
            getattr(venta, 'tipo_viaje', None) == 'INT' and venta.cantidad_apertura_usd and venta.cantidad_apertura_usd > 0
        )
        if not tiene_apertura:
            messages.error(request, "Esta venta no tiene pago de apertura.")
            return redirect(reverse('detalle_venta', kwargs={'pk': venta.pk, 'slug': venta.slug_safe}) + '?tab=abonos')
        
        if venta.modo_pago_apertura not in ['TRN', 'TAR', 'DEP']:
            messages.error(request, "Este tipo de pago no requiere comprobante.")
            return redirect(reverse('detalle_venta', kwargs={'pk': venta.pk, 'slug': venta.slug_safe}) + '?tab=abonos')
        
        # Obtener la imagen del request
        imagen = request.FILES.get('comprobante_apertura')
        if not imagen:
            messages.error(request, "Debes seleccionar una imagen del comprobante.")
            return redirect(reverse('detalle_venta', kwargs={'pk': venta.pk, 'slug': venta.slug_safe}) + '?tab=abonos')
        
        # ✅ SEGURIDAD: Validar archivo antes de guardar
        from ventas.validators import validate_uploaded_file
        from django.core.exceptions import ValidationError
        try:
            validate_uploaded_file(imagen)
        except ValidationError as e:
            messages.error(request, f"Archivo inválido: {str(e)}")
            return redirect(reverse('detalle_venta', kwargs={'pk': venta.pk, 'slug': venta.slug_safe}) + '?tab=abonos')
        
        # Guardar el comprobante
        venta.comprobante_apertura = imagen
        venta.comprobante_apertura_subido = True
        venta.comprobante_apertura_subido_en = timezone.now()
        venta.comprobante_apertura_subido_por = request.user
        # Asegurar que el estado sea EN_CONFIRMACION para que aparezca en pagos por confirmar
        # Esto es crítico: las ventas con comprobante subido deben estar en EN_CONFIRMACION
        # para que el contador las vea en "Pagos por confirmar"
        venta.estado_confirmacion = 'EN_CONFIRMACION'
        venta.save()
        
        # Verificar que el estado se mantuvo después del save (por si algún signal lo cambió)
        # Recargar desde la BD para obtener el estado actual
        venta.refresh_from_db()
        if venta.estado_confirmacion != 'EN_CONFIRMACION':
            # Si algún signal cambió el estado, forzarlo de nuevo
            venta.estado_confirmacion = 'EN_CONFIRMACION'
            venta.save(update_fields=['estado_confirmacion'])
        
        # Crear notificaciones para CONTADOR
        contadores = User.objects.filter(perfil__rol='CONTADOR')
        modo_pago_display = dict(VentaViaje.MODO_PAGO_CHOICES).get(venta.modo_pago_apertura, venta.modo_pago_apertura)
        # Para ventas internacionales mostrar USD, para nacionales MXN
        if venta.tipo_viaje == 'INT' and venta.cantidad_apertura_usd:
            monto_display = f"USD ${venta.cantidad_apertura_usd:,.2f}"
        else:
            monto_display = f"${venta.cantidad_apertura:,.2f}"
        mensaje_contador = f"{'⚠️ DÓLARES - ' if venta.tipo_viaje == 'INT' else ''}Pago de apertura pendiente de confirmación: {monto_display} ({modo_pago_display}) - Venta #{venta.pk} - Cliente: {venta.cliente.nombre_completo_display}"
        
        for contador in contadores:
            # Eliminar notificaciones previas de apertura para esta venta
            Notificacion.objects.filter(
                usuario=contador,
                venta=venta,
                abono__isnull=True,  # Notificaciones sin abono son de apertura
                tipo='PAGO_PENDIENTE',
                confirmado=False
            ).delete()
            
            Notificacion.objects.create(
                usuario=contador,
                tipo='PAGO_PENDIENTE',
                mensaje=mensaje_contador,
                venta=venta,
                confirmado=False
            )
        
        # Crear notificaciones para JEFE
        jefes = User.objects.filter(perfil__rol='JEFE')
        mensaje_jefe = f"{'⚠️ DÓLARES - ' if venta.tipo_viaje == 'INT' else ''}Pago de apertura pendiente de confirmación: {monto_display} ({modo_pago_display}) - Venta #{venta.pk}"
        for jefe in jefes:
            # Eliminar notificaciones previas
            Notificacion.objects.filter(
                usuario=jefe,
                venta=venta,
                abono__isnull=True,
                tipo='PAGO_PENDIENTE',
                confirmado=False
            ).delete()
            
            Notificacion.objects.create(
                usuario=jefe,
                tipo='APERTURA',
                mensaje=mensaje_jefe,
                venta=venta,
                confirmado=False
            )
        
        messages.success(request, "Comprobante de apertura subido exitosamente. El contador ha sido notificado.")
        return redirect(reverse('detalle_venta', kwargs={'pk': venta.pk, 'slug': venta.slug_safe}) + '?tab=abonos')


class PagosPorConfirmarView(LoginRequiredMixin, UserPassesTestMixin, TemplateView):
    """Vista para que el contador vea todos los pagos pendientes de confirmar."""
    
    template_name = 'ventas/pagos_por_confirmar.html'
    
    def test_func(self):
        """Solo CONTADOR puede acceder."""
        return perm.is_contador(self.request.user)
    
    def handle_no_permission(self):
        messages.error(self.request, "Solo el contador puede acceder a esta sección.")
        return redirect('dashboard')
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        # Obtener pagos pendientes de confirmar
        # Abonos pendientes con comprobante subido
        abonos_pendientes = AbonoPago.objects.filter(
            Q(forma_pago__in=['TRN', 'TAR', 'DEP']) &
            Q(confirmado=False) &
            Q(comprobante_subido=True)
        ).select_related('venta', 'venta__cliente', 'venta__vendedor', 'registrado_por').order_by('-fecha_pago')
        
        # Ventas con apertura pendiente
        # Para TRN/TAR/DEP: requiere comprobante subido
        # Para CRE: no requiere comprobante, solo estar en EN_CONFIRMACION
        # IMPORTANTE: Solo mostrar las que están en 'EN_CONFIRMACION' (no las confirmadas)
        # CRÍTICO: Para ventas INT, verificar cantidad_apertura_usd > 0, no solo cantidad_apertura
        # ESTRATEGIA: Consulta simplificada que verifica ambos casos de manera clara
        
        # Base: ventas en EN_CONFIRMACION (no COMPLETADO)
        base_query = VentaViaje.objects.filter(
            estado_confirmacion='EN_CONFIRMACION'
        ).exclude(
            estado_confirmacion='COMPLETADO'
        )
        
        # Para TRN/TAR/DEP: deben tener comprobante subido Y algún monto de apertura válido
        ventas_trn_tar_dep = base_query.filter(
            modo_pago_apertura__in=['TRN', 'TAR', 'DEP'],
            comprobante_apertura_subido=True
        ).filter(
            # Ventas internacionales: cantidad_apertura_usd > 0
            Q(tipo_viaje='INT', cantidad_apertura_usd__gt=0) |
            Q(tipo_viaje__in=['NAC', 'INT_MXN'], cantidad_apertura__gt=0) |
            Q(tipo_viaje__isnull=True, cantidad_apertura__gt=0)
        )
        
        # Para CRE: solo necesita estar en EN_CONFIRMACION
        ventas_cre = base_query.filter(
            modo_pago_apertura='CRE'
        )
        
        # Combinar ambas consultas
        ventas_apertura_pendiente = (ventas_trn_tar_dep | ventas_cre).distinct().select_related('cliente', 'vendedor').order_by('-fecha_creacion')
        
        # Abonos a proveedor pendientes: PENDIENTE (para aprobar) y APROBADO (para confirmar)
        abonos_proveedor_pendientes = AbonoProveedor.objects.filter(
            estado__in=['PENDIENTE', 'APROBADO']
        ).select_related('venta', 'venta__cliente', 'venta__vendedor', 'solicitud_por', 'aprobado_por').order_by('-fecha_solicitud', '-fecha_aprobacion')
        
        # Crear un objeto simple para el template que tenga los atributos abonos, aperturas y abonos_proveedor
        class PagosPendientes:
            def __init__(self, abonos, aperturas, abonos_proveedor):
                self.abonos = abonos
                self.aperturas = aperturas
                self.abonos_proveedor = abonos_proveedor
        
        context['object'] = PagosPendientes(abonos_pendientes, ventas_apertura_pendiente, abonos_proveedor_pendientes)
        context['pagos_pendientes'] = context['object']  # También disponible con este nombre
        
        # Abonos a proveedor confirmados para el historial
        abonos_proveedor_confirmados = AbonoProveedor.objects.filter(
            estado='COMPLETADO'
        ).select_related('venta', 'venta__cliente', 'confirmado_por')
        
        # Obtener pagos confirmados para el historial
        fecha_filtro = self.request.GET.get('fecha_filtro', 'mes')
        fecha_desde = self.request.GET.get('fecha_desde')
        fecha_hasta = self.request.GET.get('fecha_hasta')
        
        # Abonos confirmados - Incluir todos los confirmados
        # Primero aplicar filtros de fecha, luego ordenar
        abonos_confirmados = AbonoPago.objects.filter(
            confirmado=True
        ).select_related('venta', 'venta__cliente', 'confirmado_por')
        
        # Ventas con apertura confirmada - SOLO las que tienen modo_pago que requiere confirmación
        # y que fueron confirmadas (estado_confirmacion='COMPLETADO')
        # Para TRN/TAR/DEP: requieren comprobante subido
        # Para CRE: no requiere comprobante, solo que estado no sea EN_CONFIRMACION
        ventas_apertura_confirmada = VentaViaje.objects.filter(
            Q(estado_confirmacion='COMPLETADO') &
            (
                # Transferencia, Tarjeta, Depósito: requieren comprobante subido
                # Para NAC: cantidad_apertura > 0; para INT: cantidad_apertura_usd > 0
                (Q(modo_pago_apertura__in=['TRN', 'TAR', 'DEP']) & 
                 Q(comprobante_apertura_subido=True) &
                 (
                     Q(cantidad_apertura__gt=0) |  # Ventas nacionales
                     (Q(tipo_viaje='INT') & Q(cantidad_apertura_usd__gt=0))  # Ventas internacionales
                 )) |
                # Crédito: no requiere comprobante
                Q(modo_pago_apertura='CRE')
            )
        ).select_related('cliente', 'vendedor')
        
        # Aplicar filtros de fecha a abonos confirmados
        # IMPORTANTE: Solo usar fecha_desde y fecha_hasta si fecha_filtro es 'personalizado'
        if fecha_filtro == 'personalizado' and fecha_desde and fecha_hasta:
            try:
                fecha_desde_obj = datetime.datetime.strptime(fecha_desde, '%Y-%m-%d').date()
                fecha_hasta_obj = datetime.datetime.strptime(fecha_hasta, '%Y-%m-%d').date()
                # Filtrar por confirmado_en si existe, sino por fecha_pago
                abonos_confirmados = abonos_confirmados.filter(
                    Q(confirmado_en__date__range=[fecha_desde_obj, fecha_hasta_obj]) |
                    Q(confirmado_en__isnull=True, fecha_pago__date__range=[fecha_desde_obj, fecha_hasta_obj])
                )
                # También aplicar filtro a aperturas confirmadas por fecha de creación
                ventas_apertura_confirmada = ventas_apertura_confirmada.filter(
                    fecha_creacion__date__range=[fecha_desde_obj, fecha_hasta_obj]
                )
                # Aplicar filtro a abonos a proveedor confirmados
                abonos_proveedor_confirmados = abonos_proveedor_confirmados.filter(
                    fecha_confirmacion__date__range=[fecha_desde_obj, fecha_hasta_obj]
                )
            except ValueError:
                pass
        elif fecha_filtro == 'dia':
            # Filtrar solo los de hoy
            # Usar localdate() para obtener la fecha local de México, no UTC
            hoy = timezone.localdate()
            # Buscar por confirmado_en si existe, sino por fecha_pago
            # IMPORTANTE: Usar Q objects para que funcione correctamente con OR
            abonos_confirmados = abonos_confirmados.filter(
                Q(confirmado_en__date=hoy) | 
                Q(confirmado_en__isnull=True, fecha_pago__date=hoy)
            )
            # Para aperturas, usar fecha_creacion
            ventas_apertura_confirmada = ventas_apertura_confirmada.filter(fecha_creacion__date=hoy)
            # Para abonos a proveedor, usar fecha_confirmacion
            abonos_proveedor_confirmados = abonos_proveedor_confirmados.filter(fecha_confirmacion__date=hoy)
        elif fecha_filtro == 'semana':
            # Usar localdate() para obtener la fecha local de México, no UTC
            semana_pasada = timezone.localdate() - timedelta(days=7)
            abonos_confirmados = abonos_confirmados.filter(
                Q(confirmado_en__date__gte=semana_pasada) | Q(confirmado_en__isnull=True, fecha_pago__date__gte=semana_pasada)
            )
            ventas_apertura_confirmada = ventas_apertura_confirmada.filter(fecha_creacion__date__gte=semana_pasada)
            abonos_proveedor_confirmados = abonos_proveedor_confirmados.filter(fecha_confirmacion__date__gte=semana_pasada)
        elif fecha_filtro == 'mes':
            # Usar localdate() para obtener la fecha local de México, no UTC
            mes_pasado = timezone.localdate() - timedelta(days=30)
            abonos_confirmados = abonos_confirmados.filter(
                Q(confirmado_en__date__gte=mes_pasado) | Q(confirmado_en__isnull=True, fecha_pago__date__gte=mes_pasado)
            )
            ventas_apertura_confirmada = ventas_apertura_confirmada.filter(fecha_creacion__date__gte=mes_pasado)
            abonos_proveedor_confirmados = abonos_proveedor_confirmados.filter(fecha_confirmacion__date__gte=mes_pasado)
        
        # Ordenar abonos confirmados por confirmado_en si existe, sino por fecha_pago
        from django.db.models import Case, When, F, DateTimeField
        abonos_confirmados = abonos_confirmados.annotate(
            fecha_orden=Case(
                When(confirmado_en__isnull=False, then=F('confirmado_en')),
                default=F('fecha_pago'),
                output_field=DateTimeField()
            )
        ).order_by('-fecha_orden')
        
        # Ordenar aperturas confirmadas por fecha de creación
        ventas_apertura_confirmada = ventas_apertura_confirmada.order_by('-fecha_creacion')
        
        context['abonos_confirmados'] = abonos_confirmados[:100]  # Limitar a 100 para rendimiento
        context['ventas_apertura_confirmada'] = ventas_apertura_confirmada[:100]
        context['fecha_filtro'] = fecha_filtro
        # Solo pasar fecha_desde y fecha_hasta al contexto si es personalizado
        if fecha_filtro == 'personalizado':
            context['fecha_desde'] = fecha_desde
            context['fecha_hasta'] = fecha_hasta
        else:
            context['fecha_desde'] = None
            context['fecha_hasta'] = None
        
        return context


class ConfirmarPagoDesdeListaView(LoginRequiredMixin, UserPassesTestMixin, View):
    """Vista para confirmar un pago desde la lista de pagos por confirmar."""
    
    def test_func(self):
        """Solo CONTADOR puede confirmar pagos."""
        return perm.is_contador(self.request.user)
    
    def handle_no_permission(self):
        messages.error(self.request, "Solo el contador puede confirmar pagos.")
        return redirect('dashboard')
    
    def post(self, request, tipo, pk):
        """
        Confirma un pago. Tipo puede ser 'abono' o 'apertura'.
        """
        if tipo == 'abono':
            abono = get_object_or_404(AbonoPago, pk=pk)
            
            # Verificar que tenga comprobante subido
            if not abono.comprobante_subido:
                messages.error(request, "Este abono no tiene comprobante subido.")
                return redirect('pagos_por_confirmar')
            
            # Confirmar el abono
            abono.confirmado = True
            abono.confirmado_por = request.user
            abono.confirmado_en = timezone.now()
            abono.save()
            
            venta = abono.venta
            
            # Actualizar estado de la venta
            # El método actualizar_estado_financiero() ahora maneja correctamente
            # las aperturas confirmadas y no las cambia de estado
            venta.actualizar_estado_financiero()
            
            # Eliminar notificaciones pendientes del CONTADOR
            Notificacion.objects.filter(
                usuario=request.user,
                venta=venta,
                abono=abono,
                tipo='PAGO_PENDIENTE',
                confirmado=False
            ).delete()
            
            # Crear notificaciones para VENDEDOR y JEFE
            if venta.vendedor and venta.vendedor != request.user:
                vendedor_es_jefe = venta.vendedor.perfil.rol == 'JEFE' if hasattr(venta.vendedor, 'perfil') else False
                if not vendedor_es_jefe:
                    Notificacion.objects.create(
                        usuario=venta.vendedor,
                        tipo='PAGO_CONFIRMADO',
                        mensaje=f"✅ Tu pago ha sido confirmado por el contador. Puedes proceder con la venta #{venta.pk}",
                        venta=venta,
                        abono=abono,
                        confirmado=True
                    )
            
            # Actualizar notificaciones del JEFE (INT: mostrar USD)
            jefes = User.objects.filter(perfil__rol='JEFE')
            forma_pago_display = dict(AbonoPago.FORMA_PAGO_CHOICES).get(abono.forma_pago, abono.forma_pago)
            if venta.tipo_viaje == 'INT' and abono.monto_usd_para_display is not None:
                monto_texto_jefe = f"USD ${abono.monto_usd_para_display:,.2f}"
            else:
                monto_texto_jefe = f"${abono.monto:,.2f}"
            mensaje_jefe = f"✅ Pago confirmado por el contador: {monto_texto_jefe} ({forma_pago_display}) - Venta #{venta.pk}"
            
            for jefe in jefes:
                # Actualizar notificaciones pendientes
                notificaciones_jefe = Notificacion.objects.filter(
                    usuario=jefe,
                    venta=venta,
                    abono=abono,
                    tipo='PAGO_PENDIENTE',
                    confirmado=False
                )
                if notificaciones_jefe.exists():
                    notificaciones_jefe.update(
                        tipo='PAGO_CONFIRMADO',
                        mensaje=mensaje_jefe,
                        confirmado=True,
                        confirmado_por=request.user,
                        confirmado_en=timezone.now()
                    )
                else:
                    Notificacion.objects.create(
                        usuario=jefe,
                        tipo='PAGO_CONFIRMADO',
                        mensaje=mensaje_jefe,
                        venta=venta,
                        abono=abono,
                        confirmado=True
                    )
            
            if venta.tipo_viaje == 'INT' and abono.monto_usd_para_display is not None:
                messages.success(request, f"Abono de USD ${abono.monto_usd_para_display:,.2f} confirmado exitosamente.")
            else:
                messages.success(request, f"Abono de ${abono.monto:,.2f} confirmado exitosamente.")
            
        elif tipo == 'apertura':
            venta = get_object_or_404(VentaViaje, pk=pk)
            
            # Verificar que tenga comprobante subido (solo para TRN/TAR/DEP, no para crédito)
            if venta.modo_pago_apertura in ['TRN', 'TAR', 'DEP'] and not venta.comprobante_apertura_subido:
                messages.error(request, "Este pago de apertura no tiene comprobante subido.")
                return redirect('pagos_por_confirmar')
            
            # Confirmar la apertura: estado COMPLETADO y apertura_confirmada=True
            # (apertura_confirmada es lo que usa total_pagado y la UI para mostrar "Confirmado")
            venta.estado_confirmacion = 'COMPLETADO'
            venta.apertura_confirmada = True
            venta.save(update_fields=['estado_confirmacion', 'apertura_confirmada'])
            
            # Eliminar notificaciones pendientes del CONTADOR
            Notificacion.objects.filter(
                usuario=request.user,
                venta=venta,
                abono__isnull=True,  # Notificaciones sin abono son de apertura
                tipo='PAGO_PENDIENTE',
                confirmado=False
            ).delete()
            
            # Crear notificaciones para VENDEDOR y JEFE
            if venta.vendedor and venta.vendedor != request.user:
                vendedor_es_jefe = venta.vendedor.perfil.rol == 'JEFE' if hasattr(venta.vendedor, 'perfil') else False
                if not vendedor_es_jefe:
                    Notificacion.objects.create(
                        usuario=venta.vendedor,
                        tipo='PAGO_CONFIRMADO',
                        mensaje=f"✅ Tu pago de apertura ha sido confirmado por el contador. Puedes proceder con la venta #{venta.pk}",
                        venta=venta,
                        confirmado=True
                    )
            
            # Actualizar notificaciones del JEFE
            jefes = User.objects.filter(perfil__rol='JEFE')
            modo_pago_display = dict(VentaViaje.MODO_PAGO_CHOICES).get(venta.modo_pago_apertura, venta.modo_pago_apertura)
            if venta.modo_pago_apertura == 'CRE':
                mensaje_jefe = f"✅ Crédito confirmado por el contador ({modo_pago_display}) - Venta #{venta.pk}"
            else:
                mensaje_jefe = f"✅ Pago de apertura confirmado por el contador: ${venta.cantidad_apertura:,.2f} ({modo_pago_display}) - Venta #{venta.pk}"
            
            for jefe in jefes:
                # Actualizar notificaciones pendientes
                notificaciones_jefe = Notificacion.objects.filter(
                    usuario=jefe,
                    venta=venta,
                    abono__isnull=True,
                    tipo='PAGO_PENDIENTE',
                    confirmado=False
                )
                if notificaciones_jefe.exists():
                    notificaciones_jefe.update(
                        tipo='PAGO_CONFIRMADO',
                        mensaje=mensaje_jefe,
                        confirmado=True,
                        confirmado_por=request.user,
                        confirmado_en=timezone.now()
                    )
                else:
                    Notificacion.objects.create(
                        usuario=jefe,
                        tipo='PAGO_CONFIRMADO',
                        mensaje=mensaje_jefe,
                        venta=venta,
                        confirmado=True
                    )
            
            messages.success(request, f"Pago de apertura de ${venta.cantidad_apertura:,.2f} confirmado exitosamente.")
        else:
            messages.error(request, "Tipo de pago inválido.")
            return redirect('pagos_por_confirmar')
        
        return redirect('pagos_por_confirmar')


# ------------------- 9. COTIZACIONES -------------------
class CotizacionListView(LoginRequiredMixin, ListView):
    model = Cotizacion
    template_name = 'ventas/cotizacion_list.html'
    context_object_name = 'cotizaciones'
    paginate_by = 13  # Mismo criterio que lista de ventas; navegación por números debajo

    def get_queryset(self):
        user = self.request.user
        qs = perm.get_cotizaciones_queryset_base(Cotizacion, user, self.request).select_related('cliente', 'vendedor')
        cliente_id = self.request.GET.get('cliente')
        estado = self.request.GET.get('estado')
        if cliente_id:
            qs = qs.filter(cliente_id=cliente_id)
        if estado:
            qs = qs.filter(estado=estado)
        return qs.order_by('-creada_en')

    def get_context_data(self, **kwargs):
        from datetime import timedelta
        context = super().get_context_data(**kwargs)
        today = timezone.localdate()
        start_week = today - timedelta(days=today.weekday())
        start_month = today.replace(day=1)

        user = self.request.user
        qs_base = perm.get_cotizaciones_queryset_base(Cotizacion, user, self.request)

        def rangos(qs):
            return {
                'hoy': qs.filter(creada_en__date=today).count(),
                'semana': qs.filter(creada_en__date__gte=start_week).count(),
                'mes': qs.filter(creada_en__date__gte=start_month).count(),
            }

        context['stats_propias'] = rangos(qs_base)
        context['stats_globales'] = rangos(qs_base)
        context['busqueda_cliente'] = self.request.GET.get('cliente', '')
        context['busqueda_estado'] = self.request.GET.get('estado', '')
        return context


class CotizacionCreateView(LoginRequiredMixin, CreateView):
    model = Cotizacion
    form_class = CotizacionForm
    template_name = 'ventas/cotizacion_form.html'

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        
        # Pre-seleccionar cliente si viene en la URL (desde detalle del cliente)
        cliente_pk = self.request.GET.get('cliente_pk')
        if cliente_pk:
            if 'initial' not in kwargs:
                kwargs['initial'] = {}
            try:
                cliente = Cliente.objects.get(pk=cliente_pk)
                kwargs['initial']['cliente'] = cliente
            except Cliente.DoesNotExist:
                pass
        
        return kwargs

    def form_valid(self, form):
        form.instance.vendedor = self.request.user
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('cotizacion_detalle', kwargs={'slug': self.object.slug})


class CotizacionUpdateView(LoginRequiredMixin, UpdateView):
    model = Cotizacion
    form_class = CotizacionForm
    template_name = 'ventas/cotizacion_form.html'
    slug_field = 'slug'
    slug_url_kwarg = 'slug'

    def get_queryset(self):
        return perm.get_cotizaciones_queryset_base(Cotizacion, self.request.user, self.request).select_related('cliente', 'vendedor')

    def get_success_url(self):
        return reverse('cotizacion_detalle', kwargs={'slug': self.object.slug})


class CotizacionDetailView(LoginRequiredMixin, DetailView):
    model = Cotizacion
    template_name = 'ventas/cotizacion_detail.html'
    slug_field = 'slug'
    slug_url_kwarg = 'slug'
    context_object_name = 'cotizacion'

    def get_queryset(self):
        return perm.get_cotizaciones_queryset_base(Cotizacion, self.request.user, self.request).select_related('cliente', 'vendedor')

    def get_object(self, queryset=None):
        obj = super().get_object(queryset)
        if not obj.folio:
            obj.save(update_fields=[])  # trigger folio generation
        return obj


class CotizacionDocxView(LoginRequiredMixin, DetailView):
    model = Cotizacion
    slug_field = 'slug'
    slug_url_kwarg = 'slug'

    def get_queryset(self):
        return perm.get_cotizaciones_queryset_base(Cotizacion, self.request.user, self.request)

    def get(self, request, *args, **kwargs):
        try:
            cot = self.get_object()
        except Exception as e:
            logging.error(f"Error obteniendo cotización: {e}")
            return HttpResponse(f"Error: No se pudo obtener la cotización. {str(e)}", status=400)
        
        template_path = os.path.join(settings.BASE_DIR, 'static', 'docx', 'membrete.docx')
        if not os.path.exists(template_path):
            error_msg = f'La plantilla DOCX no fue encontrada en: {template_path}'
            logging.error(error_msg)
            return HttpResponse(error_msg, status=404)

        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.oxml.ns import qn
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as e:
            error_msg = f'python-docx no está instalado en el entorno: {e}'
            logging.error(error_msg)
            return HttpResponse(error_msg, status=500)

        def format_date(value):
            if not value:
                return '-'
            try:
                if isinstance(value, datetime.date):
                    parsed = value
                else:
                    parsed = datetime.date.fromisoformat(str(value))
                return parsed.strftime('%d/%m/%Y')
            except Exception:
                return str(value)

        def format_currency(value):
            if value in (None, '', 0):
                return '0.00'
            try:
                number = Decimal(str(value).replace(',', ''))
            except Exception:
                return str(value)
            return f"{number:,.2f}"

        doc = Document(template_path)
        section = doc.sections[0]

        MOVUMS_BLUE = RGBColor(15, 92, 192)
        MOVUMS_LIGHT_BLUE = RGBColor(92, 141, 214)
        TEXT_COLOR = RGBColor(20, 20, 20)
        MOVUMS_BLUE_CORP = RGBColor(0, 74, 142)

        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        def set_run_font(run, size=12, bold=False, color=TEXT_COLOR):
            run.font.name = 'Arial'
            run.font.size = Pt(size)
            run.bold = bold
            run.font.color.rgb = color

        def add_paragraph(doc_obj, text='', size=12, bold=False, color=TEXT_COLOR, space_before=0, space_after=0):
            paragraph = doc_obj.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(space_before)
            paragraph.paragraph_format.space_after = Pt(space_after)
            run = paragraph.add_run(text)
            set_run_font(run, size=size, bold=bold, color=color)
            return paragraph

        # Obtener fecha de cotización desde propuestas o usar fecha de creación
        fecha_cotizacion = None
        if isinstance(cot.propuestas, dict) and cot.propuestas.get('fecha_cotizacion'):
            try:
                fecha_cotizacion = datetime.date.fromisoformat(cot.propuestas['fecha_cotizacion'])
            except (ValueError, TypeError):
                fecha_cotizacion = cot.creada_en.date() if cot.creada_en else None
        else:
            fecha_cotizacion = cot.creada_en.date() if cot.creada_en else None
        
        # Título de la cotización
        titulo_paragraph = doc.add_paragraph()
        titulo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        titulo_paragraph.paragraph_format.space_before = Pt(12)
        titulo_paragraph.paragraph_format.space_after = Pt(6)
        titulo_run = titulo_paragraph.add_run(cot.titulo or 'Cotización de Viaje')
        set_run_font(titulo_run, size=16, bold=True, color=MOVUMS_BLUE_CORP)
        
        # Información del cliente
        cliente_paragraph = doc.add_paragraph()
        cliente_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cliente_paragraph.paragraph_format.space_after = Pt(6)
        cliente_label_run = cliente_paragraph.add_run('Cliente: ')
        set_run_font(cliente_label_run, size=12, bold=True, color=TEXT_COLOR)
        cliente_value_run = cliente_paragraph.add_run(cot.cliente.nombre_completo_display)
        set_run_font(cliente_value_run, size=12, bold=False, color=TEXT_COLOR)
        
        # Fecha de cotización
        fecha_paragraph = doc.add_paragraph()
        fecha_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fecha_paragraph.paragraph_format.space_after = Pt(12)
        fecha_run = fecha_paragraph.add_run(f"Fecha de Cotización: {format_date(fecha_cotizacion)}")
        set_run_font(fecha_run, size=14, bold=True, color=MOVUMS_BLUE)

        propuestas = cot.propuestas if isinstance(cot.propuestas, dict) else {}
        tipo = propuestas.get('tipo', 'vuelos')  # Default a 'vuelos' si no está en propuestas
        
        # Determinar si es traslados o renta_autos para ajustar la tabla de información
        es_traslados = tipo == 'traslados'
        es_renta_autos = tipo == 'renta_autos'
        
        # Para traslados, usar datos desde propuestas si están disponibles
        if es_traslados and propuestas.get('traslados'):
            traslados_data = propuestas.get('traslados', {})
            origen_val = traslados_data.get('desde') or cot.origen or '-'
            destino_val = traslados_data.get('hasta') or cot.destino or '-'
        elif es_renta_autos and propuestas.get('renta_autos'):
            # Para renta_autos, usar punto de origen y punto de regreso desde propuestas
            renta_autos_data = propuestas.get('renta_autos', {})
            origen_val = renta_autos_data.get('punto_origen') or '-'
            destino_val = renta_autos_data.get('punto_regreso') or '-'
        else:
            origen_val = cot.origen or '-'
            destino_val = cot.destino or '-'
        
        # Preparar información de edades de menores
        edades_menores_texto = ''
        if cot.edades_menores and cot.edades_menores.strip():
            edades_menores_texto = f" (Edades: {cot.edades_menores})"
        
        # Construir datos de la tabla según el tipo
        if es_traslados:
            # Para traslados: solo mostrar origen/destino y pasajeros
            pasajeros_texto = f"{cot.adultos or 0} Adultos / {cot.menores or 0} Menores{edades_menores_texto}"
            info_data = [
                ("Desde / Hasta", origen_val, destino_val),
                ("Pasajeros", str(cot.pasajeros) or '1', pasajeros_texto),
            ]
            info_table = doc.add_table(rows=len(info_data), cols=3)
        elif es_renta_autos:
            # Para renta_autos: solo mostrar punto de origen/regreso y pasajeros (sin fechas ni días/noches)
            pasajeros_texto = f"{cot.adultos or 0} Adultos / {cot.menores or 0} Menores{edades_menores_texto}"
            info_data = [
                ("Punto de Origen / Punto de Regreso", origen_val, destino_val),
                ("Pasajeros", str(cot.pasajeros) or '1', pasajeros_texto),
            ]
            info_table = doc.add_table(rows=len(info_data), cols=3)
        else:
            # Para otros tipos: mostrar toda la información
            pasajeros_texto = f"{cot.adultos or 0} Adultos / {cot.menores or 0} Menores{edades_menores_texto}"
            info_data = [
                ("Origen / Destino", origen_val, destino_val),
                ("Inicio / Fin", format_date(cot.fecha_inicio) if cot.fecha_inicio else '-', format_date(cot.fecha_fin) if cot.fecha_fin else '-'),
                ("Pasajeros", str(cot.pasajeros) or '1', pasajeros_texto),
                ("Viaje", f"{cot.dias or '-'} días", f"{cot.noches or '-'} noches"),
            ]
            info_table = doc.add_table(rows=len(info_data), cols=3)
        
        for row_idx, (label, v1, v2) in enumerate(info_data):
            row = info_table.rows[row_idx].cells
            label_run = row[0].paragraphs[0].add_run(label)
            set_run_font(label_run, size=14, bold=True, color=MOVUMS_BLUE)
            val1_run = row[1].paragraphs[0].add_run(v1)
            set_run_font(val1_run, size=12)
            val2_run = row[2].paragraphs[0].add_run(v2)
            set_run_font(val2_run, size=12)

        add_paragraph(doc, "", space_after=6)
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

        def agregar_subtitulo_con_vineta(texto):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            bullet_run = p.add_run('• ')
            set_run_font(bullet_run, size=14, bold=True, color=MOVUMS_BLUE_CORP)
            texto_run = p.add_run(texto)
            set_run_font(texto_run, size=14, bold=True, color=MOVUMS_BLUE_CORP)
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(2)
            return p

        def agregar_info_line(etiqueta, valor):
            if not valor:
                return
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            label_run = p.add_run(f'{etiqueta}: ')
            set_run_font(label_run, size=12, bold=True)
            value_run = p.add_run(str(valor))
            set_run_font(value_run, size=12)
            return p

        def agregar_info_inline(*pares_etiqueta_valor, separador=' | '):
            if not pares_etiqueta_valor:
                return
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            for idx, (etiqueta, valor) in enumerate(pares_etiqueta_valor):
                if not valor:
                    continue
                if idx > 0:
                    sep_run = p.add_run(separador)
                    set_run_font(sep_run, size=12)
                label_run = p.add_run(f'{etiqueta}: ')
                set_run_font(label_run, size=12, bold=True)
                value_run = p.add_run(str(valor))
                set_run_font(value_run, size=12)
            return p

        def agregar_salto_entre_secciones():
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(6)
            return spacer

        def agregar_titulo_principal(texto):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(texto)
            set_run_font(run, size=18, bold=True, color=MOVUMS_BLUE_CORP)
            return p

        if tipo == 'vuelos' and propuestas.get('vuelos'):
            for vuelo in propuestas.get('vuelos', []):
                agregar_titulo_principal("VUELO")
                agregar_subtitulo_con_vineta('Información del Vuelo')
                agregar_info_inline(
                    ('Aerolínea', vuelo.get('aerolinea') or '-'),
                    ('Salida', vuelo.get('salida') or '-'),
                    ('Regreso', vuelo.get('regreso') or '-')
                )
                agregar_info_line('Incluye', vuelo.get('incluye') or '-')
                if vuelo.get('forma_pago'):
                    agregar_info_line('Forma de Pago', vuelo.get('forma_pago'))
                total_p = doc.add_paragraph()
                total_p.paragraph_format.space_before = Pt(0)
                total_p.paragraph_format.space_after = Pt(6)
                total_run = total_p.add_run(f"Total MXN {format_currency(vuelo.get('total'))} Pesos")
                set_run_font(total_run, size=18, bold=True, color=MOVUMS_BLUE_CORP)
                total_run.font.underline = True
                agregar_salto_entre_secciones()

        elif tipo == 'hospedaje' and propuestas.get('hoteles'):
            for hotel in propuestas.get('hoteles', []):
                agregar_titulo_principal("HOSPEDAJE")
                agregar_subtitulo_con_vineta('Información del Alojamiento')
                info_table = doc.add_table(rows=4, cols=2)
                info_table.autofit = False
                for col in info_table.columns:
                    for cell in col.cells:
                        cell.width = Inches(3.25)
                nombre_cell = info_table.rows[0].cells[0]
                nombre_label = nombre_cell.paragraphs[0].add_run('Nombre: ')
                set_run_font(nombre_label, size=12, bold=True)
                nombre_val = nombre_cell.paragraphs[0].add_run(hotel.get('nombre') or 'Hotel propuesto')
                set_run_font(nombre_val, size=12)
                habitacion_cell = info_table.rows[1].cells[0]
                habitacion_label = habitacion_cell.paragraphs[0].add_run('Habitación: ')
                set_run_font(habitacion_label, size=12, bold=True)
                habitacion_val = habitacion_cell.paragraphs[0].add_run(hotel.get('habitacion') or '-')
                set_run_font(habitacion_val, size=12)
                direccion_cell = info_table.rows[2].cells[0]
                direccion_label = direccion_cell.paragraphs[0].add_run('Dirección: ')
                set_run_font(direccion_label, size=12, bold=True)
                direccion_val = direccion_cell.paragraphs[0].add_run(hotel.get('direccion') or '-')
                set_run_font(direccion_val, size=12)
                plan_cell = info_table.rows[0].cells[1]
                plan_label = plan_cell.paragraphs[0].add_run('Plan de Alimentos: ')
                set_run_font(plan_label, size=12, bold=True)
                plan_val = plan_cell.paragraphs[0].add_run(hotel.get('plan') or '-')
                set_run_font(plan_val, size=12)
                if hotel.get('forma_pago'):
                    forma_pago_cell = info_table.rows[1].cells[1]
                    forma_pago_label = forma_pago_cell.paragraphs[0].add_run('Forma de Pago: ')
                    set_run_font(forma_pago_label, size=12, bold=True)
                    forma_pago_val = forma_pago_cell.paragraphs[0].add_run(hotel.get('forma_pago'))
                    set_run_font(forma_pago_val, size=12)
                total_p = doc.add_paragraph()
                total_p.paragraph_format.space_before = Pt(0)
                total_p.paragraph_format.space_after = Pt(6)
                total_run = total_p.add_run(f"Total MXN {format_currency(hotel.get('total'))} Pesos")
                set_run_font(total_run, size=18, bold=True, color=MOVUMS_BLUE_CORP)
                total_run.font.underline = True
                agregar_salto_entre_secciones()

        elif tipo == 'paquete' and propuestas.get('paquete'):
            paquete = propuestas.get('paquete', {})
            vuelo = paquete.get('vuelo') or {}
            hotel = paquete.get('hotel') or {}
            agregar_titulo_principal("PAQUETE")
            agregar_subtitulo_con_vineta('Vuelo')
            agregar_info_inline(
                ('Aerolínea', vuelo.get('aerolinea') or '-'),
                ('Salida', vuelo.get('salida') or '-'),
                ('Regreso', vuelo.get('regreso') or '-')
            )
            agregar_info_line('Incluye', vuelo.get('incluye') or '-')
            if vuelo.get('forma_pago'):
                agregar_info_line('Forma de Pago', vuelo.get('forma_pago'))
            agregar_salto_entre_secciones()
            agregar_subtitulo_con_vineta('Hospedaje')
            hospedaje_table = doc.add_table(rows=3, cols=2)
            hospedaje_table.autofit = False
            for col in hospedaje_table.columns:
                for cell in col.cells:
                    cell.width = Inches(3.25)
            nombre_cell = hospedaje_table.rows[0].cells[0]
            nombre_label = nombre_cell.paragraphs[0].add_run('Nombre: ')
            set_run_font(nombre_label, size=12, bold=True)
            nombre_val = nombre_cell.paragraphs[0].add_run(hotel.get('nombre') or 'Hotel incluido')
            set_run_font(nombre_val, size=12)
            habitacion_cell = hospedaje_table.rows[1].cells[0]
            habitacion_label = habitacion_cell.paragraphs[0].add_run('Habitación / Plan: ')
            set_run_font(habitacion_label, size=12, bold=True)
            habitacion_val = habitacion_cell.paragraphs[0].add_run(hotel.get('habitacion') or '-')
            set_run_font(habitacion_val, size=12)
            notas_cell = hospedaje_table.rows[2].cells[0]
            notas_label = notas_cell.paragraphs[0].add_run('Notas: ')
            set_run_font(notas_label, size=12, bold=True)
            notas_val = notas_cell.paragraphs[0].add_run(hotel.get('notas') or '-')
            set_run_font(notas_val, size=12)
            if paquete.get('forma_pago'):
                agregar_info_line('Forma de Pago', paquete.get('forma_pago'))
            total_p = doc.add_paragraph()
            total_p.paragraph_format.space_before = Pt(0)
            total_p.paragraph_format.space_after = Pt(6)
            total_run = total_p.add_run(f"Total MXN {format_currency(paquete.get('total'))} Pesos")
            set_run_font(total_run, size=18, bold=True, color=MOVUMS_BLUE_CORP)
            total_run.font.underline = True
            agregar_salto_entre_secciones()
            
            # Tours del Paquete
            if paquete.get('tours') and isinstance(paquete.get('tours'), list) and len(paquete.get('tours')) > 0:
                agregar_subtitulo_con_vineta('Tours del Paquete')
                for idx, tour in enumerate(paquete.get('tours'), 1):
                    if len(paquete.get('tours')) > 1:
                        agregar_info_line('Tour', f'Tour del Paquete {idx}')
                    agregar_info_line('Nombre del Tour', tour.get('nombre') or '-')
                    
                    if tour.get('total'):
                        total_tour_p = doc.add_paragraph()
                        total_tour_p.paragraph_format.space_before = Pt(0)
                        total_tour_p.paragraph_format.space_after = Pt(6)
                        total_tour_run = total_tour_p.add_run(f"Total MXN {format_currency(tour.get('total'))} Pesos")
                        set_run_font(total_tour_run, size=16, bold=True, color=MOVUMS_BLUE_CORP)
                        total_tour_run.font.underline = True
                    
                    if tour.get('especificaciones'):
                        especificaciones = tour.get('especificaciones', '').strip()
                        if especificaciones:
                            agregar_salto_entre_secciones()
                            agregar_subtitulo_con_vineta('Especificaciones')
                            lineas = especificaciones.split('\n')
                            for linea in lineas:
                                if linea.strip():
                                    p = doc.add_paragraph()
                                    p.paragraph_format.space_after = Pt(4)
                                    run = p.add_run(linea.strip())
                                    set_run_font(run, size=12)
                    
                    if tour.get('forma_pago'):
                        agregar_salto_entre_secciones()
                        agregar_subtitulo_con_vineta('Forma de Pago')
                        agregar_info_line('Forma de Pago', tour.get('forma_pago'))
                    
                    # Agregar separador entre tours si hay más de uno
                    if idx < len(paquete.get('tours')):
                        agregar_salto_entre_secciones()
                        agregar_salto_entre_secciones()
                
                agregar_salto_entre_secciones()

        elif tipo == 'tours' and propuestas.get('tours'):
            tours_data = propuestas.get('tours', {})
            
            # Verificar si es un array de tours o un objeto único (compatibilidad)
            tours_list = []
            if isinstance(tours_data, list):
                tours_list = tours_data
            else:
                # Compatibilidad: convertir objeto único a array
                tours_list = [tours_data] if tours_data else []
            
            if tours_list:
                agregar_titulo_principal("TOURS" if len(tours_list) > 1 else "TOUR")
                
                for idx, tour in enumerate(tours_list, 1):
                    if len(tours_list) > 1:
                        agregar_subtitulo_con_vineta(f'Tour {idx}')
                    else:
                        agregar_subtitulo_con_vineta('Información del Tour')
                    
                    agregar_info_line('Nombre del Tour', tour.get('nombre') or '-')
                    
                    if tour.get('total'):
                        total_p = doc.add_paragraph()
                        total_p.paragraph_format.space_before = Pt(0)
                        total_p.paragraph_format.space_after = Pt(6)
                        total_run = total_p.add_run(f"Total MXN {format_currency(tour.get('total'))} Pesos")
                        set_run_font(total_run, size=16, bold=True, color=MOVUMS_BLUE_CORP)
                        total_run.font.underline = True
                    
                    if tour.get('especificaciones'):
                        agregar_salto_entre_secciones()
                        agregar_subtitulo_con_vineta('Especificaciones')
                        especificaciones = tour.get('especificaciones', '').strip()
                        if especificaciones:
                            lineas = especificaciones.split('\n')
                            for linea in lineas:
                                if linea.strip():
                                    p = doc.add_paragraph()
                                    p.paragraph_format.space_after = Pt(4)
                                    run = p.add_run(linea.strip())
                                    set_run_font(run, size=12)
                    
                    if tour.get('forma_pago'):
                        agregar_salto_entre_secciones()
                        agregar_subtitulo_con_vineta('Forma de Pago')
                        agregar_info_line('Forma de Pago', tour.get('forma_pago'))
                    
                    # Agregar separador entre tours si hay más de uno
                    if idx < len(tours_list):
                        agregar_salto_entre_secciones()
                        agregar_salto_entre_secciones()
                
                agregar_salto_entre_secciones()

        elif tipo == 'traslados' and propuestas.get('traslados'):
            traslados_data = propuestas.get('traslados', {})
            agregar_titulo_principal("TRASLADO")
            agregar_subtitulo_con_vineta('Información del Traslado')
            agregar_info_inline(
                ('Tipo', traslados_data.get('tipo') or '-'),
                ('Modalidad', traslados_data.get('modalidad') or '-')
            )
            agregar_info_inline(
                ('Desde', traslados_data.get('desde') or '-'),
                ('Hasta', traslados_data.get('hasta') or '-')
            )
            # Si es traslado redondo, mostrar fechas y horarios
            if traslados_data.get('modalidad') == 'REDONDO':
                if traslados_data.get('fecha_ida') or traslados_data.get('fecha_regreso'):
                    agregar_info_inline(
                        ('Fecha de Ida', format_date(traslados_data.get('fecha_ida'))),
                        ('Fecha de Regreso', format_date(traslados_data.get('fecha_regreso')))
                    )
                if traslados_data.get('hora_ida') or traslados_data.get('hora_regreso'):
                    agregar_info_inline(
                        ('Hora de Ida', traslados_data.get('hora_ida') or '-'),
                        ('Hora de Regreso', traslados_data.get('hora_regreso') or '-')
                    )
            if traslados_data.get('descripcion'):
                agregar_salto_entre_secciones()
                agregar_subtitulo_con_vineta('Descripción')
                descripcion = traslados_data.get('descripcion', '').strip()
                if descripcion:
                    lineas = descripcion.split('\n')
                    for linea in lineas:
                        if linea.strip():
                            p = doc.add_paragraph()
                            p.paragraph_format.space_after = Pt(4)
                            run = p.add_run(linea.strip())
                            set_run_font(run, size=12)
            if traslados_data.get('forma_pago'):
                agregar_salto_entre_secciones()
                agregar_subtitulo_con_vineta('Forma de Pago')
                agregar_info_line('Forma de Pago', traslados_data.get('forma_pago'))
            total_p = doc.add_paragraph()
            total_p.paragraph_format.space_before = Pt(0)
            total_p.paragraph_format.space_after = Pt(6)
            total_run = total_p.add_run(f"Total MXN {format_currency(traslados_data.get('total'))} Pesos")
            set_run_font(total_run, size=18, bold=True, color=MOVUMS_BLUE_CORP)
            total_run.font.underline = True
            agregar_salto_entre_secciones()

        elif tipo == 'renta_autos' and propuestas.get('renta_autos'):
            renta_autos_data = propuestas.get('renta_autos', {})
            agregar_titulo_principal("RENTA DE AUTOS")
            agregar_subtitulo_con_vineta('Información de la Renta')
            agregar_info_inline(
                ('Arrendadora', renta_autos_data.get('arrendadora') or '-'),
                ('Punto de Origen', renta_autos_data.get('punto_origen') or '-'),
                ('Punto de Regreso', renta_autos_data.get('punto_regreso') or '-')
            )
            if renta_autos_data.get('hora_pickup') or renta_autos_data.get('hora_devolucion'):
                agregar_info_inline(
                    ('Hora de Pickup', renta_autos_data.get('hora_pickup') or '-'),
                    ('Hora de Devolución', renta_autos_data.get('hora_devolucion') or '-')
                )
            if renta_autos_data.get('forma_pago'):
                agregar_salto_entre_secciones()
                agregar_subtitulo_con_vineta('Forma de Pago')
                agregar_info_line('Forma de Pago', renta_autos_data.get('forma_pago'))
            total_p = doc.add_paragraph()
            total_p.paragraph_format.space_before = Pt(0)
            total_p.paragraph_format.space_after = Pt(6)
            total_run = total_p.add_run(f"Total MXN {format_currency(renta_autos_data.get('total'))} Pesos")
            set_run_font(total_run, size=18, bold=True, color=MOVUMS_BLUE_CORP)
            total_run.font.underline = True
            agregar_salto_entre_secciones()

        elif tipo == 'generica' and propuestas.get('generica'):
            generica_data = propuestas.get('generica', {})
            agregar_titulo_principal("COTIZACIÓN GENÉRICA")
            if generica_data.get('contenido'):
                contenido = generica_data.get('contenido', '').strip()
                if contenido:
                    lineas = contenido.split('\n')
                    for linea in lineas:
                        if linea.strip():
                            p = doc.add_paragraph()
                            p.paragraph_format.space_after = Pt(4)
                            run = p.add_run(linea.strip())
                            set_run_font(run, size=12)

        try:
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            if buffer.getvalue() is None or len(buffer.getvalue()) == 0:
                error_msg = 'El documento generado está vacío'
                logging.error(error_msg)
                return HttpResponse(error_msg, status=500)

            filename = f"cotizacion_{cot.slug}_{timezone.now().strftime('%Y%m%d%H%M%S')}.docx"
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
        except Exception as e:
            error_msg = f'Error al generar el documento DOCX: {str(e)}'
            logging.error(error_msg, exc_info=True)
            return HttpResponse(error_msg, status=500)


class CotizacionPDFView(LoginRequiredMixin, DetailView):
    """
    Vista para generar cotizaciones en formato PDF usando WeasyPrint.
    Incluye sistema de cache para mejorar el rendimiento.
    """
    model = Cotizacion
    slug_field = 'slug'
    slug_url_kwarg = 'slug'

    def get_queryset(self):
        return perm.get_cotizaciones_queryset_base(Cotizacion, self.request.user, self.request)

    def get(self, request, *args, **kwargs):
        try:
            cot = self.get_object()
        except Exception as e:
            logging.error(f"Error obteniendo cotización: {e}")
            return HttpResponse(f"Error: No se pudo obtener la cotización. {str(e)}", status=400)
        
        # Verificar disponibilidad de WeasyPrint
        if not WEASYPRINT_AVAILABLE:
            return HttpResponse(
                "Error: WeasyPrint no está disponible. Por favor instálalo con: pip install weasyprint", 
                status=500
            )
        
        # Cache deshabilitado temporalmente para forzar regeneración con nuevos estilos
        # TODO: Re-habilitar cache una vez que los estilos estén estabilizados
        # pdf_path = self._get_cache_path(cot)
        # if os.path.exists(pdf_path):
        #     cache_mtime = os.path.getmtime(pdf_path)
        #     cot_mtime = cot.actualizada_en.timestamp() if cot.actualizada_en else 0
        #     if cache_mtime >= cot_mtime:
        #         try:
        #             with open(pdf_path, 'rb') as f:
        #                 pdf_content = f.read()
        #             return self._crear_respuesta_pdf(pdf_content, cot)
        #         except Exception as e:
        #             logging.warning(f"Error leyendo cache, regenerando: {e}")
        
        # Generar nuevo PDF
        try:
            pdf_content = self._generar_pdf(cot)
            
            # Cache deshabilitado temporalmente
            # Guardar en cache
            # try:
            #     pdf_path = self._get_cache_path(cot)
            #     os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
            #     with open(pdf_path, 'wb') as f:
            #         f.write(pdf_content)
            # except Exception as e:
            #     logging.warning(f"Error guardando en cache: {e}")
            
            return self._crear_respuesta_pdf(pdf_content, cot)
            
        except Exception as e:
            error_msg = f'Error al generar el documento PDF: {str(e)}'
            logging.error(error_msg, exc_info=True)
            return HttpResponse(error_msg, status=500)
    
    def _get_cache_path(self, cotizacion):
        """Genera la ruta del archivo cacheado."""
        cache_dir = os.path.join(settings.MEDIA_ROOT, 'cache', 'pdfs')
        # Usar slug y timestamp de actualización para invalidar cache automáticamente
        timestamp = int(cotizacion.actualizada_en.timestamp()) if cotizacion.actualizada_en else 0
        filename = f"cotizacion_{cotizacion.slug}_{timestamp}.pdf"
        return os.path.join(cache_dir, filename)
    
    def _preparar_contexto(self, cotizacion):
        """Prepara el contexto para la plantilla."""
        propuestas = cotizacion.propuestas if isinstance(cotizacion.propuestas, dict) else {}
        tipo = propuestas.get('tipo', 'vuelos')
        
        # Para cotizaciones de paquete: normalizar tours para que siempre se muestren en el PDF
        if tipo == 'paquete':
            propuestas = dict(propuestas)
            paquete = propuestas.get('paquete')
            if paquete is None:
                paquete = {}
                propuestas['paquete'] = paquete
            else:
                paquete = dict(paquete)
                propuestas['paquete'] = paquete
            
            tours_raw = paquete.get('tours') or paquete.get('Tours')
            if tours_raw is None and propuestas.get('tours'):
                tours_raw = propuestas.get('tours')
            
            if isinstance(tours_raw, list):
                tours_list = [t for t in tours_raw if isinstance(t, dict)]
            elif isinstance(tours_raw, dict) and tours_raw:
                tours_list = [tours_raw]
            else:
                tours_list = []
            
            paquete['tours'] = tours_list
        
        # Para cotizaciones de tours: normalizar propuestas['tours'] para que la tabla muestre datos
        if tipo == 'tours':
            propuestas = dict(propuestas)
            tours_raw = propuestas.get('tours') or propuestas.get('Tours') or propuestas.get('tour')
            if isinstance(tours_raw, list):
                tours_list = []
                for t in tours_raw:
                    if isinstance(t, dict):
                        tours_list.append({
                            'nombre': t.get('nombre') or t.get('nombre_tour') or '-',
                            'especificaciones': t.get('especificaciones') or '',
                            'forma_pago': t.get('forma_pago') or '',
                            'total': t.get('total'),
                        })
                propuestas['tours'] = tours_list
            elif isinstance(tours_raw, dict) and tours_raw:
                propuestas['tours'] = [{
                    'nombre': tours_raw.get('nombre') or tours_raw.get('nombre_tour') or '-',
                    'especificaciones': tours_raw.get('especificaciones') or '',
                    'forma_pago': tours_raw.get('forma_pago') or '',
                    'total': tours_raw.get('total'),
                }]
            else:
                propuestas['tours'] = []

        # Determinar template según tipo
        template_map = {
            'vuelos': 'ventas/pdf/cotizacion_vuelos_pdf.html',
            'hospedaje': 'ventas/pdf/cotizacion_hospedaje_pdf.html',
            'paquete': 'ventas/pdf/cotizacion_paquete_pdf.html',
            'tours': 'ventas/pdf/cotizacion_tours_pdf.html',
            'traslados': 'ventas/pdf/cotizacion_traslados_pdf.html',
            'renta_autos': 'ventas/pdf/cotizacion_renta_autos_pdf.html',
            'generica': 'ventas/pdf/cotizacion_generica_pdf.html',
        }
        
        template_name = template_map.get(tipo, 'ventas/pdf/cotizacion_vuelos_pdf.html')
        
        # Preparar ruta absoluta file:// para el membrete (WeasyPrint necesita URL absoluta)
        membrete_path = os.path.join(settings.BASE_DIR, 'static', 'img', 'membrete_movums.jpg')
        membrete_url = None
        if os.path.exists(membrete_path):
            # Crear URL file:// absoluta para WeasyPrint
            membrete_abs_path = os.path.abspath(membrete_path)
            # En Windows, necesitamos ajustar el formato de la ruta
            if os.name == 'nt':
                membrete_url = f"file:///{membrete_abs_path.replace(os.sep, '/')}"
            else:
                membrete_url = f"file://{membrete_abs_path}"
        
        # Obtener información del ejecutivo que realizó la cotización
        ejecutivo = None
        ejecutivo_nombre = None
        ejecutivo_telefono = None
        ejecutivo_email = None
        
        if cotizacion.vendedor:
            try:
                # Intentar obtener el Ejecutivo asociado al usuario
                ejecutivo = Ejecutivo.objects.filter(usuario=cotizacion.vendedor).first()
                
                if ejecutivo:
                    # Si hay ejecutivo, usar su nombre, teléfono y email
                    ejecutivo_nombre = ejecutivo.nombre_completo
                    if ejecutivo.telefono:
                        # Formatear teléfono a 10 dígitos (eliminar espacios, guiones, paréntesis, etc.)
                        telefono_limpio = re.sub(r'[^\d]', '', ejecutivo.telefono)
                        # Si tiene más de 10 dígitos (puede incluir código de país), tomar los últimos 10
                        if len(telefono_limpio) > 10:
                            ejecutivo_telefono = telefono_limpio[-10:]
                        elif len(telefono_limpio) == 10:
                            ejecutivo_telefono = telefono_limpio
                        else:
                            ejecutivo_telefono = ejecutivo.telefono  # Mantener original si no tiene 10 dígitos
                    if ejecutivo.email:
                        ejecutivo_email = ejecutivo.email
                else:
                    # Si no hay ejecutivo asociado, usar el nombre del usuario como fallback
                    ejecutivo_nombre = f"{cotizacion.vendedor.get_full_name() or cotizacion.vendedor.get_username()}"
                    if cotizacion.vendedor.email:
                        ejecutivo_email = cotizacion.vendedor.email
                    logger.info(f"Usuario {cotizacion.vendedor.username} no tiene Ejecutivo asociado, usando nombre de usuario")
            except Exception as e:
                logger.warning(f"Error al obtener ejecutivo para cotización {cotizacion.pk}: {e}", exc_info=True)
                # Fallback: usar el nombre del usuario
                if cotizacion.vendedor:
                    ejecutivo_nombre = f"{cotizacion.vendedor.get_full_name() or cotizacion.vendedor.get_username()}"
                    if cotizacion.vendedor.email:
                        ejecutivo_email = cotizacion.vendedor.email
        
        # Procesar fecha de cotización: convertir string ISO a objeto date si es necesario
        fecha_cotizacion_obj = None
        if isinstance(propuestas, dict) and propuestas.get('fecha_cotizacion'):
            try:
                # Si es string ISO, convertir a date
                if isinstance(propuestas['fecha_cotizacion'], str):
                    fecha_cotizacion_obj = datetime.date.fromisoformat(propuestas['fecha_cotizacion'])
                elif isinstance(propuestas['fecha_cotizacion'], (datetime.date, datetime.datetime)):
                    fecha_cotizacion_obj = propuestas['fecha_cotizacion'] if isinstance(propuestas['fecha_cotizacion'], datetime.date) else propuestas['fecha_cotizacion'].date()
            except (ValueError, TypeError, AttributeError):
                # Fallback a fecha de creación
                fecha_cotizacion_obj = cotizacion.creada_en.date() if cotizacion.creada_en else None
        else:
            # Si no hay fecha en propuestas, usar fecha de creación
            fecha_cotizacion_obj = cotizacion.creada_en.date() if cotizacion.creada_en else None
        
        contexto = {
            'cotizacion': cotizacion,
            'propuestas': propuestas,
            'tipo': tipo,
            'template_name': template_name,
            'STATIC_URL': settings.STATIC_URL,
            'membrete_url': membrete_url,  # URL absoluta file:// para WeasyPrint
            'ejecutivo': ejecutivo,
            'ejecutivo_nombre': ejecutivo_nombre,
            'ejecutivo_telefono': ejecutivo_telefono,
            'ejecutivo_email': ejecutivo_email,
            'fecha_cotizacion': fecha_cotizacion_obj,  # Objeto date para usar con filtro |date
        }
        # Lista explícita de tours del paquete para el PDF (evita problemas de visualización en template)
        if tipo == 'paquete':
            contexto['paquete_tours'] = propuestas.get('paquete', {}).get('tours', [])
        return contexto
    
    def _generar_pdf(self, cotizacion):
        """
        Genera el PDF usando WeasyPrint para mejor control de diseño y formato con HTML/CSS.
        """
        from django.template.loader import render_to_string
        from weasyprint import HTML
        from io import BytesIO
        
        # Preparar contexto (ya incluye membrete en base64)
        context = self._preparar_contexto(cotizacion)
        
        # Renderizar HTML (el CSS ya está incrustado en base_cotizacion_pdf.html)
        html_string = render_to_string(context['template_name'], context, request=self.request)
        
        # Obtener rutas de archivos estáticos para recursos como imágenes
        static_dir = os.path.join(settings.BASE_DIR, 'static')
        static_dir_abs = os.path.abspath(static_dir)
        base_url = f"file://{static_dir_abs}/"
        
        # Generar PDF (sin CSS externo, ya está en el HTML)
        html = HTML(
            string=html_string, 
            base_url=base_url
        )
        
        pdf_buffer = BytesIO()
        html.write_pdf(pdf_buffer)
        pdf_buffer.seek(0)
        
        return pdf_buffer.getvalue()
    
    def _crear_respuesta_pdf(self, pdf_content, cotizacion):
        """Crea la respuesta HTTP con el PDF."""
        filename = f"cotizacion_{cotizacion.slug}_{timezone.now().strftime('%Y%m%d%H%M%S')}.pdf"
        
        response = HttpResponse(
            pdf_content,
            content_type='application/pdf'
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        response['Content-Length'] = str(len(pdf_content))
        response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response['Pragma'] = 'no-cache'
        response['Expires'] = '0'
        
        return response


class CotizacionConvertirView(LoginRequiredMixin, View):
    def post(self, request, *args, **kwargs):
        qs = perm.get_cotizaciones_queryset_base(Cotizacion, request.user, request)
        cot = get_object_or_404(qs, slug=kwargs.get('slug'))
        
        # Extraer el total de las propuestas según el tipo de cotización
        total_cotizacion = Decimal('0.00')
        propuestas = cot.propuestas if isinstance(cot.propuestas, dict) else {}
        tipo = propuestas.get('tipo', '')
        
        def limpiar_y_convertir_total(valor):
            """Convierte un string de total (puede tener comas) a Decimal"""
            if not valor:
                return Decimal('0.00')
            try:
                # Remover comas y espacios, luego convertir a Decimal
                valor_limpio = str(valor).replace(',', '').replace('$', '').strip()
                return Decimal(valor_limpio)
            except (ValueError, InvalidOperation):
                return Decimal('0.00')
        
        # Obtener índices seleccionados del formulario
        opcion_vuelo_index = request.POST.get('opcion_vuelo_index', '')
        opcion_hotel_index = request.POST.get('opcion_hotel_index', '')
        
        # Guardar los índices en la sesión para que el formulario los use
        if opcion_vuelo_index:
            request.session[f'cotizacion_{cot.slug}_opcion_vuelo'] = opcion_vuelo_index
        if opcion_hotel_index:
            request.session[f'cotizacion_{cot.slug}_opcion_hotel'] = opcion_hotel_index
        
        if tipo == 'vuelos' and propuestas.get('vuelos'):
            # Para vuelos, usar la opción seleccionada o la primera por defecto
            vuelos = propuestas.get('vuelos', [])
            if vuelos and len(vuelos) > 0:
                # Determinar qué índice usar
                try:
                    indice = int(opcion_vuelo_index) if opcion_vuelo_index else 0
                    if indice < 0 or indice >= len(vuelos):
                        indice = 0
                except (ValueError, TypeError):
                    indice = 0
                
                # Tomar el vuelo seleccionado usando el índice correcto
                vuelo_seleccionado = vuelos[indice] if isinstance(vuelos, list) else vuelos.get(f'propuesta_{indice + 1}', {})
                if isinstance(vuelo_seleccionado, dict):
                    total_vuelo = vuelo_seleccionado.get('total')
                    if total_vuelo:
                        total_cotizacion = limpiar_y_convertir_total(total_vuelo)
                        # Log para depuración (puedes removerlo después)
                        logging.debug(f"Total del vuelo seleccionado (índice {indice}): {total_cotizacion}")
                    else:
                        # Si no hay total en el vuelo seleccionado, intentar con el primero
                        if len(vuelos) > 0 and isinstance(vuelos[0], dict) and vuelos[0].get('total'):
                            total_cotizacion = limpiar_y_convertir_total(vuelos[0].get('total'))
                            logging.debug(f"Total del primer vuelo (fallback): {total_cotizacion}")
                else:
                    # Si no se encontró el vuelo seleccionado, usar el primero
                    if len(vuelos) > 0 and isinstance(vuelos[0], dict) and vuelos[0].get('total'):
                        total_cotizacion = limpiar_y_convertir_total(vuelos[0].get('total'))
                        logging.debug(f"Total del primer vuelo (fallback - vuelo no encontrado): {total_cotizacion}")
        
        elif tipo == 'hospedaje' and propuestas.get('hoteles'):
            # Para hospedaje, usar la opción seleccionada o la primera por defecto
            hoteles = propuestas.get('hoteles', [])
            
            # 1. Asegurar índice válido
            indice = 0
            try:
                if opcion_hotel_index:
                    indice = int(opcion_hotel_index)
            except (ValueError, TypeError):
                indice = 0
            
            # 2. Obtener el hotel usando lógica dual (Lista o Diccionario)
            hotel_seleccionado = {}
            if isinstance(hoteles, list):
                if 0 <= indice < len(hoteles):
                    hotel_seleccionado = hoteles[indice]
            else:
                # Si es diccionario, buscar por claves numéricas o string 'propuesta_X'
                hotel_seleccionado = hoteles.get(f'propuesta_{indice+1}') or hoteles.get(str(indice)) or {}
            
            # 3. Extraer total (Blindado)
            if hotel_seleccionado and isinstance(hotel_seleccionado, dict):
                total_str = hotel_seleccionado.get('total')
                if total_str:
                    total_cotizacion = limpiar_y_convertir_total(total_str)
                    logging.debug(f"Total HOTEL blindado (índice {indice}): {total_cotizacion}")
        
        elif tipo == 'paquete' and propuestas.get('paquete'):
            # Para paquete, tomar el total del paquete y sumar los tours
            paquete = propuestas.get('paquete', {})
            if isinstance(paquete, dict) and paquete.get('total'):
                total_cotizacion = limpiar_y_convertir_total(paquete.get('total'))
            
            # Sumar los totales de los tours del paquete
            if isinstance(paquete, dict) and paquete.get('tours') and isinstance(paquete.get('tours'), list):
                for tour in paquete.get('tours', []):
                    if isinstance(tour, dict) and tour.get('total'):
                        total_tour = limpiar_y_convertir_total(tour.get('total'))
                        total_cotizacion += total_tour
        
        elif tipo == 'tours' and propuestas.get('tours'):
            # Para tours, puede ser un array o un objeto único (compatibilidad)
            tours = propuestas.get('tours', {})
            tours_list = []
            if isinstance(tours, list):
                tours_list = tours
            else:
                # Compatibilidad: convertir objeto único a array
                tours_list = [tours] if tours else []
            
            # Sumar todos los totales de los tours
            for tour in tours_list:
                if isinstance(tour, dict) and tour.get('total'):
                    total_tour = limpiar_y_convertir_total(tour.get('total'))
                    total_cotizacion += total_tour
        
        elif tipo == 'traslados' and propuestas.get('traslados'):
            # Para traslados, tomar el total
            traslados = propuestas.get('traslados', {})
            if isinstance(traslados, dict) and traslados.get('total'):
                total_cotizacion = limpiar_y_convertir_total(traslados.get('total'))
        
        elif tipo == 'renta_autos' and propuestas.get('renta_autos'):
            # Para renta_autos, tomar el total
            renta_autos = propuestas.get('renta_autos', {})
            if isinstance(renta_autos, dict) and renta_autos.get('total'):
                total_cotizacion = limpiar_y_convertir_total(renta_autos.get('total'))
        
        # Si no se encontró total en propuestas, usar total_estimado como fallback
        if total_cotizacion == Decimal('0.00'):
            total_cotizacion = cot.total_estimado or Decimal('0.00')
        
        # Determinar servicios según el tipo de cotización
        servicios_mapeo = {
            'vuelos': ['Vuelo'],
            'hospedaje': ['Hospedaje'],
            'tours': ['Tour y Actividades'],
            'traslados': ['Traslado'],
            'renta_autos': ['Renta de Auto'],
        }
        servicios_seleccionados = servicios_mapeo.get(tipo, [])
        
        # Para paquetes, usar 'Paquete' como servicio principal (no 'Vuelo' y 'Hospedaje' por separado)
        if tipo == 'paquete':
            servicios_seleccionados = ['Paquete']  # Cambiar a 'Paquete' para que se guarde como 'PAQ'
            # Verificar si hay tours en el paquete
            paquete = propuestas.get('paquete', {})
            if isinstance(paquete, dict) and paquete.get('tours'):
                tours = paquete.get('tours', [])
                # Verificar si hay tours (puede ser lista o objeto único)
                if isinstance(tours, list) and len(tours) > 0:
                    # Verificar que al menos un tour tenga datos
                    tiene_tours = any(
                        isinstance(tour, dict) and (
                            tour.get('nombre') or 
                            tour.get('especificaciones') or 
                            tour.get('total') or 
                            tour.get('forma_pago')
                        ) for tour in tours
                    )
                    if tiene_tours:
                        servicios_seleccionados.append('Tour y Actividades')
                elif isinstance(tours, dict) and (tours.get('nombre') or tours.get('total')):
                    # Compatibilidad: si es un objeto único con datos
                    servicios_seleccionados.append('Tour y Actividades')
        
        # Preparar edades de menores en formato lista (nombre - edad) con saltos de línea
        edades_menores_formato = ''
        if cot.edades_menores and cot.edades_menores.strip():
            # Si las edades están en formato "5, 8, 12", convertirlas a formato lista con saltos de línea
            edades = [e.strip() for e in cot.edades_menores.split(',') if e.strip()]
            if edades:
                # Crear formato con saltos de línea: "Menor 1 - 5\nMenor 2 - 8\nMenor 3 - 12"
                edades_lista = [f"Menor {i+1} - {edad}" for i, edad in enumerate(edades)]
                edades_menores_formato = '\n'.join(edades_lista)
        
        # Guardar información de la cotización en la sesión para pre-llenar el formulario
        # Asegurar que el total se guarde correctamente
        logging.debug(f"Guardando total_cotizacion en sesión: {total_cotizacion} (índice vuelo: {opcion_vuelo_index})")
        request.session['cotizacion_convertir'] = {
            'cotizacion_id': cot.pk,
            'cotizacion_slug': cot.slug,
            'cliente_id': cot.cliente.pk,
            'fecha_inicio': str(cot.fecha_inicio) if cot.fecha_inicio else None,
            'fecha_fin': str(cot.fecha_fin) if cot.fecha_fin else None,
            'total_cotizacion': str(total_cotizacion),
            'opcion_vuelo_index': opcion_vuelo_index if opcion_vuelo_index else None,
            'opcion_hotel_index': opcion_hotel_index if opcion_hotel_index else None,
            'edades_menores': edades_menores_formato,
            'servicios_seleccionados': servicios_seleccionados,
            'tipo_cotizacion': tipo,
        }
        # Forzar guardado de la sesión
        request.session.modified = True
        
        # CAMBIO ANTIGRAVITY: Pasar datos explícitos por URL para garantizar que lleguen
        # Mejorar el redirect para incluir cualquier índice (vuelo u hotel)
        indice_final = opcion_vuelo_index or opcion_hotel_index or '0'
        url_destino = reverse('crear_venta') + f'?cotizacion={cot.slug}&total_b={str(total_cotizacion)}&idx_b={indice_final}'
        return redirect(url_destino)


# ------------------- ABONOS A PROVEEDOR (VENTAS INTERNACIONALES) -------------------

class SolicitarAbonoProveedorView(LoginRequiredMixin, UserPassesTestMixin, View):
    """Vista para que un vendedor solicite un abono a proveedor."""
    
    def test_func(self):
        """Solo vendedores, jefes y contadores pueden solicitar abonos."""
        user_rol = perm.get_user_role(self.request.user, self.request)
        return perm.is_vendedor(self.request.user, self.request) or perm.has_full_access(self.request.user, self.request) or perm.is_contador(self.request.user, self.request)
    
    def handle_no_permission(self):
        messages.error(self.request, "No tienes permiso para solicitar abonos a proveedores.")
        return redirect('dashboard')
    
    def post(self, request, pk):
        """Procesa la solicitud de abono a proveedor."""
        venta = get_object_or_404(VentaViaje, pk=pk)
        
        # Única fuente de verdad: venta.puede_solicitar_abonos_proveedor (no duplicar lógica)
        puede_solicitar = venta.puede_solicitar_abonos_proveedor
        
        if not puede_solicitar:
            messages.error(request, "No se pueden solicitar abonos a proveedores para esta venta.")
            return redirect('detalle_venta', pk=venta.pk, slug=venta.slug_safe)
        
        # Ya no requerimos que la venta tenga proveedor asignado, el vendedor lo escribe libremente
        
        from .forms import SolicitarAbonoProveedorForm
        form = SolicitarAbonoProveedorForm(request.POST, venta=venta, user=request.user)
        
        if form.is_valid():
            try:
                with transaction.atomic():
                    abono = form.save(commit=False)
                    abono.venta = venta
                    abono.solicitud_por = request.user
                    abono.estado = 'PENDIENTE'
                    abono.save()
                    
                    # Solo para ventas NAC: calcular monto_usd desde monto (MXN) si hay tipo de cambio
                    if venta.tipo_viaje != 'INT':
                        if abono.tipo_cambio_aplicado and abono.tipo_cambio_aplicado > 0 and abono.monto_usd is None:
                            abono.monto_usd = (abono.monto / abono.tipo_cambio_aplicado).quantize(Decimal('0.01'))
                            abono.save(update_fields=['monto_usd'])
                        elif venta.tipo_cambio and venta.tipo_cambio > 0 and abono.monto_usd is None:
                            abono.monto_usd = (abono.monto / venta.tipo_cambio).quantize(Decimal('0.01'))
                            abono.tipo_cambio_aplicado = venta.tipo_cambio
                            abono.save(update_fields=['monto_usd', 'tipo_cambio_aplicado'])
                    
                    # Notificar a contadores (INT: mostrar USD; NAC: MXN)
                    contadores = User.objects.filter(perfil__rol='CONTADOR')
                    if venta.tipo_viaje == 'INT' and abono.monto_usd is not None:
                        msg_monto = f"USD ${abono.monto_usd:,.2f}"
                    else:
                        msg_monto = f"${abono.monto:,.2f} MXN"
                    for contador in contadores:
                        Notificacion.objects.create(
                            usuario=contador,
                            tipo='SOLICITUD_ABONO_PROVEEDOR',
                            mensaje=f"Solicitud de abono a proveedor: {msg_monto} ({venta.folio or venta.pk})",
                            venta=venta,
                            abono_proveedor=abono
                        )
                    
                    if venta.tipo_viaje == 'INT' and abono.monto_usd is not None:
                        messages.success(request, f"Solicitud de abono a {abono.proveedor} por USD ${abono.monto_usd:,.2f} enviada correctamente.")
                    else:
                        messages.success(request, f"Solicitud de abono a {abono.proveedor} por ${abono.monto:,.2f} MXN enviada correctamente.")
            except Exception as e:
                messages.error(request, f"Error al solicitar el abono: {str(e)}")
                logger.exception(f"Error al solicitar abono a proveedor: {str(e)}")
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f"{field}: {error}")
        
        return redirect(reverse('detalle_venta', kwargs={'pk': venta.pk, 'slug': venta.slug_safe}) + '?tab=logistica')


class AprobarAbonoProveedorView(LoginRequiredMixin, UserPassesTestMixin, View):
    """Vista para que un contador apruebe un abono a proveedor."""
    
    def test_func(self):
        """Solo contadores pueden aprobar abonos."""
        return perm.is_contador(self.request.user, self.request)
    
    def handle_no_permission(self):
        messages.error(self.request, "No tienes permiso para aprobar abonos a proveedores.")
        return redirect('dashboard')
    
    def post(self, request, abono_id):
        """Aprueba un abono a proveedor."""
        abono = get_object_or_404(AbonoProveedor, pk=abono_id)
        
        if abono.estado != 'PENDIENTE':
            messages.error(request, "Este abono ya no está pendiente de aprobación.")
            return redirect(reverse('detalle_venta', kwargs={'pk': abono.venta.pk, 'slug': abono.venta.slug_safe}) + '?tab=logistica')
        
        try:
            with transaction.atomic():
                abono.estado = 'APROBADO'
                abono.aprobado_por = request.user
                abono.fecha_aprobacion = timezone.now()
                abono.save()
                
                # Notificar al vendedor que solicitó
                if abono.solicitud_por:
                    Notificacion.objects.create(
                        usuario=abono.solicitud_por,
                        tipo='ABONO_PROVEEDOR_APROBADO',
                        mensaje=f"Abono a {abono.proveedor} por {'USD ' if abono.venta.tipo_viaje == 'INT' else ''}${abono.monto:,.2f}{' MXN' if abono.venta.tipo_viaje != 'INT' else ''} aprobado (Venta #{abono.venta.folio or abono.venta.pk})",
                        venta=abono.venta,
                        abono_proveedor=abono
                    )
                
                contadores = User.objects.filter(perfil__rol='CONTADOR')
                for contador in contadores:
                    Notificacion.objects.create(
                        usuario=contador,
                        tipo='ABONO_PROVEEDOR_APROBADO',
                        mensaje=f"Abono a {abono.proveedor} por {'USD ' if abono.venta.tipo_viaje == 'INT' else ''}${abono.monto:,.2f}{' MXN' if abono.venta.tipo_viaje != 'INT' else ''} aprobado, pendiente de confirmar (Venta #{abono.venta.folio or abono.venta.pk})",
                        venta=abono.venta,
                        abono_proveedor=abono
                    )
                
                messages.success(request, f"Abono a {abono.proveedor} aprobado correctamente.")
        except Exception as e:
            messages.error(request, f"Error al aprobar el abono: {str(e)}")
            logger.exception(f"Error al aprobar abono a proveedor: {str(e)}")
        
        return redirect(reverse('detalle_venta', kwargs={'pk': abono.venta.pk, 'slug': abono.venta.slug_safe}) + '?tab=logistica')


class ConfirmarAbonoProveedorView(LoginRequiredMixin, UserPassesTestMixin, View):
    """Vista para que un contador confirme un abono a proveedor con comprobante."""
    
    def test_func(self):
        """Solo contadores pueden confirmar abonos."""
        return perm.is_contador(self.request.user, self.request)
    
    def handle_no_permission(self):
        messages.error(self.request, "No tienes permiso para confirmar abonos a proveedores.")
        return redirect('dashboard')
    
    def post(self, request, abono_id):
        """Confirma un abono a proveedor con comprobante, o reemplaza el comprobante si ya está COMPLETADO."""
        abono = get_object_or_404(AbonoProveedor, pk=abono_id)
        if abono.estado not in ('APROBADO', 'COMPLETADO'):
            messages.error(request, "Este abono debe estar aprobado antes de confirmarlo.")
            return redirect(reverse('detalle_venta', kwargs={'pk': abono.venta.pk, 'slug': abono.venta.slug_safe}) + '?tab=logistica')
        # COMPLETADO: solo reemplazar comprobante; APROBADO: confirmar y marcar COMPLETADO
        es_reemplazo = abono.estado == 'COMPLETADO'
        
        from .forms import ConfirmarAbonoProveedorForm
        form = ConfirmarAbonoProveedorForm(request.POST, request.FILES, instance=abono)
        
        if form.is_valid():
            try:
                with transaction.atomic():
                    abono = form.save(commit=False)
                    if not es_reemplazo:
                        abono.estado = 'COMPLETADO'
                        abono.confirmado_por = request.user
                        abono.fecha_confirmacion = timezone.now()
                        if abono.nota_confirmacion:
                            abono.nota_solicitud = abono.nota_confirmacion
                        abono.save()
                        if abono.solicitud_por:
                            Notificacion.objects.create(
                                usuario=abono.solicitud_por,
                                tipo='ABONO_PROVEEDOR_COMPLETADO',
                                mensaje=f"Abono a {abono.proveedor} por {'USD ' if abono.venta.tipo_viaje == 'INT' else ''}${abono.monto:,.2f}{' MXN' if abono.venta.tipo_viaje != 'INT' else ''} completado (Venta #{abono.venta.folio or abono.venta.pk})",
                                venta=abono.venta,
                                abono_proveedor=abono
                            )
                        messages.success(request, f"Abono a {abono.proveedor} confirmado con comprobante.")
                    else:
                        abono.save(update_fields=['comprobante', 'nota_confirmacion'])
                        if abono.nota_confirmacion:
                            abono.nota_solicitud = abono.nota_confirmacion
                            abono.save(update_fields=['nota_solicitud'])
                        messages.success(request, "Comprobante actualizado correctamente.")
            except Exception as e:
                messages.error(request, f"Error al guardar: {str(e)}")
                logger.exception(f"Error al confirmar/reemplazar abono a proveedor: {str(e)}")
        else:
            messages.error(request, "Por favor, corrige los errores en el formulario.")
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f"{field}: {error}")
        
        return redirect(reverse('detalle_venta', kwargs={'pk': abono.venta.pk, 'slug': abono.venta.slug_safe}) + '?tab=logistica')


class EliminarComprobanteAbonoProveedorView(LoginRequiredMixin, UserPassesTestMixin, View):
    """Vista para que el contador elimine el comprobante de un abono a proveedor (COMPLETADO) y pueda subir otro."""
    
    def test_func(self):
        return perm.is_contador(self.request.user, self.request)
    
    def handle_no_permission(self):
        messages.error(self.request, "No tienes permiso para eliminar comprobantes de abonos a proveedores.")
        return redirect('dashboard')
    
    def post(self, request, abono_id):
        abono = get_object_or_404(AbonoProveedor, pk=abono_id)
        if abono.estado != 'COMPLETADO':
            messages.error(request, "Solo se puede eliminar el comprobante de un abono ya completado.")
            return redirect(reverse('detalle_venta', kwargs={'pk': abono.venta.pk, 'slug': abono.venta.slug_safe}) + '?tab=logistica')
        if not abono.comprobante:
            messages.info(request, "Este abono no tiene comprobante.")
            return redirect(reverse('detalle_venta', kwargs={'pk': abono.venta.pk, 'slug': abono.venta.slug_safe}) + '?tab=logistica')
        try:
            abono.comprobante.delete(save=False)
            abono.comprobante = None
            abono.save(update_fields=['comprobante'])
            messages.success(request, "Comprobante eliminado. Puedes subir el comprobante correcto.")
        except Exception as e:
            messages.error(request, f"Error al eliminar el comprobante: {str(e)}")
            logger.exception(f"Error al eliminar comprobante abono proveedor: {e}")
        return redirect(reverse('detalle_venta', kwargs={'pk': abono.venta.pk, 'slug': abono.venta.slug_safe}) + '?tab=logistica')


class CancelarAbonoProveedorView(LoginRequiredMixin, UserPassesTestMixin, View):
    """Vista para cancelar un abono a proveedor (solo JEFE puede modificar)."""
    
    def test_func(self):
        """Solo jefes pueden cancelar abonos."""
        user_rol = perm.get_user_role(self.request.user, self.request)
        return perm.has_full_access(self.request.user)
    
    def handle_no_permission(self):
        messages.error(self.request, "Solo los jefes pueden cancelar abonos a proveedores.")
        return redirect('dashboard')
    
    def post(self, request, abono_id):
        """Cancela un abono a proveedor."""
        abono = get_object_or_404(AbonoProveedor, pk=abono_id)
        motivo_cancelacion = request.POST.get('motivo_cancelacion', '')
        
        if abono.estado == 'COMPLETADO':
            messages.error(request, "No se puede cancelar un abono que ya está completado.")
            return redirect(reverse('detalle_venta', kwargs={'pk': abono.venta.pk, 'slug': abono.venta.slug_safe}) + '?tab=logistica')
        
        try:
            with transaction.atomic():
                abono.estado = 'CANCELADO'
                abono.cancelado_por = request.user
                abono.fecha_cancelacion = timezone.now()
                abono.motivo_cancelacion = motivo_cancelacion
                abono.save()
                
                # Notificar al vendedor que solicitó
                if abono.solicitud_por:
                    Notificacion.objects.create(
                        usuario=abono.solicitud_por,
                        tipo='ABONO_PROVEEDOR_CANCELADO',
                        mensaje=f"Abono a {abono.proveedor} por ${abono.monto:,.2f} MXN cancelado (Venta #{abono.venta.folio or abono.venta.pk})",
                        venta=abono.venta,
                        abono_proveedor=abono
                    )
                
                messages.success(request, f"Abono a {abono.proveedor} cancelado correctamente.")
        except Exception as e:
            messages.error(request, f"Error al cancelar el abono: {str(e)}")
            logger.exception(f"Error al cancelar abono a proveedor: {str(e)}")
        
        return redirect(reverse('detalle_venta', kwargs={'pk': abono.venta.pk, 'slug': abono.venta.slug_safe}) + '?tab=logistica')


class ListaAbonosProveedorView(LoginRequiredMixin, UserPassesTestMixin, TemplateView):
    """Vista para que un contador vea todas las solicitudes de abonos a proveedores pendientes."""
    template_name = 'ventas/lista_abonos_proveedor.html'
    
    def test_func(self):
        """Solo contadores y jefes pueden ver la lista de solicitudes."""
        user_rol = perm.get_user_role(self.request.user, self.request)
        return perm.is_contador(self.request.user, self.request) or perm.has_full_access(self.request.user, self.request)
    
    def handle_no_permission(self):
        messages.error(self.request, "No tienes permiso para ver las solicitudes de abonos a proveedores.")
        return redirect('dashboard')
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Filtrar solo solicitudes pendientes y aprobadas
        context['abonos_pendientes'] = AbonoProveedor.objects.filter(
            estado__in=['PENDIENTE', 'APROBADO']
        ).select_related('venta', 'solicitud_por').order_by('-fecha_solicitud')
        context['user_rol'] = perm.get_user_role(self.request.user, self.request)
        return context